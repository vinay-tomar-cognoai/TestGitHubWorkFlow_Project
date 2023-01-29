import re
import os
import shutil
import pytz
import json
import uuid
import random
import logging
import hashlib
import mimetypes
import threading
import magic
import xlwt
import pdfkit
import sys

from django.core.cache import cache
from PIL import Image
from time import sleep
from xlwt import Workbook
from docx import Document
from moviepy.editor import *
from zipfile import ZipFile
from os.path import basename
from EasyChat import settings
from django.db.models import Q
from django.utils import timezone
from django.db.models import Count
from email.mime.text import MIMEText
from datetime import datetime, timedelta
from django.shortcuts import HttpResponse
from django.utils.encoding import smart_str
from email.mime.multipart import MIMEMultipart
from LiveChatApp.utils_custom_encryption import *
from LiveChatApp.utils_validation import LiveChatInputValidation
from django.core.paginator import Page, Paginator, EmptyPage, PageNotAnInteger
from AuditTrailApp.utils import add_audit_trail
from LiveChatApp.constants import *
from DeveloperConsoleApp.utils import get_developer_console_settings, send_email_to_customer_via_awsses
from LiveChatApp.utils_payload import get_missed_chat_payload, get_offline_chat_payload, get_abandoned_chat_payload, get_agent_login_logout_payload, get_agent_not_ready_payload, get_chat_history_nps_payload, get_chat_history_payload, get_agent_performance_report_payload, get_video_call_history_payload, get_voice_call_history_payload, get_cobrowsing_history_payload


logger = logging.getLogger(__name__)


"""
function: get_random_captcha_image
It return random captcha image
"""


def get_random_captcha_image():
    files = os.listdir('EasyChatApp/static/EasyChatApp/captcha_images')
    index = random.randrange(0, len(files))
    return files[index]


"""
function: check_and_add_admin_config
input params:
    livechat_user: LiveChat user object.
output:

    This function creates livechat admin config object for new livechat admin login.

"""


def check_and_add_admin_config(livechat_user, LiveChatAdminConfig):
    logger.info("Inside check and add admin config",
                extra={'AppName': 'LiveChat'})
    try:
        if livechat_user.status == "1":
            try:
                LiveChatAdminConfig.objects.get(admin=livechat_user)
            except Exception:
                logger.info("Admin config does not exists, creating new ", extra={
                            'AppName': 'LiveChat'})
                LiveChatAdminConfig.objects.create(admin=livechat_user)

        logger.info("Exist successfully from check_and_add_admin_config.", extra={
                    'AppName': 'LiveChat'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_and_add_admin_config: %s at %s", str(e), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: add_supervisor
input params:
    user:
    supervisor_pk:
    current_user:
output:

    This function adds supervisor for a perticular agent. In case, if supervisor pk "-1", then it current user will be the supervisor for agent.

"""


def add_supervisor(user, supervisor_pk, current_user, LiveChatUser):
    try:
        if supervisor_pk != "-1":
            supervisor = LiveChatUser.objects.get(pk=int(supervisor_pk))
            supervisor.agents.add(user)
            supervisor.save()
            if supervisor != current_user:
                current_user.agents.remove(user)
                current_user.save()
        else:
            current_user.agents.add(user)
            current_user.save()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("add_supervisor: %s at %s", exc, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: save_session_details
input params:
    user:
output:
    1. This function is use to add entry in LiveChatSessionManagement based on user.
"""


def save_session_details(user, selected_status, LiveChatSessionManagement, LiveChatAgentNotReady, LiveChatUser, LiveChatAdminConfig, LiveChatConfig, Bot):
    try:
        sessions_obj = LiveChatSessionManagement.objects.filter(
            user=user, session_completed=False)[0]
        if sessions_obj.user.is_online:
            diff = timezone.now() - sessions_obj.session_ends_at
            sessions_obj.online_time += diff.seconds
            sessions_obj.session_ends_at = timezone.now()
            if sessions_obj.is_idle:
                time_diff = datetime.now(
                    timezone.utc) - sessions_obj.last_idle_time
                sessions_obj.idle_time += time_diff.seconds
            sessions_obj.is_idle = False
            agent_not_ready_obj = LiveChatAgentNotReady.objects.create(
                user=user, session_id=str(uuid.uuid4()), reason_for_offline=selected_status)
            sessions_obj.agent_not_ready.add(agent_not_ready_obj)
            sessions_obj.save()

            # he'll mark himself offline only if he's online
            if selected_status == "0":
                sessions_obj.time_marked_stop_interaction = timezone.now()
                sessions_obj.save()
            send_event_for_agent_not_ready(user, agent_not_ready_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot)
        else:
            diff = timezone.now() - sessions_obj.session_ends_at
            sessions_obj.offline_time += diff.seconds
            sessions_obj.session_ends_at = timezone.now()

            if user.current_status == "0":
                diff = timezone.now() - sessions_obj.time_marked_stop_interaction
                sessions_obj.stop_interaction_time += diff.seconds
                agent_not_ready_obj = sessions_obj.agent_not_ready.all().order_by(
                    '-not_ready_starts_at')[0]
                agent_not_ready_obj.stop_interaction_duration = diff.seconds
                sessions_obj.is_idle = False
                if selected_status != "6":
                    agent_not_ready_obj.reason_for_offline = selected_status
                agent_not_ready_obj.save()
            else:
                sessions_obj.is_idle = True
                sessions_obj.last_idle_time = datetime.now()

            sessions_obj.save()
            agent_not_ready_obj = sessions_obj.agent_not_ready.all().order_by(
                '-not_ready_starts_at')[0]
            agent_not_ready_obj.not_ready_ends_at = timezone.now()
            if selected_status == "6":
                agent_not_ready_obj.is_expired = True
            agent_not_ready_obj.save()
            send_event_for_agent_not_ready(user, agent_not_ready_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot, True)
        send_event_for_login_logout(user, sessions_obj, LiveChatUser, LiveChatAdminConfig, LiveChatConfig, Bot, True)
        send_event_for_performance_report(user, sessions_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot, True)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside save_session_details: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: save_audit_trail_data
input params:
    action:
    user:
output:
    1. This function is use to add entry in audit trail based on the action of user.
"""


def save_audit_trail_data(action, user, LiveChatAuditTrail):
    try:
        LiveChatAuditTrail.objects.create(action=action, user=user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside save_audit_trail_data: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: get_livechat_request_packet_to_channel
input params:
    session_id:
    type:
    message:
    path:
    channel:
    bot_id:
output:
    1. This function return the sample request packet sent during whatsapp and facebook integration.
"""


def get_livechat_request_packet_to_channel(session_id, type, message, path, channel, bot_id, sender_name):
    try:
        response = {
            "session_id": session_id,
            "type": type,
            "text_message": message,
            "path": path,
            "channel": channel,
            "bot_id": bot_id
        }
        if sender_name != "" and sender_name != None:
            response["agent_name"] = str(sender_name)
        return json.dumps(response)
    except Exception as err:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_livechat_request_packet_to_channel: %s at %s",
                     str(err), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: check_for_holiday
input params:
    bot_obj:
output:
    1. This function checks that request raising date is working or not.
"""


def check_for_holiday(bot_obj, LiveChatCalender, LiveChatUser):
    response = {}
    try:
        livechat_admin_obj = LiveChatUser.objects.get(
            user=bot_obj.created_by, is_deleted=False)
        if LiveChatCalender.objects.filter(event_type="2", event_date__date=timezone.now().date(), created_by=livechat_admin_obj).count():
            response["message"] = LiveChatCalender.objects.filter(
                event_type="2", event_date__date=timezone.now().date(), created_by=livechat_admin_obj)[0].auto_response
            response["status_code"] = "300"
            response["assigned_agent"] = "holiday"
            response["assigned_agent_username"] = "None"
            return True, response

        return False, response
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside check_for_holiday: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        return False, response


"""
function: check_for_non_working_hour
input params:
    bot_obj:
output:
    1. This function checks that request raising time comes within working hour or not.
"""


def check_for_non_working_hour(bot_obj, LiveChatCalender, LiveChatConfig, LiveChatUser):
    response = {}
    response[
        "message"] = "Thank you for writing to us, will respond you in next business day."
    try:
        try:
            livechat_config = LiveChatConfig.objects.get(bot=bot_obj)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Livechat config does not exists for bot obj: %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
            livechat_config = LiveChatConfig.objects.create(bot=bot_obj)
        livechat_admin_obj = LiveChatUser.objects.get(
            user=bot_obj.created_by, is_deleted=False)
        current_date = datetime.now().date()
        current_time = datetime.now().time()
        if LiveChatCalender.objects.filter(event_type="1", event_date__date=current_date, start_time__lte=current_time, end_time__gte=current_time, created_by=livechat_admin_obj).count():
            return False, response

        response["message"] = livechat_config.auto_bot_response
        response["status_code"] = "300"
        response["assigned_agent"] = "non_working_hour"
        response["assigned_agent_username"] = "None"
        return True, response
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside check_for_non_working_hour: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        return False, response


"""
function: get_number_of_day
input params:
    year:
    month:
expected output:
    1. Provides number of days in a month of a perticular year.
"""


def get_number_of_day(year, month):
    year = int(year)
    month = int(month)
    leap = 0
    if year % 400 == 0:
        leap = 1
    elif year % 100 == 0:
        leap = 0
    elif year % 4 == 0:
        leap = 1
    if month == 2:
        return 28 + leap
    list = [1, 3, 5, 7, 8, 10, 12]
    if month in list:
        return 31
    return 30


"""
function: check_and_update_based_on_event_type
input params:
    user:
    calender_obj:
    selected_event:
    start_time:
    end_time:
    description: description of event
    auto_response: auto response will be given to livechat customer on this day.
expected output:
    1. If user is admin, then this event will be updated based on event type.
"""


def check_and_update_based_on_event_type(user, calender_obj, selected_event, start_time, end_time, description, auto_response):
    try:
        if user.status == "1":
            if str(selected_event) == "1":
                calender_obj.start_time = start_time
                calender_obj.end_time = end_time
                calender_obj.event_type = "1"
            else:
                calender_obj.description = description
                calender_obj.auto_response = auto_response
                calender_obj.event_type = "2"
            calender_obj.modified_by = user
            calender_obj.modified_date = timezone.now()
            calender_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside check_and_update_based_on_event_type: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_month_list():
    return [{"key": "1", "value": "Jan"}, {"key": "2", "value": "Feb"}, {"key": "3", "value": "Mar"}, {"key": "4", "value": "Apr"}, {"key": "5", "value": "May"}, {"key": "6", "value": "Jun"}, {"key": "7", "value": "Jul"}, {"key": "8", "value": "Aug"}, {"key": "9", "value": "Sep"}, {"key": "10", "value": "Oct"}, {"key": "11", "value": "Nov"}, {"key": "12", "value": "Dec"}]


def get_year_list():
    return [{"key": "2020", "value": "2020"}, {"key": "2021", "value": "2021"}, {"key": "2022", "value": "2022"}, {"key": "2023", "value": "2023"}, {"key": "2024", "value": "2024"}, {"key": "2025", "value": "2025"}, {"key": "2026", "value": "2026"}, {"key": "2027", "value": "2027"}, {"key": "2028", "value": "2028"}, {"key": "2029", "value": "2029"}, {"key": "2030", "value": "2030"}, {"key": "2031", "value": "2031"}]


def get_priority_list():
    return [{"key": "1", "value": "1"}, {"key": "2", "value": "2"}, {"key": "3", "value": "3"}, {"key": "4", "value": "4"}, {"key": "5", "value": "5"}]


"""
function: get_allowed_livechat_user
input params:
    livechat_user:
    selected_category:
expected output:
    1. It is a helper function which two arguments livechat_user and category. It returns the all agents lying under this livechat user.
"""


def get_allowed_livechat_user(livechat_user, selected_category, bot_obj, LiveChatUser, LiveChatCategory):
    try:
        parent_user = LiveChatUser.objects.filter(
            agents__user=livechat_user.user)[0]
    except Exception:
        parent_user = livechat_user
    if int(selected_category) == -1:
        return parent_user.agents.all().filter(is_deleted=False, status="3", bots__in=[bot_obj], is_online=True)

    return parent_user.agents.all().filter(is_deleted=False, status="3", bots__in=[bot_obj], category=LiveChatCategory.objects.get(pk=int(selected_category)), is_online=True)


"""
function: get_chat_status
input params:
    chat_status
expected output:
    It is a helper function which takes a variable and return some value based on some conditions
"""


def get_chat_status(chat_status):
    if chat_status == None or chat_status == "0" or chat_status == 'None':
        chat_status = None
    elif chat_status == "1":
        chat_status = False
    else:
        chat_status = True

    return chat_status


"""
function: get_chat_duration_start_end
input params:
    chat_duration
expected output:
    It is a helper function which takes a variable and return some value based on some conditions
"""


"""
function: get_audit_objects
input params:
    query_user_obj: user objects selected during applying filter
    agent_username: agent username selected during applying filter
    chat_status: Online/Offline
    chat_duration:
    datetime_start:
    datetime_end
expected output:
    This function is used to find LiveChatCustomer objects for a given interval with given chat duration.
"""


def get_previous_session_messages(cust_obj_pk, client_id, LiveChatMISDashboard, LiveChatCustomer):
    try:
        customer_objs = LiveChatCustomer.objects.filter(
            client_id=client_id).exclude(Q(agent_id__isnull=True) & Q(followup_assignment=False)).order_by('joined_date')
        current_customer_obj = LiveChatCustomer.objects.get(pk=cust_obj_pk)
        new_index = list(customer_objs).index(current_customer_obj)
        new_customer_object = customer_objs[new_index - 1]
        new_index = new_index - 1
        messages_list_obj = []
        fetched_obj = customer_objs[new_index]
        if new_index > -1:
            messages_list_obj = LiveChatMISDashboard.objects.filter(
                livechat_customer=fetched_obj).order_by('message_time', 'pk')
        else:
            messages_list_obj = []
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("error in get_previous_session_messages %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChatApp'})
        messages_list_obj = []
    return messages_list_obj, str(new_customer_object)


def get_audit_objects(query_user_obj, chat_status, datetime_start, datetime_end, livechat_customer_list, channel_name, channel_obj, selected_category_pk, category_obj, chat_termination, livechat_followup_cust_objs):
    audit_obj_list = []
    try:
        livechat_customer_list = livechat_customer_list.filter(
            request_raised_date__range=[datetime_start, datetime_end])
        if channel_name != None and channel_name != "All":
            livechat_customer_list = livechat_customer_list.filter(
                channel=channel_obj)

        if str(selected_category_pk) != "0" and selected_category_pk != None:
            livechat_customer_list = livechat_customer_list.filter(
                category=category_obj)

        if chat_termination != None and chat_termination != "All":
            livechat_customer_list = livechat_customer_list.filter(
                chat_ended_by=chat_termination)            

        chat_status = get_chat_status(chat_status)

        if chat_status == None:
            livechat_followup_cust_objs = livechat_followup_cust_objs.filter(
                agent_id__in=query_user_obj, is_whatsapp_conversation_reinitiated=False)
            audit_obj_list = livechat_customer_list.filter(
                Q(agent_id__in=query_user_obj) | Q(pk__in=livechat_followup_cust_objs)).order_by('-joined_date')
        else:
            livechat_followup_cust_objs = livechat_followup_cust_objs.filter(
                livechat_customer__is_session_exp=chat_status, agent_id__in=query_user_obj, is_whatsapp_conversation_reinitiated=False)
            audit_obj_list = livechat_customer_list.filter(
                Q(is_session_exp=chat_status, agent_id__in=query_user_obj) | Q(pk__in=livechat_followup_cust_objs)).order_by('-joined_date')

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_audit_objects: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return audit_obj_list


def get_call_data_history_objects(agent_username, query_user_obj, datetime_end, datetime_start, agent_obj_list, call_type, LiveChatVoIPData):

    call_object_list = []
    try:
        if call_type == 'video_call':
            call_object_list = LiveChatVoIPData.objects.filter(request_datetime__range=[datetime_start, datetime_end],
                                                               call_type=call_type, agent__in=agent_obj_list, is_started=True, is_interrupted=False).order_by('-pk')
        else:
            call_object_list = LiveChatVoIPData.objects.filter(request_datetime__range=[datetime_start, datetime_end],
                                                               call_type__in=['pip', 'new_tab'], agent__in=agent_obj_list, is_started=True, is_interrupted=False).order_by('-pk')

        if agent_username != "All" and query_user_obj != None:
            call_object_list = call_object_list.filter(agent=query_user_obj)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_call_data_history_objects: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return call_object_list


"""
function: DecryptVariable
input params:
    encrypted data
expected output:
    decrypted data

This function is used to decrypt all the encrypted data.
"""


def DecryptVariable(data):
    try:
        custom_encrypt_obj = CustomEncrypt()
        data = custom_encrypt_obj.decrypt(data)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside DecryptVariable: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return data


"""
function: get_miniseconds_datetime
input params:
    datetime_obj
expected output:
    This function is used to find total number of seconds for a perticular datetime object.
"""


def get_miniseconds_datetime(datetime_obj):
    total_seconds = 0
    try:
        total_seconds = round(datetime_obj.timestamp() * 1000)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_miniseconds_datetime: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return total_seconds


"""

function: get_time
input params:
    datetime object
output:
    returns a string

This function return the time in same timzone formate as defined in settings.py. It returns
12 hour time cycle.
"""


def get_time(datetime_obj):
    import pytz
    tz = pytz.timezone(settings.TIME_ZONE)
    tz = datetime_obj.astimezone(tz)

    if tz.hour < 12:
        flag = "AM"
        hour1 = tz.hour
    elif tz.hour == 12:
        hour1 = tz.hour
        flag = "PM"
    else:
        hour1 = tz.hour - 12
        flag = "PM"

    if hour1 <= 9:
        hour2 = "0" + str(hour1)
    else:
        hour2 = str(hour1)

    if tz.minute <= 9:
        minute1 = "0" + str(tz.minute)
    else:
        minute1 = str(tz.minute)

    time1 = str(hour2) + ":" + str(minute1) + " " + flag

    return time1


"""

function: get_date
input params:
    datetime object
output:
    returns a string

This function returns 'Yesterday' for yesterday's date, day name of week for date within
that week, date for date before the week and time for today.
"""


def get_date(datetime_obj):
    import pytz
    tz = pytz.timezone(settings.TIME_ZONE)
    tz = datetime_obj.astimezone(tz)

    days_diff = (datetime.today().date() - tz.date()).days

    if days_diff == 1:
        return "Yesterday"

    if days_diff > 1 and days_diff < 8:
        return tz.date().strftime('%A')

    if days_diff >= 8:
        return tz.date().strftime('%d/%m/%Y')

    return get_time(datetime_obj)


"""
function: update_message_history_till_now
input params:
    LiveChatCustomer object

This function create LiveChatMISDashboard objects of conversation between chatbot and user.
"""


def update_message_history_till_now(livechat_cust_obj, LiveChatMISDashboard, MISDashboard):
    try:
        validation_obj = LiveChatInputValidation()
        if not livechat_cust_obj.is_chat_loaded:
            livechat_cust_obj.is_chat_loaded = True
            livechat_cust_obj.save()
            max_message_count = 10
            easychat_mis_objs = MISDashboard.objects.filter(
                creation_date=datetime.today().date(), date__lt=livechat_cust_obj.joined_date, user_id=livechat_cust_obj.easychat_user_id, bot=livechat_cust_obj.bot).order_by('-date')

            for easychat_mis_obj in easychat_mis_objs:
                LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                                    sender="Customer",
                                                    text_message=easychat_mis_obj.get_message_received(),
                                                    sender_name=livechat_cust_obj.get_complete_username(),
                                                    message_time=easychat_mis_obj.date,
                                                    attachment_file_name="",
                                                    attachment_file_path="",
                                                    is_copied_from_easychat=True)

                bot_response = easychat_mis_obj.get_bot_response_with_html_tags_intact()
                bot_response = validation_obj.replace_break_tags(bot_response)
                bot_response = validation_obj.remo_html_from_string(bot_response)
                if bot_response.find(RESPONSE_SENTENCE_SEPARATOR) != -1:
                    bot_response = bot_response.split(RESPONSE_SENTENCE_SEPARATOR)
                    for text in bot_response:
                        LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                                                    sender="Bot",
                                                                    text_message=text,
                                                                    sender_name="bot",
                                                                    message_time=easychat_mis_obj.date,
                                                                    attachment_file_name="",
                                                                    attachment_file_path="",
                                                                    is_copied_from_easychat=True)
                else:
                    LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                                        sender="Bot",
                                                        text_message=bot_response,
                                                        sender_name="bot",
                                                        message_time=easychat_mis_obj.date,
                                                        attachment_file_name="",
                                                        attachment_file_path="",
                                                        is_copied_from_easychat=True)
                max_message_count -= 2
                if max_message_count <= 0:
                    break
                                                
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_message_history_till_now: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def update_followup_lead_message_history(livechat_cust_obj, livechat_followup_cust_obj, LiveChatMISDashboard, MISDashboard):
    try:
        if not livechat_cust_obj.is_chat_loaded:
            easychat_mis_objs = MISDashboard.objects.filter(
                user_id=livechat_cust_obj.easychat_user_id, bot=livechat_cust_obj.bot).order_by('-date')

            for easychat_mis_obj in easychat_mis_objs:
                LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                                    sender="Customer",
                                                    text_message=easychat_mis_obj.get_message_received(),
                                                    sender_name=livechat_cust_obj.get_complete_username(),
                                                    message_time=easychat_mis_obj.date,
                                                    attachment_file_name="",
                                                    attachment_file_path="")
                LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                                    sender="Bot",
                                                    text_message=easychat_mis_obj.get_bot_response(),
                                                    sender_name="bot",
                                                    message_time=easychat_mis_obj.date,
                                                    attachment_file_name="",
                                                    attachment_file_path="")

            text_message = "Customer Name: " + str(livechat_cust_obj.username) + " | Agent Name: " + str(livechat_followup_cust_obj.agent_id.user.first_name) + " " + str(
                livechat_followup_cust_obj.agent_id.user.last_name) + "(" + str(livechat_followup_cust_obj.agent_id.user.username) + ")"
            LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                                sender="System",
                                                text_message=text_message,
                                                sender_name="system",
                                                message_time=timezone.now(),
                                                attachment_file_name="",
                                                attachment_file_path="")

            livechat_cust_obj.is_chat_loaded = True
            livechat_cust_obj.save(update_fields=['is_chat_loaded'])
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_followup_lead_message_history: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


############## User Validity   ###############

"""

function: is_agent_live
input params:
    agent: LiveChatUser object
output:
    returns True/False

This function returns True if last updated time of user is 30 seconds before, False otherwise.

"""


def is_agent_live(agent):
    try:
        import pytz
        import os
        tz = pytz.timezone(settings.TIME_ZONE)
        last_updated_time = agent.last_updated_time
        if last_updated_time == None:
            return False
        last_updated_time = last_updated_time.astimezone(tz)
        current_time = datetime.now().astimezone(tz)

        available_time = (current_time - last_updated_time).total_seconds()

        if available_time > 75:
            agent.is_online = False
            agent.save()
            return False

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("is_agent_live: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    logger.info("Into is_agent_live: User is valid",
                extra={'AppName': 'LiveChat'})

    return True


"""
function: get_sender_name
input params:
    customer_obj: LiveChatCustomer object
output:
    returns message sender name

This function returns name of last message sender.
"""


def get_sender_name(customer_obj, LiveChatMISDashboard):
    logger.info("Into get_sender_name....", extra={'AppName': 'LiveChat'})
    sender = "None"
    try:
        mis_objs = LiveChatMISDashboard.objects.filter(
            livechat_customer=customer_obj).order_by("-pk")

        if mis_objs:
            sender = mis_objs.first().sender_name

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_sender_name: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return sender


"""

function: get_one_previous_message
input params:
    customer_obj: LiveChatCustomer object
    LiveChatMISDashboard:
output:
    returns previous message

This function returns last message sent by user/agent

"""


def get_one_previous_message(customer_obj, admin_config, LiveChatMISDashboard):

    logger.info("Into get_one_previous_message...",
                extra={'AppName': 'LiveChat'})
    message = {
        "text_message": ""
    }

    try:
        mis_obj = LiveChatMISDashboard.objects.filter(
            livechat_customer=customer_obj, is_video_call_message=False, is_voice_call_message=False, is_cobrowsing_message=False, is_file_not_support_message=False).order_by('-pk')

        if not mis_obj:
            return message

        mis_obj = mis_obj.first()

        if customer_obj.chat_ended_by == "Customer":
            message["text_message"] = "Customer left the chat"
        elif customer_obj.chat_ended_by == "System":
            message["text_message"] = mis_obj.text_message
                
        else:
            message["sender"] = mis_obj.sender
            message["message_time"] = get_miniseconds_datetime(
                mis_obj.message_time)
            message["text_message"] = mis_obj.text_message
            message["is_guest_agent_message"] = mis_obj.is_guest_agent_message
            if mis_obj.attachment_file_name != "":
                message["attachment_name"] = mis_obj.attachment_file_name
                message["is_attachment"] = "True"
            else:
                message["is_attachment"] = "False"
        message["message_id"] = str(mis_obj.message_id)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_one_previous_message: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return message


"""

function: get_livechat_category_object
input params:
    category_pk: pk of category object
output:
    returns category object

This function returns category object. If category pk is -1, then it creates a category with title "Others" and returns this object.

"""


def get_livechat_category_object(category_pk, bot_obj, LiveChatCategory):
    category_obj = None
    try:
        if int(category_pk) == -1 and LiveChatCategory.objects.filter(title="others", bot=bot_obj, is_deleted=False).count():
            category_obj = LiveChatCategory.objects.filter(
                title="others", bot=bot_obj, is_deleted=False)[0]
        elif int(category_pk) == -1:
            category_obj = LiveChatCategory.objects.create(
                title="others", bot=bot_obj)
        else:
            category_obj = LiveChatCategory.objects.get(pk=int(category_pk))

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_livechat_category_object: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return category_obj


"""
function: get_livechat_category_using_name
input params:
    category_pk: name of category object
output:
    returns category object

This function returns category object.

"""


def get_livechat_category_using_name(category_name, bot_obj, LiveChatCategory):
    category_obj = None
    try:
        category_obj = LiveChatCategory.objects.filter(
            title=category_name, bot=bot_obj, is_deleted=False)

        if category_obj:
            return category_obj.first()
        else:
            category_obj = LiveChatCategory.objects.filter(
                title="others", bot=bot_obj, is_deleted=False)

            if category_obj:
                return category_obj.first()
            else:
                category_obj = LiveChatCategory.objects.create(
                    title="others", bot=bot_obj)

                return category_obj

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_livechat_category_using_name: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return category_obj


"""

function: get_chat_duration_list
input params:

output:
    This function returns a dictionary containing the key-value for chat duration dropdown.

"""


def get_chat_duration_list():
    return [{"key": "10000000", "value": "No Restriction"}, {"key": "15", "value": "Less than 15 Minutes"}, {"key": "30", "value": "Between 15 and 30 minutes"}, {"key": "60", "value": "Between 30 and 60 minutes"}, {"key": "61", "value": "More than 60 minutes"}]


"""

function: get_livechat_date_format
input params:
    datetime_obj:
output:
    This function returns datetime string in "'%d-%b-%Y'" format, ex. Feb 21, 2018

"""


def get_livechat_date_format(datetime_obj):
    result = "None"

    try:
        return datetime_obj.date().strftime('%d-%b-%Y')
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_livechat_date_format: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return result


"""
function: get_milliseconds_to_datetime
input params:
    milliseconds:
expected output:
    This function is used to convert milliseconds to datetime.
"""


def get_milliseconds_to_datetime(ms):
    date = datetime.now().date()
    try:
        date = get_livechat_date_format(
            datetime.fromtimestamp(int(ms) // 1000.0))

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_miniseconds_datetime: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return date


"""
function: get_agents_under_this_user
input params:
    user_obj:
output:
    1. This function returns list of livechat agents under this user_obj.
    2. If user_obj is agent then it will be empty.

"""


def get_agents_under_this_user(user_obj, add_admin_to_list=True):
    agent_obj_list = []
    if user_obj.status == "3" or (user_obj.status == "1" and add_admin_to_list):
        agent_obj_list.append(user_obj)
    try:
        for user in user_obj.agents.all():
            if user.status == "2":
                for user1 in user.agents.all():
                    agent_obj_list.append(user1)
            elif user.status == "3":
                agent_obj_list.append(user)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agents_under_this_user: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return list(set(agent_obj_list))


def get_agents_as_per_supervisors(supervisors_list, user_obj_list, LiveChatUser):
    try:
        if not isinstance(supervisors_list, list):
            supervisors_list = supervisors_list.split(',')
        agent_pks = LiveChatUser.objects.filter(pk__in=supervisors_list).values_list('agents', flat=True).exclude(agents=None)

        agent_objs = LiveChatUser.objects.filter(pk__in=agent_pks, status='3', is_deleted=False)
        return agent_objs

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agents_as_per_supervisors: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'}) 

        return user_obj_list    


def get_admin_from_active_agent(active_agent, LiveChatUser):

    try:
        if active_agent.is_allow_toggle or active_agent.status == "1":
            return active_agent

        parent_user = LiveChatUser.objects.filter(
            agents__pk=active_agent.pk)[0]
        try:
            second_parent = LiveChatUser.objects.filter(
                agents__pk=parent_user.pk)[0]
            return second_parent
        except Exception:
            return parent_user

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_admin_from_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return None


"""
function: get_livechat_user_obj_from_request

Returns LiveChatUser object from the request packet
"""


def get_livechat_user_obj_from_request(request, LiveChatUser, User):
    try:
        user_obj = LiveChatUser.objects.filter(user=request.user, is_deleted=False)[0]
        return user_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_livechat_user_obj_from_request %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return None


"""
function: generate_random_password

Generate a random password string of given length
"""


def generate_random_password():
    try:
        import string
        length = 10
        letters_and_digits = string.ascii_letters + string.digits
        result_str = ''.join((random.choice(letters_and_digits)
                              for _ in range(length)))
        return result_str
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("generate_random_password: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        return None


"""
function: send_mail
input params:
    from_email_id: Sender email id
    to_email_id: Receiver email id
    message_as_string: Email message string
    from_email_id_password: Sender email password

Send mail to given email id
"""


def send_mail(from_email_id, to_emai_id, message_as_string, from_email_id_password):
    import smtplib
    # The actual sending of the e-mail
    server = smtplib.SMTP('smtp.gmail.com:587')
    # Credentials (if needed) for sending the mail
    password = from_email_id_password
    # Start tls handshaking
    server.starttls()
    # Login to server
    server.login(from_email_id, password)
    # Send mail
    server.sendmail(from_email_id, to_emai_id, message_as_string)
    # Close session
    server.quit()


"""
function: send_password_mail
input params:
    name: Name of user
    username: Username
    email: email id of user
    password: randomly generated password for user

Send user credentials of agent on their email
"""


def send_password_mail(name, username, email, password):

    body = """
        <head>
            <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
            <title>Cogno AI</title>
            <style type="text/css" media="screen">
            </style>
        </head>
        <body>
        <div style="padding:1em;border:0.1em black solid;" class="container">
            <p>
                Dear {},
            </p>
            <p>
                We have received a request to provide you with the LiveChat User ID and Password. <br> User ID: {} <br> Password: {} <br>
                Please <a href="{}">click</a> here to login 
            </p>
            <p>&nbsp;</p>
            """
    developer_console_config = get_developer_console_settings()

    body += developer_console_config.custom_report_template_signature

    body += """</div></body>"""

    body = body.format(name, username, password,
                       settings.EASYCHAT_HOST_URL + '/chat/login/')

    email_subject = "LiveChat Login Details For " + \
        str(username)
    thread = threading.Thread(target=send_email_to_customer_via_awsses, args=(
        email, email_subject, body), daemon=True)
    thread.start()


"""
function: send_password_mail
input params:
    name: Name of user
    username: Username
    email: email id of user
    password: randomly generated password for user

Send user credentials of agent on their email
"""


def file_download(file_key, sender_username, LiveChatUser, LiveChatFileAccessManagement):
    file_access_management_obj = None
    try:
        file_access_management_obj = LiveChatFileAccessManagement.objects.get(
            key=file_key)

        is_user_authorized = check_user_authorization_for_file_access(file_access_management_obj, sender_username, LiveChatUser)
        if not is_user_authorized:
            return HttpResponse(status=404)

        path_to_file = file_access_management_obj.file_path

        filename = path_to_file.split("/")[-1]

        path_to_file = settings.BASE_DIR + path_to_file
        mime_type, _ = mimetypes.guess_type(path_to_file)

        if os.path.exists(path_to_file):
            with open(path_to_file, 'rb') as fh:
                response = HttpResponse(
                    fh.read(), status=200, content_type=mime_type)
                response['Content-Disposition'] = 'attachment; filename=%s' % smart_str(
                    str(filename))
                response['X-Sendfile'] = smart_str(path_to_file)
                return response
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error FileDownload %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return HttpResponse(status=404)


"""
function: check_user_authorization_for_file_access
input params:
    name: file_access_management object
    username: Livechat username

Checks and notifies if user is authorisied to access the file
"""


def check_user_authorization_for_file_access(file_access_management_obj, sender_username, LiveChatUser):
    try:

        user_obj = LiveChatUser.objects.filter(user__username=sender_username).first()

        if file_access_management_obj.file_access_type == "all":
            
            return True

        elif file_access_management_obj.file_access_type == "group_chat":

            accessible_users = file_access_management_obj.group.members.filter(is_removed=False).values_list('user', flat=True)

            if user_obj.pk not in accessible_users:
                return False

        elif file_access_management_obj.file_access_type == "user_group_chat":

            accessible_users = file_access_management_obj.user_group.members.all().values_list('user', flat=True)

            if user_obj.pk not in accessible_users:
                return False

        elif file_access_management_obj.file_access_type == "personal_access":

            if user_obj not in file_access_management_obj.users.all():
                return False

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error check_user_authorization_for_file_access %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return True        


"""
function: get_admin_config
input params:
    livechat_user: LiveChat user

It returns the LiveChatAdminConfig object for a particular agent.
"""


def get_admin_config(livechat_user, LiveChatAdminConfig, LiveChatUser):
    try:
        admin_user = get_admin_from_active_agent(livechat_user, LiveChatUser)
        return LiveChatAdminConfig.objects.filter(admin=admin_user)[0]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_admin_config %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return None


"""
function: get_message_history
input params:
    livechat_cust_obj: LiveChatCustomer

It returns message history of livechat_cust_obj as a json.
"""


def get_message_history(livechat_cust_obj, is_bot_side_message_history, LiveChatMISDashboard, LiveChatTranslationCache):
    message_history = []
    try:
        from LiveChatApp.utils_translation import get_translated_text

        livechat_mis_objs = LiveChatMISDashboard.objects.filter(
            livechat_customer=livechat_cust_obj).order_by('message_time', 'pk')

        for livechat_mis_obj in livechat_mis_objs:
            if (livechat_mis_obj.text_message != None and livechat_mis_obj.text_message.strip()) or (livechat_mis_obj.attachment_file_path != None and livechat_mis_obj.attachment_file_path.strip()):

                text_message = livechat_mis_obj.text_message
                translated_text = livechat_mis_obj.translated_text
                if is_bot_side_message_history and "customer name" in livechat_mis_obj.text_message.lower() and "agent name" in livechat_mis_obj.text_message.lower():
                    text_message = livechat_cust_obj.agent_id.get_agent_name() + " has joined the chat. Please ask your queries now."
                    translated_text = get_translated_text(text_message, livechat_cust_obj.customer_language.lang, LiveChatTranslationCache) if livechat_cust_obj.customer_language else ""

                message_history.append({
                    "message": text_message,
                    "sender": livechat_mis_obj.sender,
                    "sender_name": livechat_mis_obj.sender_name,
                    "time": str(get_time(livechat_mis_obj.message_time)),
                    "time_in_minisec": str(get_miniseconds_datetime(livechat_mis_obj.message_time)),
                    "attached_file_src": livechat_mis_obj.attachment_file_path,
                    "file": livechat_mis_obj.thumbnail_file_path,
                    "preview_attachment_file_path": livechat_mis_obj.preview_attachment_file_path,
                    "is_guest_agent_message": livechat_mis_obj.is_guest_agent_message,
                    "message_id": str(livechat_mis_obj.message_id),
                    "reply_message_id": livechat_mis_obj.reply_message_id,
                    "sender_username": livechat_mis_obj.sender_username,
                    "translated_text": translated_text,
                    "meeting_link": livechat_mis_obj.meeting_link,
                    "is_voice_call_message": livechat_mis_obj.is_voice_call_message,
                    "is_video_call_message": livechat_mis_obj.is_video_call_message,
                    "is_cobrowsing_message": livechat_mis_obj.is_cobrowsing_message,
                    "is_transcript_message": livechat_mis_obj.is_transcript_message,
                    "is_file_not_support_message": livechat_mis_obj.is_file_not_support_message,
                    "message_for": livechat_mis_obj.message_for,
                    "message_id": str(livechat_mis_obj.message_id),
                    "is_customer_warning_message": livechat_mis_obj.is_customer_warning_message,
                    "is_customer_report_message_notification": livechat_mis_obj.is_customer_report_message_notification,
                    "agent_name": livechat_mis_obj.sender_name,
                    "is_copied_from_easychat": livechat_mis_obj.is_copied_from_easychat,
                })

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_message_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return message_history


def get_internal_one_to_one_chat_message_history(curr_user, receiver_user, chat_index, LiveChatInternalMISDashboard, LiveChatInternalUserGroup, LiveChatInternalChatLastSeen):
    message_history = []
    all_chat_loaded = False

    try:
        user_group_obj = check_and_update_user_group(
            curr_user, receiver_user, LiveChatInternalUserGroup)
        
        last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=receiver_user, user_group=user_group_obj)

        last_seen_on_chat = 0
        if last_seen_obj:
            last_seen_on_chat = last_seen_obj.first().last_seen.timestamp() * 1000

        livechat_mis_objs = LiveChatInternalMISDashboard.objects.filter(
            sender__in=[curr_user, receiver_user], receiver__in=[receiver_user, curr_user], user_group=user_group_obj).order_by('-message_datetime', '-pk')

        max_index = min(chat_index + 50, livechat_mis_objs.count())

        if max_index >= livechat_mis_objs.count():
            all_chat_loaded = True

        livechat_mis_objs = livechat_mis_objs[chat_index:max_index]
        for livechat_mis_obj in livechat_mis_objs:
            message_history.append(get_message_history_dict(livechat_mis_obj, curr_user))

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_internal_one_to_one_chat_message_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return message_history, all_chat_loaded, last_seen_on_chat


def get_message_history_dict(livechat_mis_obj, curr_user):
    try:
        sender = "ReceiverUser"
        try:
            if livechat_mis_obj.sender.pk == curr_user.pk:
                sender = "CurrentUser"
        except Exception:
            sender = "System"

        message_dict = {
            "message_id": str(livechat_mis_obj.message_id),
            "message": livechat_mis_obj.message_info.message_text, 
            "sender": sender, 
            "sender_name": livechat_mis_obj.sender_name, 
            "time": str(get_time(livechat_mis_obj.message_datetime)), 
            "time_in_minisec": str(get_miniseconds_datetime(livechat_mis_obj.message_datetime)), 
            "attached_file_src": livechat_mis_obj.message_info.attached_file_src,
            "file": livechat_mis_obj.message_info.thumbnail_file_src, 
            "preview_attachment_file_path": livechat_mis_obj.message_info.preview_file_src, 
            "is_replied_message": livechat_mis_obj.message_info.is_replied_message, 
            "reply_message_text": livechat_mis_obj.message_info.reply_message_text,
            "reply_attached_file_src": livechat_mis_obj.message_info.reply_attached_file_src, 
            "reply_attached_file_name": livechat_mis_obj.message_info.reply_attached_file_name, 
            "reply_thumbnail_file_src": livechat_mis_obj.message_info.reply_thumbnail_file_src,
            "reply_message_time": str(get_time(livechat_mis_obj.message_info.reply_message_time))
        }

        return message_dict
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_message_history_dict %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return {}


"""
function: save_transfer_audit
input params:
    chat_transferred_by: The agent who is transferring the chat
    chat_transferred_to: The agent who is going to attain transferred the chat
    livechat_cust_obj: The customer who is being transferred
    transfer_description: The reason for transferring chat provided by chat_transferred_by

This function is used to save chat transfer audits and create tranfer chat message MIS.
"""


def save_transfer_audit(chat_transferred_by, chat_transferred_to, livechat_cust_obj, transfer_description, LiveChatTransferAudit):
    try:
        LiveChatTransferAudit.objects.create(chat_transferred_by=chat_transferred_by,
                                             chat_transferred_to=chat_transferred_to,
                                             livechat_customer=livechat_cust_obj,
                                             transfer_description=transfer_description)
        ongoing_chats = chat_transferred_by.ongoing_chats
        ongoing_chats = max(ongoing_chats - 1, 0)
        chat_transferred_by.ongoing_chats = ongoing_chats
        chat_transferred_by.save()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_transfer_audit %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: set_livechat_session_in_profile
input params:
    livechat_session_id: Livechat session for whatsapp/Facebook etc user
    channel_name: Channel name ex. WhatsApp, Facebook, Instagram
    phone: It works as user id in case of WhatsApp. It may change for other channel

This function is used to add updated livechat session in easychat Profile object
"""


def set_livechat_session_in_profile(livechat_session_id, channel_name, phone, bot_obj, Profile):
    try:
        if channel_name != "WhatsApp":
            return
        profile_obj = Profile.objects.get(user_id=phone, bot=bot_obj)
        profile_obj.livechat_connected = True
        profile_obj.livechat_session_id = livechat_session_id
        profile_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_transfer_audit %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: get_agent_token
input params:
    username: username of agent

This function is used to generate hash of current agent. And this has is used for creating room for current agent.
"""


def get_agent_token(username):
    return "74a65650-2902-4d43-8eb4-d68fa8eeb4dd123"


"""
function: get_blacklisted_keyword_for_current_agent
input params:
    agent_obj: current user
    LiveChatBlackListKeyword: 

returns this function provide all the canned response for agent_obj. agent_obj can be agent/supervisor/admin or livechatonlyadmin
"""


def get_blacklisted_keyword_for_current_agent(agent_obj, LiveChatBlackListKeyword, LiveChatUser):
    blacklisted_keyword = []
    try:
        if agent_obj.status == "1" or agent_obj.is_allow_toggle:
            blacklisted_keyword = LiveChatBlackListKeyword.objects.filter(
                agent_id=agent_obj, blacklist_keyword_for="agent")
            return blacklisted_keyword

        if agent_obj.status == "2":
            blacklisted_keyword = LiveChatBlackListKeyword.objects.filter(agent_id=agent_obj, blacklist_keyword_for="agent") | LiveChatBlackListKeyword.objects.filter(
                agent_id=LiveChatUser.objects.filter(agents__user=agent_obj.user)[0], blacklist_keyword_for="agent")
            return blacklisted_keyword
        parent_users = LiveChatUser.objects.filter(agents__user=agent_obj.user)
        livechat_admin_obj = None
        if LiveChatUser.objects.filter(agents__user=parent_users[0].user).count():
            livechat_admin_obj = LiveChatUser.objects.filter(
                agents__user=parent_users[0].user)[0]

        blacklisted_keyword = LiveChatBlackListKeyword.objects.filter(
            agent_id__in=parent_users, blacklist_keyword_for="agent") | LiveChatBlackListKeyword.objects.filter(agent_id=livechat_admin_obj, blacklist_keyword_for="agent")
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_blacklisted_keyword_for_current_agent: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return blacklisted_keyword


"""
function: get_canned_response_for_current_agent
input params:
    user_obj: current user
    CannedResponse

returns this function provide all the canned response for user_obj. user_obj can be agent/supervisor/admin or livechatonlyadmin
"""


def get_canned_response_for_current_agent(user_obj, CannedResponse, LiveChatUser):
    canned_response_list = []
    try:
        if user_obj.status == "1" or user_obj.is_allow_toggle:
            canned_response_list = CannedResponse.objects.filter(
                agent_id=user_obj, is_deleted=False)
            return canned_response_list

        if user_obj.status == "2":
            parent_user = LiveChatUser.objects.filter(
                agents__user=user_obj.user)[0]
            canned_response_list = CannedResponse.objects.filter(
                agent_id__in=[user_obj, parent_user], is_deleted=False)
        else:
            parent_users = LiveChatUser.objects.filter(
                agents__user=user_obj.user)
            livechat_admin_obj = None
            if LiveChatUser.objects.filter(agents__user=parent_users[0].user).count():
                livechat_admin_obj = LiveChatUser.objects.filter(
                    agents__user=parent_users[0].user)[0]

            canned_response_list = CannedResponse.objects.filter(agent_id__in=[
                                                                 user_obj, livechat_admin_obj], is_deleted=False) | CannedResponse.objects.filter(agent_id__in=parent_users, is_deleted=False)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_canned_response_for_current_agent: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return canned_response_list


"""
function: check_if_livechat_only_admin
input params:
    user_obj: current user
    LiveChatUser: 

returns this function checks if current user is livechat only admin or not.
"""


def check_if_livechat_only_admin(user_obj, LiveChatUser):
    try:
        if user_obj.is_livechat_only_admin:
            return LiveChatUser.objects.filter(livechat_only_admin__user=user_obj.user)[0]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_if_livechat_only_admin: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return user_obj


"""
function: get_total_number_of_active_livechat_users_under_admin
input params:
    user_obj: current user
    LiveChatUser:

returns the total numnber of active users under a LiveChat Admin including Supervisors as well as agents.
"""


def get_total_number_of_active_livechat_users_under_admin(user_obj, LiveChatUser):
    admin_user = get_admin_from_active_agent(user_obj, LiveChatUser)
    user_objs_under_admin = admin_user.agents.filter(is_deleted=False)  # can have both agents as well as supervisors
    total_users_under_admin = 0
    for user_obj in user_objs_under_admin:
        if user_obj.agents.filter(is_deleted=False).count() == 0:
            total_users_under_admin = total_users_under_admin + 1
        else:
            total_users_under_admin = total_users_under_admin + 1 + user_obj.agents.filter(is_deleted=False).count()

    return total_users_under_admin


"""
function: check_if_user_creation_allowed
input params:
    user_obj: current user
    LiveChatUser:
    LiveChatAdminConfig
    no_of_user_expected_to_be_created (optional)

returns boolean which define if new user(s) can be created or not
"""


def check_if_user_creation_allowed(user_obj, LiveChatUser, LiveChatAdminConfig, no_of_user_expected_to_be_created=1):
    try:
        livechat_admin_config_obj = get_admin_config(user_obj, LiveChatAdminConfig, LiveChatUser)
        max_users_allowed_to_be_created = livechat_admin_config_obj.max_users_allowed_to_be_created
        current_number_of_active_users = get_total_number_of_active_livechat_users_under_admin(user_obj, LiveChatUser)

        creation_allowed = (current_number_of_active_users + no_of_user_expected_to_be_created) <= max_users_allowed_to_be_created

        return creation_allowed

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_if_agent_creation_allowed: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    
    return False


def check_agent_has_bot_assign_category(user_obj, bot_obj, livechat_category):
    if bot_obj in user_obj.bots.all():
        user_obj.category.add(livechat_category)
        user_obj.save()


def create_image_thumbnail(filename):
    thumbnail_file_name = ""

    try:
        original_file = Image.open(
            LIVECHAT_SECURED_FILES_ATTACHMENT_PATH + filename)
        original_file.thumbnail((80, 80))

        thumbnail_file_name = filename.split(
            '.')[0] + '_thumbnail.' + filename.split('.')[1]
        original_file.save(
            LIVECHAT_SECURED_FILES_ATTACHMENT_PATH + thumbnail_file_name)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("create_image_thumbnail: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return thumbnail_file_name


def create_video_thumbnail(filename):
    thumbnail_file_name = ""

    try:
        clip = VideoFileClip(settings.SECURE_MEDIA_ROOT +
                             'LiveChatApp/attachment/' + filename)
        duration = clip.duration
        frame_at_second = min(5, duration)
        frame = clip.get_frame(frame_at_second)
        thumbnail = Image.fromarray(frame)

        thumbnail_file_name = filename.split('.')[0] + '_thumbnail.png'

        thumbnail.save(
            LIVECHAT_SECURED_FILES_ATTACHMENT_PATH + thumbnail_file_name)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("create_video_thumbnail: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return thumbnail_file_name


"""
function: get_masked_data

input_params:
    message: text message.
output_params:
    masked_data: Used to mask all confidential data such as PAN, Adhar, Account number and account balance etc.
"""


def get_masked_data(message):
    final_string = message
    try:
        # pan
        pan_pattern = re.findall(
            r"[a-zA-Z]{3}[pP][a-zA-Z][1-9][0-9]{3}[a-zA-Z]", message)
        for item in pan_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # email id
        email_pattern = re.findall(r'[\w.+-]+@[\w-]+\.[\w.-]+', message)
        for item in email_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # dates
        # Regex date format: dd/mm/yyyy
        date_format_ddmmyyyy = re.findall(
            r"[\d]{1,2}/[\d]{1,2}/[\d]{2,4}", message)
        for item in date_format_ddmmyyyy:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # Regex date format: dd-mm-yy
        date_format_ddmmyyyy_two = re.findall(
            r"[\d]{1,2}-[\d]{1,2}-[\d]{2}", message)
        for item in date_format_ddmmyyyy_two:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # Regex date format: dd AUG YYYY
        date_format_ddmonthnameyyyy = re.findall(
            r"[\d]{1,2} [ADFJMNOS]\w* [\d]{2,4}", message)
        for item in date_format_ddmonthnameyyyy:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # Regex date format: AUG dd YYYY
        date_format_monthnameddyyyy = re.findall(
            r"[ADFJMNOS]\w* [\d]{1,2} [\d]{2,4}", message)
        for item in date_format_monthnameddyyyy:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # mobile number, account number, aadhar number (10-12 digits number
        # should have space before and after)
        mobile_pattern = re.findall(r"\b[0-9]{10,12}\b", message)
        for item in mobile_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # age and address (1-3 digits number
        # should have space before and after)
        age_pattern = re.findall(r"\b[0-9]{1,3}\b", message)
        for item in age_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # customer_id (any string that contains atleast 1 digit)
        id_pattern = re.findall(r"\b[A-Za-z0-9]*\d[A-Za-z0-9]*\b", message)
        for item in id_pattern:
            reg = r"\b" + str(item) + r"\b"
            message = re.sub(reg, hashlib.sha256(
                item.encode()).hexdigest(), message)

        final_string = message
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_masked_data: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return final_string


"""
function: mask_pii_data

input_params:
    message: customer_obj
output_params:
    masking of confidential data such as PAN, Adhar, Account number and account balance etc in messages of this customer.
"""


def mask_pii_data(customer_obj, LiveChatMISDashboard):
    try:
        livechat_mis = LiveChatMISDashboard.objects.filter(
            livechat_customer=customer_obj, sender='Customer')

        for mis in livechat_mis:
            message = mis.text_message
            message = get_masked_data(message)
            mis.text_message = message
            mis.save()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("mask_pii_data: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def send_otp_mail(email_ids, otp, username, bot_obj):
    body = """
    <head>
      <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
      <title>Cogno AI</title>
      <style type="text/css" media="screen">
      </style>
    </head>
    <body>

    <div style="padding:1em;border:0.1em black solid;" class="container">
        <p>
            """ + str(username) + """ is trying to turn off the PII Masking settings for the bot """ + str(bot_obj.name) + """ - """ + str(bot_obj.pk) + """. Code to accept it is <b>""" + str(otp) + """<b>.
        </p>
    </div>

    </body>"""

    # Create message container - the correct MIME type is
    # multipart/alternative here!

    for email_id in email_ids:
        send_email_to_customer_via_awsses(email_id, "PII Change Request", body)


def paginate(obj_list, item_count, page):
    try:
        total_audits = obj_list.count()
    except Exception:
        # obj_list is a python list
        total_audits = len(obj_list)

    paginator = Paginator(obj_list, item_count)

    try:
        obj_list = paginator.page(page)
    except PageNotAnInteger:
        obj_list = paginator.page(1)
    except EmptyPage:
        obj_list = paginator.page(paginator.num_pages)

    if page != None:
        start_point = item_count * (int(page) - 1) + 1
        end_point = min(item_count *
                        int(page), total_audits)
    else:
        start_point = 1
        end_point = min(item_count, total_audits)

    return total_audits, obj_list, start_point, end_point


"""
function: get_trailing_list
input_params:
    current_history_id: the id related to the history on which the user currently is(modal state open). 
    audit_obj_list: audit_obj_list complete 
    AUDIT_TRAIL_ITEM_COUNT: Item count per page,
    audit_obj_list_final: audit_object_list after pagination on current page,
    page: page open in back ground of modal,
    total_pages: total no. of pages.
Returns the complete list of maximum (AUDIT_TRAIL_ITEM_COUNT*3) objects of audit_objs.
"""


def get_trailing_list(current_history_id, audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, audit_obj_list_final, page, total_pages):
    try:
        current_index = -1
        current_page = -1
        if current_history_id != '':
            for index, item in enumerate(audit_obj_list):
                if item.session_id == uuid.UUID(current_history_id):
                    current_index = index
                    break
            current_page = int((current_index + 1) /
                               AUDIT_TRAIL_ITEM_COUNT) + 1

        audit_obj_next_list_final = []
        audit_obj_prev_list_final = []

        if current_page != -1:
            audit_obj_prev_list_final, audit_obj_list_final, audit_obj_next_list_final = get_audit_obj_session_lists(
                current_page, audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, audit_obj_list_final, current_page, total_pages)
        else:
            audit_obj_prev_list_final, audit_obj_list_final, audit_obj_next_list_final = get_audit_obj_session_lists(
                page, audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, audit_obj_list_final, current_page, total_pages)

        trailing_list = list(audit_obj_prev_list_final) + \
            list(audit_obj_list_final) + list(audit_obj_next_list_final)
        return trailing_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_trailing_list: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""
function: get_audit_obj_session_lists
input_params:
    page: page open in back ground of modal,
    audit_obj_list: audit_obj_list complete 
    AUDIT_TRAIL_ITEM_COUNT: Item count per page,
    audit_obj_list_final: audit_object_list after pagination on current page,
    currently_opened_page: -1 if no modal history is opened and other non negative integer(page no.) if a particular history is open
    total_pages: total no. of pages.
Returns the list of items on previous page to current page, list of items on current page and list of items on the next page of current page
"""


def get_audit_obj_session_lists(page, audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, audit_obj_list_final, currently_opened_page, total_pages):
    try:
        if int(page) == total_pages:
            audit_obj_next_list_final = []
        else:
            total_audits, audit_obj_next_list_final, start_point, end_point = paginate(
                audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, int(page) + 1)

        if currently_opened_page != -1:
            total_audits, audit_obj_list_final, start_point, end_point = paginate(
                audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, int(page))

        if int(page) == 1:
            audit_obj_prev_list_final = []
        else:
            total_audits, audit_obj_prev_list_final, start_point, end_point = paginate(
                audit_obj_list, AUDIT_TRAIL_ITEM_COUNT, int(page) - 1)

        return audit_obj_prev_list_final, audit_obj_list_final, audit_obj_next_list_final
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_audit_obj_session_lists: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def fetch_all_agents_under_this_user(user_obj):
    user_obj_list = user_obj.agents.none()
    try:
        user_obj_list = user_obj.agents.all().filter(is_deleted=False).order_by('-pk')
        supervisor_obj_list = user_obj_list.filter(status="2")
        for supervisors in supervisor_obj_list:
            user_obj_list = user_obj_list | supervisors.agents.all().filter(
                is_deleted=False).order_by('-pk')
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("mask_pii_data: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return user_obj_list.distinct()


def export_today_data(user, report_type, filter_param, LiveChatCustomer, LiveChatSessionManagement, LiveChatAgentNotReady, LiveChatFollowupCustomer):
    try:
        # adding new report types "8" for abandoned chats and "9" for customer declined chats
        if report_type == "2":
            thread = threading.Thread(target=export_today_chat_history, args=(
                user, filter_param, LiveChatCustomer, LiveChatFollowupCustomer), daemon=True)
            thread.start()
        elif report_type == "6":
            thread = threading.Thread(target=export_today_missed_chats_report, args=(
                user, filter_param, LiveChatCustomer), daemon=True)
            thread.start()
        elif report_type == "5":
            thread = threading.Thread(target=export_today_login_logout_report, args=(
                user, filter_param, LiveChatSessionManagement), daemon=True)
            thread.start()
        elif report_type == "0":
            thread = threading.Thread(target=export_today_agent_not_ready_report, args=(
                user, filter_param, LiveChatAgentNotReady), daemon=True)
            thread.start()
        elif report_type == "1":
            thread = threading.Thread(target=export_today_agent_performance_report, args=(
                user, filter_param, LiveChatSessionManagement), daemon=True)
            thread.start()
        elif report_type == "4":
            thread = threading.Thread(target=export_today_hourly_interaction_report, args=(
                user, filter_param, LiveChatCustomer), daemon=True)
            thread.start()
        elif report_type == "3":
            thread = threading.Thread(target=export_today_daily_interaction_report, args=(
                user, filter_param, LiveChatCustomer), daemon=True)
            thread.start()
        elif report_type == "8":
            thread = threading.Thread(target=export_today_offline_chats_report, args=(
                user, filter_param, LiveChatCustomer), daemon=True)
            thread.start()
        elif report_type == "9":
            thread = threading.Thread(target=export_today_abandoned_chats_report, args=(
                user, filter_param, LiveChatCustomer), daemon=True)
            thread.start()
        return True

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_data: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        return False


def get_html_body_for_mailer_report(username, start_date_obj, end_date_obj, domain, export_zip_file_path, report_name="Report"):
    
    date_format = "%Y-%m-%d"
    start_date = datetime.strptime(start_date_obj, date_format).date()
    end_date = datetime.strptime(end_date_obj, date_format).date()
    
    body = f"""
            <head>
                <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                <title>Cogno AI</title>
                <style type="text/css" media="screen">
                </style>
            </head>
            <body>
            <div style="padding:1em;border:0.1em black solid;" class="container">
                <p>
                    Dear {username},
                </p>
                <p>
                    We have received a request to provide you with the LiveChat {report_name} report from {start_date} to {end_date}. Please click on the link below to download the file.
                </p>
                <a href="{domain}/{export_zip_file_path}">click here</a>
                <p>&nbsp;</p>
    """
    developer_console_config = get_developer_console_settings()
    body += developer_console_config.custom_report_template_signature 
    body += """</div></body>"""

    return body


def export_custom_hourly_interaction_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatCustomer, zip_obj):

    message_history_datadump_log = open(
        "log/livechat_hourly_interaction_report_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]

        agent_list = get_agents_under_this_user(user_obj)
        query_bot_obj_list = user_obj.bots.all()
        tz = pytz.timezone(settings.TIME_ZONE)

        for temp_date in list_of_date:

            interaction_objects = []
            interaction_object_list = LiveChatCustomer.objects.filter(agent_id__in=agent_list,
                                                                      request_raised_date=temp_date.date(),
                                                                      ).order_by('-joined_date')

            for bot_obj in query_bot_obj_list:
                
                current_datetime = temp_date.replace(hour=0, minute=0, second=0).astimezone(tz)
                start_time_obj = current_datetime.time().replace(hour=0, minute=0, second=0)

                for itr in range(1, 25):
                    if itr != 24:
                        end_time_obj = current_datetime.time().replace(hour=itr, minute=0, second=0)
                    else:
                        end_time_obj = current_datetime.time().replace(hour=itr - 1, minute=59, second=59)
                    interaction_count = interaction_object_list.filter(bot=bot_obj, 
                                                                            joined_date__time__range=[
                                                                                start_time_obj, end_time_obj]
                                                                       ).count()
                    
                    interaction_objects.append({"date": current_datetime, "start_time": start_time_obj,
                                                "end_time": end_time_obj, "bot_name": bot_obj.name, "frequency": interaction_count})
                    
                    start_time_obj = current_datetime.time().replace(hour=itr % 24, minute=0, second=1)

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Bot Name")
            sheet1.write(0, 2, "Start Time")
            sheet1.write(0, 3, "End Time")
            sheet1.write(0, 4, "Interaction Count")
            row = 1
            for index in range(0, len(interaction_objects)):
                sheet1.write(row, 0, str(interaction_objects[index]["date"].date()))
                sheet1.write(row, 1, interaction_objects[index]["bot_name"])
                sheet1.write(row, 2, interaction_objects[index]["start_time"].strftime("%H:%M:%S"))
                sheet1.write(row, 3, interaction_objects[index]["end_time"].strftime("%H:%M:%S"))
                sheet1.write(row, 4, interaction_objects[index]["frequency"])
                row += 1

            file_path = LIVECHAT_HOURLY_INTERACTION_REPORT + \
                        str(username) + \
                       HOURLY_INTERACTION_REPORT + str(temp_date.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()

        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Hourly Interactions")
        email_subject = f"LiveChat Hourly Interactions Report For {str(username)}" 

        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_hourly_interaction_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_daily_interaction_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatCustomer, zip_obj):

    message_history_datadump_log = open(
        "log/livechat_daily_interaction_report_dump.log", "a")

    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]

        agent_list = get_agents_under_this_user(user_obj)

        for temp_date in list_of_date:

            interaction_objects = []

            interaction_object_list = LiveChatCustomer.objects.filter(
                agent_id__in=agent_list, request_raised_date=temp_date)

            query_bot_obj_list = user_obj.bots.all()

            for bot_obj in query_bot_obj_list:
                interaction_objs_for_a_bot_obj = interaction_object_list.filter(bot=bot_obj)
                interaction_objects.append({
                    "date": temp_date, 
                    "bot_name": bot_obj.name, 
                    "frequency": interaction_objs_for_a_bot_obj.count()
                })

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Bot Name")
            sheet1.write(0, 2, "Interaction Count")
            row = 1
            for index in range(0, len(interaction_objects)):
                sheet1.write(row, 0, interaction_objects[index]["date"].strftime("%d-%m-%Y"))
                sheet1.write(row, 1, interaction_objects[index]["bot_name"])
                sheet1.write(row, 2, interaction_objects[index]["frequency"])
                row += 1

            file_path = LIVECHAT_DAILY_INTERACTION_REPORT_PATH + \
                        str(username) + \
                       DAILY_INTERACTION_REPORT_PATH + str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()

        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Daily Interactions")
        email_subject = f"LiveChat Daily Interactions Report For {str(username)}"\

        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_daily_interaction_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_missed_chat_history_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatCustomer, zip_obj):

    message_history_datadump_log = open(
        "log/livechat_offline_message_report_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        time_zone = pytz.timezone(settings.TIME_ZONE)
        for temp_date in list_of_date:

            message_history_yesterday = LiveChatCustomer.objects.filter(
                last_appearance_date__date=temp_date, is_denied=True, is_system_denied=False, bot__in=user_obj.bots.all(), category__in=user_obj.category.all())

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Created On")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Message")
            sheet1.write(0, 5, "Channel")
            sheet1.write(0, 6, "Category")
            row = 1
            for message in message_history_yesterday:
                sheet1.write(row, 0, str((message.joined_date.astimezone(time_zone)).strftime(
                    "%d/%m/%Y %I:%M %p")))
                sheet1.write(row, 1, message.username)
                sheet1.write(row, 2, message.phone)
                sheet1.write(row, 3, message.email)
                sheet1.write(row, 4, message.message)
                sheet1.write(row, 5, message.channel.name)
                if message.category.title:
                    sheet1.write(row, 6, message.category.title)
                else:
                    sheet1.write(row, 6, 'None')
                row += 1

            file_path = LIVECHAT_MISSED_CHATS_REPORT_PATH + \
                        str(username) + \
                       MISSED_CHATS_REPORT_PATH + str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()

        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Missed Chat History")
        email_subject = f"LiveChat Missed Chats History For {str(username)}"
        for email_id in email_list: 
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_missed_chat_history_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_offline_chat_history_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatCustomer, zip_obj):
    
    message_history_datadump_log = open(
        "log/livechat_total_declined_chats_report_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        time_zone = pytz.timezone(settings.TIME_ZONE)
        for temp_date in list_of_date:

            message_history_yesterday = LiveChatCustomer.objects.filter(
                last_appearance_date__date=temp_date, is_system_denied=True, bot__in=user_obj.bots.all(), category__in=user_obj.category.all())

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Created On")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Message")
            sheet1.write(0, 5, "Channel")
            sheet1.write(0, 6, "Category")
            row = 1
            for message in message_history_yesterday:
                sheet1.write(row, 0, str((message.joined_date.astimezone(time_zone)).strftime(
                    "%d/%m/%Y %I:%M %p")))
                sheet1.write(row, 1, message.username)
                sheet1.write(row, 2, message.phone)
                sheet1.write(row, 3, message.email)
                sheet1.write(row, 4, message.message)
                sheet1.write(row, 5, message.channel.name)
                try:
                    sheet1.write(row, 6, message.category.title)
                except Exception:
                    sheet1.write(row, 6, 'None')
                    pass
                row += 1

            file_path = LIVECHAT_OFFLINE_CHATS_REPORT + \
                        str(username) + \
                       OFFLINE_CHATS_REPORT + str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()

        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Offline Chat History")
        email_subject = f"LiveChat Offline Chats History For {str(username)}" 
            
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_offline_chat_history_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()
        

def export_custom_abandoned_chat_history_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatCustomer, zip_obj):
    
    message_history_datadump_log = open(
        "log/livechat_total_declined_chats_report_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        time_zone = pytz.timezone(settings.TIME_ZONE)
        for temp_date in list_of_date:
            
            abandoned_chat_history_objs = LiveChatCustomer.objects.filter(
                last_appearance_date__date=temp_date, agent_id=None, abruptly_closed=True, bot__in=user_obj.bots.all(), category__in=user_obj.category.all())

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Created On")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Message")
            sheet1.write(0, 5, "Channel")
            sheet1.write(0, 6, "Category")
            sheet1.write(0, 7, "Wait Time")
            row = 1
            for abandoned_chat_history_obj in abandoned_chat_history_objs.iterator():
                sheet1.write(row, 0, str((abandoned_chat_history_obj.joined_date.astimezone(time_zone)).strftime(
                    "%d/%m/%Y, %I:%M:%p")))
                sheet1.write(row, 1, abandoned_chat_history_obj.username)
                sheet1.write(row, 2, abandoned_chat_history_obj.phone)
                sheet1.write(row, 3, abandoned_chat_history_obj.email)
                sheet1.write(row, 4, abandoned_chat_history_obj.message)
                sheet1.write(row, 5, abandoned_chat_history_obj.channel.name)
                if abandoned_chat_history_obj.category.title:
                    sheet1.write(row, 6, abandoned_chat_history_obj.category.title)
                else:
                    sheet1.write(row, 6, 'None')
                sheet1.write(row, 7, abandoned_chat_history_obj.get_wait_time())
                row += 1

            file_path = LIVECHAT_ABANDONED_CHATS_REPORT_PATH + \
                        str(username) + \
                        ABANDONED_CHATS_REPORT_PATH + str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()
        
        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Abandoned Chat History")
        email_subject = f"LiveChat Abandoned Chats History For {str(username)}" 
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_abandoned_chat_history_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_voice_call_history_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatVoIPData, zip_obj):
    
    message_history_datadump_log = open(
        "log/voip_history_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        agent_list = get_agents_under_this_user(user_obj)
        time_zone = pytz.timezone(settings.TIME_ZONE)
        for temp_date in list_of_date:

            voip_history_objs = LiveChatVoIPData.objects.filter(request_datetime__date=temp_date, call_type__in=['pip', 'new_tab'], agent__in=agent_list, is_completed=True, is_interrupted=False).order_by("-request_datetime")
            
            test_wb = Workbook()
            
            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Customer Name")
            sheet1.write(0, 1, "Agent Username")
            sheet1.write(0, 2, "Start Date-Time")
            sheet1.write(0, 3, "End Date-Time")
            sheet1.write(0, 4, "Total Duration")
            sheet1.write(0, 5, "Meeting Status")

            row = 1
            for voip_data in voip_history_objs.iterator():

                sheet1.write(row, 0, voip_data.customer.get_username())
                sheet1.write(row, 1, voip_data.agent.user.username)

                sheet1.write(row, 2, voip_data.start_datetime.astimezone(time_zone).strftime('%d-%b-%Y, %I:%M %p'))
                sheet1.write(row, 3, voip_data.end_datetime.astimezone(time_zone).strftime('%d-%b-%Y, %I:%M %p'))

                sheet1.write(row, 4, voip_data.get_call_duration())
                sheet1.write(row, 5, "Completed")

                row += 1
            
            file_path = "livechat-voip-history/" + \
                str(username) + "/voip_history_" + \
                str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()
        
        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "VOIP History")
        email_subject = f"LiveChat VOIP History For {str(username)}" 
        
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_voice_call_history_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_video_call_history_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatVoIPData, zip_obj):
    
    message_history_datadump_log = open(
        "log/vc_history_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        agent_list = get_agents_under_this_user(user_obj)
        time_zone = pytz.timezone(settings.TIME_ZONE)
        for temp_date in list_of_date:
            
            vc_history_objs = LiveChatVoIPData.objects.filter(request_datetime__date=temp_date, call_type='video_call', agent__in=agent_list, is_completed=True, is_interrupted=False).order_by("-request_datetime")

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Customer Name")
            sheet1.write(0, 1, "Agent Username")
            sheet1.write(0, 2, "Start Date-Time")
            sheet1.write(0, 3, "End Date-Time")
            sheet1.write(0, 4, "Total Duration")
            sheet1.write(0, 5, "Meeting Status")

            row = 1
            for vc_data in vc_history_objs.iterator():
    
                sheet1.write(row, 0, vc_data.customer.get_username())
                sheet1.write(row, 1, vc_data.agent.user.username)

                sheet1.write(row, 2, vc_data.start_datetime.astimezone(time_zone).strftime('%d-%b-%Y, %I:%M %p'))
                sheet1.write(row, 3, vc_data.end_datetime.astimezone(time_zone).strftime('%d-%b-%Y, %I:%M %p'))

                sheet1.write(row, 4, vc_data.get_call_duration())
                sheet1.write(row, 5, "Completed")

                row += 1

            file_path = "livechat-vc-history/" + \
                str(username) + "/vc_history_" + \
                str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()
        
        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]
    
        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "VC History")
        email_subject = f"LiveChat VC History For {str(username)}"
            
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_video_call_history_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_agent_not_ready_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatAgentNotReady, zip_obj):
    
    message_history_datadump_log = open(
        "log/livechat_agent_not_ready_report_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        agent_list = get_agents_under_this_user(user_obj)
        for temp_date in list_of_date:
            
            agent_not_ready_objs = LiveChatAgentNotReady.objects.filter(
                not_ready_starts_at__date=temp_date, user__in=agent_list)

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Agent Name")
            sheet1.write(0, 2, "Agent Username")
            sheet1.write(0, 3, "Reason")
            sheet1.write(0, 4, "Stop Interaction Duration")
            sheet1.write(0, 5, "Start Time")
            sheet1.write(0, 6, "End Time")
            sheet1.write(0, 7, "Offline Time")
            row = 1
            for agent_not_ready_obj in agent_not_ready_objs.iterator():
                sheet1.write(row, 0, str(
                    agent_not_ready_obj.not_ready_starts_at.date().strftime("%d-%m-%Y")))
                sheet1.write(row, 1, str(
                    agent_not_ready_obj.user.user.first_name) + " " + str(agent_not_ready_obj.user.user.last_name))
                sheet1.write(row, 2, agent_not_ready_obj.user.user.username)
                sheet1.write(row, 3, agent_not_ready_obj.get_reason_for_offline())
                sheet1.write(row, 4, agent_not_ready_obj.get_stop_interaction_duration())
                sheet1.write(
                    row, 5, str((agent_not_ready_obj.not_ready_starts_at.astimezone(pytz.timezone(settings.TIME_ZONE))).strftime("%H:%M:%S")))
                sheet1.write(row, 6, str((agent_not_ready_obj.not_ready_ends_at.astimezone(pytz.timezone(settings.TIME_ZONE))).strftime("%H:%M:%S")))
                sheet1.write(row, 7, agent_not_ready_obj.get_offline_duration())
                row += 1

            file_path = "livechat-agent-not-ready-report/" + \
                str(username) + \
                "/agent_not_ready_report_" + str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()
        
        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Agent Not Ready")
        email_subject = f"LiveChat Agent Not Ready Report For {str(username)}"
            
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_agent_not_ready_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_login_logout_report(user_obj, username, list_of_date, requested_data, export_zip_file_path, LiveChatSessionManagement, zip_obj):
    
    message_history_datadump_log = open(
        "log/livechat_login_logout_report_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        agent_list = get_agents_under_this_user(user_obj)
        for temp_date in list_of_date:
            
            message_history_objs = LiveChatSessionManagement.objects.filter(
                session_starts_at__date=temp_date, user__in=agent_list).order_by('-session_starts_at')

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Agent Username")
            sheet1.write(0, 1, "Date")
            sheet1.write(0, 2, "Duration")
            sheet1.write(0, 3, "Agent Name")
            sheet1.write(0, 4, "Login Time")
            sheet1.write(0, 5, "Logout Time")
            sheet1.write(0, 6, "Total Offline Time")
            sheet1.write(0, 7, "Total Online Time")
            sheet1.write(0, 8, "Not Ready Duration")
            row = 1
            for message in message_history_objs:
                sheet1.write(row, 0, message.user.user.username)
                sheet1.write(row, 1, str(
                    message.session_starts_at.date().strftime("%d-%m-%Y")))
                sheet1.write(row, 2, str(
                    message.get_session_duration()))
                sheet1.write(row, 3, message.get_name())
                sheet1.write(row, 4, str(
                    (message.session_starts_at.astimezone(pytz.timezone(settings.TIME_ZONE))).strftime("%H:%M:%S")))
                sheet1.write(row, 5, str(
                    (message.session_ends_at.astimezone(pytz.timezone(settings.TIME_ZONE))).strftime("%H:%M:%S")))
                sheet1.write(row, 6, str(
                    message.get_total_offline_time()))
                sheet1.write(row, 7, str(
                    message.get_total_online_time()))
                sheet1.write(
                    row, 8, str(message.get_session_stop_interaction_duration()))
                row += 1

            file_path = "livechat-login-logout-report/" + \
                        str(username) + \
                        "/login_logout_report_" + str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)
            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()
        
        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Login-Logout")
        email_subject = f"LiveChat Login Logout History For {str(username)}"
            
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_login_logout_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_custom_chat_history(user_obj, username, list_of_date, requested_data, LiveChatCustomer, LiveChatFollowupCustomer, LiveChatConfig, zip_obj):
    
    message_history_datadump_log = open(
        "log/message_history_dump.log", "a")
    today = datetime.now()
    try:
        domain = settings.EASYCHAT_HOST_URL
        email_str = requested_data["email"]
        email_list = [email.strip()
                        for email in email_str.split(",") if email != ""]
        agent_list = get_agents_under_this_user(user_obj)
        for temp_date in list_of_date:
            
            livechat_followup_cust_objs = LiveChatFollowupCustomer.objects.filter(
                completed_date__date=temp_date, agent_id__in=agent_list, is_completed=True).values('livechat_customer')
            message_history_objs = LiveChatCustomer.objects.filter(
                Q(agent_id__in=agent_list, joined_date__date=temp_date) | Q(pk__in=livechat_followup_cust_objs)).order_by('-joined_date')

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Chat NPS Score")
            sheet1.write(0, 5, "Chat NPS Feedback")
            sheet1.write(0, 6, "Interaction Start Date-Time")
            sheet1.write(0, 7, "Interaction End Date-Time")
            sheet1.write(0, 8, "Interaction Duration")
            sheet1.write(0, 9, "Category")
            sheet1.write(0, 10, "Agent Name")
            sheet1.write(0, 11, "User ID")
            sheet1.write(0, 12, "Channel")
            sheet1.write(0, 13, "Session ID")
            sheet1.write(0, 14, "Wait Time")
            sheet1.write(0, 15, "FTR")
            sheet1.write(0, 16, "Chat Termination")
            sheet1.write(0, 17, "Cobrowsing NPS")
            sheet1.write(0, 18, "Cobrowsing NPS Comment")

            # For handling single cell Multi-line
            algn1 = xlwt.Alignment()
            algn1.wrap = 1
            style1 = xlwt.XFStyle()
            style1.alignment = algn1
            
            row = 1
            for message in message_history_objs:
                config_obj = LiveChatConfig.objects.get(bot=message.bot)
                is_original_information_in_reports_enabled = config_obj.is_original_information_in_reports_enabled
                joined_date = message.joined_date.strftime("%d/%m/%Y")
                sheet1.write(row, 0, joined_date)
                if is_original_information_in_reports_enabled and message.original_username != "" and message.username != message.original_username:
                    sheet1.write(row, 1, str(message.username) + "(original - " + str(message.original_username) + ")")
                else:
                    sheet1.write(row, 1, message.username)
                if is_original_information_in_reports_enabled and message.original_phone != "" and message.phone != message.original_phone:
                    sheet1.write(row, 2, str(message.phone) + "(original - " + str(message.original_phone) + ")")
                else:
                    sheet1.write(row, 2, message.phone)
                if is_original_information_in_reports_enabled and message.original_email != "" and message.email != message.original_email:
                    sheet1.write(row, 3, str(message.email) + "(original - " + str(message.original_email) + ")")
                else:
                    sheet1.write(row, 3, message.email)
                if message.rate_value == -1:
                    sheet1.write(row, 4, "NA")
                else:
                    sheet1.write(row, 4, message.rate_value)
                
                if message.nps_text_feedback == "":
                    sheet1.write(row, 5, "NA")
                else:
                    sheet1.write(row, 5, message.nps_text_feedback)

                start_time = (message.joined_date + timedelta(hours=5, minutes=30)).strftime(
                    "%d/%m/%Y, %I:%M:%p")
                sheet1.write(row, 6, start_time)
                end_time = (message.last_appearance_date + timedelta(hours=5, minutes=30)).strftime(
                    "%d/%m/%Y, %I:%M:%p")
                if message.followup_assignment:
                    livechat_followup_cust_obj = LiveChatFollowupCustomer.objects.get(livechat_customer=message)
                    end_time = (livechat_followup_cust_obj.completed_date + timedelta(hours=5, minutes=30)).strftime(
                        "%d/%m/%Y, %I:%M:%p")
                sheet1.write(
                    row, 7, end_time)
                sheet1.write(row, 8, message.get_chat_duration())
                try:
                    sheet1.write(row, 9, message.category.title)
                except Exception:
                    sheet1.write(row, 9, 'None')
                    pass
                try:
                    sheet1.write(
                        row, 10, message.get_agent_name())
                    sheet1.write(
                        row, 11, message.get_agent_username())
                except Exception:
                    sheet1.write(row, 10, 'None')
                    sheet1.write(row, 11, 'None')
                    pass
                channel_name = message.channel.name
                if message.followup_assignment:
                    channel_name = message.channel.name + ' (Follow Up)'
                sheet1.write(row, 12, channel_name)
                sheet1.write(row, 13, str(message.session_id))
                sheet1.write(row, 14, message.get_wait_time())
                sheet1.write(row, 15, message.get_agent_first_time_response_time())
                sheet1.write(row, 16, message.chat_ended_by)

                cobrowsing_nps_rating, cobrowsing_nps_comment = message.get_cobrowsing_nps_data()
                sheet1.write(row, 17, cobrowsing_nps_rating, style1)
                sheet1.write(row, 18, cobrowsing_nps_comment, style1)
                row += 1

            file_path = "livechat-chat-history/" + \
                str(username) + "/chat_history_" + \
                str(temp_date) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)
            zip_obj.write(settings.MEDIA_ROOT + file_path, basename(file_path))
        zip_obj.close()
        export_zip_file_path = 'secured_files/LiveChatApp/' + LIVECHAT_CHAT_HISTORY_PATH +\
                              username + "/ChatHistoryCustom.zip"
        from LiveChatApp.models import LiveChatFileAccessManagement
            
        export_zip_file_path = "/" + export_zip_file_path
        file_access_management_obj = LiveChatFileAccessManagement.objects.create(
            file_path=export_zip_file_path, is_public=False)
        file_access_management_obj.file_access_type = "personal_access"
        file_access_management_obj.users.add(user_obj)
        file_access_management_obj.save()

        export_zip_file_path = 'livechat/download-file/' + \
            str(file_access_management_obj.key) + '/LiveChatHistory.zip'

        start_date_obj = requested_data["startdate"]
        end_date_obj = requested_data["enddate"]

        body = get_html_body_for_mailer_report(str(username), start_date_obj, end_date_obj, domain, export_zip_file_path, "Chat History")
        email_subject = f"LiveChat Chat History For {str(username)}"
            
        for email_id in email_list:
            send_email_to_customer_via_awsses(email_id, email_subject, body)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_custom_chat_history: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        message_history_datadump_log.write(str(today) + ": failed: " + str(exc))
    message_history_datadump_log.close()


def export_today_chat_history(user, filter_param, LiveChatCustomer, LiveChatFollowupCustomer):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/message_history_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-chat-history'):
            os.makedirs(settings.MEDIA_ROOT + 'livechat-chat-history')
        try:
            from LiveChatApp.models import LiveChatConfig
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]
            current_time = today.strftime("%H:%M:%S")

            agent_list = get_agents_under_this_user(user)
            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-chat-history/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-chat-history/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-chat-history/" + \
                str(user.user.username) + "/chat_history_" + \
                str(today.date()) + "_" + str(current_time) + ".xls"

            livechat_followup_cust_objs = LiveChatFollowupCustomer.objects.filter(
                completed_date__date=today.date(), agent_id__in=agent_list, is_completed=True).values('livechat_customer')
            message_history_objs = LiveChatCustomer.objects.filter(
                Q(agent_id__in=agent_list, joined_date__date=today.date()) | Q(pk__in=livechat_followup_cust_objs)).order_by('-joined_date')
            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Web-Chat Feedback")
            sheet1.write(0, 5, "Interaction Start Date-Time")
            sheet1.write(0, 6, "Interaction End Date-Time")
            sheet1.write(0, 7, "Interaction Duration")
            sheet1.write(0, 8, "Category")
            sheet1.write(0, 9, "Agent Name")
            sheet1.write(0, 10, "User ID")
            sheet1.write(0, 11, "Channel")
            sheet1.write(0, 12, "Session ID")
            sheet1.write(0, 13, "Wait Time")
            sheet1.write(0, 14, "FTR")
            sheet1.write(0, 15, "Chat Termination")

            row = 1
            for message in message_history_objs:
                config_obj = LiveChatConfig.objects.get(bot=message.bot)
                is_original_information_in_reports_enabled = config_obj.is_original_information_in_reports_enabled
                joined_date = message.joined_date.strftime("%d/%m/%Y")
                sheet1.write(row, 0, joined_date)
                if is_original_information_in_reports_enabled and message.original_username != "" and message.username != message.original_username:
                    sheet1.write(row, 1, str(
                        message.username) + "(original - " + str(message.original_username) + ")")
                else:
                    sheet1.write(row, 1, message.username)
                if is_original_information_in_reports_enabled and message.original_phone != "" and message.phone != message.original_phone:
                    sheet1.write(row, 2, str(message.phone) +
                                 "(original - " + str(message.original_phone) + ")")
                else:
                    sheet1.write(row, 2, message.phone)
                if is_original_information_in_reports_enabled and message.original_email != "" and message.email != message.original_email:
                    sheet1.write(row, 3, str(message.email) +
                                 "(original - " + str(message.original_email) + ")")
                else:
                    sheet1.write(row, 3, message.email)
                if message.rate_value == -1:
                    sheet1.write(row, 4, "NA")
                else:
                    sheet1.write(row, 4, message.rate_value)
                start_time = (message.joined_date + timedelta(hours=5, minutes=30)).strftime(
                    "%m/%d/%Y, %I:%M:%p")
                sheet1.write(row, 5, start_time)
                end_time = (message.last_appearance_date + timedelta(hours=5, minutes=30)).strftime(
                    "%m/%d/%Y, %I:%M:%p")
                if message.followup_assignment:
                    livechat_followup_cust_obj = LiveChatFollowupCustomer.objects.get(
                        livechat_customer=message)
                    end_time = (livechat_followup_cust_obj.completed_date + timedelta(hours=5, minutes=30)).strftime(
                        "%d/%m/%Y, %I:%M:%p")
                sheet1.write(
                    row, 6, end_time)
                sheet1.write(row, 7, message.get_chat_duration())
                try:
                    sheet1.write(row, 8, message.category.title)
                except Exception:
                    sheet1.write(row, 8, 'None')
                try:
                    sheet1.write(
                        row, 9, message.get_agent_name())
                    sheet1.write(
                        row, 10, message.get_agent_username())
                except Exception:
                    sheet1.write(row, 9, 'None')
                    sheet1.write(row, 10, 'None')
                channel_name = message.channel.name
                if message.followup_assignment:
                    channel_name = message.channel.name + ' (Follow Up)'
                sheet1.write(row, 11, channel_name)
                sheet1.write(row, 12, str(message.session_id))
                sheet1.write(row, 13, message.get_wait_time())
                sheet1.write(row, 14, message.get_agent_first_time_response_time())
                sheet1.write(row, 15, message.chat_ended_by)
                row += 1

            file_path = "livechat-chat-history/" + \
                str(user.user.username) + "/chat_history_" + \
                str(today.date()) + "_" + str(current_time) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            if not os.path.exists('secured_files/LiveChatApp/livechat-chat-history/' + str(user.user.username)):
                os.makedirs('secured_files/LiveChatApp/livechat-chat-history/' + str(user.user.username))
                    
            export_zip_file_path = "secured_files/LiveChatApp/livechat-chat-history/" + \
                str(user.user.username) + "/ChatHistoryToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-chat-history/" + \
                    str(user.user.username) + "/chat_history_" + \
                    str(today.date()) + "_" + str(current_time) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("Chat History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()
            
            from LiveChatApp.models import LiveChatFileAccessManagement
            
            export_zip_file_path = "/" + export_zip_file_path
            file_access_management_obj = LiveChatFileAccessManagement.objects.create(
                file_path=export_zip_file_path, is_public=False)
            file_access_management_obj.file_access_type = "personal_access"
            file_access_management_obj.users.add(user)
            file_access_management_obj.save()

            export_zip_file_path = 'livechat/download-file/' + \
                str(file_access_management_obj.key) + '/LiveChatHistory.zip'

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Chat History for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat Chat History For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Chat History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_chat_history: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        

def generate_transcript_html(customer_obj, message_objs, LiveChatTranslationCache, LiveChatFileAccessManagement):

    domain = settings.EASYCHAT_HOST_URL
    start_date_time = (customer_obj.joined_date + timedelta(hours=5, minutes=30)).strftime(
        "%d/%m/%Y, %I:%M %p")
    start_date = start_date_time.split(",")[0]
    start_time = start_date_time.split(",")[1]
    
    chat_duration = "-"
    if customer_obj.get_chat_duration():
        chat_duration = customer_obj.get_chat_duration()
    html_template = """<html>
    <head>
    <title>Chat Transcript </title>
    <meta charset="UTF-8">
    </head>
    <body>
    """
    
    body_template = """<h2> Chat Transcript</h2><br>"""
    body_template += """Date of Conversation:&nbsp""" + start_date + """<br>"""
    body_template += """Time of Conversation:""" + start_time + """<br>"""
    body_template += """Duration:&nbsp""" + chat_duration + """<br><br>"""

    from LiveChatApp.utils_translation import get_translated_text
    for message in message_objs:
        if message.sender == "System":
            if "customer name" in message.text_message.lower() and "agent name" in message.text_message.lower():
                text_message = customer_obj.agent_id.get_agent_name() + " has joined the chat. Please ask your queries now."
                translated_text = get_translated_text(text_message, customer_obj.customer_language.lang, LiveChatTranslationCache) if customer_obj.customer_language else ""
                body_template += """<p>""" + message.sender_name + """:&nbsp""" + translated_text + """</p>"""
            elif (message.is_video_call_message or message.is_voice_call_message or message.is_cobrowsing_message) and message.message_for == 'customer':
                if message.translated_text:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.translated_text + """</p>"""
                else:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.text_message + """</p>"""
            elif message.is_customer_warning_message:
                if message.translated_text:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.translated_text + """</p>"""
                else:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.text_message + """</p>"""
            elif TRANSCRIPT_MAIL_TEXT in message.text_message:
                translated_text = get_translated_text(message.text_message, customer_obj.customer_language.lang, LiveChatTranslationCache) if customer_obj.customer_language else ""
                body_template += """<p>""" + message.sender_name + """:&nbsp""" + translated_text + """</p>"""
        if message.sender == "Bot":
            body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.text_message + """</p>"""
        if message.sender == "Agent":
            if message.attachment_file_path:
                get_file_access_management_obj_key = message.attachment_file_path.split("/")[3]
                get_file_access_management_obj = LiveChatFileAccessManagement.objects.filter(
                    key=get_file_access_management_obj_key, is_public=False).first()
                get_file_access_management_obj.file_access_type = "all"
                get_file_access_management_obj.is_public_to_all_user = True
                get_file_access_management_obj.save()
                body_template += """<p>""" + message.sender_name + """:&nbsp""" + """<a href='""" + domain + message.attachment_file_path + """'>""" + message.attachment_file_name + """</a></p>"""
            if message.text_message:
                if message.translated_text:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.translated_text + """</p>"""
                else:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.text_message + """</p>"""
        if message.sender == "Customer":
            if message.attachment_file_path:
                if customer_obj.channel.name not in ["Web", "Android", "iOS"]:
                    if domain in message.attachment_file_path: 
                        body_template += """<p>""" + message.sender_name + """:&nbsp""" + """<a href='""" + message.attachment_file_path + """' target="_blank">""" + message.attachment_file_name + """</a></p>"""
                    else:
                        body_template += """<p>""" + message.sender_name + """:&nbsp""" + """<a href='""" + domain + message.attachment_file_path + """' target="_blank">""" + message.attachment_file_name + """</a></p>"""
                else:
                    body_template += """<p>""" + message.sender_name + """:&nbsp""" + """<a href='""" + message.preview_attachment_file_path + """'>""" + message.attachment_file_name + """</a></p>"""
            if message.text_message:
                body_template += """<p>""" + message.sender_name + """:&nbsp""" + message.text_message + """</p>"""
    
    html_template += body_template
    html_template += """</body>
    </html>
    """
    return html_template
            
    
def export_chat_transcript(username, email_id, customer_obj, message_objs, LiveChatTranslationCache, LiveChatFileAccessManagement):
    try:
        domain = settings.EASYCHAT_HOST_URL
        html_template = generate_transcript_html(customer_obj, message_objs, LiveChatTranslationCache, LiveChatFileAccessManagement)

        today = datetime.now()
        current_time = today.strftime("%H:%M:%S")
        
        body = """
            <head>
                <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                <title>Cogno AI</title>
                <style type="text/css" media="screen">
                </style>
            </head>
            <body>

            <div style="padding:1em;border:0.1em black solid;" class="container">
                
            <p>
            Hi! <br> You can find the transcript of your conversation by clicking
            <a href="{}/{}">here</a>
            </p>
            <p>Thanks and Regards</p>
            """
            
        body += """</div></body>"""
        
        if not os.path.exists('secured_files/LiveChatApp/livechat-transcript/' + str(username)):
            os.makedirs('secured_files/LiveChatApp/livechat-transcript/' + str(username))
            
        export_file_path_pdf = "secured_files/LiveChatApp/livechat-transcript/" + \
            str(username) + "/transcript_" + str(today.date()) + "_" + str(current_time) + ".pdf"

        try:
            pdfkit.from_string(html_template, export_file_path_pdf, css=settings.LIVECHAT_ROOT + 'static/LiveChatApp/css/pdf.css')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Transcript Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                'AppName': 'LiveChat'})
        
        export_file_path_pdf = "/" + export_file_path_pdf
        file_access_management_obj = LiveChatFileAccessManagement.objects.create(
            file_path=export_file_path_pdf, is_public=False)
        file_access_management_obj.file_access_type = "all"
        file_access_management_obj.is_public_to_all_user = True
        file_access_management_obj.save()

        export_zip_file_path = 'livechat/download-file/' + \
            str(file_access_management_obj.key) + '/LiveChatTranscript.pdf'

        body = body.format(domain, export_zip_file_path)
        email_subject = "Your Transcript " 
        send_email_to_customer_via_awsses(
            email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_chat_transcript: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_agent_performance_report(user, filter_param, LiveChatSessionManagement):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_agent_performance_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-agent-performance-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-agent-performance-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)
            if not os.path.exists(settings.MEDIA_ROOT + LIVECHAT_AGENT_PERFORMANCE_REPORT_FOLDER + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + LIVECHAT_AGENT_PERFORMANCE_REPORT_FOLDER +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + LIVECHAT_AGENT_PERFORMANCE_REPORT_FOLDER + \
                str(user.user.username) + "/agent_performance_report_" + \
                str(today.date()) + ".xls"

            message_history_objs = LiveChatSessionManagement.objects.filter(
                user__in=agent_list, session_starts_at__date=today.date()).order_by('-session_starts_at')

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Agent Name")
            sheet1.write(0, 1, "Login Date-time")
            sheet1.write(0, 2, "Logout Date-time")
            sheet1.write(0, 3, "Login Duration")
            sheet1.write(0, 4, "Average Handle Time")
            sheet1.write(0, 5, "Average FTR")
            sheet1.write(0, 6, "No Response Count")
            sheet1.write(0, 7, "Idle Time")
            sheet1.write(0, 8, "Not Ready Count")
            sheet1.write(0, 9, "Not Ready Duration")
            sheet1.write(0, 10, "Interaction Count")
            sheet1.write(0, 11, "Total Interaction Duration")
            sheet1.write(0, 12, "Self Assigned Chat Count")
            sheet1.write(0, 13, "Transferred Chat Recieved Count")
            sheet1.write(0, 14, "Transferred Chat Made Count")
            sheet1.write(0, 15, "Total Group Chat Requests")
            sheet1.write(0, 16, "Accepted Group Chats")
            sheet1.write(0, 17, "Declined Group Chats")
            sheet1.write(0, 18, "No Accept/Reject Group Chat")
            sheet1.write(0, 19, "Total Group Chat Duration")
            sheet1.write(0, 20, "Average NPS")

            row = 1
            for message in message_history_objs.iterator():
                sheet1.write(row, 0, message.get_name())
                sheet1.write(row, 1, str((message.session_starts_at.astimezone(
                    pytz.timezone(settings.TIME_ZONE))).strftime("%d/%m/%Y, %I:%M:%p")))
                if message.session_completed:
                    sheet1.write(row, 2, str((message.session_ends_at.astimezone(
                        pytz.timezone(settings.TIME_ZONE))).strftime("%d/%m/%Y, %I:%M:%p")))
                else:
                    sheet1.write(row, 2, '-')
                sheet1.write(
                    row, 3, message.get_session_duration())
                sheet1.write(
                    row, 4, message.get_average_handle_time())
                sheet1.write(
                    row, 5, message.get_average_first_time_response_time())
                sheet1.write(row, 6, message.get_no_response_count())
                sheet1.write(row, 7, message.get_idle_duration())
                sheet1.write(row, 8, message.get_not_ready_count())
                sheet1.write(
                    row, 9, message.get_session_stop_interaction_duration())
                sheet1.write(
                    row, 10, message.get_interaction_count())
                sheet1.write(
                    row, 11, message.get_interaction_duration())
                sheet1.write(row, 12, message.get_self_assigned_chat())
                sheet1.write(
                    row, 13, message.get_total_transferred_chat_received())
                sheet1.write(
                    row, 14, message.get_total_transferred_chat_made())
                sheet1.write(row, 15, message.get_total_group_chat_request())
                sheet1.write(row, 16, message.get_total_group_chat_accept())
                sheet1.write(row, 17, message.get_total_group_chat_reject())
                sheet1.write(
                    row, 18, message.get_total_group_chat_no_response())
                sheet1.write(row, 19, message.get_total_group_chat_duration())
                sheet1.write(row, 20, message.get_average_nps())
                row += 1

            file_path = LIVECHAT_AGENT_PERFORMANCE_REPORT_FOLDER + \
                str(user.user.username) + "/agent_performance_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/" + LIVECHAT_AGENT_PERFORMANCE_REPORT_FOLDER + \
                str(user.user.username) + "/AgentPerformanceReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = LIVECHAT_AGENT_PERFORMANCE_REPORT_FOLDER + \
                    str(user.user.username) + "/agent_performance_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("AgentPerformanceReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Agent Performance Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat agent performance report For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("AgentPerformanceReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_agent_performance_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_daily_interaction_report(user, filter_param, LiveChatCustomer):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_daily_interaction_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-daily-interaction-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-daily-interaction-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)
            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-daily-interaction-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-daily-interaction-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-daily-interaction-report/" + \
                str(user.user.username) + "/daily_interaction_report_" + \
                str(today.date()) + ".xls"

            interaction_objects = []

            interaction_object_list = LiveChatCustomer.objects.filter(
                agent_id__in=agent_list, request_raised_date=today.date())

            query_bot_obj_list = user.bots.all()

            for bot_obj in query_bot_obj_list:
                interaction_objs_for_a_bot_obj = interaction_object_list.filter(bot=bot_obj)
                interaction_objects.append({
                    "date": today, 
                    "bot_name": bot_obj.name, 
                    "frequency": interaction_objs_for_a_bot_obj.count()
                })

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Bot Name")
            sheet1.write(0, 2, "Interaction Count")

            row = 1
            for index in range(0, len(interaction_objects)):
                sheet1.write(row, 0, interaction_objects[index]["date"].strftime("%d-%m-%Y"))
                sheet1.write(row, 1, interaction_objects[
                             index]["bot_name"])
                sheet1.write(row, 2, interaction_objects[
                             index]["frequency"])
                row += 1

            file_path = "livechat-daily-interaction-report/" + \
                str(user.user.username) + "/daily_interaction_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-daily-interaction-report/" + \
                str(user.user.username) + "/DailyInteractionReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-daily-interaction-report/" + \
                    str(user.user.username) + "/daily_interaction_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("DailyInteractionReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Daily Interaction Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                      """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat DailyInteractionReport For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("DailyInteractionReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_daily_interaction_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_hourly_interaction_report(user, filter_param, LiveChatCustomer):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_hourly_interaction_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-hourly-interaction-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-hourly-interaction-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)
            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-hourly-interaction-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-hourly-interaction-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-hourly-interaction-report/" + \
                str(user.user.username) + "/hourly_interaction_report_" + \
                str(today.date()) + ".xls"

            interaction_objects = []

            interaction_object_list = LiveChatCustomer.objects.filter(
                agent_id__in=agent_list, last_appearance_date__date=today.date())

            query_bot_obj_list = user.bots.all()

            for bot_obj in query_bot_obj_list:
                start_time_obj = today + \
                    timedelta(hours=5, minutes=30)
                start_time_obj = today.time().replace(hour=0, minute=0, second=0)
                for itr in range(1, 25):
                    end_time_obj = today.time().replace(hour=itr % 24, minute=0, second=0)
                    interaction_count = interaction_object_list.filter(bot=bot_obj, request_raised_date=today.date(
                    ), joined_date__time__range=[start_time_obj, end_time_obj]).count()
                    interaction_objects.append({"date": today, "start_time": start_time_obj,
                                                "end_time": end_time_obj, "bot_name": bot_obj.name, "frequency": interaction_count})
                    start_time_obj = today.time().replace(hour=itr % 24, minute=0, second=1)

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Bot Name")
            sheet1.write(0, 2, "Start Time")
            sheet1.write(0, 3, "End Time")
            sheet1.write(0, 4, "Interaction Count")

            row = 1
            for index in range(0, len(interaction_objects)):
                sheet1.write(row, 0, str(
                    interaction_objects[index]["date"].date().strftime(DATE_DD_MM_YYYY)))
                sheet1.write(row, 1, interaction_objects[
                             index]["bot_name"])
                sheet1.write(row, 2, str(
                    interaction_objects[index]["start_time"]))
                sheet1.write(row, 3, str(
                    interaction_objects[index]["end_time"]))
                sheet1.write(row, 4, interaction_objects[
                             index]["frequency"])
                row += 1

            file_path = "livechat-hourly-interaction-report/" + \
                str(user.user.username) + "/hourly_interaction_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-hourly-interaction-report/" + \
                str(user.user.username) + "/HourlyInteractionReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-hourly-interaction-report/" + \
                    str(user.user.username) + "/hourly_interaction_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("HourlyInteractionReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Hourly Interaction Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat HourlyInteractionReport For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("HourlyInteractionReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_hourly_interaction_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_agent_not_ready_report(user, filter_param, LiveChatAgentNotReady):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_agent_not_ready_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-agent-not-ready-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-agent-not-ready-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)
            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-agent-not-ready-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-agent-not-ready-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-agent-not-ready-report/" + \
                str(user.user.username) + "/agent_not_ready_report_" + \
                str(today.date()) + ".xls"

            message_history_objs = LiveChatAgentNotReady.objects.filter(
                not_ready_starts_at__date=today.date(), user__in=agent_list)
            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Date")
            sheet1.write(0, 1, "Agent Name")
            sheet1.write(0, 2, "Agent Username")
            sheet1.write(0, 3, "Reason")
            sheet1.write(0, 4, "Stop Interaction Duration")
            sheet1.write(0, 5, "Start Time")
            sheet1.write(0, 6, "End Time")
            sheet1.write(0, 7, "Offline Time")
            row = 1
            for message in message_history_objs:
                message.not_ready_starts_at = message.not_ready_starts_at + \
                    timedelta(hours=5, minutes=30)
                message.not_ready_ends_at = message.not_ready_ends_at + \
                    timedelta(hours=5, minutes=30)
                sheet1.write(row, 0, str(
                    message.not_ready_starts_at.date().strftime(DATE_DD_MM_YYYY)))
                sheet1.write(row, 1, str(
                    message.user.user.first_name) + " " + str(message.user.user.last_name))
                sheet1.write(row, 2, message.user.user.username)
                sheet1.write(row, 3, message.get_reason_for_offline())
                sheet1.write(row, 4, message.get_stop_interaction_duration())
                sheet1.write(
                    row, 5, str(message.not_ready_starts_at.time().strftime("%H:%M:%S")))
                sheet1.write(row, 6, str(
                    message.not_ready_ends_at.time().strftime("%H:%M:%S")))
                sheet1.write(row, 7, message.get_offline_duration())
                row += 1

            file_path = "livechat-agent-not-ready-report/" + \
                str(user.user.username) + "/agent_not_ready_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-agent-not-ready-report/" + \
                str(user.user.username) + "/AgentNotReadyReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-agent-not-ready-report/" + \
                    str(user.user.username) + "/agent_not_ready_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("AgentNotReadyReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Agent Not Ready Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat AgentNotReadyReport For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("AgentNotReadyReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_agent_not_ready_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_missed_chats_report(user, filter_param, LiveChatCustomer):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_offline_message_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-missed-chats-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-missed-chats-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-missed-chats-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-missed-chats-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-missed-chats-report/" + \
                str(user.user.username) + "/missed_chats_report_" + \
                str(today.date()) + ".xls"
            category_list = user.category.all()
            message_history_objs = LiveChatCustomer.objects.filter(
                last_appearance_date__date=today.date(), is_denied=True, is_system_denied=False, bot__in=user.bots.all(), category__in=category_list)

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Created On")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Message")
            sheet1.write(0, 5, "Channel")
            sheet1.write(0, 6, "Category")

            row = 1
            for message in message_history_objs:
                sheet1.write(row, 0, str((message.joined_date + timedelta(hours=5, minutes=30)).strftime(
                    DATE_DD_MM_YYYY_TIME_HH_MIN_P)))
                sheet1.write(row, 1, message.username)
                sheet1.write(row, 2, message.phone)
                sheet1.write(row, 3, message.email)
                sheet1.write(row, 4, message.message)
                sheet1.write(row, 5, message.channel.name)
                try:
                    sheet1.write(row, 6, message.category.title)
                except Exception:
                    sheet1.write(row, 6, 'None')
                row += 1

            file_path = "livechat-missed-chats-report/" + \
                str(user.user.username) + "/missed_chats_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-missed-chats-report/" + \
                str(user.user.username) + "/MissedChatsReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-missed-chats-report/" + \
                    str(user.user.username) + "/missed_chats_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("MissedChatsReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Missed Chats Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat Missed Chats History For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("MissedChatsReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_missed_chats_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_offline_chats_report(user, filter_param, LiveChatCustomer):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_abandoned_chats_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-offline-chats-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-offline-chats-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-offline-chats-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-offline-chats-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-offline-chats-report/" + \
                str(user.user.username) + "/offline_chats_report_" + \
                str(today.date()) + ".xls"
            category_list = user.category.all()
            message_history_objs = LiveChatCustomer.objects.filter(
                last_appearance_date__date=today.date(), is_system_denied=True, bot__in=user.bots.all(), category__in=category_list)

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Created On")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Message")
            sheet1.write(0, 5, "Channel")
            sheet1.write(0, 6, "Category")

            row = 1
            for message in message_history_objs:
                sheet1.write(row, 0, str((message.joined_date + timedelta(hours=5, minutes=30)).strftime(
                    DATE_DD_MM_YYYY_TIME_HH_MIN_P)))
                sheet1.write(row, 1, message.username)
                sheet1.write(row, 2, message.phone)
                sheet1.write(row, 3, message.email)
                sheet1.write(row, 4, message.message)
                sheet1.write(row, 5, message.channel.name)
                try:
                    sheet1.write(row, 6, message.category.title)
                except Exception:
                    sheet1.write(row, 6, 'None')
                row += 1

            file_path = "livechat-offline-chats-report/" + \
                str(user.user.username) + "/offline_chats_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-offline-chats-report/" + \
                str(user.user.username) + "/OfflineChatsReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-offline-chats-report/" + \
                    str(user.user.username) + "/offline_chats_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("OfflineChatsReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Offline Chats Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat Offline Chats History For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("OfflineChatsReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_offline_chats_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_abandoned_chats_report(user, filter_param, LiveChatCustomer):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_total_declined_chats_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-abandoned-chats-report'):
            os.makedirs(settings.MEDIA_ROOT +
                        'livechat-abandoned-chats-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-abandoned-chats-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-abandoned-chats-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-abandoned-chats-report/" + \
                str(user.user.username) + "/abandoned_chats_report_" + \
                str(today.date()) + ".xls"
            category_list = user.category.all()
            message_history_objs = LiveChatCustomer.objects.filter(
                last_appearance_date__date=today.date(), agent_id=None, abruptly_closed=True, bot__in=user.bots.all(), category__in=category_list)

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Created On")
            sheet1.write(0, 1, "Customer Name")
            sheet1.write(0, 2, "Mobile Number")
            sheet1.write(0, 3, "Email")
            sheet1.write(0, 4, "Message")
            sheet1.write(0, 5, "Channel")
            sheet1.write(0, 6, "Category")
            sheet1.write(0, 7, "Wait Time")
            row = 1
            for message in message_history_objs:
                sheet1.write(row, 0, str((message.joined_date + timedelta(hours=5, minutes=30)).strftime(
                    DATE_DD_MM_YYYY_TIME_HH_MIN_P)))
                sheet1.write(row, 1, message.username)
                sheet1.write(row, 2, message.phone)
                sheet1.write(row, 3, message.email)
                sheet1.write(row, 4, message.message)
                sheet1.write(row, 5, message.channel.name)
                try:
                    sheet1.write(row, 6, message.category.title)
                except Exception:
                    sheet1.write(row, 6, 'None')
                sheet1.write(row, 7, message.get_wait_time())
                row += 1

            file_path = "livechat-abandoned-chats-report/" + \
                str(user.user.username) + "/abandoned_chats_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-abandoned-chats-report/" + \
                str(user.user.username) + "/AbandonedChatsReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-abandoned-chats-report/" + \
                    str(user.user.username) + "/abandoned_chats_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("AbandonedChatsReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Abandoned Chats Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat Abandoned Chats Report For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Abandoned Chats Report Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_abandoned_chats_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def export_today_login_logout_report(user, filter_param, LiveChatSessionManagement):
    try:
        domain = settings.EASYCHAT_HOST_URL
        message_history_datadump_log = open(
            "log/livechat_login_logout_report_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-login-logout-report'):
            os.makedirs(settings.MEDIA_ROOT + 'livechat-login-logout-report')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)
            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-login-logout-report/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-login-logout-report/' +
                            str(user.user.username))

            file_path = settings.MEDIA_ROOT + "livechat-login-logout-report/" + \
                str(user.user.username) + "/login_logout_report_" + \
                str(today.date()) + ".xls"

            message_history_objs = LiveChatSessionManagement.objects.filter(
                session_starts_at__date=today.date(), user__in=agent_list).order_by('-session_starts_at')

            test_wb = Workbook()

            sheet1 = test_wb.add_sheet("Sheet1")

            sheet1.write(0, 0, "Agent Username")
            sheet1.write(0, 1, "Date")
            sheet1.write(0, 2, "Duration")
            sheet1.write(0, 3, "Agent Name")
            sheet1.write(0, 4, "Login Time")
            sheet1.write(0, 5, "Logout Time")
            sheet1.write(0, 6, "Total Offline Time")
            sheet1.write(0, 7, "Total Online Time")

            row = 1
            for message in message_history_objs:
                message.session_starts_at = message.session_starts_at + \
                    timedelta(hours=5, minutes=30)
                message.session_ends_at = message.session_ends_at + \
                    timedelta(hours=5, minutes=30)
                sheet1.write(row, 0, message.user.user.username)
                sheet1.write(row, 1, str(
                    message.session_starts_at.date().strftime(DATE_DD_MM_YYYY)))
                sheet1.write(row, 2, str(
                    message.get_session_duration()))
                sheet1.write(row, 3, message.get_name())
                sheet1.write(row, 4, str(
                    message.session_starts_at.time().strftime("%H:%M:%S")))
                sheet1.write(row, 5, str(
                    message.session_ends_at.time().strftime("%H:%M:%S")))
                sheet1.write(row, 6, str(
                    message.get_total_offline_time()))
                sheet1.write(row, 7, str(
                    message.get_total_online_time()))
                row += 1

            file_path = "livechat-login-logout-report/" + \
                str(user.user.username) + "/login_logout_report_" + \
                str(today.date()) + ".xls"
            test_wb.save(settings.MEDIA_ROOT + file_path)

            export_zip_file_path = "files/livechat-login-logout-report/" + \
                str(user.user.username) + "/LoginLogoutReportToday-" + \
                str(today.time().replace(microsecond=0)) + ".zip"

            zip_obj = ZipFile(export_zip_file_path, 'w')

            try:
                file_path = "livechat-login-logout-report/" + \
                    str(user.user.username) + "/login_logout_report_" + \
                    str(today.date()) + '.xls'
                zip_obj.write(settings.MEDIA_ROOT +
                              file_path, basename(file_path))
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("LoginLogoutReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                    'AppName': 'LiveChat'})

            zip_obj.close()

            body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Login Logout Report for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
            developer_console_config = get_developer_console_settings()

            body += developer_console_config.custom_report_template_signature

            body += """</div></body>"""

            start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)
            for email_id in email_list:
                body = body.format(email_id, str(start_date),
                                   domain, export_zip_file_path)
                email_subject = "LiveChat LoginLogoutReport For " + \
                    str(email_id)
                send_email_to_customer_via_awsses(
                    email_id, email_subject, body)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("LoginLogoutReport Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_login_logout_report: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def check_for_system_commands(code, bot_pk, LiveChatConfig, Bot):
    try:
        config_obj = LiveChatConfig.objects.get(
            bot=Bot.objects.get(pk=int(bot_pk)))
        system_commands = json.loads(
            config_obj.system_commands.replace("'", '"'))

        for command in system_commands:
            if command in code:
                return True
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_for_system_commands: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyChat', 'user_id': 'None', 'source': 'None', 'channel': 'None', 'bot_id': 'None'})

    return False


def open_file(file_dir, method):
    try:
        file_dir = settings.SECURE_MEDIA_ROOT + file_dir

        if '..' in file_dir:
            logger.error("user is trying to access this file: %s", str(file_dir), extra={
                'AppName': 'EasyChat', 'user_id': 'None', 'source': 'None', 'channel': 'None', 'bot_id': 'None'})
            return

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("open_file: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyChat', 'user_id': 'None', 'source': 'None', 'channel': 'None', 'bot_id': 'None'})

    return open(file_dir, method)


def get_agent_token_based_on_username(username):
    try:
        access_token = username.strip().lower()
        access_token = hashlib.md5(access_token.encode()).hexdigest()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("error in get_agent_token  %s %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return str(access_token)


def get_user_obj_dict(user_obj, user, user_group_obj, LiveChatInternalMISDashboard, LiveChatInternalChatLastSeen):
    try:
        user_obj_dict = {}
        username = user_obj.user.username
        full_name = str(user_obj.user.first_name) + \
            " " + str(user_obj.user.last_name)
        websocket_token = get_agent_token_based_on_username(username)
        user_obj_dict["username"] = username
        user_obj_dict["websocket_token"] = websocket_token
        user_obj_dict["full_name"] = full_name.strip()
        user_obj_dict['is_user_group'] = False

        last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=user, user_group=user_group_obj)

        unread_message_count = 0
        if last_seen_obj:
            last_seen_obj = last_seen_obj.first()
            unread_message_count = LiveChatInternalMISDashboard.objects.filter(sender=user_obj, receiver=user, user_group=user_group_obj, message_datetime__gt=last_seen_obj.last_seen).count()
        else:
            unread_message_count = LiveChatInternalMISDashboard.objects.filter(sender=user_obj, receiver=user, user_group=user_group_obj).count()

        user_obj_dict['unread_message_count'] = unread_message_count
        last_message_obj = LiveChatInternalMISDashboard.objects.filter(
            sender__in=[user, user_obj], receiver__in=[user, user_obj], user_group=user_group_obj).order_by('-pk')

        if last_message_obj:
            user_obj_dict['last_msg_text'] = last_message_obj[0].message_info.message_text
            user_obj_dict['filename'] = last_message_obj[0].message_info.attached_file_name
            user_obj_dict['sender'] = last_message_obj[0].sender_name
            user_obj_dict['last_msg_time'] = get_date(
                last_message_obj[0].message_datetime)
            user_obj_dict['is_chat_started'] = True
        else:
            user_obj_dict['last_msg_text'] = 'No messages yet'
            user_obj_dict['filename'] = ''
            user_obj_dict['sender'] = ''
            user_obj_dict['last_msg_time'] = ''
            user_obj_dict['is_chat_started'] = user_group_obj.is_chat_started

        return user_obj_dict
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_user_obj_dict: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat', 'user_id': 'None', 'source': 'None', 'channel': 'None', 'bot_id': 'None'})


def check_and_create_bulk_user_groups(user_objs, user, LiveChatInternalUserGroup):
    try:
        for user_obj in user_objs:
            if not user_obj or user_obj.is_deleted:
                continue

            check_and_update_user_group(
                user_obj, user, LiveChatInternalUserGroup)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_dictonary_of_user_objects: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat', 'user_id': 'None', 'source': 'None', 'channel': 'None', 'bot_id': 'None'})


def check_and_update_user_group(user_obj, user, LiveChatInternalUserGroup):
    try:
        user_group_obj = LiveChatInternalUserGroup.objects.filter(is_converted_into_group=False, members__in=[
                                                                  user_obj, user]).annotate(num_attr=Count('members')).filter(num_attr=len([user_obj, user]))

        if not user_group_obj:
            user_group_obj = LiveChatInternalUserGroup.objects.create()

            user_group_obj.members.add(user_obj)
            user_group_obj.members.add(user)
        else:
            user_group_obj = user_group_obj.first()

        return user_group_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_and_update_user_group: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat'})


def get_group_chat_list(group_chat_objs, user, LiveChatInternalMISDashboard, LiveChatInternalChatGroupMembers, LiveChatInternalChatLastSeen):
    try:
        group_chat_list = []
        count = 0
        for group_chat_obj in group_chat_objs:
            group_member_obj = LiveChatInternalChatGroupMembers.objects.get(
                group=group_chat_obj, user=user)

            if group_member_obj.is_deleted:
                continue

            last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=user, group=group_chat_obj)

            unread_message_count = 0

            if group_member_obj.is_removed:
                remove_datetime = group_member_obj.remove_datetime
                last_msg_obj = LiveChatInternalMISDashboard.objects.filter(
                    group=group_chat_obj, message_datetime__lte=remove_datetime).order_by('-pk')

                if last_seen_obj:
                    last_seen_obj = last_seen_obj.first()
                    unread_message_count = LiveChatInternalMISDashboard.objects.filter(group=group_chat_obj, message_datetime__lte=remove_datetime, message_datetime__gt=last_seen_obj.last_seen).exclude(sender=user).count()
                else:
                    unread_message_count = LiveChatInternalMISDashboard.objects.filter(group=group_chat_obj, message_datetime__lte=remove_datetime).exclude(sender=user).count()

            elif group_member_obj.has_left:
                leave_datetime = group_member_obj.left_datetime
                last_msg_obj = LiveChatInternalMISDashboard.objects.filter(
                    group=group_chat_obj, message_datetime__lte=leave_datetime).order_by('-pk')
    
                if last_seen_obj:
                    last_seen_obj = last_seen_obj.first()
                    unread_message_count = LiveChatInternalMISDashboard.objects.filter(group=group_chat_obj, message_datetime__lte=leave_datetime, message_datetime__gt=last_seen_obj.last_seen).exclude(sender=user).count()
                else:
                    unread_message_count = LiveChatInternalMISDashboard.objects.filter(group=group_chat_obj, message_datetime__lte=leave_datetime).exclude(sender=user).count()
                    
            else:
                last_msg_obj = LiveChatInternalMISDashboard.objects.filter(
                    group=group_chat_obj).order_by('-pk')

                if last_seen_obj:
                    last_seen_obj = last_seen_obj.first()
                    unread_message_count = LiveChatInternalMISDashboard.objects.filter(group=group_chat_obj, message_datetime__gt=last_seen_obj.last_seen).exclude(sender=user).count()
                else:
                    unread_message_count = LiveChatInternalMISDashboard.objects.filter(group=group_chat_obj).exclude(sender=user).count()
                    
            last_msg_dict = {}
            if last_msg_obj:
                last_msg_dict['text'] = last_msg_obj[0].message_info.message_text
                last_msg_dict['filename'] = last_msg_obj[0].message_info.attached_file_name
                last_msg_dict['sender'] = last_msg_obj[0].sender_name

                if last_msg_obj[0].sender:
                    last_msg_dict['sender_username'] = last_msg_obj[0].sender.user.username
                else:
                    last_msg_dict['sender_username'] = 'Deleted User'

                last_msg_dict['time'] = get_date(
                    last_msg_obj[0].message_datetime)

            else:
                last_msg_dict['text'] = 'No messages yet'
                last_msg_dict['filename'] = ''
                last_msg_dict['sender'] = group_chat_obj.group_name
                last_msg_dict['sender_username'] = ''
                last_msg_dict['time'] = get_date(
                    group_chat_obj.created_datetime)

            group_chat_list.append({
                'id': str(group_chat_obj.group_id),
                'name': group_chat_obj.group_name.strip() if group_chat_obj.group_name.strip() != "" else "NA",
                'desc': group_chat_obj.group_description,
                'icon_path': group_chat_obj.icon_path,
                'created_date': str(group_chat_obj.created_datetime.strftime('%d/%m/%y')),
                'is_deleted': group_chat_obj.is_deleted,
                'admin': group_chat_obj.created_by.user.username,
                'last_message': last_msg_dict,
                'unread_message_count': unread_message_count
            })

            group_members = []
            for group_member in group_chat_obj.members.all():
                last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=group_member.user, group=group_chat_obj)

                last_seen_on_chat = 0
                if last_seen_obj:
                    last_seen_on_chat = last_seen_obj.first().last_seen.timestamp() * 1000
                
                if group_member.user:
                    last_seen_time = 0
                    if group_member.user.last_updated_time:
                        last_seen_time = group_member.user.last_updated_time.timestamp() * 1000

                    group_members.append({
                        'name': group_member.user.user.username,
                        'is_supervisor': group_member.user.status == '2',
                        'is_admin': group_member.user.status == '1',
                        'is_removed': group_member.is_removed,
                        'has_left': group_member.has_left,
                        'is_deleted': group_member.is_deleted,
                        'last_seen_time': last_seen_time,
                        'last_seen_on_chat': last_seen_on_chat
                    })

            group_chat_list[count]['members'] = group_members
            count += 1

        return group_chat_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_group_chat_list: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat'})

    return []


def get_user_group_dict(user_group_obj, livechat_user, LiveChatInternalMISDashboard, LiveChatInternalChatLastSeen):
    try:
        user_group_dict = {}

        last_msg_obj = LiveChatInternalMISDashboard.objects.filter(
            user_group=user_group_obj).order_by('-pk')

        last_msg_dict = {}
        if last_msg_obj:
            last_msg_dict['text'] = last_msg_obj[0].message_info.message_text
            last_msg_dict['filename'] = last_msg_obj[0].message_info.attached_file_name
            last_msg_dict['sender'] = last_msg_obj[0].sender_name
            last_msg_dict['sender_username'] = last_msg_obj[0].sender.user.username
            last_msg_dict['time'] = get_date(
                last_msg_obj[0].message_datetime)

        else:
            last_msg_dict['text'] = 'No messages yet'
            last_msg_dict['filename'] = ''
            last_msg_dict['sender'] = str(user_group_obj.chat_belong_to)
            last_msg_dict['sender_username'] = ''
            last_msg_dict['time'] = get_date(
                user_group_obj.created_datetime)
        
        last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=livechat_user, user_group=user_group_obj)

        unread_message_count = 0
        if last_seen_obj:
            last_seen_obj = last_seen_obj.first()
            unread_message_count = LiveChatInternalMISDashboard.objects.filter(user_group=user_group_obj, message_datetime__gt=last_seen_obj.last_seen).exclude(sender=livechat_user).count()
        else:
            unread_message_count = LiveChatInternalMISDashboard.objects.filter(user_group=user_group_obj).exclude(sender=livechat_user).count()

        user_group_dict = {
            'is_user_group': True,
            'id': str(user_group_obj.group_id),
            'full_name': f'{str(user_group_obj.chat_belong_to.user.first_name)} {str(user_group_obj.chat_belong_to.user.last_name)}',
            'username': str(user_group_obj.chat_belong_to),
            'created_date': str(user_group_obj.created_datetime.strftime('%d/%m/%y')),
            'last_message': last_msg_dict,
            'unread_message_count': unread_message_count,
        }

        group_members = []
        for group_member in user_group_obj.members.all():
            last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=group_member, user_group=user_group_obj)

            last_seen_on_chat = 0
            if last_seen_obj:
                last_seen_on_chat = last_seen_obj.first().last_seen.timestamp() * 1000

            if group_member:
                last_seen_time = 0
                if group_member.last_updated_time:
                    last_seen_time = group_member.last_updated_time.timestamp() * 1000

                group_members.append({
                    'name': group_member.user.username,
                    'is_supervisor': group_member.status == '2',
                    'is_admin': group_member.status == '1',
                    'last_seen_time': last_seen_time,
                    'last_seen_on_chat': last_seen_on_chat
                })

        user_group_dict['members'] = group_members

        return user_group_dict
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_user_group_dict: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat'})


def get_user_group_list(user_group_objs, livechat_user, LiveChatInternalMISDashboard, LiveChatInternalChatLastSeen):
    try:
        group_chat_list = []
        for user_group in user_group_objs:
            if user_group.is_converted_into_group:
                if user_group.chat_belong_to == None:
                    continue

                user_dict = get_user_group_dict(user_group, livechat_user, LiveChatInternalMISDashboard, LiveChatInternalChatLastSeen)
            else:
                if user_group.members.count() != 2:
                    continue

                user_objs = user_group.members.all()
                if livechat_user == user_objs.first():
                    user_obj = user_objs[1]
                else:
                    user_obj = user_objs[0]

                user_dict = get_user_obj_dict(user_obj, livechat_user, user_group, LiveChatInternalMISDashboard, LiveChatInternalChatLastSeen)
            
            if isinstance(user_dict, dict):
                group_chat_list.append(user_dict)
                
        return group_chat_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_user_group_list: %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat'})

    return []


def get_agents_and_supervisor_under_this_user(user_obj):
    agent_obj_list = []
    if user_obj.status == "3":
        agent_obj_list.append(user_obj)

    try:
        for user in user_obj.agents.filter(is_deleted=False):
            if user.status == "2":
                agent_obj_list.append(user)
                for user1 in user.agents.filter(is_deleted=False):
                    agent_obj_list.append(user1)
            elif user.status == "3":
                agent_obj_list.append(user)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agents_under_this_user: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return list(set(agent_obj_list))


def get_other_supervisors(user_obj, admin):
    supervisor_list = []
    try:
        for user in admin.agents.filter(is_deleted=False, status="2"):
            if user != user_obj:
                supervisor_list.append(user)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_other_supervisors: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return list(set(supervisor_list))


def get_internal_group_message_history(group_obj, curr_user, chat_index, LiveChatInternalMISDashboard, LiveChatInternalChatGroupMembers):
    message_history = []
    all_chat_loaded = False

    try:
        group_member_obj = LiveChatInternalChatGroupMembers.objects.get(
            group=group_obj, user=curr_user)

        if group_member_obj.is_removed:
            remove_datetime = group_member_obj.remove_datetime
            livechat_mis_objs = LiveChatInternalMISDashboard.objects.filter(
                group=group_obj, message_datetime__lte=remove_datetime).order_by('-message_datetime', '-pk')
        elif group_member_obj.has_left:
            leave_datetime = group_member_obj.left_datetime
            livechat_mis_objs = LiveChatInternalMISDashboard.objects.filter(
                group=group_obj, message_datetime__lte=leave_datetime).order_by('-message_datetime', '-pk')
        else:
            livechat_mis_objs = LiveChatInternalMISDashboard.objects.filter(
                group=group_obj).order_by('-message_datetime', '-pk')

        max_index = min(chat_index + 50, livechat_mis_objs.count())

        if max_index >= livechat_mis_objs.count():
            all_chat_loaded = True

        livechat_mis_objs = livechat_mis_objs[chat_index:max_index]
        for livechat_mis_obj in livechat_mis_objs:
            text_message = livechat_mis_obj.message_info.message_text
            attachment_file_path = livechat_mis_obj.message_info.attached_file_src
            preview_file_path = livechat_mis_obj.message_info.preview_file_src
            thumbnail_file_path = livechat_mis_obj.message_info.thumbnail_file_src

            if livechat_mis_obj.sender_name.lower() == 'system':
                sender = 'System'
            elif livechat_mis_obj.sender == curr_user:
                sender = "CurrentUser"
            else:
                sender = "ReceiverUser"

            if livechat_mis_obj.sender:
                sender_username = livechat_mis_obj.sender.user.username
            else:
                sender_username = 'Deleted User'

            message_history.append({"message": text_message, "sender": sender, "sender_name": livechat_mis_obj.sender_name, "time": str(get_time(
                livechat_mis_obj.message_datetime)), "time_in_minisec": str(get_miniseconds_datetime(livechat_mis_obj.message_datetime)), "attached_file_src": attachment_file_path,
                "file": thumbnail_file_path, "preview_attachment_file_path": preview_file_path, "sender_username": sender_username})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_internal_group_message_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return message_history, all_chat_loaded


def get_internal_user_group_message_history(user_group_obj, curr_user, chat_index, LiveChatInternalMISDashboard):
    message_history = []
    all_chat_loaded = False

    try:
        livechat_mis_objs = LiveChatInternalMISDashboard.objects.filter(
            user_group=user_group_obj).order_by('-message_datetime', '-pk')

        max_index = min(chat_index + 50, livechat_mis_objs.count())

        if max_index >= livechat_mis_objs.count():
            all_chat_loaded = True

        livechat_mis_objs = livechat_mis_objs[chat_index:max_index]
        for livechat_mis_obj in livechat_mis_objs:
            text_message = livechat_mis_obj.message_info.message_text
            attachment_file_path = livechat_mis_obj.message_info.attached_file_src
            preview_file_path = livechat_mis_obj.message_info.preview_file_src
            thumbnail_file_path = livechat_mis_obj.message_info.thumbnail_file_src

            if livechat_mis_obj.sender_name.lower() == 'system':
                sender = 'System'
            elif livechat_mis_obj.sender == curr_user:
                sender = "CurrentUser"
            else:
                sender = "ReceiverUser"

            message_history.append({"message": text_message, "sender": sender, "sender_name": livechat_mis_obj.sender_name, "time": str(get_time(
                livechat_mis_obj.message_datetime)), "time_in_minisec": str(get_miniseconds_datetime(livechat_mis_obj.message_datetime)), "attached_file_src": attachment_file_path,
                "file": thumbnail_file_path, "preview_attachment_file_path": preview_file_path, "sender_username": livechat_mis_obj.sender.user.username})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_internal_user_group_message_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return message_history, all_chat_loaded


def ensure_element_tree(xlrd):
    try:
        xlrd.xlsx.ensure_elementtree_imported(False, None)
        xlrd.xlsx.Element_has_iter = True
    except Exception:
        pass


def parse_audit_object_list(audit_objs):
    try:
        tz = pytz.timezone(settings.TIME_ZONE)

        audit_obj_list = []
        for audit_obj in audit_objs:

            is_followup_lead = False
            if audit_obj.followup_assignment:
                is_followup_lead = True

            audit_obj_list.append({
                'pk': str(audit_obj.pk),
                'username': audit_obj.get_username(),
                'agent_username': audit_obj.get_agent_username(),
                'channel': audit_obj.channel.name,
                'source': audit_obj.get_source_name_from_choice_field(),
                'closing_category': audit_obj.get_closing_category_title(),
                'rate_value': audit_obj.rate_value,
                'rate_text': audit_obj.nps_text_feedback,
                'form_filled': audit_obj.form_filled,
                'joined_date': audit_obj.joined_date.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'last_appearance_date': audit_obj.get_last_appearance_date(),
                'chat_duration': audit_obj.get_chat_duration(),
                'wait_time': audit_obj.get_wait_time(),
                'is_session_exp': audit_obj.is_session_exp,
                'chat_ended_by': audit_obj.chat_ended_by,
                'is_followup_lead': is_followup_lead,
                'previous_channel': audit_obj.previous_channel,
                'first_time_response': audit_obj.get_agent_first_time_response_time(),
            })

        return audit_obj_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_audit_object_list %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return []


def parse_livechat_cust_objs(cust_objs, LiveChatConfig):
    try:
        tz = pytz.timezone(settings.TIME_ZONE)

        cust_obj_list = []
        for cust_obj in cust_objs:
            config_obj = LiveChatConfig.objects.get(bot=cust_obj.bot)
            cust_obj_list.append({
                'pk': str(cust_obj.pk),
                'username': cust_obj.get_username(),
                'channel': cust_obj.channel.name,
                'category': cust_obj.category.title,
                'joined_date': cust_obj.joined_date.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'timestamp': cust_obj.joined_date.astimezone(tz).strftime('%m/%d/%Y, %I:%M:%S %p'),
                'queue_timer': config_obj.queue_timer,
            })

        return cust_obj_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_livechat_cust_objs %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return []


def parse_and_get_followup_leads(cust_objs):
    try:
        tz = pytz.timezone(settings.TIME_ZONE)

        cust_obj_list = []
        for cust_obj in cust_objs:
            cust_obj_list.append({
                'pk': str(cust_obj.livechat_customer.pk),
                'username': cust_obj.livechat_customer.get_username(),
                'email': cust_obj.livechat_customer.email,
                'phone': cust_obj.livechat_customer.phone,
                'category': cust_obj.livechat_customer.category.get_category_name(),
                'assigned_date': cust_obj.livechat_customer.joined_date.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'source': cust_obj.get_lead_source(),
                'agent': cust_obj.agent_id.get_agent_name(),
                'channel': cust_obj.livechat_customer.channel.name,
                'bot_pk': cust_obj.livechat_customer.bot.pk,
            })

        return cust_obj_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_and_get_followup_leads %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return []


def parse_and_get_reported_blocked_users(cust_objs):
    try:
        cust_obj_list = []
        for cust_obj in cust_objs:

            time_diff = "0"
            if cust_obj.is_blocked:
                time_diff = get_time_left_to_unblock(cust_obj.blocked_date, cust_obj.block_duration) - timezone.now()
                days, seconds = time_diff.days, time_diff.seconds
                hours = days * 24 + seconds // 3600
                if days > 0:
                    time_diff = str(days) + " day/s" 
                else:
                    if hours > 0:
                        time_diff = str(hours) + " hour/s"
                    else: 
                        time_diff = "1 hour/s"

                if check_if_user_should_be_unblocked(cust_obj):
                    continue

            cust_obj_list.append({
                'pk': str(cust_obj.livechat_customer.pk),
                'username': cust_obj.livechat_customer.get_username(),
                'email': cust_obj.livechat_customer.email,
                'phone': cust_obj.livechat_customer.phone,
                'agent': cust_obj.livechat_customer.agent_id.get_agent_name(),
                'channel': cust_obj.livechat_customer.channel.name,
                'ip_address': cust_obj.livechat_customer.ip_address,
                'client_id': cust_obj.livechat_customer.client_id,
                'time_diff': time_diff,
            })

        return cust_obj_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_and_get_reported_blocked_users %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return []


def get_time_left_to_unblock(blocked_date, block_duration):
    try: 

        if block_duration == "1":
            return blocked_date + timedelta(minutes=30)
        elif block_duration == "2":
            return blocked_date + timedelta(minutes=60)
        elif block_duration == "3":
            return blocked_date + timedelta(days=1)
        elif block_duration == "4":
            return blocked_date + timedelta(days=7)
        elif block_duration == "5":
            return blocked_date + timedelta(days=14)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_time_left_to_unblock %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return timezone.now()


def check_if_user_should_be_unblocked(cust_obj):
    try:

        if timezone.now() >= get_time_left_to_unblock(cust_obj.blocked_date, cust_obj.block_duration):

            cust_obj.is_completed = True
            cust_obj.save()

            return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error check_if_user_should_be_unblocked %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return False


def parse_voip_history_object_list(voip_data_objs, LiveChatFileAccessManagement):
    try:
        tz = pytz.timezone(settings.TIME_ZONE)

        voip_data_obj_list = []
        for voip_data_obj in voip_data_objs:
            voip_data = {
                'pk': str(voip_data_obj.pk),
                'username': voip_data_obj.customer.get_username(),
                'agent_username': voip_data_obj.agent.user.username,
                'start_date_time': voip_data_obj.start_datetime.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'end_date_time': voip_data_obj.end_datetime.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'call_duration': voip_data_obj.get_call_duration(),
                'meeting_status': "Completed" if voip_data_obj.is_completed else "Ongoing"
            }

            if voip_data["meeting_status"] == "Ongoing":
                voip_data["end_date_time"] = '-'
                voip_data["call_duration"] = '-'

            is_recording_available, file_path = get_recording_available(
                voip_data_obj, LiveChatFileAccessManagement)

            voip_data["is_recording_available"] = is_recording_available
            voip_data["file_path"] = file_path

            voip_data_obj_list.append(voip_data)

        return voip_data_obj_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_audit_object_list %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return []


def get_recording_available(voip_obj, LiveChatFileAccessManagement):
    is_recording_available = False
    file_path = ""

    try:

        if voip_obj.merged_file_path != "":
            is_recording_available = True

            file_pk = voip_obj.merged_file_path
            file_obj = LiveChatFileAccessManagement.objects.filter(pk=file_pk)

            if file_obj:
                file_path = '/livechat/download-file/' + \
                    str(file_pk) + '/' + file_obj[0].file_path.split('/')[-1]
            else:
                is_recording_available = False

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_recording_available %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return is_recording_available, file_path


def get_audit_trail_pagination_data(audit_objs: Page):
    try:
        pagination_data = {
            'has_other_pages': audit_objs.has_other_pages(),
            'page_range': audit_objs.paginator.page_range.stop,
            'number': audit_objs.number,
            'num_pages': audit_objs.paginator.num_pages,
        }
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_audit_trail_pagination_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        pagination_data = {
            'has_other_pages': False,
        }

    return pagination_data


def wrap_do_not_translate_keywords(text):
    try:
        text = text.replace(text, "<span translate='no'>" + text + "</span>")
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error wrap_do_not_translate_keywords %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return text


def get_livechat_notification(channel, agent_fullname, customer_language, end_chat_message, EasyChatTranslationCache):
    from LiveChatApp.utils_translation import get_translated_text
    
    if customer_language and customer_language.lang != 'en':
        agent_fullname = wrap_do_not_translate_keywords(agent_fullname.strip())
    message = ""
    remove_whitespace = False
    try:
        if end_chat_message:
            if channel == "Instagram" or channel == "Twitter" or channel == "Web" or channel == "GoogleRCS" or channel == "Telegram":
                message = "Your chat with " + agent_fullname.strip() + " has ended."
            else:
                remove_whitespace = True
                message = "Your chat with *" + agent_fullname.strip() + "* has ended."
        else:
            if channel == "Instagram" or channel == "Twitter" or channel == "GoogleRCS" or channel == "Telegram":
                message = agent_fullname.strip() + " has joined the chat. Please ask your queries now. If you wish to end your conversation, type " + END_CHAT_PRESERVER + "."
            else:
                remove_whitespace = True
                message = "*" + agent_fullname.strip() + "* has joined the chat. Please ask your queries now. If you wish to end your conversation, type " + END_CHAT_PRESERVER + "."
                
        if customer_language and customer_language.lang != 'en':
            message = get_translated_text(message, customer_language.lang, EasyChatTranslationCache, remove_whitespace, True)
            message = message.replace(END_CHAT_MATCHER, " " + STOP_CONVERSATION_LIST[customer_language.lang] + " ")
        else:
            message = message.replace(END_CHAT_PRESERVER, STOP_CONVERSATION_LIST[customer_language.lang])
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_livechat_notification %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return message


def send_voip_history_email(email_list, export_zip_file_path):
    try:
        domain = settings.EASYCHAT_HOST_URL
        today = datetime.now()
        body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat VOIP History for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
        developer_console_config = get_developer_console_settings()

        body += developer_console_config.custom_report_template_signature

        body += """</div></body>"""

        start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)

        for email_id in email_list:
            body = body.format(email_id, str(start_date),
                               domain, export_zip_file_path)
            email_subject = "LiveChat VOIP History For " + \
                str(email_id)
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_voip_history_email: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_today_voip_history_zip_file_path(user, agent_list, LiveChatVoIPData):
    try:
        today = datetime.now()
        tz = pytz.timezone(settings.TIME_ZONE)

        voip_history_objs = LiveChatVoIPData.objects.filter(request_datetime__date=today.date(
        ), call_type__in=['pip', 'new_tab'], agent__in=agent_list, is_completed=True, is_interrupted=False).order_by("-request_datetime")

        test_wb = Workbook()

        sheet1 = test_wb.add_sheet("Sheet1")

        sheet1.write(0, 0, "Customer Name")
        sheet1.write(0, 1, "Agent Username")
        sheet1.write(0, 2, "Start Date-Time")
        sheet1.write(0, 3, "End Date-Time")
        sheet1.write(0, 4, "Total Duration")
        sheet1.write(0, 5, "Meeting Status")

        row = 1
        for voip_data in voip_history_objs:

            sheet1.write(row, 0, voip_data.customer.get_username())
            sheet1.write(row, 1, voip_data.agent.user.username)

            sheet1.write(row, 2, voip_data.start_datetime.astimezone(
                tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P))
            sheet1.write(row, 3, voip_data.end_datetime.astimezone(
                tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P))

            sheet1.write(row, 4, voip_data.get_call_duration())
            sheet1.write(row, 5, "Completed")

            row += 1

        file_path = "livechat-voip-history/" + \
            str(user.user.username) + "/voip_history_" + \
            str(today.date()) + ".xls"

        test_wb.save(settings.MEDIA_ROOT + file_path)

        export_zip_file_path = "files/livechat-voip-history/" + \
            str(user.user.username) + "/VOIPHistoryToday-" + \
            str(today.time().replace(microsecond=0)) + ".zip"

        zip_obj = ZipFile(export_zip_file_path, 'w')

        try:
            file_path = "livechat-voip-history/" + \
                str(user.user.username) + "/voip_history_" + \
                str(today.date()) + '.xls'
            zip_obj.write(settings.MEDIA_ROOT +
                          file_path, basename(file_path))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("VOIP History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                'AppName': 'LiveChat'})

        zip_obj.close()

        return export_zip_file_path
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get__today_voip_history_zip_file_path: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        return ""


def export_today_voip_history_data(user, filter_param, LiveChatVoIPData):
    try:

        message_history_datadump_log = open(
            "log/voip_history_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-voip-history'):
            os.makedirs(settings.MEDIA_ROOT + 'livechat-voip-history')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)

            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-voip-history/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-voip-history/' +
                            str(user.user.username))

            export_zip_file_path = get_today_voip_history_zip_file_path(
                user, agent_list, LiveChatVoIPData)

            send_voip_history_email(email_list, export_zip_file_path)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("VOIP History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_voip_history: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def generate_livechat_api_document(username):

    def docx_replace_regex(doc_obj, regex, replace_text):
        for para in doc_obj.paragraphs:
            if regex.search(para.text):
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for index in range(len(inline)):
                    if regex.search(inline[index].text):
                        text = regex.sub(replace_text, inline[index].text)
                        inline[index].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace_text)

    def edit_document(document_path, target_path, metadata_dict):
        document_obj = Document(document_path)

        for metadata in metadata_dict:
            docx_replace_regex(document_obj, re.compile(
                r"" + metadata + ""), metadata_dict[metadata])

        document_obj.save(target_path)

    def get_metadata_dict():
        livechat_integration_url = settings.EASYCHAT_HOST_URL + \
            "/livechat/external/" + document_details["url_suffix"] + "/"

        metadata_dict = {
            "livechat_integration_url": livechat_integration_url
        }

        return metadata_dict

    def get_base_document_path(document_details):
        original_file_name = document_details["original_file_name"]
        base_document_path = "files/templates/livechat-docs-template/" + original_file_name
        return base_document_path

    def get_document_folder_path():
        target_document_folder = f'{LIVECHAT_DOCUMENTS_PATH}/{username}'

        create_directory(target_document_folder)
        return target_document_folder

    def get_target_document_path(document_details):
        target_document_folder = get_document_folder_path()
        original_file_name = document_details["original_file_name"]
        target_document_path = target_document_folder + "/" + original_file_name
        return target_document_path

    try:
        for document_type, document_details in LIVECHAT_DOCUMENTS.items():
            base_document_path = get_base_document_path(document_details)
            target_document_path = get_target_document_path(document_details)

            if document_type == "analytics-collection":
                with open(base_document_path, "r") as base_data, open(target_document_path, "w") as target_data:
                    target_data.write(base_data.read())
                continue

            metadata_dict = get_metadata_dict()
            edit_document(base_document_path,
                          target_document_path, metadata_dict)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("generate_livechat_api_document! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})
        return None


def create_directory(directory_path, remove_first=False):
    if remove_first and os.path.exists(directory_path):
        try:
            shutil.rmtree(directory_path)
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_directory %s at %s", str(e), str(
                exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)


def export_today_vc_history_data(user, filter_param, LiveChatVoIPData):
    try:

        message_history_datadump_log = open(
            "log/vc_history_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-vc-history'):
            os.makedirs(settings.MEDIA_ROOT + 'livechat-vc-history')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)

            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-vc-history/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-vc-history/' +
                            str(user.user.username))

            export_zip_file_path = get_today_vc_history_zip_file_path(
                user, agent_list, LiveChatVoIPData)

            send_vc_history_email(email_list, export_zip_file_path)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("VOIP History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_voip_history: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_today_vc_history_zip_file_path(user, agent_list, LiveChatVoIPData):
    try:
        today = datetime.now()
        tz = pytz.timezone(settings.TIME_ZONE)

        vc_history_objs = LiveChatVoIPData.objects.filter(request_datetime__date=today.date(
        ), call_type='video_call', agent__in=agent_list, is_completed=True, is_interrupted=False).order_by("-request_datetime")

        test_wb = Workbook()

        sheet1 = test_wb.add_sheet("Sheet1")

        sheet1.write(0, 0, "Customer Name")
        sheet1.write(0, 1, "Agent Username")
        sheet1.write(0, 2, "Start Date-Time")
        sheet1.write(0, 3, "End Date-Time")
        sheet1.write(0, 4, "Total Duration")
        sheet1.write(0, 5, "Meeting Status")

        row = 1
        for vc_data in vc_history_objs:

            sheet1.write(row, 0, vc_data.customer.get_username())
            sheet1.write(row, 1, vc_data.agent.user.username)

            sheet1.write(row, 2, vc_data.start_datetime.astimezone(
                tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P))
            sheet1.write(row, 3, vc_data.end_datetime.astimezone(
                tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P))

            sheet1.write(row, 4, vc_data.get_call_duration())
            sheet1.write(row, 5, "Completed")

            row += 1

        file_path = "livechat-vc-history/" + \
            str(user.user.username) + "/vc_history_" + \
            str(today.date()) + ".xls"

        test_wb.save(settings.MEDIA_ROOT + file_path)

        export_zip_file_path = "files/livechat-vc-history/" + \
            str(user.user.username) + "/VCHistoryToday-" + \
            str(today.time().replace(microsecond=0)) + ".zip"

        zip_obj = ZipFile(export_zip_file_path, 'w')

        try:
            file_path = "livechat-vc-history/" + \
                str(user.user.username) + "/vc_history_" + \
                str(today.date()) + '.xls'
            zip_obj.write(settings.MEDIA_ROOT +
                          file_path, basename(file_path))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("VC History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                'AppName': 'LiveChat'})

        zip_obj.close()

        return export_zip_file_path
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get__today_vc_history_zip_file_path: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        return ""


def send_vc_history_email(email_list, export_zip_file_path):
    try:
        domain = settings.EASYCHAT_HOST_URL
        today = datetime.now()
        body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat VC History for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
        developer_console_config = get_developer_console_settings()

        body += developer_console_config.custom_report_template_signature

        body += """</div></body>"""

        start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)

        for email_id in email_list:
            body = body.format(email_id, str(start_date),
                               domain, export_zip_file_path)
            email_subject = "LiveChat VC History For " + \
                str(email_id)
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_voip_history_email: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def clean_dynamic_form_data(form_filled):
    try:
        validation_obj = LiveChatInputValidation()

        for form_data in form_filled:
            if form_data and form_data["value"] != "*No Data filled*" and (form_data["type"] == '1' or form_data["type"] == '5'):
                form_data["value"] = validation_obj.remo_html_from_string(
                    form_data["value"])
                form_data["value"] = validation_obj.remo_unwanted_characters(
                    form_data["value"])

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("clean_dynamic_form_data: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return form_filled


def get_cobrowsing_status(cust_obj, LiveChatCobrowsingData):
    cobrowsing_status = 'none'
    cobrowsing_meeting_id = None
    try:
        cobrowsing_request = LiveChatCobrowsingData.objects.filter(
            customer=cust_obj, is_notification_displayed=False)

        if cobrowsing_request:
            cobrowsing_request = cobrowsing_request.order_by('-pk').first()
            cobrowsing_status = 'initiated'

            if cobrowsing_request.is_accepted:
                cobrowsing_status = 'accepted'

            if cobrowsing_request.is_rejected:
                cobrowsing_status = 'rejected'

            if cobrowsing_request.is_started:
                cobrowsing_status = 'ongoing'

            if cobrowsing_request.is_completed:
                cobrowsing_status = 'completed'
                cobrowsing_request.is_notification_displayed = True
                cobrowsing_request.save()

            cobrowsing_meeting_id = str(cobrowsing_request.meeting_id)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_cobrowsing_status: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return cobrowsing_status, cobrowsing_meeting_id


def get_cobrowsing_info_based_agent(agent_obj, LiveChatCobrowsingData):
    cobrowsing_info = {
        'status': 'none',
        'meeting_id': 'none',
        'session_id': 'none',
    }

    try:
        cobrowsing_request = LiveChatCobrowsingData.objects.filter(
            agent=agent_obj)

        if cobrowsing_request:
            cobrowsing_request = cobrowsing_request.order_by('-pk').first()

            if not cobrowsing_request.customer.is_session_exp:
                cobrowsing_info["status"] = 'initiated'
                cobrowsing_info["meeting_id"] = str(
                    cobrowsing_request.meeting_id)
                cobrowsing_info["session_id"] = str(
                    cobrowsing_request.customer.session_id)

                if cobrowsing_request.is_accepted:
                    cobrowsing_info['cobrowse_session_id'] = cobrowsing_request.cobrowse_session_id
                    cobrowsing_info["status"] = 'accepted'

                if cobrowsing_request.is_rejected:
                    cobrowsing_info["status"] = 'rejected'

                if cobrowsing_request.is_started:
                    cobrowsing_info["status"] = 'ongoing'

                if cobrowsing_request.is_completed:
                    cobrowsing_info["status"] = 'none'
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_cobrowsing_info_based_agent: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return cobrowsing_info


def get_cobrowsing_info_based_guest_agent(agent_obj, LiveChatCobrowsingData, LiveChatCustomer):
    cobrowsing_guest_info = []

    try:
        customer_objs = LiveChatCustomer.objects.filter(
            guest_agents__in=[agent_obj], is_session_exp=False)
        cobrowsing_requests = LiveChatCobrowsingData.objects.filter(
            customer__in=customer_objs, is_accepted=True, is_completed=False)

        for cobrowsing_request in cobrowsing_requests:
            cobrowsing_info = {
                'status': 'initiated',
                'meeting_id': str(cobrowsing_request.meeting_id),
                'session_id': str(cobrowsing_request.customer.session_id),
                'primary_agent_name': cobrowsing_request.agent.get_agent_name(),
                'joined': True if agent_obj in cobrowsing_request.guest_agents.all() else False,
            }

            if cobrowsing_request.is_accepted:
                cobrowsing_info['cobrowse_session_id'] = str(
                    cobrowsing_request.cobrowse_session_id)
                cobrowsing_info["status"] = 'accepted'

            if cobrowsing_request.is_rejected:
                cobrowsing_info["status"] = 'rejected'

            if cobrowsing_request.is_started:
                cobrowsing_info["status"] = 'ongoing'

            if cobrowsing_request.is_completed:
                cobrowsing_info["status"] = 'none'

            cobrowsing_guest_info.append(cobrowsing_info)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_cobrowsing_info_based_guest_agent: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return cobrowsing_guest_info


def get_cobrowsing_request_text(bot_obj, LiveChatConfig):
    try:
        config_obj = LiveChatConfig.objects.get(bot=bot_obj)

        text = config_obj.cobrowse_request_text

        return text
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_cobrowsing_request_text: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return COBROWSING_REQUEST_TEXT


def get_cobrowsing_data_history_objects(agent_username, query_user_obj, datetime_end, datetime_start, agent_obj_list, LiveChatCobrowsingData):

    cobrowsing_object_list = []
    try:
        cobrowsing_object_list = LiveChatCobrowsingData.objects.filter(request_datetime__date__range=[datetime_start, datetime_end],
                                                                       agent__in=agent_obj_list, is_started=True, is_interrupted=False).order_by('-pk')

        if agent_username != "All" and query_user_obj != None:
            cobrowsing_object_list = cobrowsing_object_list.filter(
                agent=query_user_obj)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_cobrowsing_data_history_objects: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return cobrowsing_object_list


def parse_cobrowsing_history_object_list(cobrowsing_data_objs):
    try:
        tz = pytz.timezone(settings.TIME_ZONE)

        cobrowsing_data_obj_list = []
        for cobrowsing_data_obj in cobrowsing_data_objs:
            cobrowsing_data = {
                'pk': str(cobrowsing_data_obj.pk),
                'username': cobrowsing_data_obj.customer.get_username(),
                'agent_username': cobrowsing_data_obj.agent.user.username,
                'start_date_time': cobrowsing_data_obj.start_datetime.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'end_date_time': cobrowsing_data_obj.end_datetime.astimezone(tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P),
                'call_duration': cobrowsing_data_obj.get_duration(),
                'meeting_status': "Completed" if cobrowsing_data_obj.is_completed else "Ongoing"
            }

            if cobrowsing_data["meeting_status"] == "Ongoing":
                cobrowsing_data["end_date_time"] = '-'
                cobrowsing_data["call_duration"] = '-'

            cobrowsing_data_obj_list.append(cobrowsing_data)

        return cobrowsing_data_obj_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_cobrowsing_history_object_list %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return []


def send_cobrowsing_history_email(email_list, export_zip_file_path):
    try:
        domain = settings.EASYCHAT_HOST_URL
        today = datetime.now()
        body = """
                <head>
                  <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
                  <title>Cogno AI</title>
                  <style type="text/css" media="screen">
                  </style>
                </head>
                <body>

                <div style="padding:1em;border:0.1em black solid;" class="container">
                    <p>
                        Dear {},
                    </p>
                    <p>
                        We have received a request to provide you with the LiveChat Cobrowsing History for {}. Please click on the link below to download the file.
                    </p>
                    <a href="{}/{}">click here</a>
                    <p>&nbsp;</p>
                    """
        developer_console_config = get_developer_console_settings()

        body += developer_console_config.custom_report_template_signature

        body += """</div></body>"""

        start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)

        for email_id in email_list:
            body = body.format(email_id, str(start_date),
                               domain, export_zip_file_path)
            email_subject = "LiveChat Cobrowsing History For " + \
                str(email_id)
            send_email_to_customer_via_awsses(email_id, email_subject, body)

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_cobrowsing_history_email: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_today_cobrowsing_history_zip_file_path(user, agent_list, LiveChatCobrowsingData):
    try:
        today = datetime.now()
        tz = pytz.timezone(settings.TIME_ZONE)

        cobrowsing_history_objs = LiveChatCobrowsingData.objects.filter(request_datetime__date=today.date(
        ), agent__in=agent_list, is_completed=True, is_interrupted=False).order_by("-request_datetime")

        test_wb = Workbook()

        sheet1 = test_wb.add_sheet("Sheet1")

        sheet1.write(0, 0, "Customer Name")
        sheet1.write(0, 1, "Agent Username")
        sheet1.write(0, 2, "Start Date-Time")
        sheet1.write(0, 3, "End Date-Time")
        sheet1.write(0, 4, "Total Duration")
        sheet1.write(0, 5, "Meeting Status")

        row = 1
        for cobrowsing_history_obj in cobrowsing_history_objs:

            sheet1.write(
                row, 0, cobrowsing_history_obj.customer.get_username())
            sheet1.write(row, 1, cobrowsing_history_obj.agent.user.username)

            sheet1.write(row, 2, cobrowsing_history_obj.start_datetime.astimezone(
                tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P))
            sheet1.write(row, 3, cobrowsing_history_obj.end_datetime.astimezone(
                tz).strftime(DATE_DD_MMM_YYYY_TIME_HH_MIN_P))

            sheet1.write(row, 4, cobrowsing_history_obj.get_duration())
            sheet1.write(row, 5, "Completed")

            row += 1

        file_path = "livechat-cobrowsing-history/" + \
            str(user.user.username) + "/cobrowsing_history_" + \
            str(today.date()) + ".xls"

        test_wb.save(settings.MEDIA_ROOT + file_path)

        export_zip_file_path = "files/livechat-cobrowsing-history/" + \
            str(user.user.username) + "/CobrowsingHistoryToday-" + \
            str(today.time().replace(microsecond=0)) + ".zip"

        zip_obj = ZipFile(export_zip_file_path, 'w')

        try:
            file_path = "livechat-cobrowsing-history/" + \
                str(user.user.username) + "/cobrowsing_history_" + \
                str(today.date()) + '.xls'
            zip_obj.write(settings.MEDIA_ROOT +
                          file_path, basename(file_path))
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Cobrowsing History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                'AppName': 'LiveChat'})

        zip_obj.close()

        return export_zip_file_path
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get__today_cobrowsing_history_zip_file_path: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        return ""


def export_today_cobrowsing_history_data(user, filter_param, LiveChatCobrowsingData):
    try:

        message_history_datadump_log = open(
            "log/cobrowsing_history_dump.log", "a")
        if not os.path.exists(settings.MEDIA_ROOT + 'livechat-cobrowsing-history'):
            os.makedirs(settings.MEDIA_ROOT + 'livechat-cobrowsing-history')
        try:
            filter_param = json.loads(filter_param)
            today = datetime.now()
            email_str = filter_param["email"]
            email_list = [email.strip()
                          for email in email_str.split(",") if email != ""]

            agent_list = get_agents_under_this_user(user)

            if not os.path.exists(settings.MEDIA_ROOT + 'livechat-cobrowsing-history/' + str(user.user.username)):
                os.makedirs(settings.MEDIA_ROOT + 'livechat-cobrowsing-history/' +
                            str(user.user.username))

            export_zip_file_path = get_today_cobrowsing_history_zip_file_path(
                user, agent_list, LiveChatCobrowsingData)

            send_cobrowsing_history_email(email_list, export_zip_file_path)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("VOIP History Datadump! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                         'AppName': 'LiveChat'})
            message_history_datadump_log.write(
                str(today) + ": failed: " + str(e))
        message_history_datadump_log.close()
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_today_cobrowsing_history: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_ip_address(request):
    try:

        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')

        if x_forwarded_for:
            ip_address = x_forwarded_for.split(',')[0]
        else:
            ip_address = request.META.get('REMOTE_ADDR')
        return ip_address
        
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_ip_address: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def is_blacklist_keyword_present(word, user_obj, blacklist_for, LiveChatBlackListKeyword, LiveChatUser):
    try:

        admin_obj = get_admin_from_active_agent(user_obj, LiveChatUser)
        word = word.lower().strip()
        keyword_check_for_agent = False

        if LiveChatBlackListKeyword.objects.filter(word=word, agent_id=user_obj, blacklist_keyword_for=blacklist_for).count() or LiveChatBlackListKeyword.objects.filter(word=word, agent_id=admin_obj, blacklist_keyword_for=blacklist_for).count():
            keyword_check_for_agent = True

        if blacklist_for == "agent":
            return keyword_check_for_agent

        elif blacklist_for == "customer":
            if keyword_check_for_agent or LiveChatBlackListKeyword.objects.filter(word=word.lower().strip(), agent_id__in=user_obj.agents.all(), blacklist_keyword_for=blacklist_for).count():
                return True
            else:
                return False

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("is_blacklist_keyword_present: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return False    


def get_customer_blacklisted_keywords(agent_obj, LiveChatBlackListKeyword, LiveChatUser):
    customer_blacklisted_keyword = []
    try:
        if agent_obj.status == "1" or agent_obj.status == "2" or agent_obj.is_allow_toggle:

            admin_obj = get_admin_from_active_agent(agent_obj, LiveChatUser)
            customer_blacklisted_keyword = LiveChatBlackListKeyword.objects.filter(agent_id=admin_obj, blacklist_keyword_for="customer") | LiveChatBlackListKeyword.objects.filter(
                agent_id__in=admin_obj.agents.all(), blacklist_keyword_for="customer")    

            return customer_blacklisted_keyword

        parent_users = LiveChatUser.objects.filter(agents__user=agent_obj.user)
        livechat_admin_obj = None
        supervisor_objs = []
        if LiveChatUser.objects.filter(agents__user=parent_users[0].user).count():
            livechat_admin_obj = LiveChatUser.objects.filter(
                agents__user=parent_users[0].user)[0]
            supervisor_objs = livechat_admin_obj.agents.filter(status="2", is_deleted=False)

        customer_blacklisted_keyword = LiveChatBlackListKeyword.objects.filter(agent_id__in=parent_users, blacklist_keyword_for="customer") | LiveChatBlackListKeyword.objects.filter(agent_id=livechat_admin_obj, blacklist_keyword_for="customer") | LiveChatBlackListKeyword.objects.filter(agent_id__in=supervisor_objs, blacklist_keyword_for="customer")

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_customer_blacklisted_keywords: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return customer_blacklisted_keyword


def get_chat_escalation_status(customer_obj, LiveChatMISDashboard, LiveChatReportedCustomer):
    chat_escalation_status = "safe"
    try:
        reported_customer_obj = LiveChatReportedCustomer.objects.filter(livechat_customer=customer_obj)

        if reported_customer_obj.exists():
            
            reported_customer_obj = reported_customer_obj[0]

            if reported_customer_obj.is_completed:
                chat_escalation_status = "completed"
            else:
                if reported_customer_obj.is_reported:
                    chat_escalation_status = "reported"
                else:
                    if LiveChatMISDashboard.objects.filter(livechat_customer=customer_obj, message_contains_blacklisted_keyword=True, message_time__gt=reported_customer_obj.chat_escalation_report_ignored_time).count():
                        chat_escalation_status = "to_be_reported"
                    else:
                        chat_escalation_status = "warned" 
        
        else:
            if LiveChatMISDashboard.objects.filter(livechat_customer=customer_obj, message_contains_blacklisted_keyword=True, message_time__gt=customer_obj.chat_escalation_warn_ignored_time).count():
                chat_escalation_status = "to_be_warned" 

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_chat_escalation_status: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        chat_escalation_status = "safe"

    return chat_escalation_status


def check_is_customer_blocked(livechat_cust_obj, LiveChatReportedCustomer):
    try:
        if livechat_cust_obj.channel.name == "Web":
            blocked_ip_address = list(LiveChatReportedCustomer.objects.filter(livechat_customer__bot=livechat_cust_obj.bot, is_blocked=True, is_completed=False).values_list("livechat_customer__ip_address", flat=True))
            if livechat_cust_obj.ip_address in blocked_ip_address:
                return True
        else:
            blocked_client_ids = list(LiveChatReportedCustomer.objects.filter(livechat_customer__bot=livechat_cust_obj.bot, is_blocked=True, is_completed=False).values_list("livechat_customer__client_id", flat=True))
            if livechat_cust_obj.client_id in blocked_client_ids:
                return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_is_customer_blocked: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return False    


def get_last_seen_object_based_user_group(livechat_user, user_group_obj, LiveChatInternalChatLastSeen):
    try:
        last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=livechat_user, user_group=user_group_obj)

        if last_seen_obj:
            return last_seen_obj.first()
        
        last_seen_obj = LiveChatInternalChatLastSeen.objects.create(user=livechat_user, user_group=user_group_obj)

        return last_seen_obj

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_last_seen_object_based_user_group: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_last_seen_object_based_group(livechat_user, group_obj, LiveChatInternalChatLastSeen):
    try:
        last_seen_obj = LiveChatInternalChatLastSeen.objects.filter(user=livechat_user, group=group_obj)

        if last_seen_obj:
            return last_seen_obj.first()
        
        last_seen_obj = LiveChatInternalChatLastSeen.objects.create(user=livechat_user, group=group_obj)

        return last_seen_obj

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_last_seen_object_based_group: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def store_keyword_for_livechat_intent(bot_pk, whatsapp_reinitiating_keyword, previous_reinitiating_keyword, Bot):
    try:
        bot_obj = Bot.objects.get(pk=int(bot_pk))

        livechat_training_data = json.loads(bot_obj.livechat_default_intent.training_data)
        exiting_keyword_position = None

        for key, value in livechat_training_data.items():
            if value == previous_reinitiating_keyword:
                exiting_keyword_position = key
                break

        if exiting_keyword_position:
            livechat_training_data[exiting_keyword_position] = whatsapp_reinitiating_keyword
        else:
            last_data_key = list(livechat_training_data.keys())[-1]
            livechat_training_data[int(last_data_key) + 1] = whatsapp_reinitiating_keyword

        bot_obj.livechat_default_intent.training_data = json.dumps(livechat_training_data)
        bot_obj.livechat_default_intent.save()
        
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("store_keyword_for_livechat_intent: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def check_if_whatsapp_keyword_present_in_intent(bot_pk, is_whatsapp_reinitiation_enabled, whatsapp_reinitiating_keyword, Bot, Intent, LiveChatConfig):
    try:
        bot_obj = Bot.objects.get(pk=int(bot_pk))
        livechat_config_obj = LiveChatConfig.objects.get(bot=bot_obj)

        if is_whatsapp_reinitiation_enabled and whatsapp_reinitiating_keyword.lower() != livechat_config_obj.whatsapp_reinitiating_keyword.lower():

            intents = Intent.objects.filter(bots__in=[bot_obj], is_deleted=False)
            for intent in intents:
                training_data = json.loads(intent.training_data)

                for key, value in training_data.items():
                    if value.lower() == whatsapp_reinitiating_keyword.lower():
                        return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_if_whatsapp_keyword_present_in_intent: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return False
    

def remove_agent_from_prev_groups(user_obj, new_user_obj, supervisor_obj, LiveChatInternalChatGroup, LiveChatInternalChatGroupMembers, LiveChatInternalMessageInfo, LiveChatInternalMISDashboard):
    try:
        group_member_objs = LiveChatInternalChatGroupMembers.objects.filter(user=new_user_obj, is_removed=False, is_deleted=False, has_left=False)
        supervisor_member_objs = LiveChatInternalChatGroupMembers.objects.filter(user=supervisor_obj, is_removed=False, is_deleted=False, has_left=False)

        group_objs = LiveChatInternalChatGroup.objects.filter(members__in=group_member_objs).exclude(members__in=supervisor_member_objs)

        for group_obj in group_objs:
            group_member_obj = LiveChatInternalChatGroupMembers.objects.filter(
                user=new_user_obj, group=group_obj).first()
            
            if group_member_obj:
                remove_datetime = timezone.now()
                group_member_obj.is_removed = True

                group_member_obj.remove_datetime = remove_datetime
                group_member_obj.save()
                msg_text = 'removed ' + new_user_obj.user.username + ' from the chat'
                msg_info_obj = LiveChatInternalMessageInfo.objects.create(
                    message_text=msg_text)
                LiveChatInternalMISDashboard.objects.create(
                    message_datetime=remove_datetime, sender_name='System', sender=user_obj, group=group_obj, message_info=msg_info_obj)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("remove_agent_from_prev_groups: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_phone_number_and_country_code(phone, channel, to_check_edited_number=False):
    try:
        if channel not in ["Web", "WhatsApp", "GoogleRCS"] and not to_check_edited_number:
            return 'None', "", True

        import phonenumbers

        if channel == "GoogleRCS":
            phone = phone.split("_")[2]

        if channel == "WhatsApp":
            phone = "+" + phone

        phone_number_obj = phonenumbers.parse(phone, None)
        is_valid_number = phonenumbers.is_valid_number(phone_number_obj)

        return str(phone_number_obj.national_number), str(phone_number_obj.country_code), is_valid_number

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("validate_and_get_country_code: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        phone, "", False


def get_export_users_status(livechat_user_obj, user_objs, LiveChatUser, LiveChatFileAccessManagement):
    try:

        export_path = ""
        export_status = "export_request_failed"

        if user_objs.count() <= 100:

            export_path = export_livechat_users(livechat_user_obj, user_objs, False, LiveChatUser, LiveChatFileAccessManagement)

            if export_path != "":
                export_status = "export_request_completed"

        else:

            thread = threading.Thread(target=export_livechat_users, args=(
                livechat_user_obj, user_objs, True, LiveChatUser, LiveChatFileAccessManagement), daemon=True)
            thread.start()

            export_status = "export_request_saved"      

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_export_users_status: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return export_status, export_path    


def export_livechat_users(livechat_user_obj, user_objs, to_be_mailed, LiveChatUser, LiveChatFileAccessManagement):
    try:

        if not os.path.exists(settings.SECURE_MEDIA_ROOT + 'LiveChatApp/livechat-user-details/' + str(livechat_user_obj.user.username)):
            os.makedirs(settings.SECURE_MEDIA_ROOT + 'LiveChatApp/livechat-user-details/' +
                        str(livechat_user_obj.user.username))

        today = datetime.now()
        current_time = today.strftime("%H:%M:%S")

        workbook = Workbook()
        sheet = workbook.add_sheet("Sheet1")

        style = xlwt.XFStyle()
        font = xlwt.Font()
        font.bold = True
        style.font = font

        sheet.write(0, 0, "FirstName", style=style)
        sheet.write(0, 1, "LastName", style=style)
        sheet.write(0, 2, "Phone No.", style=style)
        sheet.write(0, 3, "Email ID", style=style)
        sheet.write(0, 4, "Password", style=style)
        sheet.write(0, 5, "Status ID (2 for Supervisor, 3 for Agent)", style=style)
        sheet.write(0, 6, "Bot ID", style=style)
        sheet.write(0, 7, "Category Name", style=style)
        sheet.write(0, 8, "Maximum Customers Allowed", style=style)
        sheet.write(0, 9, "Supervisor (if any, leave blank for supervisor)", style=style)
        sheet.write(0, 10, "Profile Creation Timestamp", style=style)
        sheet.write(0, 11, "Last Active Timestamp", style=style)

        row = 1
        for user_obj in user_objs:
            sheet.write(row, 0, user_obj.user.first_name)
            sheet.write(row, 1, user_obj.user.last_name)
            sheet.write(row, 2, user_obj.phone_number)
            sheet.write(row, 3, user_obj.user.email)
            sheet.write(row, 4, '-')
            sheet.write(row, 5, user_obj.status)

            bots = str(list(user_obj.bots.filter(is_deleted=False).values_list('pk', flat=True)))[1:-1]
            sheet.write(row, 6, bots)

            categories = str(list(user_obj.category.filter(is_deleted=False).values_list('title', flat=True)))[1:-1]
            sheet.write(row, 7, categories)

            if user_obj.status == '3':
                sheet.write(row, 8, user_obj.get_max_customer_count())
                sheet.write(row, 9, user_obj.get_supervisor())
            else:
                sheet.write(row, 8, '-')
                sheet.write(row, 9, '-')

            sheet.write(row, 10, (user_obj.user.date_joined + timedelta(hours=5, minutes=30)).strftime("%m/%d/%Y %H:%M"))

            last_updated_time = user_obj.last_updated_time
            if user_obj.status == '3' and last_updated_time:
                sheet.write(row, 11, (last_updated_time + timedelta(hours=5, minutes=30)).strftime("%m/%d/%Y %H:%M"))
            else:
                sheet.write(row, 11, '-')

            row += 1

        file_name = "user_details_" + \
            str(today.date()) + "_" + str(current_time) + ".xls"
        file_path = "LiveChatApp/livechat-user-details/" + \
            str(livechat_user_obj.user.username) + '/' + file_name
        workbook.save(settings.SECURE_MEDIA_ROOT + file_path)

        file_path = '/secured_files/' + file_path

        file_access_management_obj = LiveChatFileAccessManagement.objects.create(
            file_path=file_path, is_public=False)

        export_path = '/livechat/download-file/' + \
            str(file_access_management_obj.key) + '/' + file_name

        if not to_be_mailed:

            return export_path

        body = """
            <head>
              <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
              <title>Cogno AI</title>
              <style type="text/css" media="screen">
              </style>
            </head>
            <body>

            <div style="padding:1em;border:0.1em black solid;" class="container">
                <p>
                    Dear {},
                </p>
                <p>
                    We have received a request to provide you with the LiveChat User Details for {}. Please click on the link below to download the file.
                </p>
                <a href="{}/{}">click here</a>
                <p>&nbsp;</p>

                """

        developer_console_config = get_developer_console_settings()

        body += developer_console_config.custom_report_template_signature

        body += """</div></body>"""

        start_date = datetime.strftime(today.date(), DATE_DD_MM_YYYY)

        email_id = livechat_user_obj.user.email
        body = body.format(email_id, str(start_date),
                           settings.EASYCHAT_HOST_URL, export_path[1:])
        email_subject = "LiveChat User Details For " + \
            str(email_id)
        send_email_to_customer_via_awsses(
            email_id, email_subject, body)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("export_livechat_users: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        if not to_be_mailed:

            return ""


def update_file_access_type(chat_type, group_id, user_group_id, receiver_username, sender_username, file_access_management_obj, LiveChatUser, LiveChatInternalChatGroup, LiveChatInternalUserGroup):
    try:

        if chat_type == "group_chat":
            group = LiveChatInternalChatGroup.objects.filter(group_id=group_id).first()

            if group:
                file_access_management_obj.file_access_type = "group_chat"
                file_access_management_obj.group = group
                file_access_management_obj.save()

        elif chat_type == "user_group_chat":
            user_group = LiveChatInternalUserGroup.objects.filter(group_id=user_group_id).first()

            if user_group:
                file_access_management_obj.file_access_type = "user_group_chat"
                file_access_management_obj.user_group = user_group
                file_access_management_obj.save()

        elif chat_type == "user_chat":
            user_obj = LiveChatUser.objects.filter(user__username=sender_username).first()
            user = LiveChatUser.objects.filter(user__username=receiver_username).first()
            
            user_group = check_and_update_user_group(user_obj, user, LiveChatInternalUserGroup)

            if user_group:
                file_access_management_obj.file_access_type = "user_group_chat"
                file_access_management_obj.user_group = user_group
                file_access_management_obj.save()      

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_file_access_type: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return file_access_management_obj


def get_livechat_cronjob_tracker_obj(cronjob_id, LiveChatCronjobTracker):
    try:
        livechat_cronjob_tracker_obj = LiveChatCronjobTracker.objects.filter(cronjob_id=cronjob_id).first()
        if livechat_cronjob_tracker_obj and livechat_cronjob_tracker_obj.is_object_expired():
            livechat_cronjob_tracker_obj.delete()
            return None
        return livechat_cronjob_tracker_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in get_livechat_cronjob_tracker_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return None


def create_livechat_cronjob_tracker_obj(cronjob_id, LiveChatCronjobTracker):
    try:
        LiveChatCronjobTracker.objects.create(cronjob_id=cronjob_id)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in create_livechat_cronjob_tracker_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def delete_livechat_cronjob_tracker_obj(cronjob_id, LiveChatCronjobTracker):
    try:
        livechat_cronjob_tracker_obj = LiveChatCronjobTracker.objects.filter(cronjob_id=cronjob_id).first()
        if livechat_cronjob_tracker_obj:
            livechat_cronjob_tracker_obj.delete()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in delete_livechat_cronjob_tracker_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def send_event_for_report_creation(customer_obj, LiveChatUser, LiveChatAdminConfig, LiveChatConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        if customer_obj.agent_id == None:
            return

        livechat_user = customer_obj.agent_id
        parent_user, second_user = get_parent_from_active_agent(livechat_user, LiveChatUser)
        data = get_chat_history_payload(livechat_user, customer_obj, parent_user, second_user)
        admin_config = get_livechat_admin_config_obj(customer_obj.bot, LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_report_creation! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})
        

def send_event_for_nps(customer_obj, LiveChatUser, LiveChatAdminConfig, LiveChatConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        if customer_obj.agent_id == None:
            return

        livechat_user = customer_obj.agent_id
        parent_user, second_user = get_parent_from_active_agent(livechat_user, LiveChatUser)
        data = get_chat_history_nps_payload(livechat_user, customer_obj, parent_user, second_user)
        admin_config = get_livechat_admin_config_obj(customer_obj.bot, LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_nps! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})


def get_parent_from_active_agent(active_agent, LiveChatUser):

    parent_user = None
    second_parent = None
    try:
        if active_agent.is_allow_toggle or active_agent.status == "1":
            parent_user = active_agent
            return parent_user, second_parent

        parent_user = LiveChatUser.objects.filter(
            agents__pk=active_agent.pk)[0]
        try:
            second_parent = LiveChatUser.objects.filter(
                agents__pk=parent_user.pk)[0]
        except Exception:
            second_parent = None

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_admin_from_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

    return parent_user, second_parent


def send_event_for_login_logout(livechat_user, livechat_session_management, LiveChatUser, LiveChatAdminConfig, LiveChatConfig, Bot, update=False):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        if livechat_user.status != "2":
            parent_user, second_user = get_parent_from_active_agent(livechat_user, LiveChatUser)
            data = get_agent_login_logout_payload(livechat_session_management, parent_user, second_user, update)
            admin_config = get_livechat_admin_config_obj(livechat_user.bots.all()[0], LiveChatConfig, LiveChatAdminConfig, Bot)
            send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_login_logout! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})


def send_event_for_performance_report(livechat_user, livechat_session_management, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot, update=False):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        if livechat_user.status != "2":
            parent_user, second_user = get_parent_from_active_agent(livechat_user, LiveChatUser)
            data = get_agent_performance_report_payload(livechat_session_management, parent_user, second_user, update)
            admin_config = get_livechat_admin_config_obj(livechat_user.bots.all()[0], LiveChatConfig, LiveChatAdminConfig, Bot)
            send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_performance_report! %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'LiveChat'})
        

def send_event_for_agent_not_ready(livechat_user, agent_not_ready_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot, update=False):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        if livechat_user.status != "2":
            parent_user, second_user = get_parent_from_active_agent(livechat_user, LiveChatUser)
            data = get_agent_not_ready_payload(agent_not_ready_obj, parent_user, second_user, update)
            admin_config = get_livechat_admin_config_obj(livechat_user.bots.all()[0], LiveChatConfig, LiveChatAdminConfig, Bot)
            send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_login_logout! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})
        

def send_event_for_video_call_history(meeting_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        parent_user, second_user = get_parent_from_active_agent(meeting_obj.agent, LiveChatUser)
        data = get_video_call_history_payload(meeting_obj, parent_user, second_user)
        admin_config = get_livechat_admin_config_obj(meeting_obj.agent.bots.all()[0], LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_video_call_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})
        

def send_event_for_voice_call_history(meeting_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        parent_user, second_user = get_parent_from_active_agent(meeting_obj.agent, LiveChatUser)
        data = get_voice_call_history_payload(meeting_obj, parent_user, second_user)
        admin_config = get_livechat_admin_config_obj(meeting_obj.agent.bots.all()[0], LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_voice_call_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'LiveChat'})


def send_event_for_abandoned_chat(customer_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        list_of_user = list(LiveChatUser.objects.filter(bots=customer_obj.bot).filter(Q(status="2") | Q(status="1")).values_list('user__username', flat=True))
        data = get_abandoned_chat_payload(customer_obj, list_of_user)
        admin_config = get_livechat_admin_config_obj(customer_obj.bot, LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_abandoned_chat! %s %s", str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


"""

This function provides the file access url, it may be a LiveChat mailer file or not, depends upon is_mailer_file.

"""


def get_livechat_secure_file_path(LiveChatFileAccessManagement, file_path, is_public=False, is_mailer_file=False, file_name='LiveChatFile'):
    try:
        if is_mailer_file:
            file_access_object = LiveChatFileAccessManagement.objects.create(
                file_path=file_path, is_public=is_public, is_mailer_report=True)
        else:
            file_access_object = LiveChatFileAccessManagement.objects.create(
                file_path=file_path, is_public=is_public)
        file_key = file_access_object.key
        file_url = settings.EASYCHAT_HOST_URL + \
            '/livechat/download-file/' + str(file_key) + '/' + file_name

        return file_url
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_livechat_secure_file_path! %s %s", str(e), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})

        return ""


def is_valid_livechat_file_access_request(file_management_obj):
    try:
        if file_management_obj.is_mailer_report and not file_management_obj.is_obj_time_limit_exceeded():
            return True
        else:
            return False
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("is_valid_livechat_file_access_request! %s %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
        return False

        
def send_event_for_cobrowsing_history(cobrowsing_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        parent_user, second_user = get_parent_from_active_agent(cobrowsing_obj.agent, LiveChatUser)
        data = get_cobrowsing_history_payload(cobrowsing_obj, parent_user, second_user)
        admin_config = get_livechat_admin_config_obj(cobrowsing_obj.agent.bots.all()[0], LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_cobrowsing_history! %s %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def send_event_for_offline_chat(livechat_cust_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        list_of_user = list(LiveChatUser.objects.filter(bots=livechat_cust_obj.bot).filter(Q(status="2") | Q(status="1")).values_list('user__username', flat=True))
        data = get_offline_chat_payload(livechat_cust_obj, list_of_user)
        admin_config = get_livechat_admin_config_obj(livechat_cust_obj.bot, LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_offline_chat! %s %s", str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def send_event_for_missed_chat(livechat_cust_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot):
    from EasyChat.kafka_producer import send_packet_into_kafka_producer
    try:
        list_of_user = list(LiveChatUser.objects.filter(bots=livechat_cust_obj.bot).filter(Q(status="2") | Q(status="1")).values_list('user__username', flat=True))
        data = get_missed_chat_payload(livechat_cust_obj, list_of_user)
        admin_config = get_livechat_admin_config_obj(livechat_cust_obj.bot, LiveChatConfig, LiveChatAdminConfig, Bot)
        send_packet_into_kafka_producer(settings.KAFKA_CONFIG["livechat_report_topic"], data, admin_config)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("send_event_for_missed_chat! %s %s", str(e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def handle_agent_not_ready(username, LiveChatUser, LiveChatSessionManagement, LiveChatConfig, LiveChatAdminConfig, Bot):
    try:
        user_obj = LiveChatUser.objects.filter(
            user__username=username, is_deleted=False).first()
        if user_obj:
            sessions_obj = LiveChatSessionManagement.objects.filter(
                user=user_obj).order_by('-session_starts_at').first()
            agent_not_ready_obj = sessions_obj.agent_not_ready.all().order_by(
                '-not_ready_starts_at').first()

        if agent_not_ready_obj and not agent_not_ready_obj.is_expired:
            diff = timezone.now() - sessions_obj.session_ends_at
            sessions_obj.offline_time += diff.seconds
            sessions_obj.session_ends_at = timezone.now()
            sessions_obj.is_idle = False

            if user_obj.current_status == "0":
                diff = timezone.now() - sessions_obj.time_marked_stop_interaction
                sessions_obj.stop_interaction_time += diff.seconds
                agent_not_ready_obj.stop_interaction_duration = diff.seconds
            else:
                sessions_obj.last_idle_time = datetime.now()

            agent_not_ready_obj.not_ready_ends_at = timezone.now()
            agent_not_ready_obj.is_expired = True
            sessions_obj.save()
            agent_not_ready_obj.save()
            send_event_for_agent_not_ready(
                user_obj, agent_not_ready_obj, LiveChatUser, LiveChatConfig, LiveChatAdminConfig, Bot, True)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("handle_agent_not_ready! %s %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def is_kafka_enabled():
    return settings.IS_REPORT_GENERATION_VIA_KAFKA_ENABLED


def create_export_request(user_obj, report_type, requested_data, LiveChatDataExportRequest):
    try:
        LiveChatDataExportRequest.objects.create(
            user=user_obj, report_type=report_type, filter_param=json.dumps(requested_data))
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("create_export_request! %s %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def check_if_sentence_valid(sentence, channel):
    try:
        char_limit = 3000
        if(channel == 'Facebook'):
            char_limit = LIVECHAT_FACEBOOK_CHAR_LIMIT
        elif(channel == 'Instagram'):
            char_limit = LIVECHAT_INSTAGRAM_CHAR_LIMIT

        return (len(sentence) and len(sentence) < char_limit)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("check_if_sentence_valid! %s %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def create_customer_details_system_message(livechat_cust_obj, LiveChatMISDashboard):
    try:
        text_message = "Customer Name: " + str(livechat_cust_obj.username) + " | Agent Name: " + str(livechat_cust_obj.agent_id.user.first_name) + " " + str(
            livechat_cust_obj.agent_id.user.last_name) + "(" + str(livechat_cust_obj.agent_id.user.username) + ")"
        
        LiveChatMISDashboard.objects.create(livechat_customer=livechat_cust_obj,
                                            sender="System",
                                            text_message=text_message,
                                            sender_name="system",
                                            message_time=timezone.now(),
                                            attachment_file_name="",
                                            attachment_file_path="")

    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("create_customer_details_system_message: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})


def get_livechat_admin_config_obj(bot, LiveChatConfig, LiveChatAdminConfig, Bot):
    try:
        admin_config = cache.get("LiveChatAdminConfig_" + str(bot.pk))
        if not admin_config:
            admin_config = set_livechat_admin_config_obj(bot, LiveChatConfig, LiveChatAdminConfig, Bot)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_livechat_admin_config_obj: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return admin_config


def set_livechat_admin_config_obj(bot, LiveChatConfig, LiveChatAdminConfig, Bot):
    try:
        config_obj = LiveChatConfig.objects.filter(bot=bot)
        admin_config = LiveChatAdminConfig.objects.filter(livechat_config__in=config_obj)[0]
        cache.set("LiveChatAdminConfig_" + str(bot.pk), admin_config, settings.CACHE_TIME)
    except Exception as exc:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("set_livechat_admin_config_obj: %s at %s", str(exc), str(
            exc_tb.tb_lineno), extra={'AppName': 'LiveChat'})
    return admin_config
