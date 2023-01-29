import base64
import threading

from Crypto.Cipher import AES
from django.shortcuts import render, redirect, HttpResponse

# Django REST framework
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from rest_framework.views import APIView
from rest_framework.authentication import SessionAuthentication, \
    BasicAuthentication
from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils import timezone
import xlrd
from xlrd import open_workbook
from xlwt import XFStyle, Font, Workbook

from EasyTMSApp.html_parser import strip_html_tags
from EasyTMSApp.constants import *
from EasyTMSApp.send_email import send_password_over_email, send_agent_query_message
from EasyTMSApp.utils_client_server_signal import send_data_from_server_to_client
from AuditTrailApp.utils import add_audit_trail
from docx import Document
from django.utils.encoding import smart_str

import json
from datetime import datetime, timedelta
from django.utils.dateformat import DateFormat
from django.utils.formats import get_format
from EasyTMSApp.utils_custom_encryption import *
from EasyTMSApp.utils_parse import *
from django.db.models import Q, Count
import operator

import uuid
import sys
from django.conf import settings
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import pytz
import magic
import mimetypes
import os
import hashlib
import re
import shutil

from DeveloperConsoleApp.utils import get_developer_console_settings, send_email_to_customer_via_awsses

# Logger
import logging

logger = logging.getLogger(__name__)


def logout_all(username, UserSession, Session):
    try:
        logger.info("In logout_all", extra={
            'AppName': 'EasyTMS', 'user_id': str(username)})
        user_session_objs = UserSession.objects.filter(user__username=username)
        for user_session_obj in user_session_objs:
            delete_user_session(user_session_obj, Session)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In logout_all: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS', 'user_id': str(username)})


def delete_user_session(user_session, Session):
    try:
        logger.info("In delete_user_session", extra={'AppName': 'EasyTMS',
                                                     'user_id': user_session.user.username})
        session_objs = Session.objects.filter(pk=user_session.session_key)
        user_session.delete()
        if session_objs:
            session_objs.delete()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In delete_user_session: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS', 'user_id': str(user_session.user.username)})


def get_number_of_day(year, month):
    year = int(year)
    month = int(month)
    leap = 0
    if year % 400 == 0:
        leap = 1
    elif year % 100 == 0:
        leap = 0
    elif year % 4 == 0:
        leap = 1
    if month == 2:
        return 28 + leap
    list = [1, 3, 5, 7, 8, 10, 12]
    if month in list:
        return 31
    return 30


# def send_mail(from_email_id, to_emai_id, message_as_string, from_email_id_password):
#     import smtplib
#     # The actual sending of the e-mail
#     server = smtplib.SMTP('smtp.gmail.com:587')
#     # Credentials (if needed) for sending the mail
#     password = from_email_id_password
#     # Start tls handshaking
#     server.starttls()
#     # Login to server
#     server.login(from_email_id, password)
#     # Send mail
#     server.sendmail(from_email_id, to_emai_id, message_as_string)
#     # Close session
#     server.quit()


def generate_email(email_id, start_date, end_date, filepath):

    config_obj = get_developer_console_settings()

    domain = settings.EASYCHAT_HOST_URL
    body = """
            <head>
              <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
              <title>Cogno AI</title>
              <style type="text/css" media="screen">
              </style>
            </head>
            <body>

            <div style="padding:1em;border:0.1em black solid;" class="container">
                <p>
                    Hi,
                </p>
                <p>
                    We have received a request to provide you with the Cogno TMS agent performance report from {} to {}. Please click on the link below to download the file.
                </p>
                <a href="{}/{}">click here</a>
                <p>&nbsp;</p>"""

    body += config_obj.custom_report_template_signature

    body += """</div></body>"""

    body = body.format(str(start_date), str(
        end_date), domain, filepath)

    send_email_to_customer_via_awsses(
        email_id, "Cogno TMS Agent Performance", body)


def remo_html_from_string(raw_str):
    try:
        regex_cleaner = re.compile('<.*?>')
        cleaned_raw_str = re.sub(regex_cleaner, '', raw_str)
        return cleaned_raw_str
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In remo_html_from_string: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})
        return raw_str


def remo_special_tag_from_string(raw_str):
    try:
        cleaned_raw_str = raw_str.replace(
            "+", "").replace("|", "").replace("-", "").replace("=", "")
        return cleaned_raw_str
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In remo_special_tag_from_string: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})
        return raw_str


def add_selected_supervisor(selected_supervisor_pk_list, active_agent, agent, Agent):
    try:
        logger.info("inside add supervisor ", extra={'AppName': 'EasyTMS'})

        for current_supervisor_pk in selected_supervisor_pk_list:
            current_supervisor_pk = int(current_supervisor_pk)
            if current_supervisor_pk == -1:
                current_supervisor = active_agent
            else:
                current_supervisor = Agent.objects.get(
                    pk=current_supervisor_pk)

            current_supervisor.agents.add(agent)
            current_supervisor.save()

        logger.info("successfully exited add supervisor ",
                    extra={'AppName': 'EasyTMS'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In add_selected_supervisor: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})


def add_bot_to_agent(agent, selected_bot_pk_list, Bot):
    try:
        logger.info("Inside add support language",
                    extra={'AppName': 'EasyTMS'})
        for bot_pk in selected_bot_pk_list:
            bot_obj = Bot.objects.get(pk=int(bot_pk))
            agent.bots.add(bot_obj)
        agent.save()
        logger.info("Secussefully exited add bot to agent",
                    extra={'AppName': 'EasyTMS'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_bot_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def add_ticket_category_to_agent(agent, selected_ticket_category_pk_list, TicketCategory):
    try:
        logger.info("Inside add_ticket_category_to_agent",
                    extra={'AppName': 'EasyTMS'})
        for ticket_category_pk in selected_ticket_category_pk_list:
            bot_pk, category_pk = ticket_category_pk.split("-")
            ticket_category_obj = TicketCategory.objects.get(
                pk=int(category_pk), bot__pk=int(bot_pk))
            if ticket_category_obj not in agent.ticket_categories.all():
                agent.ticket_categories.add(ticket_category_obj)
        agent.save()
        logger.info("Successfully exited from add_ticket_category_to_agent", extra={
            'AppName': 'EasyTMS'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_ticket_category_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def add_ticket_category_to_agent_by_pk_list(agent, selected_ticket_category_pk_list, TicketCategory):
    try:
        logger.info("Inside add_ticket_category_to_agent",
                    extra={'AppName': 'EasyTMS'})
        for ticket_category_pk in selected_ticket_category_pk_list:
            ticket_category_obj = TicketCategory.objects.get(
                pk=int(ticket_category_pk))
            if ticket_category_obj not in agent.ticket_categories.all():
                agent.ticket_categories.add(ticket_category_obj)
        agent.save()
        logger.info("Successfully exited from add_ticket_category_to_agent", extra={
            'AppName': 'EasyTMS'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_ticket_category_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def generate_password(password_prefix):
    DIGITS = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
    SYMBOLS = ['@', '#', '$', '&']

    password_digits = ""
    for itr in range(0, 4):
        password_digits = password_digits + random.choice(DIGITS)
    password = password_prefix + \
        random.choice(SYMBOLS) + password_digits + random.choice(SYMBOLS)
    return password


def update_resend_password_counter(agent):
    try:
        logger.info("Inside update_resend_password_counter", extra={
            'AppName': 'EasyTMS'})
        current_date = timezone.now().date()

        if agent.last_password_resend_date != current_date:
            agent.last_password_resend_date = timezone.now()
            agent.resend_password_count = RESEND_PASSWORD_THRESHOLD

        if agent.resend_password_count >= 0:
            agent.resend_password_count -= 1
        agent.save()

        logger.info("Successfully Exiting update_resend_password_counter", extra={
            'AppName': 'EasyTMS'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error archive_cobrowse_objects %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def get_mapped_agents(active_agent, Agent, is_active=True, is_account_active=True, include_supervisor=True,
                      include_self=True, is_absent=False):
    try:
        agents = active_agent.agents.filter(pk=-1)

        if active_agent.role == "supervisor":
            agents |= active_agent.agents.all()
        elif active_agent.role == "admin":
            agents |= active_agent.agents.filter(role='agent')

            if is_account_active is not None:
                supervisors = active_agent.agents.filter(
                    role='supervisor', is_account_active=is_account_active)
            else:
                supervisors = active_agent.agents.filter(role='supervisor')

            for supervisor in supervisors:
                agents = agents | supervisor.agents.all()
                if include_supervisor:
                    agents = agents | Agent.objects.filter(pk=supervisor.pk)

        if include_self:
            agents |= Agent.objects.filter(pk=active_agent.pk)

        if is_active is not None:
            agents = agents.filter(is_active=is_active)

        if is_absent is not None:
            agents = agents.filter(is_absent=is_absent)

        if is_account_active is not None:
            agents = agents.filter(is_account_active=is_account_active)

        return agents.distinct()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_mapped_agents: %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return []


def get_active_agent_obj(request, User, Agent):
    active_agent = None
    try:
        active_user = User.objects.filter(
            username=request.user.username).first()
        active_agent = Agent.objects.filter(
            user=active_user, is_account_active=True).first()
        return active_agent
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_active_agent_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return active_agent


def get_access_token_from_bot(bot_id, Bot, Agent, TMSAccessToken):
    bot_obj = Bot.objects.filter(pk=int(bot_id))
    agent_obj = Agent.objects.filter(role="admin", bots__in=bot_obj)[0]
    access_token_obj = TMSAccessToken.objects.get(agent=agent_obj)
    return access_token_obj


def get_admin_from_active_agent(active_agent, Agent):
    try:
        admin_obj = None
        if active_agent.role == "admin":
            admin_obj = active_agent
        elif active_agent.role == "supervisor":
            admin_obj = Agent.objects.filter(
                role="admin", agents__pk=active_agent.pk).first()
        else:
            supervisor_obj = Agent.objects.filter(
                role="supervisor", agents__pk=active_agent.pk).first()
            if supervisor_obj:
                admin_obj = Agent.objects.filter(
                    role="admin", agents__pk=supervisor_obj.pk).first()
            if admin_obj is None:
                admin_obj = Agent.objects.filter(
                    role="admin", agents__pk=active_agent.pk).first()
        return admin_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_admin_from_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return None


'''

method : get_supervisor_from_active_agent

if active_agent is agent then it's supervisor is returned
if active_agent is supervisor then it self is returned
if active_agent is admin then it self is returned

'''


def get_supervisor_from_active_agent(active_agent, Agent):
    try:
        if active_agent.role in ["supervisor", "admin"]:
            return active_agent

        supervisor = Agent.objects.filter(
            agents__pk=active_agent.pk, role="supervisor").exclude(pk=active_agent.pk).first()

        return supervisor
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_supervisor_from_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return None


def get_access_token_obj(active_agent, Agent, TMSAccessToken):
    try:
        admin_agent = get_admin_from_active_agent(active_agent, Agent)
        access_token = TMSAccessToken.objects.filter(agent=admin_agent).first()
        return access_token
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_access_token_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return None


def get_category_objs(active_agent, Agent):
    category_objs = None
    try:
        supervisor_obj = get_supervisor_from_active_agent(active_agent, Agent)
        category_objs = supervisor_obj.category.filter(
            bot__is_deleted=False, bot__is_tms_allowed=True)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_category_objs %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return category_objs


def get_relevant_agent_list(ticket_obj, Agent):
    agent_list = None
    try:
        admin_agent = ticket_obj.access_token.agent
        agent_list = Agent.objects.filter(pk=admin_agent.pk)

        agent_list |= admin_agent.agents.filter(
            role="supervisor", bots__in=[ticket_obj.bot])

        if ticket_obj.agent:
            agent_list |= Agent.objects.filter(pk=ticket_obj.agent.pk)

        agent_list = agent_list.distinct()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_relevant_agent_list %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return agent_list


def get_ticket_objs(active_agent, Agent, Ticket, TMSAccessToken, is_resolved=False):
    try:
        access_token = get_access_token_obj(
            active_agent, Agent, TMSAccessToken)

        if is_resolved == None:
            ticket_objs = Ticket.objects.filter(access_token=access_token)
        else:
            ticket_objs = Ticket.objects.filter(
                access_token=access_token, is_resolved=is_resolved)

        if active_agent.role == "agent":
            ticket_objs = ticket_objs.filter(agent=active_agent)
        elif active_agent.role == "supervisor":
            ticket_objs = ticket_objs.filter(bot__in=active_agent.bots.all())
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_ticket_objs %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return ticket_objs


def create_ticket_priority_obj(name, access_token, TicketPriority):
    ticket_priority_obj = TicketPriority.objects.filter(
        name=name, access_token=access_token).first()
    if ticket_priority_obj == None:
        ticket_priority_obj = TicketPriority.objects.create(
            name=name, access_token=access_token)

    access_token.ticket_priorities.add(ticket_priority_obj)
    access_token.save()

    return ticket_priority_obj


def create_ticket_status_obj(name, access_token, TicketStatus):
    ticket_status_obj = TicketStatus.objects.filter(
        name=name, access_token=access_token).first()
    if ticket_status_obj == None:
        ticket_status_obj = TicketStatus.objects.create(
            name=name, access_token=access_token)

    access_token.ticket_statuses.add(ticket_status_obj)
    access_token.save()

    return ticket_status_obj


def parse_ticket_details(ticket_obj):
    def get_bot(ticket_obj):
        data = None
        try:
            data = {
                'pk': ticket_obj.bot.pk,
                'name': ticket_obj.bot.name,
                'bot_display_name': ticket_obj.bot.bot_display_name,
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_bot %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_bot_channel(ticket_obj):
        data = None
        try:
            data = {
                'pk': ticket_obj.bot_channel.pk,
                'name': ticket_obj.bot_channel.name,
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_bot_channel %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_ticket_category(ticket_obj):
        data = None
        try:
            data = {
                'pk': ticket_obj.ticket_category.pk,
                'ticket_category': ticket_obj.ticket_category.ticket_category,
                'is_deleted': ticket_obj.ticket_category.is_deleted,
                'ticket_period': ticket_obj.ticket_category.ticket_period,
                'created_datetime': get_formatted_date(ticket_obj.ticket_category.created_datetime),
                'bot': get_bot(ticket_obj),
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_ticket_category %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_ticket_status(ticket_obj):
        data = None
        try:
            data = {
                'pk': ticket_obj.ticket_status.pk,
                'name': ticket_obj.ticket_status.name,
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_ticket_status %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_ticket_priority(ticket_obj):
        data = None
        try:
            data = {
                'pk': ticket_obj.ticket_priority.pk,
                'name': ticket_obj.ticket_priority.name,
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_ticket_priority %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_formatted_date(date_obj):
        data = None
        try:
            est = pytz.timezone(settings.TIME_ZONE)
            data = date_obj.astimezone(
                est).strftime(YEAR_TIME_AM_PM_FORMAT)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_formatted_date %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_agent_details(agent_obj):
        agent_detail = None
        try:
            agent_detail = {
                "pk": agent_obj.pk,
                "agent_username": agent_obj.user.username,
                "name": get_agent_name(agent_obj),
                "role": agent_obj.role,
                "level": agent_obj.level,
                "phone_number": agent_obj.phone_number,
                "is_absent": agent_obj.is_absent,
                "is_active": agent_obj.is_active,
                "is_account_active": agent_obj.is_account_active,
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_agent_details %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return agent_detail

    ticket_data = None
    try:
        access_token_key = None
        if ticket_obj.access_token:
            access_token_key = str(ticket_obj.access_token.key)

        customer_name = ticket_obj.customer_name if ticket_obj.customer_name else ""
        customer_email_id = ticket_obj.customer_email_id if ticket_obj.customer_email_id else ""
        customer_mobile_number = ticket_obj.customer_mobile_number if ticket_obj.customer_mobile_number else ""
        customer_information = json.loads(
            ticket_obj.customer_information) if ticket_obj.customer_information else {}

        agent_name = ""
        agent_details = get_agent_details(ticket_obj.agent)
        if agent_details:
            agent_name = agent_details['name']

        ticket_data = {
            'ticket_id': str(ticket_obj.ticket_id),
            'access_token_key': str(access_token_key),
            'agent_name': agent_name,
            'agent_details': agent_details,
            'issue_date_time': get_formatted_date(ticket_obj.issue_date_time),
            'query_description': ticket_obj.query_description,
            'query_attachment': ticket_obj.query_attachment,
            'bot': get_bot(ticket_obj),
            'bot_channel': get_bot_channel(ticket_obj),
            'ticket_category': get_ticket_category(ticket_obj),
            'ticket_status': get_ticket_status(ticket_obj),
            'ticket_priority': get_ticket_priority(ticket_obj),
            'is_resolved': ticket_obj.is_resolved,
            'resolved_date_time': get_formatted_date(ticket_obj.resolved_date_time),
            'customer_name': customer_name,
            'customer_email_id': customer_email_id,
            'customer_mobile_number': customer_mobile_number,
            'customer_information': customer_information,
        }

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_ticket_details %s at %s for ticket id %s",
                     str(e), str(exc_tb.tb_lineno), str(ticket_obj.ticket_id), extra={'AppName': 'EasyTMS'})

    return ticket_data


def auto_assign_agent(category_obj, bot_obj, access_token_obj, Agent, Bot, Ticket):
    try:
        agent_dict = {}

        admin_agent = access_token_obj.agent

        active_agents = get_mapped_agents(admin_agent, Agent, is_active=True, is_account_active=True,
                                          include_supervisor=False, include_self=False, is_absent=False)
        active_agents = active_agents.filter(
            bots__in=[bot_obj], ticket_categories__in=[category_obj])

        agent_dict = {}
        for agent in active_agents:
            agent_dict[agent.user.username] = []
            agent_dict[agent.user.username].append(0)
        if not agent_dict:
            logger.info("Active Agents are not available",
                        extra={'AppName': 'EasyTMS'})
            return None

        not_resolved_tickt_objs = Ticket.objects.filter(
            agent__in=active_agents, is_resolved=False)

        for ticket_obj in not_resolved_tickt_objs:
            if ticket_obj.agent.user.username in agent_dict:
                agent_dict[ticket_obj.agent.user.username][0] += 1

        agent_min_lead_entry = min(
            agent_dict.items(), key=lambda item: (item[1][0]))

        logger.info("agent_min_lead_entry: " +
                    str(agent_min_lead_entry), extra={'AppName': 'EasyTMS'})

        agent_username = agent_min_lead_entry[0]
        agent_active_leads_count = agent_min_lead_entry[1][0]

        logger.info("agent_dict : " + str(agent_dict),
                    extra={'AppName': 'EasyTMS'})

        logger.info("Agent with min lead: " + str(agent_username) + ":" +
                    str(agent_active_leads_count), extra={'AppName': 'EasyTMS'})

        selected_agent = Agent.objects.filter(
            user__username=agent_username).first()

        return selected_agent
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Agent autoassign error: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return None


def change_ticket_status(name, ticket_obj, TicketStatus):
    if not ticket_obj or name == "":
        return False

    try:
        ticket_status_obj = TicketStatus.objects.filter(
            access_token=ticket_obj.access_token,
            name=name,
        ).first()
        if ticket_status_obj:
            ticket_obj.ticket_status = ticket_status_obj
            ticket_obj.save()

        return True
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Agent change_ticket_status error: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return False


def get_agent_console_meta_data(agent_obj):
    meta_data = {"lead_data_cols": DEFAULT_LEAD_TABLE_COL}
    try:
        if agent_obj.role == "agent":
            meta_data["lead_data_cols"] = [
                obj for obj in DEFAULT_LEAD_TABLE_COL if obj["name"] != "agent_name"]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Agent change_ticket_status error: %s at %s",
                     e, str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return json.dumps(meta_data)


def validate_user_new_password(active_agent, new_password, old_password, AgentPasswordHistory):
    agent_password_history_objs = AgentPasswordHistory.objects.filter(
        agent=active_agent).order_by('-datetime')[:5]
    old_password_list = [
        item.password_hash for item in agent_password_history_objs]

    new_password_hash = hashlib.sha256(new_password.encode()).hexdigest()
    old_password_hash = hashlib.sha256(old_password.encode()).hexdigest()

    if old_password_hash not in old_password_list:
        old_password_list.append(old_password_hash)

    if new_password_hash in old_password_list:
        return "SIMILAR_TO_OLD_PASSWORD"

    return "VALID"


def check_malicious_file_from_filename(filename, allowed_files_list=None):
    if allowed_files_list == None:
        allowed_files_list = [
            "png", "jpg", "jpeg", "bmp", "gif", "tiff", "exif", "jfif", "webm", "mpg", "jpe",
            "mp2", "mpeg", "mpe", "mpv", "ogg", "mp4", "m4p", "m4v", "avi", "wmv", "mov", "qt", "doc", "docx",
            "flv", "swf", "avchd", "mp3", "aac", "pdf", "xls", "xlsx", "json", "xlsm", "xlt", "xltm", "zip", "xlb"
        ]

    logger.info("In check_malicious_file_from_filename: %s ", str(filename), extra={
        'AppName': 'EasyTMS'})

    try:
        if len(filename.split('.')) != 2:
            return True

        file_extension = filename.split('.')[-1].lower().strip()

        if file_extension not in allowed_files_list:
            return True
        else:
            return False
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In check_malicious_file_from_filename: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})
        return True


def check_malicious_file(uploaded_file):
    return check_malicious_file_from_filename(uploaded_file.name)


def check_malicious_file_from_content(base64_data, allowed_files_list=None):

    # Untill we find any concrete solution for detecting file type
    return False
    # decoded = base64.b64decode(base64_data)
    # mime_type = magic.from_buffer(decoded, mime=True)
    # file_ext = mimetypes.guess_extension(mime_type)
    # return check_malicious_file_from_filename(file_ext, allowed_files_list)


def save_ticket_audit_trail(ticket, agent, action_type, description, TicketAuditTrail):
    ticket_audit_trail_obj = TicketAuditTrail.objects.create(ticket=ticket,
                                                             agent=agent,
                                                             action_type=action_type,
                                                             description=description)

    return ticket_audit_trail_obj


def send_action_info_to_agent(agent, action_name, action_info):
    try:
        data = {
            "action_name": action_name,
            "action_info": action_info
        }

        send_data_from_server_to_client(
            "data_changed", data, agent.user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error notify_agent_data_change_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return False


def send_notification_to_agent(agent, notification_message):
    try:
        data = {
            "notification_message": notification_message
        }

        send_data_from_server_to_client(
            "notification", data, agent.user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error send_notification_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return False


def create_user_notification(agent_obj, ticket_obj, description, UserNotification):
    try:
        user_notification_obj = UserNotification.objects.create(
            agent=agent_obj,
            ticket=ticket_obj,
            description=description
        )
        return user_notification_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_user_notification %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return None


def process_request_data(request_data, access_token):
    from EasyTMSApp.models import TicketCategory, Agent, TicketPriority, TicketStatus, TMSAccessToken
    from EasyChatApp.models import Channel, Bot

    def get_request_value(key):
        value = ""
        try:
            value = request_data[key]
        except Exception:
            pass
        return value

    customer_name = get_request_value("customer_name")
    query_description = get_request_value("issue_description")
    query_attachment = get_request_value("query_attachment")
    customer_mobile_number = get_request_value("customer_mobile_number")
    customer_email_id = get_request_value("customer_email_id")

    agent_username = get_request_value("agent_username")
    bot_id = get_request_value("bot_id")
    ticket_category = get_request_value("ticket_category")
    ticket_priority = get_request_value("ticket_priority")
    channel_name = get_request_value("channel_name")

    processed_data = {}

    processed_data["customer_name"] = customer_name
    processed_data["query_description"] = query_description
    processed_data["query_attachment"] = query_attachment
    processed_data["customer_mobile_number"] = customer_mobile_number
    processed_data["customer_email_id"] = customer_email_id
    processed_data["access_token_obj"] = access_token

    try:
        possible_agents = get_mapped_agents(access_token.agent, Agent, is_active=None, is_account_active=True,
                                            include_supervisor=True, include_self=True, is_absent=None)
        agent_obj = possible_agents.filter(
            user__username=agent_username).first()
        if agent_obj == None:
            agent_obj = access_token.agent
        ticket_status_obj = TicketStatus.objects.filter(
            name="PENDING", access_token=access_token).first()
    except Exception:
        agent_obj = None
        ticket_status_obj = TicketStatus.objects.filter(
            name="UNASSIGNED", access_token=access_token).first()
    processed_data["agent_obj"] = agent_obj
    processed_data["ticket_status_obj"] = ticket_status_obj

    bot_obj = None
    try:
        bot_obj = access_token.agent.bots.filter(pk=int(bot_id)).first()
    except Exception:
        pass
    processed_data["bot_obj"] = bot_obj

    category_obj = None
    try:
        category_obj = TicketCategory.objects.get(
            bot=bot_obj, ticket_category__iexact=ticket_category)
    except Exception:
        category_obj = TicketCategory.objects.filter(
            bot=bot_obj, ticket_category__iexact="OTHERS").first()
    processed_data["category_obj"] = category_obj

    try:
        priority_obj = TicketPriority.objects.get(
            name__iexact=ticket_priority, access_token=access_token)
    except Exception:
        priority_obj = TicketPriority.objects.filter(
            name="LOW", access_token=access_token).first()
    processed_data["priority_obj"] = priority_obj

    try:
        bot_channel_obj = Channel.objects.filter(
            name__iexact=str(channel_name)).first()
    except Exception:
        bot_channel_obj = None
    processed_data["bot_channel_obj"] = bot_channel_obj

    return processed_data


def create_ticket_from_processed_data(processed_data):
    from EasyTMSApp.models import Ticket

    ticket_obj = Ticket.objects.create(
        customer_name=processed_data["customer_name"],
        query_description=processed_data["query_description"],
        query_attachment=processed_data["query_attachment"],
        customer_mobile_number=processed_data["customer_mobile_number"],
        customer_email_id=processed_data["customer_email_id"],
        access_token=processed_data["access_token_obj"],
        agent=processed_data["agent_obj"],
        ticket_status=processed_data["ticket_status_obj"],
        bot=processed_data["bot_obj"],
        ticket_category=processed_data["category_obj"],
        ticket_priority=processed_data["priority_obj"],
        bot_channel=processed_data["bot_channel_obj"],
    )

    return str(ticket_obj.ticket_id)


def get_ticket_data(ticket_obj):
    def get_bot(ticket_obj):
        data = None
        try:
            data = {
                'pk': ticket_obj.bot.pk,
                'name': ticket_obj.bot.bot_display_name,
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_bot %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_bot_channel(ticket_obj):
        data = ""
        try:
            data = ticket_obj.bot_channel.name
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_bot_channel %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_ticket_category(ticket_obj):
        data = ""
        try:
            data = ticket_obj.ticket_category.ticket_category
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_ticket_category %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_ticket_status(ticket_obj):
        data = ""
        try:
            data = ticket_obj.ticket_status.name
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_ticket_status %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_ticket_priority(ticket_obj):
        data = None
        try:
            data = ticket_obj.ticket_priority.name
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_ticket_priority %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    def get_formatted_date(date_obj):
        data = None
        try:
            est = pytz.timezone(settings.TIME_ZONE)
            data = date_obj.astimezone(
                est).strftime("%d %b %Y %I:%M:%S %p")
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_formatted_date %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return data

    ticket_data = None
    try:
        access_token_key = None
        if ticket_obj.access_token:
            access_token_key = str(ticket_obj.access_token.key)

        customer_name = ticket_obj.customer_name if ticket_obj.customer_name else ""
        customer_email_id = ticket_obj.customer_email_id if ticket_obj.customer_email_id else ""
        customer_mobile_number = ticket_obj.customer_mobile_number if ticket_obj.customer_mobile_number else ""

        agent_username = ""
        if ticket_obj.agent:
            agent_username = ticket_obj.agent.user.username

        ticket_data = {
            'ticket_id': str(ticket_obj.ticket_id),
            'access_token_key': access_token_key,
            'agent_username': agent_username,
            'issue_date_time': get_formatted_date(ticket_obj.issue_date_time),
            'query_description': ticket_obj.query_description,
            'query_attachment': ticket_obj.query_attachment,
            'bot': get_bot(ticket_obj),
            'bot_channel': get_bot_channel(ticket_obj),
            'ticket_category': get_ticket_category(ticket_obj),
            'ticket_status': get_ticket_status(ticket_obj),
            'ticket_priority': get_ticket_priority(ticket_obj),
            'is_resolved': ticket_obj.is_resolved,
            'resolved_date_time': get_formatted_date(ticket_obj.resolved_date_time),
            'customer_name': customer_name,
            'customer_email_id': customer_email_id,
            'customer_mobile_number': customer_mobile_number
        }

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_ticket_data %s at %s for ticket id %s",
                     str(e), str(exc_tb.tb_lineno), str(ticket_obj.ticket_id), extra={'AppName': 'EasyTMS'})

    return ticket_data


def get_ticket_activity_data(audit_trail_obj):
    def get_agent_info(agent_obj):
        agent_detail = {}
        try:
            agent_detail = {
                "agent_username": agent_obj.user.username,
                "role": agent_obj.role
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error get_agent_info %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        return agent_detail

    audit_trail_data = {}
    try:
        audit_trail_data = {
            "ticket_id": audit_trail_obj.ticket.ticket_id,
            "agent_detail": get_agent_info(audit_trail_obj.agent),
            "action_type": audit_trail_obj.action_type,
            "description": audit_trail_obj.description,
            "datetime": get_formatted_date(audit_trail_obj.datetime),
        }
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_ticket_activity_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return audit_trail_data


def create_directory(directory_path, remove_first=False):
    if remove_first and os.path.exists(directory_path):
        try:
            shutil.rmtree(directory_path)
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_directory %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)


def generate_crm_api_document(access_token):

    def docx_replace_regex(doc_obj, regex, replace_text):
        for para in doc_obj.paragraphs:
            if regex.search(para.text):
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for index in range(len(inline)):
                    if regex.search(inline[index].text):
                        text = regex.sub(replace_text, inline[index].text)
                        inline[index].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace_text)

    def edit_document(document_path, target_path, metadata_dict):
        document_obj = Document(document_path)

        for metadata in metadata_dict:
            docx_replace_regex(document_obj, re.compile(
                r"" + metadata + ""), metadata_dict[metadata])

        document_obj.save(target_path)

    def get_document_folder_path():
        target_document_folder = "secured_files/EasyTMSApp/CRMDocuments/" + \
            str(access_token.key)
        create_directory(target_document_folder)
        return target_document_folder

    def get_base_document_path(document_details):
        original_file_name = document_details["original_file_name"]
        base_document_path = "files/templates/tms-crm-template/" + original_file_name
        return base_document_path

    def get_target_document_path(document_details):
        target_document_folder = get_document_folder_path()
        original_file_name = document_details["original_file_name"]
        target_document_path = target_document_folder + "/" + original_file_name
        return target_document_path

    def get_metadata_dict():
        crm_integration_url = settings.EASYCHAT_HOST_URL + \
            "/tms/crm/" + document_details["url_suffix"] + "/"

        metadata_dict = {
            "crm_integration_url": crm_integration_url
        }

        return metadata_dict

    try:
        for document_type, document_details in CRM_DOCUMENTS.items():
            base_document_path = get_base_document_path(document_details)
            target_document_path = get_target_document_path(document_details)
            metadata_dict = get_metadata_dict()
            edit_document(base_document_path,
                          target_document_path, metadata_dict)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("generate_crm_api_document! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyTMS'})
        return None


def get_crm_integration_model(request, CRMIntegrationModel):
    try:
        authorization_header = request.META['HTTP_AUTHORIZATION']
        auth_token = authorization_header.replace("Bearer ", "")

        min_datetime = datetime.now() - timedelta(hours=1)
        integration_obj = CRMIntegrationModel.objects.filter(
            auth_token=auth_token, is_expired=False, datetime__gte=min_datetime).first()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_crm_integration_model! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyTMS'})
        integration_obj = None

    return integration_obj


def ticket_resolved_action(active_agent, ticket_obj, TicketStatus, TicketAuditTrail, UserNotification, Agent):
    ticket_obj.is_resolved = True
    ticket_obj.resolved_date_time = datetime.now()

    old_ticket_status = ticket_obj.ticket_status

    if old_ticket_status == None or (old_ticket_status and old_ticket_status.name != "RESOLVED"):
        change_ticket_status("RESOLVED", ticket_obj, TicketStatus)

        try:
            action_type = "STATUS_CHANGED"
            if old_ticket_status == None:
                description = active_agent.user.username + \
                    " set status to " + "RESOLVED".title()
            else:
                description = active_agent.user.username + " changed status from " + \
                    old_ticket_status.name.title() + " to " + "RESOLVED".title()
            save_ticket_audit_trail(
                ticket_obj, active_agent, action_type, description, TicketAuditTrail)

            agent_objs = get_relevant_agent_list(ticket_obj, Agent)

            for agent_obj in agent_objs:

                if agent_obj != active_agent:
                    create_user_notification(
                        agent_obj, ticket_obj, description, UserNotification)
                    send_action_info_to_agent(agent_obj, action_name="new_user_notification", action_info={
                        "send_notification": True,
                        "notification_message": description,
                        "ticket_id": ticket_obj.ticket_id
                    })
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error on Audit Trail Save %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    ticket_obj.save()


def generate_agent_query(agent, ticket, access_token, ticket_audit_trail, extra_information, AgentQuery):
    try:
        agent_query_obj = AgentQuery.objects.create(
            agent=agent,
            ticket=ticket,
            ticket_audit_trail=ticket_audit_trail,
            extra_information=extra_information
        )

        agent_query_info = parse_agent_query_related_info(agent_query_obj)

        thread_whatsapp = threading.Thread(
            target=send_agent_query_over_whatsapp,
            args=(
                access_token,
                agent_query_info,
            ),
            daemon=True
        )
        thread_whatsapp.start()

        thread_email = threading.Thread(
            target=send_agent_query_over_email,
            args=(
                access_token,
                agent_query_info,
            ),
            daemon=True
        )

        thread_email.start()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error generate_agent_query %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def get_agent_query(agent_query_id):
    from EasyTMSApp.models import AgentQuery
    response = {}
    agent_message = ""
    try:
        agent_customer_query_obj = AgentQuery.objects.filter(
            pk=str(agent_query_id)).first()
        agent_message = agent_customer_query_obj.ticket_audit_trail.description
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_agent_query %s at %s", str(e), str(
            exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    response["agent_message"] = str(agent_message)
    return response


def save_customer_reply_on_agent_query(ticket_id, customer_message, attachment):
    from EasyTMSApp.models import TicketAuditTrail, Ticket, Agent, UserNotification
    from EasyTMSApp.utils import save_ticket_audit_trail, create_user_notification, send_action_info_to_agent

    ticket_obj = Ticket.objects.get(ticket_id=str(ticket_id))

    if attachment == None or attachment == "None":
        attachment = ""

    action_type = "CUSTOMER_COMMENT"
    description = {
        "customer_message": customer_message,
        "attachment": attachment,
        "attachment_name": "Click to download"
    }
    description = json.dumps(description)
    save_ticket_audit_trail(ticket_obj, None, action_type,
                            description, TicketAuditTrail)

    try:
        agent_objs = get_relevant_agent_list(ticket_obj, Agent)
        description = "Customer " + ticket_obj.customer_name + " added comment."
        for agent_obj in agent_objs:
            create_user_notification(
                agent_obj, ticket_obj, description, UserNotification)
            send_action_info_to_agent(agent_obj, action_name="new_user_notification", action_info={
                "send_notification": True,
                "notification_message": description,
                "ticket_id": ticket_obj.ticket_id
            })
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_customer_reply_on_agent_query %s at %s", str(e), str(exc_tb.tb_lineno),
                     extra={'AppName': 'EasyTMS'})


def mark_agent_customer_query_resolved(agent_query_id):
    from EasyTMSApp.models import AgentQuery

    agent_customer_query_obj = AgentQuery.objects.filter(
        pk=str(agent_query_id)).first()
    agent_customer_query_obj.is_active = False
    agent_customer_query_obj.save()


def check_agent_message_sent_to_customer(audit_trail_obj):
    from EasyTMSApp.models import AgentQuery

    agent_query_obj = None
    try:
        agent_query_obj = AgentQuery.objects.filter(
            ticket_audit_trail=audit_trail_obj).filter().first()
    except Exception:
        pass

    return agent_query_obj != None


def parse_ticket_audit_trail(audit_trail_obj):
    audit_trail_data = parse_ticket_audit_trail_model(audit_trail_obj)
    try:
        audit_trail_data["is_sent_to_customer"] = check_agent_message_sent_to_customer(
            audit_trail_obj)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_ticket_audit_trail %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return audit_trail_data


def file_download(file_key, FileAccessManagement):
    try:
        file_access_management_obj = FileAccessManagement.objects.get(
            key=file_key)

        path_to_file = file_access_management_obj.file_path
        original_file_name = file_access_management_obj.original_file_name

        if original_file_name is None:
            original_file_name = path_to_file.split("/")[-1]

        path_to_file = settings.BASE_DIR + path_to_file
        mime_type, _ = mimetypes.guess_type(path_to_file)

        if os.path.exists(path_to_file):
            with open(path_to_file, 'rb') as fh:
                response = HttpResponse(
                    fh.read(), status=200, content_type=mime_type)
                response['Content-Disposition'] = 'attachment; filename="%s"' % smart_str(
                    str(original_file_name))
                return response
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error FileDownload %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return HttpResponse(status=404)


def get_requested_data_for_daily():
    # Last day data
    date_format = YYYY_MM_DD_FORMAT

    start_date = datetime.now() - timedelta(days=1)
    start_date = start_date.strftime(date_format)
    end_date = start_date

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def get_requested_data_for_week():
    # Last 7 days data
    date_format = YYYY_MM_DD_FORMAT

    start_date = datetime.now() - timedelta(days=7)
    start_date = start_date.strftime(date_format)
    end_date = datetime.now() - timedelta(days=1)
    end_date = end_date.strftime(date_format)

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def get_requested_data_for_month():
    # Last 30 days data
    date_format = YYYY_MM_DD_FORMAT

    start_date = datetime.now() - timedelta(days=30)
    start_date = start_date.strftime(date_format)
    end_date = datetime.now() - timedelta(days=1)
    end_date = end_date.strftime(date_format)

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def get_requested_data_custom(start_date, end_date):
    date_format = YYYY_MM_DD_FORMAT

    start_date = start_date.strftime(date_format)
    end_date = end_date.strftime(date_format)

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def add_general_analytics_in_excel_sheet(data_dump_wb, active_agent, start_date, end_date):
    def create_new_sheet():
        new_sheet = data_dump_wb.add_sheet("Genral Analytics")

        new_sheet.write(0, 0, "Date")
        new_sheet.col(0).width = 256 * 20
        new_sheet.write(0, 1, "Total Ticket Generated")
        new_sheet.col(1).width = 256 * 20
        new_sheet.write(0, 2, "Resolved")
        new_sheet.col(2).width = 256 * 20
        new_sheet.write(0, 3, "Pending")
        new_sheet.col(3).width = 256 * 20
        new_sheet.write(0, 4, "Unassigned")
        new_sheet.col(4).width = 256 * 20

        return new_sheet

    try:
        from EasyTMSApp.models import Ticket, TMSAccessToken, TicketStatus, Agent

        access_token = get_access_token_obj(
            active_agent, Agent, TMSAccessToken)

        ticket_objs = get_ticket_objs(
            active_agent, Agent, Ticket, TMSAccessToken, None)

        status__resolved = TicketStatus.objects.filter(
            access_token=access_token, name__iexact="RESOLVED").first()
        status__pending = TicketStatus.objects.filter(
            access_token=access_token, name__iexact="PENDING").first()
        status__unassigned = TicketStatus.objects.filter(
            access_token=access_token, name__iexact="UNASSIGNED").first()

        total_sum_count = {
            'total_ticket_generated': 0,
            'total_ticket_resolved': 0,
            'total_ticket_pending': 0,
            'total_ticket_unassigned': 0,
        }

        sheet = create_new_sheet()

        index = 1
        while start_date <= end_date:
            excel_sheet_date_format = start_date.strftime("%d-%m-%Y")

            filtered_ticket_objs = ticket_objs.filter(
                issue_date_time__date=start_date)

            total_ticket_generated = filtered_ticket_objs.count()
            total_ticket_resolved = filtered_ticket_objs.filter(
                ticket_status=status__resolved).count()
            total_ticket_pending = filtered_ticket_objs.filter(
                ticket_status=status__pending).count()
            total_ticket_unassigned = filtered_ticket_objs.filter(
                ticket_status=status__unassigned).count()

            sheet.write(index, 0, excel_sheet_date_format)
            sheet.write(index, 1, total_ticket_generated)
            sheet.write(index, 2, total_ticket_resolved)
            sheet.write(index, 3, total_ticket_pending)
            sheet.write(index, 4, total_ticket_unassigned)

            total_sum_count['total_ticket_generated'] += total_ticket_generated
            total_sum_count['total_ticket_resolved'] += total_ticket_resolved
            total_sum_count['total_ticket_pending'] += total_ticket_pending
            total_sum_count['total_ticket_unassigned'] += total_ticket_unassigned

            index += 1
            start_date = start_date + timedelta(1)

        sheet.write(index, 0, "Sum")
        sheet.write(index, 1, total_sum_count['total_ticket_generated'])
        sheet.write(index, 2, total_sum_count['total_ticket_resolved'])
        sheet.write(index, 3, total_sum_count['total_ticket_pending'])
        sheet.write(index, 4, total_sum_count['total_ticket_unassigned'])

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("add_general_analytics_in_excel_sheet %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})


def add_agent_wise_analytics_in_excel_sheet(data_dump_wb, active_agent, start_date, end_date):
    def create_new_sheet():
        new_sheet = data_dump_wb.add_sheet("Agent Wise Analytics")

        new_sheet.write(0, 0, "Agent")
        new_sheet.col(0).width = 256 * 20
        new_sheet.write(0, 1, "Total Ticket Assigned")
        new_sheet.col(1).width = 256 * 20
        new_sheet.write(0, 2, "Resolved")
        new_sheet.col(2).width = 256 * 20
        new_sheet.write(0, 3, "Pending")
        new_sheet.col(3).width = 256 * 20

        return new_sheet

    try:
        from EasyTMSApp.models import Ticket, TMSAccessToken, TicketStatus, Agent

        access_token = get_access_token_obj(
            active_agent, Agent, TMSAccessToken)

        ticket_objs = get_ticket_objs(
            active_agent, Agent, Ticket, TMSAccessToken, None)

        ticket_objs = ticket_objs.filter(
            issue_date_time__date__gte=start_date,
            issue_date_time__date__lte=end_date,
        )

        agents = get_mapped_agents(
            active_agent,
            Agent,
            is_active=None,
            is_account_active=None,
            include_supervisor=True,
            include_self=True,
            is_absent=None
        )

        status__resolved = TicketStatus.objects.filter(
            access_token=access_token, name__iexact="RESOLVED").first()
        status__pending = TicketStatus.objects.filter(
            access_token=access_token, name__iexact="PENDING").first()

        total_sum_count = {
            'total_ticket_generated': 0,
            'total_ticket_resolved': 0,
            'total_ticket_pending': 0,
        }

        sheet = create_new_sheet()

        index = 1
        for agent in agents:
            filtered_ticket_objs = ticket_objs.filter(agent=agent)

            total_ticket_generated = filtered_ticket_objs.count()
            total_ticket_resolved = filtered_ticket_objs.filter(
                ticket_status=status__resolved).count()
            total_ticket_pending = filtered_ticket_objs.filter(
                ticket_status=status__pending).count()

            sheet.write(index, 0, agent.user.username)
            sheet.write(index, 1, total_ticket_generated)
            sheet.write(index, 2, total_ticket_resolved)
            sheet.write(index, 3, total_ticket_pending)

            total_sum_count['total_ticket_generated'] += total_ticket_generated
            total_sum_count['total_ticket_resolved'] += total_ticket_resolved
            total_sum_count['total_ticket_pending'] += total_ticket_pending

            index += 1
            start_date = start_date + timedelta(1)

        sheet.write(index, 0, "Sum")
        sheet.write(index, 1, total_sum_count['total_ticket_generated'])
        sheet.write(index, 2, total_sum_count['total_ticket_resolved'])
        sheet.write(index, 3, total_sum_count['total_ticket_pending'])

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("add_agent_wise_analytics_in_excel_sheet %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})


def get_custom_analytics_export(requested_data, agent):
    def get_folder_path():
        folder_path = "EasyTMSApp/Analytics/" + str(agent.user.username) + "/"
        return folder_path

    def get_file_name(extension):
        file_name = "analytics_" + \
            str(start_date) + "_to_" + str(end_date) + extension
        return file_name

    def get_relative_file_path(extension=".xls"):
        return "secured_files/" + get_folder_path() + get_file_name(extension)

    def get_absolute_file_path(extension=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not os.path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        return absolute_folder_path + get_file_name(extension)

    try:
        logger.info("Inside get_custom_analytics_export",
                    extra={'AppName': 'EasyTMS'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = YYYY_MM_DD_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if os.path.exists(absolute_file_path) and end_date < yesterdays_date:
            return relative_file_path

        analytics_workbook = Workbook()

        add_general_analytics_in_excel_sheet(
            analytics_workbook, agent, start_date, end_date)
        add_agent_wise_analytics_in_excel_sheet(
            analytics_workbook, agent, start_date, end_date)

        analytics_workbook.save(absolute_file_path)
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_analytics_export! %s %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyTMS'})

        return "None"


def create_excel_wrong_user_data(wrong_data_list):
    filename = None
    try:
        logger.info("In create_excel_wrong_user_data",
                    extra={'AppName': 'EasyTMS'})

        wrong_data_workbook = Workbook()
        wrong_data_sheet = wrong_data_workbook.add_sheet(
            "Wrong Information Sheet")

        style = XFStyle()

        font = Font()
        font.bold = True
        style.font = font

        wrong_data_sheet.write(0, 0, "Row No", style=style)
        wrong_data_sheet.col(0).width = 256 * 10
        wrong_data_sheet.write(0, 1, "Detail", style=style)
        wrong_data_sheet.col(1).width = 256 * 40
        wrong_data_sheet.write(0, 2, "Name", style=style)
        wrong_data_sheet.col(2).width = 256 * 20
        wrong_data_sheet.write(0, 3, "Email", style=style)
        wrong_data_sheet.col(3).width = 256 * 20
        wrong_data_sheet.write(0, 4, "Mobile number", style=style)
        wrong_data_sheet.col(4).width = 256 * 20
        wrong_data_sheet.write(0, 5, "User type", style=style)
        wrong_data_sheet.col(5).width = 256 * 20
        wrong_data_sheet.write(0, 6, "Bot ID", style=style)
        wrong_data_sheet.col(6).width = 256 * 20
        wrong_data_sheet.write(0, 7, "Ticket Category", style=style)
        wrong_data_sheet.col(7).width = 256 * 20
        wrong_data_sheet.write(0, 8, "Supervisor Email/Name", style=style)
        wrong_data_sheet.col(8).width = 256 * 20

        row_number = 1
        for wrong_data in wrong_data_list:
            data_list = [wrong_data["row_num"], wrong_data["detail"]]
            data_list += wrong_data["data_array"]
            col_number = 0
            for data in data_list:
                wrong_data_sheet.write(row_number, col_number, data)
                col_number += 1
            row_number += 1

        current_datetime = str(
            datetime.now().astimezone(pytz.timezone(settings.TIME_ZONE)).strftime("%d%b%Y%I:%M%p"))
        filename = "files/TMSWrongUserData/Wrong_Data_Excel_" + \
            str(current_datetime) + ".xls"
        wrong_data_workbook.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_wrong_user_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
        filename = None
    return filename


def add_users_from_excel_document(file_path, active_agent):
    from EasyTMSApp.models import Agent, TicketCategory
    from EasyChatApp.models import Bot, User

    def process_user_first_last_name(first_last_name):
        first_last_name = first_last_name.strip()

        if len(first_last_name) == 0:
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Name Empty"})
        if not re.search(reg_name, first_last_name):
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Name Wrong Format"})

        return first_last_name

    def process_user_email(email_id):
        email_id = email_id.strip()

        if len(email_id) == 0:
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Email Empty"})
        if not re.search(reg_email, email_id):
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Email Wrong Format"})
        if Agent.objects.filter(user__email=email_id, is_account_active=True).count() > 0:
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Email Already Exists"})

        return email_id

    def process_user_mobile(phone_number):
        phone_number = phone_number.strip()
        if phone_number == "":
            phone_number = None
        if phone_number is not None:
            INVALID_MOBILE_NUMBER_MESSAGE = "Invalid Mobile Number"
            if len(phone_number) < 10:
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": INVALID_MOBILE_NUMBER_MESSAGE})
            try:
                phone_number = str(int(float(phone_number)))
            except Exception:
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": INVALID_MOBILE_NUMBER_MESSAGE})
            if not re.search(reg_mob, phone_number):
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": INVALID_MOBILE_NUMBER_MESSAGE})
            if Agent.objects.filter(phone_number=phone_number, is_account_active=True).count() > 0:
                logger.info(
                    str(phone_number) + " already exists for a user", extra={'AppName': 'EasyTMS'})
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Mobile Number Already Exists"})

        return phone_number

    def process_user_type(agent_type):
        agent_type = agent_type.strip()
        if len(agent_type) == 0:
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Invalid User Type"})

        agent_type = agent_type.lower()
        if agent_type != "agent" and agent_type != "supervisor":
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Invalid User Type"})

        return agent_type

    def process_agent_bot_ids_input(bot_ids_raw_input, parent_agent_obj):
        pk_list = []
        try:
            logger.info("bot_ids_raw_input = " +
                        str(bot_ids_raw_input), extra={"AppName": "EasyTMS"})
            bot_ids_raw_input = str(bot_ids_raw_input).strip()

            if bot_ids_raw_input == "":
                for bot_obj in parent_agent_obj.bots.all():
                    pk_list.append(bot_obj.pk)
            else:
                for bot_id in bot_ids_raw_input.split(","):
                    bot_id = str(bot_id).strip()
                    bot_id = float(bot_id)
                    bot_id = int(bot_id)
                    bot_obj = parent_agent_obj.bots.get(pk=int(bot_id))
                    pk_list.append(bot_obj.pk)

            logger.info("pk_list = " + str(pk_list),
                        extra={"AppName": "EasyTMS"})
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error process_agent_bot_ids_input %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Invalid Bot ID"})

        return pk_list

    def process_ticket_category_input(ticket_category_raw_input, parent_agent_obj, extracted_bot_pk_list):
        pk_list = []

        try:
            logger.info("ticket_category_raw_input =." +
                        str(ticket_category_raw_input) + ".", extra={"AppName": "EasyTMS"})
            ticket_category_raw_input = str(ticket_category_raw_input).strip()

            if ticket_category_raw_input == "":
                for ticket_category_obj in parent_agent_obj.ticket_categories.all():
                    pk_list.append(ticket_category_obj.pk)
            else:
                for bot_id_ticket_category in ticket_category_raw_input.split(","):
                    logger.info("bot_id_ticket_category " +
                                str(bot_id_ticket_category), extra={"AppName": "EasyTMS"})
                    bot_id_ticket_category = str(bot_id_ticket_category)
                    bot_pk, ticket_category = bot_id_ticket_category.split("-")

                    bot_pk = int(float(str(bot_pk).strip()))
                    ticket_category = str(ticket_category).strip()

                    if bot_pk in extracted_bot_pk_list:
                        ticket_category_obj = parent_agent_obj.ticket_categories.filter(
                            ticket_category__iexact=ticket_category, bot__pk=int(bot_pk)).first()
                        pk_list.append(ticket_category_obj.pk)
                    else:
                        wrong_data_list.append(
                            {
                                "row_num": row_num,
                                "data_array": row_values,
                                "detail": "Bot does not belongs to this agent : {} of {}".format(bot_pk, ticket_category)
                            })

            logger.info("pk_list = " + str(pk_list),
                        extra={"AppName": "EasyTMS"})
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error process_ticket_category_input %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Invalid Ticket Category"})

        return pk_list

    def process_supervisor(target_supervisor_username, agent_type):
        if len(target_supervisor_username) == 0:
            wrong_data_list.append(
                {"row_num": row_num, "data_array": row_values, "detail": "Supervisor Name Empty"})

        if agent_type == "agent":
            target_supervisor_obj = active_agent.agents.all().filter(
                role="supervisor", is_account_active=True, user__username__iexact=target_supervisor_username).first()

            if target_supervisor_obj is None:
                if active_agent.user.username.lower() == target_supervisor_username.lower():
                    target_supervisor_obj = active_agent
                else:
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": "Wrong Supervisor Name"})
        else:
            target_supervisor_obj = active_agent

        return target_supervisor_obj

    method_response = {"status": 500, "message": "Exception Occurred"}
    try:
        logger.info("Inside add_users_from_excel_document",
                    extra={'AppName': 'EasyTMS'})

        wrong_data_list = []

        try:
            xlrd.xlsx.ensure_elementtree_imported(False, None)
            xlrd.xlsx.Element_has_iter = True
        except Exception:
            pass

        data_workbook = open_workbook(file_path)
        data_sheet = data_workbook.sheet_by_name(
            data_workbook.sheet_names()[0])

        reg_name = '^[^\s][a-zA-Z ]+$'
        reg_email = '^[a-zA-Z0-9._-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,4}$'
        reg_mob = '^[6-9]{1}[0-9]{9}$'

        total_rows = data_sheet.nrows
        for row_num in range(1, total_rows):
            error_list_size_at_start = len(wrong_data_list)

            logger.info("row: " + str(row_num), extra={'AppName': 'EasyTMS'})

            row_values = data_sheet.row_values(row_num)

            logger.info("row_values: " + str(row_values),
                        extra={'AppName': 'EasyTMS'})

            user_name = str(row_values[0]).strip()
            user_name = process_user_first_last_name(user_name)

            user_email = str(row_values[1]).strip()
            user_email = process_user_email(user_email)

            user_mobile = str(row_values[2]).strip()
            user_mobile = process_user_mobile(user_mobile)

            user_type = str(row_values[3]).strip()
            user_type = process_user_type(user_type)

            supervisor_username = str(row_values[6]).strip()
            supervisor_obj = process_supervisor(supervisor_username, user_type)

            bot_ids_input = str(row_values[4]).strip()
            bot_pk_list = process_agent_bot_ids_input(
                bot_ids_input, supervisor_obj)

            ticket_category_input = str(row_values[5]).strip()
            ticket_category_pk_list = process_ticket_category_input(
                ticket_category_input, supervisor_obj, bot_pk_list)

            logger.info("wrong_data_list = " + str(wrong_data_list),
                        extra={"AppName": "EasyTMS"})

            error_list_size_at_end = len(wrong_data_list)
            if error_list_size_at_end > error_list_size_at_start:
                continue

            logger.info("user_email = " + str(user_email),
                        extra={"AppName": "EasyTMS"})

            password = generate_password("TMS")

            try:
                user = User.objects.get(username=user_email)
                user.first_name = user_name
                user.email = user_email
                user.save()
            except Exception:
                user = User.objects.create(first_name=user_name,
                                           email=user_email,
                                           username=user_email,
                                           status="2",
                                           role="bot_builder")
                user.set_password(password)
                user.save()

            platform_url = settings.EASYCHAT_HOST_URL

            thread = threading.Thread(target=send_password_over_email, args=(
                user_email, user_name, password, platform_url,), daemon=True)
            thread.start()

            user.set_password(password)
            user.save()

            tms_agent = Agent.objects.create(user=user,
                                             phone_number=user_mobile,
                                             role=user_type, )

            supervisor_obj.agents.add(tms_agent)
            supervisor_obj.save()
            add_bot_to_agent(tms_agent, bot_pk_list, Bot)
            add_ticket_category_to_agent_by_pk_list(
                tms_agent, ticket_category_pk_list, TicketCategory)
            logger.info("Added user user_name: " + str(user_name),
                        extra={'AppName': 'EasyTMS'})

        if total_rows <= 1:
            method_response["status"] = 301
            method_response["message"] = "File is empty."
        elif wrong_data_list:
            method_response["status"] = 302
            method_response["message"] = "Some Data in wrong format found"
            method_response["file_path"] = create_excel_wrong_user_data(
                wrong_data_list)
            logger.info("Some wrong data found %s", str(
                wrong_data_list), extra={'AppName': 'EasyTMS'})
        else:
            method_response["status"] = 200
            method_response["message"] = "users created"
            logger.info("Successfully exited from add_users_from_excel_document", extra={
                        'AppName': 'EasyTMS'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_users_from_excel_document %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})
    return method_response


def create_excel_user_details(active_agent, Agent):
    try:
        logger.info("In create_excel_user_details",
                    extra={'AppName': 'EasyAssist'})

        agents = get_mapped_agents(
            active_agent,
            Agent,
            is_active=None,
            is_account_active=True,
            include_supervisor=False,
            include_self=False,
            is_absent=None
        )

        agents = agents.order_by('-agent_creation_datetime')
        current_datetime = datetime.now()

        support_history_wb = Workbook()
        sheet = support_history_wb.add_sheet("Agent Details")

        sheet.write(0, 0, "DateTime")
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, "Name")
        sheet.col(1).width = 256 * 20
        sheet.write(0, 2, "Email")
        sheet.col(2).width = 256 * 20
        sheet.write(0, 3, "Mobile Number")
        sheet.col(3).width = 256 * 20
        sheet.write(0, 4, "Status")
        sheet.col(4).width = 256 * 20
        sheet.write(0, 5, "Creation DateTime")
        sheet.col(0).width = 256 * 20

        index = 1
        for agent in agents:
            sheet.write(index, 0, str(current_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(YEAR_TIME_AM_PM_FORMAT)))
            sheet.write(index, 1, agent.user.first_name)
            sheet.write(index, 2, agent.user.username)
            sheet.write(index, 3, agent.phone_number)
            if agent.is_absent:
                sheet.write(index, 4, "Absent")
            else:
                sheet.write(index, 4, "Present")
            sheet.write(index, 5, str(agent.agent_creation_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(YEAR_TIME_AM_PM_FORMAT)))

            index += 1

        filename = "secured_files/EasyTMSApp/AgentDetails/agents_" + \
            str(active_agent.user.username) + ".xls"

        support_history_wb.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_user_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_supervisor_details(active_agent):
    try:
        logger.info("In create_excel_supervisor_details",
                    extra={'AppName': 'EasyAssist'})

        supervisors = active_agent.agents.filter(
            is_account_active=True, role="supervisor")
        supervisors = supervisors.order_by('-agent_creation_datetime')

        current_datetime = datetime.now()

        support_history_wb = Workbook()
        sheet = support_history_wb.add_sheet("Supervisor Details")

        sheet.write(0, 0, "DateTime")
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, "Name")
        sheet.col(1).width = 256 * 20
        sheet.write(0, 2, "Email")
        sheet.col(2).width = 256 * 20
        sheet.write(0, 3, "Mobile Number")
        sheet.col(3).width = 256 * 20
        sheet.write(0, 4, "Status")
        sheet.col(4).width = 256 * 20
        sheet.write(0, 5, "Creation DateTime")
        sheet.col(0).width = 256 * 20

        index = 1
        for supervisor in supervisors:
            sheet.write(index, 0, str(current_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(YEAR_TIME_AM_PM_FORMAT)))
            sheet.write(index, 1, supervisor.user.first_name)
            sheet.write(index, 2, supervisor.user.username)
            sheet.write(index, 3, supervisor.phone_number)
            if supervisor.is_absent:
                sheet.write(index, 4, "Absent")
            else:
                sheet.write(index, 4, "Present")
            sheet.write(index, 5, str(supervisor.agent_creation_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(YEAR_TIME_AM_PM_FORMAT)))
            index += 1

        filename = "secured_files/EasyTMSApp/AgentDetails/supervisors_" + \
            str(active_agent.user.username) + ".xls"
        support_history_wb.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_user_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def send_agent_query_over_whatsapp(access_token, agent_query_info):
    try:
        function_dictionary = {}
        processor_code = str(
            access_token.get_whatsapp_api_processor_obj().function)
        exec(processor_code, function_dictionary)
        function_dictionary['f']({
            "agent_comment": agent_query_info["query_info"]["message"],
            "ticket_id": agent_query_info["ticket_info"]["ticket_id"],
            "customer_mobile_number": agent_query_info["customer_info"]["customer_mobile_number"],
        })
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error send_agent_query_over_whatsapp message send %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def send_agent_query_over_email(access_token, agent_query_info):
    try:
        send_agent_query_message({
            "agent_comment": agent_query_info["query_info"]["message"],
            "ticket_id": agent_query_info["ticket_info"]["ticket_id"],
            "customer_email_id": agent_query_info["customer_info"]["customer_email_id"],
        })
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error send_agent_query_over_email message send %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})


def process_agent_comment(comment_html):
    try:
        comment_html = comment_html.replace("&lt;", "")
        comment_html = comment_html.replace("&gt;", "")
        return comment_html
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error process_agent_comment message send %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyTMS'})

    return ""
