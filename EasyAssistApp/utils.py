from django.conf import settings
from django.utils import timezone
from django.db.models import Q, Count, Sum
from django.utils.encoding import smart_str
from django.shortcuts import HttpResponse
from django import forms
from EasyAssistApp.encrypt import CustomEncrypt, generate_random_key
from EasyAssistApp.html_parser import strip_html_tags
from EasyAssistApp.constants import *
from EasyAssistApp.send_email import send_password_over_email, send_drop_link_over_email, send_reported_bug_over_email, send_cobrowse_proxy_drop_link_over_email
from EasyAssistApp.utils_client_server_signal import send_data_from_server_to_client
from AuditTrailApp.utils import add_audit_trail
from EasyAssistApp.jaas_jwt import JaaSJwtBuilder
from EasyAssistApp.utils_validation import *
from django.core.paginator import Page, Paginator, EmptyPage, PageNotAnInteger

from DeveloperConsoleApp.utils import *
from DeveloperConsoleApp.utils_aws_s3 import *

import hashlib
import requests
import sys
import json
import operator
import logging
import re
import ast
import random
import array
import datetime
import pytz
import threading
import os
import uuid
import mimetypes
from zipfile import ZipFile
import geocoder
import shutil
import magic
import base64

from urllib.parse import urljoin, urlparse, urlencode
from random import randint
from xlwt import Workbook, XFStyle, Font, Formula, easyxf
import xlrd
from xlrd import open_workbook
from os import path
from os.path import basename
from docx import Document

logger = logging.getLogger(__name__)


def is_valid_domain(request, origin, cobrowse_access_token_obj):
    is_valid_domain = cobrowse_access_token_obj.is_valid_domain(origin)
    is_proxy_cobrowsing = request.COOKIES.get("is_proxy_cobrowsing")

    if is_proxy_cobrowsing:
        is_valid_domain = True

    return is_valid_domain


class UrlShortenTinyurl:
    URL = "https://api.tinyurl.com/create"

    def shorten(self, url_long):
        try:
            console_config = get_developer_console_settings()
            api_token = "yihhXDx75Mg9ZEd9lhnXszwuPQZlcPw1dEsWza7mhtrpVj7x8ksSUTur42IO"
            if console_config:
                api_token = console_config.tiny_url_api_token

            data = {
                "url": url_long,
                "domain": "tinyurl.com",
                "alias": "",
                "tags": "",
            }
            url = self.URL + "?api_token=" + api_token
            res = requests.post(url, data=data)

            json_response = res.json()
            try:
                tiny_url = json_response["data"]["tiny_url"]
                return tiny_url
            except Exception as e:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                logger.error("Error shorten %s at %s",
                             str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
                return "Error"

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error shorten %s at %s",
                         str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return ""


def convert_seconds_to_hours_minutes(seconds):
    try:
        hour = seconds // 3600
        seconds %= 3600
        minutes = seconds // 60

        duration = ""
        if hour > 0:
            duration = str(hour) + (" hours " if hour > 1 else " hour ")
        if duration == "" or minutes > 0:
            duration += str(minutes) + \
                (" minutes" if minutes > 1 else " minute")
        return duration.strip()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In convert_seconds_to_hours_minutes: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyAssist'})
        return "0 minute"


def get_formatted_time_from_seconds(seconds):
    try:
        hour = seconds // 3600
        seconds %= 3600
        minutes = seconds // 60
        seconds %= 60

        duration = ""
        if hour > 0:
            duration = str(hour) + (" hours " if hour > 1 else " hour ")
        if minutes > 0:
            duration += str(minutes) + \
                (" minutes " if minutes > 1 else " minute ")
        if duration == "" or hour == 0:
            duration += str(seconds) + \
                (" seconds" if seconds > 1 else " second")
        return duration.strip()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In convert_seconds_to_hours_minutes: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyAssist'})
        return "0 minute"


def delete_user_session(user_session, Session):
    try:
        logger.info("In delete_user_session", extra={'AppName': 'EasyAssist',
                                                     'user_id': user_session.user.username})
        session_objs = Session.objects.filter(pk=user_session.session_key)
        user_session.delete()
        if session_objs:
            session_objs.delete()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In delete_user_session: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
            'AppName': 'EasyAssist', 'user_id': str(user_session.user.username)})


def is_online(username, UserSession):
    try:
        logger.info("In is_online", extra={
                    'AppName': 'EasyAssist', 'user_id': str(username)})
        user_session_objs = UserSession.objects.filter(user__username=username)
        logger.info("user sessions count: " + str(user_session_objs.count()), extra={
            'AppName': 'EasyAssist', 'user_id': str(username)})
        if user_session_objs:
            return True
        else:
            return False
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In is_online: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
    return False


def logout_all(username, UserSession, Session):
    try:
        logger.info("In logout_all", extra={
                    'AppName': 'EasyAssist', 'user_id': str(username)})
        user_session_objs = UserSession.objects.filter(user__username=username)
        for user_session_obj in user_session_objs:
            delete_user_session(user_session_obj, Session)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In delete_user_session: %s at %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist', 'user_id': str(username)})


def check_and_update_lead(primary_value_list, meta_data, cobrowse_io, CobrowseCapturedLeadData):
    try:
        cobrowse_io.captured_lead.clear()
        cobrowse_access_token = cobrowse_io.access_token
        for primary_id_obj in primary_value_list:
            primary_id = primary_id_obj['value']
            primary_id = primary_id.strip().lower()
            primary_value = hashlib.md5(primary_id.encode()).hexdigest()

            lead_obj = CobrowseCapturedLeadData.objects.filter(
                primary_value=primary_value, session_id=cobrowse_io.session_id).first()
            if lead_obj is None:
                logger.info("lead for " + str(cobrowse_io.session_id) +
                            " not found. Creating a new lead object.", extra={'AppName': 'EasyAssist'})

                primary_label = primary_id_obj['label']
                search_field = cobrowse_access_token.search_fields.filter(
                    tag_label=primary_label).first()
                lead_obj = CobrowseCapturedLeadData.objects.create(
                    primary_value=primary_value, session_id=cobrowse_io.session_id)

                if search_field:
                    lead_obj.search_field = search_field
                    lead_obj.save()
            else:
                logger.info("lead for " + str(cobrowse_io.session_id) +
                            " found.", extra={'AppName': 'EasyAssist'})

            cobrowse_io.captured_lead.add(lead_obj)
        cobrowse_io.meta_data = meta_data
        cobrowse_io.last_update_datetime = timezone.now()
        cobrowse_io.is_active = True
        cobrowse_io.save()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error check_and_update_lead %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def add_supported_language(cobrowse_user, supported_language_list, LanguageSupport):
    try:
        logger.info("Inside add support language",
                    extra={'AppName': 'EasyAssist'})
        for language_pk in supported_language_list:
            language_obj = LanguageSupport.objects.get(pk=int(language_pk))
            cobrowse_user.supported_language.add(language_obj)
        cobrowse_user.save()
        logger.info("Secussefully exited add support language",
                    extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_supported_language %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def add_supported_product(cobrowse_user, supported_product_list, ProductCategory):
    try:
        logger.info("Inside add supported product",
                    extra={'AppName': 'EasyAssist'})
        for product_pk in supported_product_list:
            product_obj = ProductCategory.objects.get(pk=int(product_pk))
            cobrowse_user.product_category.add(product_obj)
        cobrowse_user.save()
        logger.info("Secussefully exited add supported product",
                    extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_supported_product %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def add_selected_supervisor(selected_supervisor_pk_list, active_agent, cobrowse_agent, CobrowseAgent):
    logger.info("inside add supervisor ", extra={'AppName': 'EasyAssist'})
    
    if cobrowse_agent.role == "admin_ally":
        for current_supervisor_pk in selected_supervisor_pk_list:
            current_supervisor_pk = int(current_supervisor_pk)
            if current_supervisor_pk == -1:
                active_agent.agents.add(cobrowse_agent)
            else:
                active_agent.agents.add(cobrowse_agent)
                current_supervisor = CobrowseAgent.objects.get(
                    pk=current_supervisor_pk)
                cobrowse_agent.agents.add(current_supervisor)
            cobrowse_agent.save()
    else:
        for current_supervisor_pk in selected_supervisor_pk_list:
            current_supervisor_pk = int(current_supervisor_pk)
            if current_supervisor_pk == -1:
                current_supervisor = active_agent
            else:
                current_supervisor = CobrowseAgent.objects.get(
                    pk=current_supervisor_pk)
            current_supervisor.agents.add(cobrowse_agent)
            current_supervisor.save()
    logger.info("successfully exited add supervisor ",
                extra={'AppName': 'EasyAssist'})


def reset_agents_language(cobrowse_user):
    try:
        agents = get_list_agents_under_admin(cobrowse_user, None, None)
        agents += get_list_supervisor_under_admin(cobrowse_user, None)

        for agent in agents:
            for language in agent.supported_language.all():
                if language not in cobrowse_user.supported_language.filter(is_deleted=False):
                    agent.supported_language.remove(language)
                    agent.save()
    except Exception:
        return


def save_language_support(cobrowse_user, supported_language, LanguageSupport):
    try:
        for language_obj in cobrowse_user.supported_language.all():
            language_obj.is_deleted = True
            language_obj.save()

        supported_language_list = supported_language.split(",")
        for index, language in enumerate(supported_language_list):
            language = str(language).strip()
            language = remo_html_from_string(strip_html_tags(language))
            if len(language) < 2:
                continue
            language = language[0].upper() + language[1:].lower()

            language_obj = cobrowse_user.supported_language.filter(
                title=language).first()

            if language_obj == None:
                language_obj = LanguageSupport.objects.create(title=language)

            language_obj.index = index
            language_obj.is_deleted = False
            language_obj.save()

            cobrowse_user.supported_language.add(language_obj)

        cobrowse_user.save()

        reset_agents_language(cobrowse_user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_language_support %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def save_product_category(cobrowse_user, product_category, ProductCategory):
    try:
        for category_obj in cobrowse_user.product_category.all():
            category_obj.is_deleted = True
            category_obj.save()

        product_category_list = product_category.split(",")
        for index, category in enumerate(product_category_list):
            category = str(category).strip()
            category = remo_html_from_string(strip_html_tags(category))
            if len(category) < 2:
                continue
            category = category[0].upper() + category[1:].lower()
            category_obj = cobrowse_user.product_category.filter(
                title=category).first()

            if category_obj == None:
                category_obj = ProductCategory.objects.create(title=category)

            category_obj.index = index
            category_obj.is_deleted = False
            category_obj.save()

            cobrowse_user.product_category.add(category_obj)

        cobrowse_user.save()

        reset_agents_product_category(cobrowse_user)
        logger.info("Successfully exited from save_product_category",
                    extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_product_category %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def reset_agents_product_category(cobrowse_user):
    try:
        logger.info("Inside reset_agents_product_category",
                    extra={'AppName': 'EasyAssist'})
        agents = get_list_agents_under_admin(cobrowse_user, None, None)
        agents += get_list_supervisor_under_admin(cobrowse_user, None)
        for agent in agents:
            for category in agent.product_category.all():
                if category not in cobrowse_user.product_category.filter(is_deleted=False):
                    agent.product_category.remove(category)
                    agent.save()
        logger.info("Successfully exited from reset_agents_product_category", extra={
                    'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error reset_agents_product_category %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return


def add_product_category_to_user(cobrowse_user, selected_product_category_pk_list, ProductCategory):
    try:
        logger.info("Inside add_product_category_to_user",
                    extra={'AppName': 'EasyAssist'})
        for product_category_pk in selected_product_category_pk_list:
            product_category_obj = ProductCategory.objects.get(
                pk=int(product_category_pk))
            if product_category_obj not in cobrowse_user.product_category.all():
                cobrowse_user.product_category.add(product_category_obj)
        cobrowse_user.save()
        logger.info("Successfully exited from add_product_category_to_user", extra={
                    'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_product_category_to_user %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_agents_for_product_and_language(agents_list=[], product_category_list=[], language_list=[], is_language_enabled=True, is_product_category_enabled=True):
    filtered_agents_list = []
    product_category_set = None
    language_set = None
    try:
        if len(product_category_list) != 0:
            product_category_set = set(product_category_list)
        if len(language_list) != 0:
            language_set = set(language_list)

        if(len(agents_list) != 0):
            for agent in agents_list:
                if is_product_category_enabled:
                    agent_products_set = set(agent.product_category.all())
                    if product_category_set != None and product_category_set.issubset(agent_products_set):
                        if is_language_enabled:
                            agent_languages_set = set(
                                agent.supported_language.all())
                            if language_set != None and language_set.issubset(agent_languages_set):
                                filtered_agents_list.append(agent)
                        else:
                            filtered_agents_list.append(agent)

                else:
                    if is_language_enabled:
                        agent_languages_set = set(
                            agent.supported_language.all())
                        if language_set != None and language_set.issubset(agent_languages_set):
                            filtered_agents_list.append(agent)

        return filtered_agents_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_agents_for_product_and_language %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return filtered_agents_list


def get_active_agent_obj(request, CobrowseAgent):
    active_agent = None
    try:
        active_agent = CobrowseAgent.objects.filter(user=request.user, is_account_active=True).first()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.warning("Error get_active_agent_obj %s at %s",
                       str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return active_agent


def get_list_of_unresolved_screencast(ScreenCast, agent_obj):
    return ScreenCast.objects.filter(status="unresolved")


def get_list_of_resolved_screencast(ScreenCast, agent_obj):
    return ScreenCast.objects.filter(status="resolved")


def perform_masking(screen_cast_obj, Field):
    try:
        input_html_elements = json.loads(screen_cast_obj.input_html_elements)
        for html_element in input_html_elements:
            for md5_html_element_str in input_html_elements[html_element]:
                url = screen_cast_obj.website_url
                try:
                    field_obj = Field.objects.get(
                        form__url=url, field_id=md5_html_element_str)
                    name = field_obj.name
                    if field_obj.is_masked:
                        current_value = input_html_elements[
                            html_element][md5_html_element_str]["value"]
                        input_html_elements[html_element][md5_html_element_str]["value"] = "*" * len(
                            current_value)
                    input_html_elements[html_element][
                        md5_html_element_str]["name"] = name
                    input_html_elements[html_element][
                        md5_html_element_str]["is_valid_field"] = True
                    input_html_elements[html_element][md5_html_element_str][
                        "group_color"] = field_obj.group_color
                    input_html_elements[html_element][md5_html_element_str][
                        "field_order"] = field_obj.field_order
                except Exception:
                    input_html_elements[html_element][
                        md5_html_element_str]["is_valid_field"] = True
                    input_html_elements[html_element][
                        md5_html_element_str]["group_color"] = ""
                    input_html_elements[html_element][
                        md5_html_element_str]["field_order"] = 0

        return input_html_elements
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error perform_masking %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return None


def save_remote_file(url):
    try:
        response = requests.get(url, timeout=5)
        filename = url.split("/")[-1]
        with open(settings.MEDIA_ROOT + "static/" + filename, "wb") as temp_file:
            temp_file.write(response.content)
            temp_file.close()
        return "https://easyassist.allincall.in/files/static/" + filename
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_remote_file %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return url


def get_file_type(url):
    try:
        return str(url).split("/")[-1].split(".")[-1]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_file_type %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return None


def convert_relative_path_to_absolute(text, address):
    try:
        import re
        req_css_links = re.findall(r'url\((.*?)\)', text)
        content = text
        for req_css_link in req_css_links:
            css_link = req_css_link.replace("\"", "")
            css_link = css_link.replace("'", "")
            updated_css_link = urljoin(address, str(css_link))
            file_type = get_file_type(updated_css_link)
            if file_type in ["woff2", "woff", "tff", "png", "jpg", "jpeg", "jpe"]:
                updated_css_link = save_remote_file(updated_css_link)
            elif updated_css_link != address:
                updated_css_link = "https://easyassist.allincal.in/easy-assist/files/" + updated_css_link

            content = content.replace(
                "url(" + req_css_link + ")", "url(" + updated_css_link + ")")
        return content
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error convert_relative_path_to_absolute %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def convert_absolute_path_in_html(html, address):
    try:
        import re
        req_css_links = re.findall(r'url\((.*?)\)', html)
        content = html
        for req_css_link in req_css_links:
            css_link = req_css_link.replace("\"", "")
            css_link = css_link.replace("'", "")
            updated_css_link = urljoin(address, str(css_link))
            content = content.replace(
                "url(" + req_css_link + ")", "url(" + updated_css_link + ")")

        return content
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error convert_absolute_path_in_html %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def get_least_loaded_agent(User, CobrowseAgent, CobrowseIO):
    agent_obj = None
    try:
        agent_dict = {}
        for agent in CobrowseAgent.objects.filter(role="agent"):
            agent_dict[agent.user.username] = 0

        cobrowseio_obj_list = CobrowseIO.objects.filter(~Q(agent=None)).values(
            'agent__user__username').annotate(total=Count('agent')).order_by("total")

        for cobrowseio_obj in cobrowseio_obj_list:
            agent_dict[cobrowseio_obj["agent__user__username"]
                       ] = cobrowseio_obj["total"]

        agent_username = min(agent_dict.items(), key=operator.itemgetter(1))[0]
        user = User.objects.get(username=agent_username)
        agent_obj = CobrowseAgent.objects.get(user=user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_least_loaded_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    logger.info("get_least_loaded_agent: %s", agent_obj,
                extra={'AppName': 'EasyAssist'})
    return agent_obj


def get_cobrowse_access_token_obj(request, CobrowseAccessToken):
    cobrowse_access_token = None
    try:
        access_token = request.META.get('HTTP_X_ACCESSTOKEN')
        cobrowse_access_token = CobrowseAccessToken.objects.get(
            key=access_token)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_cobrowse_access_token_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return cobrowse_access_token


def get_request_origin(request):
    try:
        origin = request.META.get('HTTP_ORIGIN')
        origin = urlparse(origin)
        return origin.netloc
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_request_origin %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def get_list_agents_under_admin(admin_user, is_active=True, is_account_active=True):
    try:
        cobrowse_agents = []
        agents = []
        supervisors = []
        if is_active == None:
            if is_account_active == None:
                agents = list(admin_user.agents.all().filter(
                    role="agent"))
            else:
                agents = list(admin_user.agents.all().filter(
                    role="agent", is_account_active=is_account_active))
        else:
            if is_account_active == None:
                agents = list(admin_user.agents.all().filter(
                    role="agent", is_active=is_active))
            else:
                agents = list(admin_user.agents.all().filter(
                    role="agent", is_active=is_active, is_account_active=is_account_active))

        admin_ally_supervisors = []
        if admin_user.role == "admin":
            admin_allies = []
            if is_active == None:
                if is_account_active == None:
                    admin_allies = list(admin_user.agents.all().filter(
                        role="admin_ally"))
                else:
                    admin_allies = list(admin_user.agents.all().filter(
                        role="admin_ally", is_account_active=is_account_active))
            else:
                if is_account_active == None:
                    admin_allies = list(admin_user.agents.all().filter(
                        role="admin_ally", is_active=is_active))
                else:
                    admin_allies = list(admin_user.agents.all().filter(
                        role="admin_ally", is_active=is_active, is_account_active=is_account_active))

            for admin_ally in admin_allies:
                admin_ally_supervisors += get_list_supervisor_under_admin(
                    admin_ally, None)
            supervisors += list(set(admin_ally_supervisors))
            if is_account_active == None:
                supervisors += admin_user.agents.all().filter(
                    is_switch_allowed=True)
            else:
                supervisors += admin_user.agents.all().filter(
                    is_switch_allowed=True, is_account_active=is_account_active)

        cobrowse_agents += agents
        if is_account_active == None:
            supervisors += admin_user.agents.all().filter(
                role="supervisor")
        else:
            supervisors += admin_user.agents.all().filter(
                role="supervisor", is_account_active=is_account_active)

        if is_account_active == None:
            supervisors += admin_user.agents.all().filter(
                is_switch_allowed=True)
        else:
            supervisors += admin_user.agents.all().filter(
                is_switch_allowed=True, is_account_active=is_account_active)
        supervisors = list(set(supervisors))
        for supervisor in supervisors:
            if is_active == None:
                if is_account_active == None:
                    cobrowse_agents += list(supervisor.agents.all().filter(
                        role="agent"))
                else:
                    cobrowse_agents += list(supervisor.agents.all().filter(
                        role="agent", is_account_active=is_account_active))
            else:
                if is_account_active == None:
                    cobrowse_agents += list(supervisor.agents.all().filter(
                        role="agent", is_active=is_active))
                else:
                    cobrowse_agents += list(supervisor.agents.all().filter(
                        role="agent", is_active=is_active, is_account_active=is_account_active))

            if supervisor.is_switch_allowed:
                if is_active == None or (is_active == False and supervisor.is_active == False) or (is_active == True and supervisor.is_active == True):
                    cobrowse_agents += [supervisor]
                # elif is_active == False and supervisor.is_active == False:
                #     cobrowse_agents += [supervisor]
                # elif is_active == True and supervisor.is_active == True:
                #     cobrowse_agents += [supervisor]

        if admin_user.is_switch_allowed:
            if is_active == None or (is_active == False and admin_user.is_active == False) or (is_active == True and admin_user.is_active == True):
                cobrowse_agents += [admin_user]
            # elif is_active == False and admin_user.is_active == False:
            #     cobrowse_agents += [admin_user]
            # elif is_active == True and admin_user.is_active == True:
            #     cobrowse_agents += [admin_user]

        return list(set(cobrowse_agents))
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_list_agents_under_admin: %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return []


def get_list_supervisor_under_admin(admin_user, is_active=True, is_account_active=True):
    try:
        agents = []
        if is_account_active:
            if is_active == None:
                agents = admin_user.agents.all().filter(
                    role="supervisor", is_account_active=True)
            else:
                agents = admin_user.agents.all().filter(
                    role="supervisor", is_active=is_active, is_account_active=True)

            if is_active == None:
                agents |= admin_user.agents.all().filter(
                    role="supervisor", is_account_active=True)
            else:
                agents |= admin_user.agents.all().filter(
                    role="supervisor", is_active=is_active, is_account_active=True)

        else:
            if is_active == None:
                agents = admin_user.agents.all().filter(
                    role="supervisor")
            else:
                agents = admin_user.agents.all().filter(
                    role="supervisor", is_active=is_active)

            if is_active == None:
                agents |= admin_user.agents.all().filter(
                    role="supervisor")
            else:
                agents |= admin_user.agents.all().filter(
                    role="supervisor", is_active=is_active)

        return list(set(agents))
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_list_supervisor_under_admin: %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return []


def get_list_admin_ally(admin_user, is_active=True, is_account_active=True):
    try:
        agents = []
        if is_account_active:
            if is_active == None:
                agents = admin_user.agents.all().filter(
                    role="admin_ally", is_account_active=True)
            else:
                agents = admin_user.agents.all().filter(
                    role="admin_ally", is_active=is_active, is_account_active=True)

            if is_active == None:
                agents |= admin_user.agents.all().filter(
                    role="admin_ally", is_account_active=True)
            else:
                agents |= admin_user.agents.all().filter(
                    role="admin_ally", is_active=is_active, is_account_active=True)

        else:
            if is_active == None:
                agents = admin_user.agents.all().filter(
                    role="admin_ally")
            else:
                agents = admin_user.agents.all().filter(
                    role="admin_ally", is_active=is_active)

            if is_active == None:
                agents |= admin_user.agents.all().filter(
                    role="admin_ally")
            else:
                agents |= admin_user.agents.all().filter(
                    role="admin_ally", is_active=is_active)

        return list(set(agents))
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_list_admin_ally: %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return []

# This function is Deprecated


def is_agent_under_me(agent_obj, agent_name):

    def is_direct_mapped(target_agent_obj):
        agent_objs = target_agent_obj.agents.filter(
            user__username__iexact=agent_name)

        if agent_objs.count() > 0:
            return True
        else:
            return False

    try:
        is_under_me = False
        if agent_obj.role == "admin":
            supervisors = agent_obj.agents.filter(role="supervisor")

            for supervisor in supervisors:
                if is_direct_mapped(supervisor):
                    is_under_me = True

                if supervisor.user.username.lower() == agent_name.lower():
                    is_under_me = True

                if is_under_me:
                    break

        if is_direct_mapped(agent_obj):
            is_under_me = True

        if agent_obj.user.username.lower() == agent_name.lower():
            is_under_me = True

        return is_under_me
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("is_agent_under_me: %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return False


def is_agent_under_admin(agent_obj, agent_name, CobrowseAgent):
    try:
        cobrowse_agent_obj = CobrowseAgent.objects.filter(
            user__username__iexact=agent_name).first()
        return cobrowse_agent_obj.get_access_token_obj() == agent_obj.get_access_token_obj()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("is_agent_under_admin: %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return False


def get_admin_from_active_agent(active_agent, CobrowseAgent):

    try:
        if active_agent.is_switch_allowed and active_agent.role == "admin":
            return active_agent

        parent_user = CobrowseAgent.objects.filter(
            agents__pk=active_agent.pk)[0]
        try:
            second_parent = CobrowseAgent.objects.filter(
                agents__pk=parent_user.pk)[0]
            return second_parent
        except Exception:
            return parent_user

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_admin_from_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def get_supervisor_from_active_agent(active_agent, CobrowseAgent):
    try:
        return list(CobrowseAgent.objects.filter(Q(role="supervisor") | Q(is_switch_allowed=True)).filter(is_account_active=True, agents__in=[active_agent]))
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_supervisor_from_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return []


def filter_free_active_agent(agent_list, cobrowse_io_obj, for_meeting=False, support_levels=None):
    free_agent_list = []
    try:
        for agent in agent_list:
            if support_levels and (agent.support_level not in support_levels):
                continue

            if (for_meeting == False and agent.is_cobrowsing_active == False) or (for_meeting == True and agent.is_cognomeet_active == False):
                if cobrowse_io_obj.access_token.allow_language_support:
                    if cobrowse_io_obj.supported_language in agent.supported_language.filter(is_deleted=False):
                        if cobrowse_io_obj.access_token.choose_product_category:
                            if cobrowse_io_obj.product_category in agent.product_category.filter(is_deleted=False):
                                free_agent_list.append(agent)
                        else:
                            free_agent_list.append(agent)
                else:
                    if cobrowse_io_obj.access_token.choose_product_category:
                        if cobrowse_io_obj.product_category in agent.product_category.filter(is_deleted=False):
                            free_agent_list.append(agent)
                    else:
                        free_agent_list.append(agent)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error filter_free_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        free_agent_list = agent_list

    return free_agent_list


def generate_md5(str):
    return hashlib.md5(str.encode()).hexdigest()


def save_audit_trail(cobrowse_agent, action, description, CobrowsingAuditTrail):
    try:
        CobrowsingAuditTrail.objects.create(agent=cobrowse_agent,
                                            action=action,
                                            action_description=description)
        logger.info("Audit Trail Saved Successfully.",
                    extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_audit_trail %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def generate_random_password():
    # maximum length of password needed
    # this can be changed to suit your password length
    MAX_LEN = 12

    # declare arrays of the character that we need in out password
    # Represented as chars to enable easy string concatenation
    DIGITS = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
    LOCASE_CHARACTERS = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h',
                         'i', 'j', 'k', 'm', 'n', 'o', 'p', 'q',
                         'r', 's', 't', 'u', 'v', 'w', 'x', 'y',
                         'z']

    UPCASE_CHARACTERS = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H',
                         'I', 'J', 'K', 'M', 'N', 'O', 'p', 'Q',
                         'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y',
                         'Z']

    SYMBOLS = ['@', '#', '$', '%', '=', ':', '?', '.', '/', '|', '~', '>',
               '*', '(', ')']

    # combines all the character arrays above to form one array
    COMBINED_LIST = DIGITS + UPCASE_CHARACTERS + LOCASE_CHARACTERS + SYMBOLS

    # randomly select at least one character from each character set above
    rand_digit = random.choice(DIGITS)
    rand_upper = random.choice(UPCASE_CHARACTERS)
    rand_lower = random.choice(LOCASE_CHARACTERS)
    rand_symbol = random.choice(SYMBOLS)

    # combine the character randomly selected above
    # at this stage, the password contains only 4 characters but
    # we want a 12-character password
    temp_pass = rand_digit + rand_upper + rand_lower + rand_symbol

    # now that we are sure we have at least one character from each
    # set of characters, we fill the rest of
    # the password length by selecting randomly from the combined
    # list of character above.
    for value in range(MAX_LEN - 4):
        temp_pass = temp_pass + random.choice(COMBINED_LIST)
        # convert temporary password into array and shuffle to
        # prevent it from having a consistent pattern
        # where the beginning of the password is predictable

        temp_pass_list = [char for char in temp_pass]
        random.shuffle(temp_pass_list)

    # traverse the temporary password array and append the chars
    # to form the password
    password = ""
    for value in temp_pass_list:
        password = password + value

    # print out password
    return password


def generate_password(password_prefix):
    DIGITS = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
    SYMBOLS = ['@', '#', '$', '&']

    password_digits = ""
    for value in range(0, 4):
        password_digits = password_digits + random.choice(DIGITS)
    password = password_prefix + \
        random.choice(SYMBOLS) + password_digits + random.choice(SYMBOLS)
    return password


def delete_expired_cobrowse_middleware_token(CobrowsingMiddlewareAccessToken):
    try:
        current_date = (datetime.datetime.now() - datetime.timedelta(1)).date()
        CobrowsingMiddlewareAccessToken.objects.filter(
            timestamp__date__lte=current_date).delete()
        CobrowsingMiddlewareAccessToken.objects.filter(
            is_expired=True).delete()
        logger.info(
            "Expired Cobrowsing Middleware Access Token deleted successfully.", extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error delete_expired_cobrowse_middleware_token %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def create_cobrowse_middleware_token(CobrowsingMiddlewareAccessToken):
    return CobrowsingMiddlewareAccessToken.objects.create()


def is_valid_cobrowse_middleware_token(token, CobrowsingMiddlewareAccessToken):
    is_valid = False
    try:
        access_token = CobrowsingMiddlewareAccessToken.objects.get(token=token)
        if access_token.is_expired:
            is_valid = False
        else:
            access_token.is_expired = True
            access_token.save()
            is_valid = True
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error is_valid_cobrowse_middleware_token %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return is_valid


def random_with_n_digits(number):
    range_start = 10**(number - 1)
    range_end = (10**number) - 1
    return randint(range_start, range_end)


def extract_authorization_params(request):
    try:
        custom_encrypt_obj = CustomEncrypt()
        auth_token = request.META["HTTP_AUTHORIZATION"].split(" ")[-1]
        auth_token = custom_encrypt_obj.decrypt(auth_token)
        return tuple(auth_token.split(":"))
    except Exception:
        return None


def save_agent_closing_comments_cobrowseio(cobrowse_io_obj,
                                           agent_obj,
                                           comments,
                                           CobrowseAgentComment,
                                           comment_desc="",
                                           subcomments=""):
    try:
        CobrowseAgentComment.objects.create(cobrowse_io=cobrowse_io_obj,
                                            agent=agent_obj,
                                            agent_comments=comments,
                                            agent_subcomments=subcomments,
                                            comment_desc=comment_desc)
        logger.info("Cobrowsing session closing remarks saved successfully.", extra={
                    'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_agent_closing_comments_cobrowseio %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_visited_page_title_list_with_agent(active_agent, CobrowseIO):
    try:
        if active_agent.role == "admin":
            active_admin_user = active_agent
        else:
            active_admin_user = active_agent.get_access_token_obj().agent

        cobrowse_io_objs = CobrowseIO.objects.filter(is_archived=True,
                                                     access_token__agent=active_admin_user).filter(~Q(title=None))

        query_pages = list(
            set(list(cobrowse_io_objs.values_list('title', flat=True))))

        return query_pages
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_visited_page_title_list_with_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return []


def cobrowseio_agent_comments_manually_converted_leads(cobrowse_io_objs, CobrowseAgentComment):
    agent_comments_dict = {}
    try:
        if len(cobrowse_io_objs) != 0:
            for cobrowse_io in cobrowse_io_objs:
                agent_comments_dict[str(cobrowse_io.pk)] = {
                    "session_archived_cause": cobrowse_io.session_archived_cause,
                    "prev_agent_comments": "",
                    "prev_agent_subcomments": "",
                    "prev_comment_desc": "",
                    "agent_comments": "",
                    "agent_subcomments": "",
                    "comment_desc": ""
                }
                agent_comments_objs = CobrowseAgentComment.objects.filter(
                    cobrowse_io=cobrowse_io).order_by("datetime")
                if agent_comments_objs:
                    if cobrowse_io.session_archived_cause != "UNATTENDED":
                        agent_comment = agent_comments_objs.first()

                        if agent_comment.agent_comments != None:
                            agent_comments_dict[str(cobrowse_io.pk)][
                                "agent_comments"] = agent_comment.agent_comments

                        if agent_comment.agent_subcomments != None:
                            agent_comments_dict[str(cobrowse_io.pk)][
                                "agent_subcomments"] = agent_comment.agent_subcomments

                        if agent_comment.comment_desc != None:
                            agent_comments_dict[str(cobrowse_io.pk)][
                                "comment_desc"] = agent_comment.comment_desc

                    else:
                        if agent_comments_objs.count() > 1:
                            agent_comments = agent_comments_objs[:2]

                            if agent_comments[0].agent_comments != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "prev_agent_comments"] = agent_comments[0].agent_comments

                            if agent_comments[0].agent_subcomments != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "prev_agent_subcomments"] = agent_comments[0].agent_subcomments

                            if agent_comments[0].comment_desc != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "prev_comment_desc"] = agent_comments[0].comment_desc

                            if agent_comments[1].agent_comments != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "agent_comments"] = agent_comments[1].agent_comments

                            if agent_comments[1].agent_subcomments != None:
                                agent_comments_dict[str(cobrowse_io.pk)]["agent_subcomments"] = agent_comments[
                                    1].agent_subcomments

                            if agent_comments[1].comment_desc != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "comment_desc"] = agent_comments[1].comment_desc

                        elif agent_comments_objs.count() == 1:
                            agent_comments = agent_comments_objs[0]

                            if agent_comments.agent_comments != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "agent_comments"] = agent_comments.agent_comments

                            if agent_comments.agent_subcomments != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "agent_subcomments"] = agent_comments.agent_subcomments

                            if agent_comments.comment_desc != None:
                                agent_comments_dict[str(cobrowse_io.pk)][
                                    "comment_desc"] = agent_comments.comment_desc

            return agent_comments_dict

    except Exception as e:
        agent_comments_dict = {}
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_cobrowseio_manually_converted_leads %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return agent_comments_dict


def get_cobrowseio_agent_comments(cobrowse_io_objs, CobrowseAgent, CobrowseIO, CobrowseAgentComment):
    agent_comments_dict = {}
    try:
        if len(cobrowse_io_objs) != 0:
            for cobrowse_io in cobrowse_io_objs:
                agent_comments_dict[str(cobrowse_io.pk)] = {
                    "agent_comments": "",
                    "agent_subcomments": "",
                    "comment_desc": "",
                    "is_helpful": str(cobrowse_io.is_helpful)
                }
                agent_comments_objs = CobrowseAgentComment.objects.filter(
                    cobrowse_io=cobrowse_io).order_by("datetime")
                if agent_comments_objs:
                    agent_comment = agent_comments_objs.first()

                    if agent_comment.agent_comments != None:
                        agent_comments_dict[str(cobrowse_io.pk)][
                            "agent_comments"] = agent_comment.agent_comments

                    if agent_comment.agent_subcomments != None:
                        agent_comments_dict[str(cobrowse_io.pk)][
                            "agent_subcomments"] = agent_comment.agent_subcomments

                    if agent_comment.comment_desc != None:
                        agent_comments_dict[str(cobrowse_io.pk)][
                            "comment_desc"] = agent_comment.comment_desc

            return agent_comments_dict

    except Exception as e:
        agent_comments_dict = {}
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_cobrowseio_agent_comments %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return agent_comments_dict


def save_system_audit_trail(category, description, cobrowse_io, cobrowse_access_token, SystemAuditTrail, sender=None):
    try:
        SystemAuditTrail.objects.create(category=category,
                                        description=description,
                                        cobrowse_io=cobrowse_io,
                                        cobrowse_access_token=cobrowse_access_token,
                                        sender=sender)
        logger.info("System Audit Trail saved successfully",
                    extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_system_audit_trail %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def create_excel_sales_support_history(user, cobrowse_io_objs, access_token_obj):
    filename = None
    try:
        logger.info("create_excel_sales_support_history",
                    extra={'AppName': 'EasyAssist'})

        support_history_wb = Workbook()
        sheet1 = support_history_wb.add_sheet("Support History")

        sheet1.write(0, 0, CUSTOMER_NAME)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, "Agent Email ID")
        sheet1.col(2).width = 256 * 35
        sheet1.write(0, 3, "Cobrowsing Start Date & Time")
        sheet1.col(3).width = 256 * 25
        sheet1.write(0, 4, "Session End Date & Time")
        sheet1.col(4).width = 256 * 25
        sheet1.write(0, 5, "Time Spent")
        sheet1.col(5).width = 256 * 15
        sheet1.write(0, 6, "Lead Status")
        sheet1.col(6).width = 256 * 15
        sheet1.write(0, 7, "Title")
        sheet1.col(7).width = 256 * 60
        sheet1.write(0, 8, "Session ID")
        sheet1.col(8).width = 256 * 45
        sheet1.write(0, 9, AGENT_REMARKS)
        sheet1.col(9).width = 256 * 80
        if access_token_obj is not None and access_token_obj.allow_language_support:
            sheet1.write(0, 10, "Language")
            sheet1.col(10).width = 256 * 25

        index = 1
        for cobrowse_io_obj in cobrowse_io_objs:

            if not cobrowse_io_obj.is_lead:
                sheet1.write(index, 0, cobrowse_io_obj.full_name)
                sheet1.write(index, 1, cobrowse_io_obj.mobile_number)
            else:
                sheet1.write(
                    index, 0, cobrowse_io_obj.get_sync_data_value("name"))
                sheet1.write(
                    index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))
            sheet1.write(index, 2, cobrowse_io_obj.agent.user.username)
            sheet1.write(index, 3, str(cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%Y-%m-%d %I:%M %p")))
            sheet1.write(index, 4, str(cobrowse_io_obj.last_agent_update_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%Y-%m-%d %I:%M %p")))
            sheet1.write(index, 5, str(cobrowse_io_obj.total_time_spent()))

            if cobrowse_io_obj.is_archived == True and cobrowse_io_obj.is_helpful == True:
                sheet1.write(index, 6, "Converted")
            else:
                sheet1.write(index, 6, "Not Converted")
            sheet1.write(index, 7, cobrowse_io_obj.title)
            sheet1.write(index, 8, str(cobrowse_io_obj.session_id))
            sheet1.write(index, 9, cobrowse_io_obj.agent_comments)
            if access_token_obj is not None and access_token_obj.allow_language_support:
                if cobrowse_io_obj.supported_language is not None:
                    sheet1.write(
                        index, 10, cobrowse_io_obj.supported_language.title)
                else:
                    sheet1.write(index, 10, "")
            index += 1

        if not os.path.exists('secured_files/EasyAssistApp/CobrowseSupportHistory'):
            os.makedirs('secured_files/EasyAssistApp/CobrowseSupportHistory')

        filename = "secured_files/EasyAssistApp/CobrowseSupportHistory/" + \
            str(user.username) + ".xls"
        support_history_wb.save(filename)

        if get_save_in_s3_bucket_status():
            s3_bucket_upload_file_by_file_path(
                filename, str(user.username) + ".xls")
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_sales_support_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_unattended_leads_datails(user, cobrowse_io_objs, access_token_obj):
    filename = None
    try:
        logger.info("create_excel_unattended_leads_datails",
                    extra={'AppName': 'EasyAssist'})

        support_history_wb = Workbook()
        sheet1 = support_history_wb.add_sheet("Support History")

        sheet1.write(0, 0, CUSTOMER_NAME)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, "Agent Email ID")
        sheet1.col(2).width = 256 * 35
        sheet1.write(0, 3, COBROWSING_REQUEST_DATETIME)
        sheet1.col(3).width = 256 * 30
        sheet1.write(0, 4, "Title")
        sheet1.col(4).width = 256 * 80
        sheet1.write(0, 5, AGENT_REMARKS)
        sheet1.col(5).width = 256 * 80
        if access_token_obj is not None and access_token_obj.allow_language_support:
            sheet1.write(0, 6, "Language")
            sheet1.col(6).width = 256 * 25

        index = 1
        for cobrowse_io_obj in cobrowse_io_objs:
            if not cobrowse_io_obj.is_lead:
                sheet1.write(index, 0, cobrowse_io_obj.full_name)
                sheet1.write(index, 1, cobrowse_io_obj.mobile_number)
            else:
                sheet1.write(
                    index, 0, cobrowse_io_obj.get_sync_data_value("name"))
                sheet1.write(
                    index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

            sheet1.write(index, 2, cobrowse_io_obj.agent.user.username)
            sheet1.write(index, 3, str(cobrowse_io_obj.request_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime("%Y-%m-%d %I:%M %p")))
            sheet1.write(index, 4, cobrowse_io_obj.title)
            sheet1.write(index, 5, cobrowse_io_obj.agent_comments)
            if access_token_obj is not None and access_token_obj.allow_language_support:
                if cobrowse_io_obj.supported_language is not None:
                    sheet1.write(
                        index, 6, cobrowse_io_obj.supported_language.title)
                else:
                    sheet1.write(index, 6, "")
            index += 1

        if not os.path.exists('secured_files/EasyAssistApp/UnattendedLeads'):
            os.makedirs('secured_files/EasyAssistApp/UnattendedLeads')

        filename = "secured_files/EasyAssistApp/UnattendedLeads/" + \
            str(user.username) + ".xls"
        support_history_wb.save(filename)

        if get_save_in_s3_bucket_status():
            s3_bucket_upload_file_by_file_path(
                filename, str(user.username) + ".xls")

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_unattended_leads_datails %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_declined_leads_datails(user, cobrowse_io_objs, access_token_obj):
    filename = None
    try:
        logger.info("create_excel_declined_leads_datails",
                    extra={'AppName': 'EasyAssist'})

        support_history_wb = Workbook()
        sheet1 = support_history_wb.add_sheet("Support History")

        sheet1.write(0, 0, CUSTOMER_NAME)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, "Agent Email ID")
        sheet1.col(2).width = 256 * 35
        sheet1.write(0, 3, COBROWSING_REQUEST_DATETIME)
        sheet1.col(3).width = 256 * 30
        sheet1.write(0, 4, "Title")
        sheet1.col(4).width = 256 * 80
        sheet1.write(0, 5, AGENT_REMARKS)
        sheet1.col(5).width = 256 * 80
        if access_token_obj is not None and access_token_obj.allow_language_support:
            sheet1.write(0, 6, "Language")
            sheet1.col(6).width = 256 * 25

        index = 1
        for cobrowse_io_obj in cobrowse_io_objs:
            if not cobrowse_io_obj.is_lead:
                sheet1.write(index, 0, cobrowse_io_obj.full_name)
                sheet1.write(index, 1, cobrowse_io_obj.mobile_number)
            else:
                sheet1.write(
                    index, 0, cobrowse_io_obj.get_sync_data_value("name"))
                sheet1.write(
                    index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

            sheet1.write(index, 2, cobrowse_io_obj.agent.user.username)
            sheet1.write(index, 3, str(cobrowse_io_obj.request_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime("%Y-%m-%d %I:%M %p")))
            sheet1.write(index, 4, cobrowse_io_obj.title)
            sheet1.write(index, 5, cobrowse_io_obj.agent_comments)
            if access_token_obj is not None and access_token_obj.allow_language_support:
                if cobrowse_io_obj.supported_language is not None:
                    sheet1.write(
                        index, 6, cobrowse_io_obj.supported_language.title)
                else:
                    sheet1.write(index, 6, "")
            index += 1

        if not os.path.exists('secured_files/EasyAssistApp/DeclinedLeads'):
            os.makedirs('secured_files/EasyAssistApp/DeclinedLeads')

        filename = "secured_files/EasyAssistApp/DeclinedLeads/" + \
            str(user.username) + ".xls"
        support_history_wb.save(filename)

        if get_save_in_s3_bucket_status():
            s3_bucket_upload_file_by_file_path(
                filename, str(user.username) + ".xls")

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_declined_leads_datails %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_followup_leads_datails(user, cobrowse_io_objs, access_token_obj):
    filename = None
    try:
        logger.info("create_excel_followup_leads_datails",
                    extra={'AppName': 'EasyAssist'})

        support_history_wb = Workbook()
        sheet1 = support_history_wb.add_sheet("Support History")

        sheet1.write(0, 0, CUSTOMER_NAME)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, COBROWSING_REQUEST_DATETIME)
        sheet1.col(2).width = 256 * 30
        sheet1.write(0, 3, "Title")
        sheet1.col(3).width = 256 * 80
        if access_token_obj is not None and access_token_obj.choose_product_category:
            sheet1.write(0, 4, "Category")
            sheet1.col(4).width = 256 * 30
        if access_token_obj is not None and access_token_obj.allow_language_support:
            sheet1.write(0, 5, "Language")
            sheet1.col(5).width = 256 * 25

        index = 1
        for cobrowse_io_obj in cobrowse_io_objs:
            if not cobrowse_io_obj.is_lead:
                sheet1.write(index, 0, cobrowse_io_obj.full_name)
                sheet1.write(index, 1, cobrowse_io_obj.mobile_number)
            else:
                sheet1.write(
                    index, 0, cobrowse_io_obj.get_sync_data_value("name"))
                sheet1.write(
                    index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

            sheet1.write(index, 2, str(cobrowse_io_obj.request_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime("%Y-%m-%d %I:%M %p")))
            sheet1.write(index, 3, cobrowse_io_obj.title)
            if access_token_obj is not None and access_token_obj.choose_product_category:
                if cobrowse_io_obj.product_category is not None:
                    sheet1.write(
                        index, 4, cobrowse_io_obj.product_category.title)
                else:
                    sheet1.write(index, 4, "")
            if access_token_obj is not None and access_token_obj.allow_language_support:
                if cobrowse_io_obj.supported_language is not None:
                    sheet1.write(
                        index, 5, cobrowse_io_obj.supported_language.title)
                else:
                    sheet1.write(index, 5, "")
            index += 1

        if not os.path.exists('secured_files/EasyAssistApp/FollowUpLeads'):
            os.makedirs('secured_files/EasyAssistApp/FollowUpLeads')

        filename = "secured_files/EasyAssistApp/FollowUpLeads/" + \
            str(user.username) + ".xls"
        support_history_wb.save(filename)

        if get_save_in_s3_bucket_status():
            s3_bucket_upload_file_by_file_path(
                filename, str(user.username) + ".xls")

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_followup_leads_datails %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_meeting_support_history(user, audit_trail_objs, CobrowsingFileAccessManagement):
    filename = None
    try:
        logger.info("create_excel_sales_support_history",
                    extra={'AppName': 'EasyAssist'})

        if not os.path.exists('secured_files/EasyAssistApp/MeetingSupportHistory'):
            os.makedirs('secured_files/EasyAssistApp/MeetingSupportHistory')

        meeting_support_history_wb = Workbook()
        sheet1 = meeting_support_history_wb.add_sheet(
            "Meeting Support History")

        sheet1.write(0, 0, CUSTOMER_NAME)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, "Agent Email ID")
        sheet1.col(2).width = 256 * 35
        sheet1.write(0, 3, "Meeting Start Date & Time")
        sheet1.col(3).width = 256 * 25
        sheet1.write(0, 4, "Meeting End Date & Time")
        sheet1.col(4).width = 256 * 25
        sheet1.write(0, 5, "Time Spent")
        sheet1.col(5).width = 256 * 15
        sheet1.write(0, 6, "Meeting Status")
        sheet1.col(6).width = 256 * 15
        sheet1.write(0, 7, "Meeting ID")
        sheet1.col(7).width = 256 * 45
        sheet1.write(0, 8, "Meeting Notes")
        sheet1.col(8).width = 256 * 45
        sheet1.write(0, 9, "Customer Location")
        sheet1.col(9).width = 256 * 90
        sheet1.write(0, 10, "Attachment")
        sheet1.col(10).width = 256 * 45

        index = 1

        for audit_trail_obj in audit_trail_objs:
            client_location_details = json.loads(
                audit_trail_obj.client_location_details)
            client_address = []
            for location_obj in client_location_details['items']:
                client_name = location_obj['client_name']
                longitude = location_obj['longitude']
                latitude = location_obj['latitude']
                if latitude == 'None' or longitude == 'None':
                    continue

                if "address" in location_obj and location_obj["address"] != "None":
                    client_address.append(
                        client_name + " - " + location_obj["address"])
                else:
                    try:
                        location = geocoder.google(
                            [latitude, longitude], method="reverse", key=GOOGLE_GEOCODER_KEY)
                        client_address.append(
                            client_name + " - " + location.address)
                        location_obj["address"] = location.address
                    except Exception:
                        logger.warning("Error in create_excel_meeting_support_history, client's address not found",
                                       extra={'AppName': 'EasyAssist'})

            screenshots = json.loads(audit_trail_obj.meeting_screenshot)
            audit_trail_obj.client_location_details = json.dumps(
                client_location_details)
            audit_trail_obj.save()

            attachment_link = ""
            if len(screenshots['items']) > 0:
                zip_file_path = "secured_files/EasyAssistApp/MeetingSupportHistory/attachment_" + \
                    uuid.uuid4().hex + ".zip"

                zip_obj = ZipFile(zip_file_path, 'w')

                for screenshot_obj in screenshots['items']:
                    file_key = screenshot_obj['screenshot']
                    try:
                        file_access_management_obj = CobrowsingFileAccessManagement.objects.get(
                            key=file_key, is_public=False)

                        path_to_file = file_access_management_obj.file_path
                        file_name = path_to_file.split(os.sep)[-1]
                        zip_obj.write(path_to_file[1:], file_name)
                    except Exception:
                        logger.warning("Error in create_excel_meeting_support_history, Invalid file_key",
                                       extra={'AppName': 'EasyAssist'})
                zip_obj.close()

                if get_save_in_s3_bucket_status():
                    key = s3_bucket_upload_file_by_file_path(
                        zip_file_path, file_name)
                    s3_file_path = s3_bucket_download_file(
                        key, 'EasyAssistApp/MeetingSupportHistory/', '.zip')
                    zip_file_path = s3_file_path.split("EasyChat/", 1)[1]

                zip_file_path = "/" + zip_file_path
                zip_file_obj = CobrowsingFileAccessManagement.objects.create(
                    file_path=zip_file_path, is_public=False)
                attachment_link = settings.EASYCHAT_HOST_URL + \
                    '/easy-assist/download-file/' + str(zip_file_obj.key)

            sheet1.write(index, 0, audit_trail_obj.cobrowse_video.full_name)
            sheet1.write(
                index, 1, audit_trail_obj.cobrowse_video.mobile_number)

            sheet1.write(
                index, 2, audit_trail_obj.cobrowse_video.agent.user.username)
            sheet1.write(index, 3, str(audit_trail_obj.agent_joined.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%Y-%m-%d - %I:%M %p")))
            sheet1.write(index, 4, str(audit_trail_obj.meeting_ended.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%Y-%m-%d - %I:%M %p")))
            sheet1.write(index, 5, str(
                audit_trail_obj.get_readable_meeting_duration()))

            if audit_trail_obj.cobrowse_video.is_expired == True:
                sheet1.write(index, 6, "Completed")
            else:
                sheet1.write(index, 6, "Scheduled")
            sheet1.write(index, 7, str(
                audit_trail_obj.cobrowse_video.meeting_id))

            if audit_trail_obj.agent_notes == None:
                sheet1.write(index, 8, "")
            else:
                sheet1.write(index, 8, audit_trail_obj.agent_notes)

            sheet1.write(index, 9, " | ".join(client_address))
            sheet1.write(index, 10, attachment_link)

            index += 1

        filename = "secured_files/EasyAssistApp/MeetingSupportHistory/" + \
            str(user.username) + ".xls"
        meeting_support_history_wb.save(filename)

        if get_save_in_s3_bucket_status():
            s3_bucket_upload_file_by_file_path(
                filename, str(user.username) + ".xls")

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_meeting_support_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_screen_recording_history(user, screen_recording_objs, CobrowsingFileAccessManagement):
    filename = None
    try:
        logger.info("create_excel_screen_recording_history",
                    extra={'AppName': 'EasyAssist'})

        if not os.path.exists('secured_files/EasyAssistApp/ScreenRecordingHistory'):
            os.makedirs('secured_files/EasyAssistApp/ScreenRecordingHistory')

        meeting_support_history_wb = Workbook()
        sheet1 = meeting_support_history_wb.add_sheet(
            "Meeting Support History")

        sheet1.write(0, 0, CUSTOMER_NAME)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, "Agent Email ID")
        sheet1.col(2).width = 256 * 35
        sheet1.write(0, 3, "Recording Start Date & Time")
        sheet1.col(3).width = 256 * 25
        sheet1.write(0, 4, "Recording Duration")
        sheet1.col(4).width = 256 * 20

        index = 1

        for screen_recording_obj in screen_recording_objs:

            sheet1.write(index, 0, screen_recording_obj.cobrowse_io.full_name)
            sheet1.write(
                index, 1, screen_recording_obj.cobrowse_io.mobile_number)

            sheet1.write(
                index, 2, screen_recording_obj.agent.user.username)
            sheet1.write(index, 3, str(screen_recording_obj.recording_started.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%Y-%m-%d - %I:%M %p")))
            sheet1.write(index, 4, str(
                screen_recording_obj.get_screen_recording_duration()))

            index += 1

        filename = "secured_files/EasyAssistApp/ScreenRecordingHistory/" + \
            str(user.username) + ".xls"
        meeting_support_history_wb.save(filename)

        if get_save_in_s3_bucket_status():
            s3_bucket_upload_file_by_file_path(
                filename, str(user.username) + ".xls")

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_screen_recording_history %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_easyassist_customer_details(active_agent, easyassist_customer_objs):
    filename = None
    try:
        logger.info("create_excel_easyassist_customer_details",
                    extra={'AppName': 'EasyAssist'})

        easyassist_customer_details_wb = Workbook()
        sheet1 = easyassist_customer_details_wb.add_sheet("Customer Details")

        sheet1.write(0, 0, "Date & Time of Clicking Request for Support")
        sheet1.col(0).width = 256 * 50
        sheet1.write(0, 1, CUSTOMER_NAME)
        sheet1.col(1).width = 256 * 20
        sheet1.write(0, 2, CUSTOMER_MOBILE_NUMBER)
        sheet1.col(2).width = 256 * 25
        sheet1.write(0, 3, "Pop-up Type")
        sheet1.col(3).width = 256 * 50

        index = 1

        for easyassist_customer in easyassist_customer_objs:
            sheet1.write(index, 0, str(easyassist_customer.request_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime("%d-%b-%Y, %I:%M %p")))
            sheet1.write(index, 1, easyassist_customer.full_name)
            sheet1.write(index, 2, get_masked_data_if_hashed(
                easyassist_customer.mobile_number))
            if easyassist_customer.lead_initiated_by == "greeting_bubble":
                sheet1.write(index, 3, "Greeting bubble")
            elif easyassist_customer.lead_initiated_by == "icon":
                sheet1.write(index, 3, "Icon")
            elif easyassist_customer.lead_initiated_by == "floating_button":
                sheet1.write(index, 3, "Floating button")
            elif easyassist_customer.lead_initiated_by == "exit_intent":
                sheet1.write(index, 3, "Exit Intent")
            elif easyassist_customer.lead_initiated_by == "inactivity_popup":
                sheet1.write(index, 3, "Inactivity Pop-up")

            index += 1

        folder_path = EASYASSISTAPP_SECURED_FILES_PATH + "CustomerDetails/"

        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        filename = folder_path + "EasyAssist_CustomerDetails_" + \
            str(active_agent.user.username) + ".xls"
        easyassist_customer_details_wb.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_customer_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_easyassist_conversions_by_url(cobrowsing_type, active_agent, cobrowseio_converted_by_url_objs):
    filename = None
    try:
        logger.info("create_excel_easyassist_conversions_by_url",
                    extra={'AppName': 'EasyAssist'})

        easyassist_converted_by_url_wb = Workbook()
        sheet1 = easyassist_converted_by_url_wb.add_sheet(
            "Customer_converted_through_URL")

        sheet1.write(0, 0, "Date & Time")
        sheet1.col(0).width = 256 * 25
        sheet1.write(0, 1, "Session ID")
        sheet1.col(1).width = 256 * 45
        sheet1.write(0, 2, "Customer Converted through URL")
        sheet1.col(2).width = 256 * 80

        index = 1

        for cobrowse_io_obj in cobrowseio_converted_by_url_objs:
            sheet1.write(index, 0, str(cobrowse_io_obj.lead_converted_url_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime("%d-%b-%Y, %I:%M %p")))
            sheet1.write(index, 1, str(cobrowse_io_obj.session_id))
            sheet1.write(index, 2, cobrowse_io_obj.lead_converted_url)

            index += 1

        folder_path = EASYASSISTAPP_SECURED_FILES_PATH + "URL_Conversions/"

        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        filename_prefix = folder_path + "EasyAssist_Inbound_URL_Conversions"
        if cobrowsing_type == "outbound":
            filename_prefix = folder_path + "EasyAssist_Outbound_URL_Conversions"

        elif cobrowsing_type == "reverse":
            filename_prefix = folder_path + "EasyAssist_Reverse_URL_Conversions"

        filename = filename_prefix + ".xls"
        easyassist_converted_by_url_wb.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_conversions_by_url %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_user_details(active_agent):
    filename = None
    try:
        logger.info("In create_excel_user_details",
                    extra={'AppName': 'EasyAssist'})
        support_history_wb = Workbook()
        
        def create_new_sheet():
            global supervisors_col, account_status_col, last_active_date_col, creation_date_col,\
                deactivation_date_col, activation_date_col
            nonlocal wb_index
            sheet = support_history_wb.add_sheet("Agent Details-" + str(wb_index))
            
            style = easyxf(HEADING_FONT_STYLE)
            
            sheet.write(0, 0, "Date & Time", style=style)
            sheet.col(0).width = 256 * 20
            sheet.write(0, 1, "Agent Name", style=style)
            sheet.col(1).width = 256 * 20
            sheet.write(0, 2, "Agent Email ID", style=style)
            sheet.col(2).width = 256 * 20
            sheet.write(0, 3, "Mobile Number", style=style)
            sheet.col(3).width = 256 * 20
            sheet.write(0, 4, "Activity Status", style=style)
            sheet.col(4).width = 256 * 20
            
            prev_col = 4
            
            if active_agent.role == "admin_ally":
                supervisors_col = prev_col + 1
                sheet.write(0, supervisors_col, "Supervisors", style=style)
                sheet.col(supervisors_col).width = 256 * 20
                prev_col = supervisors_col
            
            account_status_col = prev_col + 1
            sheet.write(0, account_status_col, "Account Status", style=style)
            sheet.col(account_status_col).width = 256 * 20
            prev_col = account_status_col
            
            last_active_date_col = prev_col + 1
            sheet.write(0, last_active_date_col, "Last Active Date", style=style)
            sheet.col(last_active_date_col).width = 256 * 20
            prev_col = last_active_date_col
            
            creation_date_col = prev_col + 1
            sheet.write(0, creation_date_col, "Creation Date", style=style)
            sheet.col(creation_date_col).width = 256 * 20
            prev_col = creation_date_col
            
            deactivation_date_col = prev_col + 1
            sheet.write(0, deactivation_date_col, "Deactivation Date", style=style)
            sheet.col(deactivation_date_col).width = 256 * 20
            prev_col = deactivation_date_col 
            
            activation_date_col = prev_col + 1 
            sheet.write(0, activation_date_col, "Activation Date", style=style)
            sheet.col(activation_date_col).width = 256 * 20
            prev_col = activation_date_col + 1

            return sheet

        agents = get_list_agents_under_admin(active_agent, is_active=None, is_account_active=None)
        
        admin_allys = get_list_admin_ally(
            active_agent, is_active=None, is_account_active=None)
            
        supervisors = []
        for admin_ally in admin_allys:
            supervisors += list(admin_ally.agents.filter(role="supervisor"))
        for supervisor in supervisors:
            agents += list(supervisor.agents.filter(role="agent"))
    
        agents = list(set(agents))
        
        agents.sort(
            key=lambda agent: agent.agent_creation_datetime, reverse=True)
        current_datetime = datetime.datetime.now()
        wb_index = 1
        access_token_obj = active_agent.get_access_token_obj()
        sheet = create_new_sheet()
        
        sheet.set_panes_frozen('1')
        sheet.set_horz_split_pos(1)  
        sheet.set_vert_split_pos(0) 
    
        active_status_style = easyxf(ACTIVE_STATUS_STYLE)
        inactive_status_style = easyxf(INACTIVE_STATUS_STYLE)
    
        index = 1
        for agent in agents:
            if index > 50000:
                wb_index += 1
                sheet = create_new_sheet()
                index = 1

            sheet.write(index, 0, str(current_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            sheet.write(index, 1, agent.user.first_name)
            sheet.write(index, 2, agent.user.username)
            
            if agent.mobile_number:
                sheet.write(index, 3, agent.mobile_number)
            else:
                sheet.write(index, 3, "-")

            if agent.is_active and agent.is_cobrowsing_active == False and agent.is_cognomeet_active == False and access_token_obj.allow_agent_to_customer_cobrowsing == False:
                sheet.write(index, 4, "Available")
            elif agent.is_active:
                if access_token_obj.allow_agent_to_customer_cobrowsing == False and agent.is_cobrowsing_active or agent.is_cognomeet_active:
                    sheet.write(index, 4, "Busy")
                else:
                    sheet.write(index, 4, "Online")
            else:
                sheet.write(index, 4, "Offline")

            if active_agent.role == "admin_ally":
                sheet.write(index, supervisors_col, agent.get_supervisors())
            
            if agent.is_account_active:
                sheet.write(index, account_status_col, "Active", style=active_status_style)
            else:
                sheet.write(index, account_status_col, "Inactive", style=inactive_status_style)
                
            last_agent_active_datetime_str = str(agent.last_agent_active_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT))
            
            agent_creation_datetime_str = str(agent.agent_creation_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT))
            
            if last_agent_active_datetime_str == agent_creation_datetime_str:
                sheet.write(index, last_active_date_col, '-')
            else:
                sheet.write(index, last_active_date_col, last_agent_active_datetime_str)
                
            sheet.write(index, creation_date_col, agent_creation_datetime_str)

            if agent.agent_deactivation_datetime:
                sheet.write(index, deactivation_date_col, str(agent.agent_deactivation_datetime.astimezone(
                    pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            else:
                sheet.write(index, deactivation_date_col, '-')
        
            if agent.agent_activation_datetime:
                sheet.write(index, activation_date_col, str(agent.agent_activation_datetime.astimezone(
                    pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            else:
                sheet.write(index, activation_date_col, '-')
            index += 1

        folder_path = EASYASSISTAPP_SECURED_FILES_PATH + "AgentDetails/"

        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            
        filename = folder_path + "EasyAssist_Agent_Details_" + \
            str(active_agent.user.username) + ".xls"
        support_history_wb.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_user_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_supervisor_details(active_agent):
    filename = None
    try:
        logger.info("In create_excel_supervisor_details",
                    extra={'AppName': 'EasyAssist'})

        supervisors = get_list_supervisor_under_admin(
            active_agent, is_active=None, is_account_active=None)

        admin_allys = get_list_admin_ally(
            active_agent, is_active=None, is_account_active=None)
        
        for admin_ally in admin_allys:
            supervisors += list(admin_ally.agents.filter(role="supervisor"))

        supervisors = list(set(supervisors))
        
        supervisors.sort(
            key=lambda supervisor: supervisor.agent_creation_datetime, reverse=True)
        current_datetime = datetime.datetime.now()

        admin_allys.sort(
            key=lambda admin_allys: admin_allys.agent_creation_datetime, reverse=True)

        support_history_wb = Workbook()
        sheet1 = support_history_wb.add_sheet("Admin_Ally and Supervisors")
    
        style = easyxf(HEADING_FONT_STYLE)

        sheet1.write(0, 0, "Date & Time", style=style)
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, "Agent Name", style=style)
        sheet1.col(1).width = 256 * 20
        sheet1.write(0, 2, "Agent Email ID", style=style)
        sheet1.col(2).width = 256 * 20
        sheet1.write(0, 3, "Role", style=style)
        sheet1.col(3).width = 256 * 20
        sheet1.write(0, 4, "Mobile Number", style=style)
        sheet1.col(4).width = 256 * 20
        sheet1.write(0, 5, "Account Status", style=style)
        sheet1.col(5).width = 256 * 20
        sheet1.write(0, 6, "Creation Date", style=style)
        sheet1.col(6).width = 256 * 20
        sheet1.write(0, 7, "Deactivation Date", style=style)
        sheet1.col(7).width = 256 * 20
        sheet1.write(0, 8, "Activation Date", style=style)
        sheet1.col(8).width = 256 * 20
        
        sheet1.set_panes_frozen('1')
        sheet1.set_horz_split_pos(1)  
        sheet1.set_vert_split_pos(0)
        
        active_status_style = easyxf(ACTIVE_STATUS_STYLE)
        inactive_status_style = easyxf(INACTIVE_STATUS_STYLE)
        
        index = 1
        for admin_ally in admin_allys:
            sheet1.write(index, 0, str(current_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            sheet1.write(index, 1, admin_ally.user.first_name)
            sheet1.write(index, 2, admin_ally.user.username)
            sheet1.write(index, 3, "Admin Ally")
            
            if admin_ally.mobile_number:
                sheet1.write(index, 4, admin_ally.mobile_number)
            else:
                sheet1.write(index, 4, "-")
            
            if admin_ally.is_account_active:
                sheet1.write(index, 5, "Active", style=active_status_style)
            else:
                sheet1.write(index, 5, "Inactive", style=inactive_status_style)
                
            sheet1.write(index, 6, str(admin_ally.agent_creation_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            
            if admin_ally.agent_deactivation_datetime:
                sheet1.write(index, 7, str(admin_ally.agent_deactivation_datetime.astimezone(
                    pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            else:
                sheet1.write(index, 7, '-')
                
            if admin_ally.agent_activation_datetime:
                sheet1.write(index, 8, str(admin_ally.agent_activation_datetime.astimezone(
                    pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            else:
                sheet1.write(index, 8, '-')
            
            index += 1

        for supervisor in supervisors:
            sheet1.write(index, 0, str(current_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            sheet1.write(index, 1, supervisor.user.first_name)
            sheet1.write(index, 2, supervisor.user.username)
            sheet1.write(index, 3, supervisor.role.capitalize())
            
            if supervisor.mobile_number:
                sheet1.write(index, 4, supervisor.mobile_number)
            else:
                sheet1.write(index, 4, "-")
            
            if supervisor.is_account_active:
                sheet1.write(index, 5, "Active", style=active_status_style)
            else:
                sheet1.write(index, 5, "Inactive", style=inactive_status_style)
                
            sheet1.write(index, 6, str(supervisor.agent_creation_datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            if supervisor.agent_deactivation_datetime != None:
                sheet1.write(index, 7, str(supervisor.agent_deactivation_datetime.astimezone(
                    pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            else:
                sheet1.write(index, 7, '-')
            if supervisor.agent_activation_datetime != None:
                sheet1.write(index, 8, str(supervisor.agent_activation_datetime.astimezone(
                    pytz.timezone(ASIA_KOLKATA)).strftime(TIME_FORMAT)))
            else:
                sheet1.write(index, 8, '-')
            index += 1

        folder_path = EASYASSISTAPP_SECURED_FILES_PATH + 'AgentDetails/'
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        filename = folder_path + \
            str(active_agent.user.username) + ".xls"
        support_history_wb.save(filename)
        if get_save_in_s3_bucket_status():
            key = s3_bucket_upload_file_by_file_path(
                filename, str(active_agent.user.username) + ".xls")
            s3_file_path = s3_bucket_download_file(
                key, 'EasyAssistApp/AgentDetails/', '.xls')
            return s3_file_path.split("EasyChat/", 1)[1]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_user_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_wrong_user_data(wrong_data_list):
    filename = None
    try:
        logger.info("In create_excel_wrong_user_data",
                    extra={'AppName': 'EasyAssist'})

        wrong_data_workbook = Workbook()
        wrong_data_sheet = wrong_data_workbook.add_sheet(
            "Wrong Information Sheet")

        style = XFStyle()

        font = Font()
        font.bold = True
        style.font = font

        wrong_data_sheet.write(0, 0, "Row No", style=style)
        wrong_data_sheet.col(0).width = 256 * 10
        wrong_data_sheet.write(0, 1, "Detail", style=style)
        wrong_data_sheet.col(1).width = 256 * 40
        wrong_data_sheet.write(0, 2, "Name", style=style)
        wrong_data_sheet.col(2).width = 256 * 20
        wrong_data_sheet.write(0, 3, "Email ID", style=style)
        wrong_data_sheet.col(3).width = 256 * 20
        wrong_data_sheet.write(0, 4, "Mobile number", style=style)
        wrong_data_sheet.col(4).width = 256 * 20
        wrong_data_sheet.write(0, 5, "User type", style=style)
        wrong_data_sheet.col(5).width = 256 * 20
        wrong_data_sheet.write(0, 6, "Product Category", style=style)
        wrong_data_sheet.col(6).width = 256 * 20
        wrong_data_sheet.write(0, 7, "Supported languages", style=style)
        wrong_data_sheet.col(7).width = 256 * 20
        wrong_data_sheet.write(0, 8, "Support Level", style=style)
        wrong_data_sheet.col(8).width = 256 * 20
        wrong_data_sheet.write(0, 9, "Supervisor Email/Name", style=style)
        wrong_data_sheet.col(9).width = 256 * 20
        wrong_data_sheet.write(
            0, 10, "Assign follow-up leads(Yes/No)", style=style)
        wrong_data_sheet.col(10).width = 256 * 20

        row_number = 1
        for wrong_data in wrong_data_list:
            data_list = []
            data_list.append(wrong_data["row_num"])
            data_list.append(wrong_data["detail"])
            data_list += wrong_data["data_array"]
            col_number = 0
            for data in data_list:
                wrong_data_sheet.write(row_number, col_number, data)
                col_number += 1
            row_number += 1

        current_datetime = str(datetime.datetime.now().astimezone(
            pytz.timezone(ASIA_KOLKATA)).strftime("%d%b%Y%I:%M%p"))

        folder_path = EASYASSISTAPP_SECURED_FILES_PATH + "EasyassistWrongUserData/"

        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        filename = folder_path + "Wrong_Data_Excel_" + \
            str(current_datetime) + ".xls"
        wrong_data_workbook.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_wrong_user_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def add_users_from_excel_document(file_path, active_agent, User, CobrowseAgent, LanguageSupport, ProductCategory, CobrowsingFileAccessManagement):
    response = {}
    response["status"] = 500
    response["message"] = "Exception Occured"
    try:
        logger.info("Inside add_users_from_excel_document",
                    extra={'AppName': 'EasyAssist'})
        cobrowse_access_token = active_agent.get_access_token_obj()

        wrong_data_list = []
        
        try:
            xlrd.xlsx.ensure_elementtree_imported(False, None)
            xlrd.xlsx.Element_has_iter = True
        except Exception:
            pass

        workbook1 = open_workbook(file_path)
        sheet1 = workbook1.sheet_by_name(workbook1.sheet_names()[0])

        regName = '^[^\s][a-zA-Z ]+$'
        regLanguage = '^[^\s][a-zA-Z]+$'
        regEmail = '^[a-zA-Z0-9._-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,4}$'
        regMob = '^[6-9]{1}[0-9]{9}$'

        total_rows = sheet1.nrows
        for row_num in range(1, total_rows):
            invalid_data_found = False
            logger.info("row: " + str(row_num),
                        extra={'AppName': 'EasyAssist'})
            row_values = sheet1.row_values(row_num)
            logger.info("row_values: " + str(row_values),
                        extra={'AppName': 'EasyAssist'})

            user_type = str(row_values[3]).strip()
            user_type = user_type.lower()

            if user_type == "agent" and active_agent.is_agent_creation_limit_exhausted():
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": AGENT_CREATION_LIMIT_EXHAUST_ERROR})
                continue

            if user_type == "supervisor" and active_agent.is_supervisor_creation_limit_exhausted():
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": SUPERVISOR_CREATION_LIMIT_EXHAUST_ERROR})
                continue
            
            user_name = str(row_values[0]).strip()
            if len(user_name) == 0:
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Name Empty"})
                continue
            if not re.search(regName, user_name):
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Name Wrong Format"})
                continue
            logger.info("Valid user_name: " + str(user_name),
                        extra={'AppName': 'EasyAssist'})

            user_email = str(row_values[1]).strip()
            if len(user_email) == 0:
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Email Empty"})
                continue
            if not re.search(regEmail, user_email):
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Email Wrong Format"})
                continue
            if CobrowseAgent.objects.filter(user__email=user_email, is_account_active=True).count() > 0:
                logger.info(str(user_email) + " is already an existing user",
                            extra={'AppName': 'EasyAssist'})
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Email Already Exists"})
                continue
            logger.info("Valid user_email: " + str(user_email),
                        extra={'AppName': 'EasyAssist'})

            user_mobile = str(row_values[2]).strip()
            if user_mobile == "":
                user_mobile = None
            if user_mobile is not None:
                if len(user_mobile) < 10:
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": INVALID_MOBILE_NUMBER})
                    continue
                try:
                    user_mobile = str(int(float(user_mobile)))
                except Exception:
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": INVALID_MOBILE_NUMBER})
                    continue
                if not re.search(regMob, user_mobile):
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": INVALID_MOBILE_NUMBER})
                    continue
                if CobrowseAgent.objects.filter(mobile_number=user_mobile, is_account_active=True).count() > 0:
                    logger.info(
                        str(user_mobile) + " already exists for a user", extra={'AppName': 'EasyAssist'})
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": "Mobile Number Already Exists"})
                    continue
            logger.info("Valid user_mobile: " + str(user_mobile),
                        extra={'AppName': 'EasyAssist'})

            if len(user_type) == 0:
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Invalid User Type"})
                continue
            if active_agent.role == "admin_ally" and user_type == "admin_ally":
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Admin ally creation is not allowed from admin ally account."})
                continue
            if user_type != "agent" and user_type != "supervisor" and user_type != "admin_ally":
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Invalid User Type"})
                continue
            logger.info("Valid user_type: " + str(user_type),
                        extra={'AppName': 'EasyAssist'})

            product_category_list = str(row_values[4]).strip()
            product_category_pk_list = []
            # if cobrowse_access_token.choose_product_category == True:
            supported_product_category_list = product_category_list.split(",")
            for product_category in supported_product_category_list:
                product_category = str(product_category).strip()
                if product_category == "" and cobrowse_access_token.choose_product_category == False:
                    continue
                if len(product_category) < 2:
                    wrong_data_list.append({"row_num": row_num, "data_array": row_values,
                                            "detail": "Product category length should be greater than one"})
                    invalid_data_found = True
                    break
                if len(product_category) > 20:
                    wrong_data_list.append({"row_num": row_num, "data_array": row_values,
                                            "detail": "Product category length should be less than 20"})
                    invalid_data_found = True
                    break
                product_category = product_category[0].upper(
                ) + product_category[1:].lower()
                product_category_obj = active_agent.product_category.filter(
                    title=product_category).first()
                if product_category_obj == None:
                    total_product_category = active_agent.product_category.all().count()
                    product_category_obj = ProductCategory.objects.create(
                        title=product_category)
                    product_category_obj.index = total_product_category
                    product_category_obj.is_deleted = False
                    product_category_obj.save()
                    active_agent.product_category.add(product_category_obj)
                    active_agent.save()
                if product_category_obj:
                    product_category_pk_list.append(product_category_obj.pk)
                else:
                    product_category_pk_list.clear()
                    break

            if invalid_data_found:
                continue
            # if len(product_category_pk_list) == 0:
            #     continue
            logger.info("Valid product_category_list: " + str(product_category_list),
                        extra={'AppName': 'EasyAssist'})

            language_list = str(row_values[5]).strip()
            selected_language_pk_list = []
            # if cobrowse_access_token.allow_language_support == True:
            supported_language_list = language_list.split(",")
            for language in supported_language_list:
                language = str(language).strip()
                if language == "" and cobrowse_access_token.allow_language_support == False:
                    continue
                if len(language) < 2:
                    wrong_data_list.append({"row_num": row_num, "data_array": row_values,
                                            "detail": "Support language length should be greater than one"})
                    invalid_data_found = True
                    break
                if len(language) > 20:
                    wrong_data_list.append({"row_num": row_num, "data_array": row_values,
                                            "detail": "Support language length should be less than 20"})
                    invalid_data_found = True
                    break
                if not re.search(regLanguage, language):
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": "Invalid Support Language"})
                    invalid_data_found = True
                    break
                language = language[0].upper() + language[1:].lower()
                language_obj = active_agent.supported_language.filter(
                    title=language).first()
                if language_obj == None:
                    total_language_count = active_agent.supported_language.all().count()
                    language_obj = LanguageSupport.objects.create(
                        title=language)
                    language_obj.index = total_language_count
                    language_obj.is_deleted = False
                    language_obj.save()

                    active_agent.supported_language.add(language_obj)
                    active_agent.save()
                if language_obj:
                    selected_language_pk_list.append(language_obj.pk)
                else:
                    selected_language_pk_list.clear()
                    break

            if invalid_data_found:
                continue

            # if len(selected_language_pk_list) == 0:
            #     continue
            logger.info("Valid language_list: " + str(language_list),
                        extra={'AppName': 'EasyAssist'})

            support_level = str(row_values[6]).strip()
            if len(support_level) > 0:
                support_level = support_level[0].upper() + support_level[1:]
            if len(support_level) == 0 or support_level not in COBROWSING_AGENT_SUPPORT_DICT:
                wrong_data_list.append(
                    {"row_num": row_num, "data_array": row_values, "detail": "Wrong Support Level"})
                continue
            logger.info("Valid support_level: " + str(support_level),
                        extra={'AppName': 'EasyAssist'})

            supervisor_username = str(row_values[7]).strip()
            supervisor_obj = None
            if user_type == "agent":
                if len(supervisor_username) == 0:
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": "Supervisor Name Empty"})
                    continue
                supervisor_obj = active_agent.agents.all().filter(
                    role="supervisor", is_account_active=True, user__username=supervisor_username).first()
                if supervisor_obj is None:
                    if active_agent.user.username == supervisor_username:
                        supervisor_obj = active_agent
                    else:
                        wrong_data_list.append(
                            {"row_num": row_num, "data_array": row_values, "detail": "Wrong Supervisor Name"})
                        continue
            else:
                supervisor_obj = active_agent
            logger.info("Valid supervisor_username: " +
                        str(supervisor_username), extra={'AppName': 'EasyAssist'})

            if user_type == "admin_ally":
                selected_supervisor_pk_list = ["-1"]
            else:
                selected_supervisor_pk_list = [str(supervisor_obj.pk)]

            category_matched = check_for_supervisor_category_language_match(active_agent, user_type, selected_supervisor_pk_list, selected_language_pk_list,
                                                                            product_category_pk_list, CobrowseAgent)
            if len(category_matched):
                if category_matched["matched_error"] == "language":
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": "Invalid Support Language"})
                else:
                    wrong_data_list.append(
                        {"row_num": row_num, "data_array": row_values, "detail": "Invalid Product Category"})
                continue

            if cobrowse_access_token.enable_followup_leads_tab:
                assign_followup_lead = str(row_values[8]).strip()
                if user_type == "agent":
                    if len(assign_followup_lead) == 0:
                        wrong_data_list.append(
                            {"row_num": row_num, "data_array": row_values, "detail": "Please enter Yes or No in Assign follow-up leads column for the entered agent"})
                        continue
                    assign_followup_lead = assign_followup_lead.lower()
                    if assign_followup_lead != "yes" and assign_followup_lead != "no":
                        wrong_data_list.append(
                            {"row_num": row_num, "data_array": row_values, "detail": "Please enter Yes or No in Assign follow-up leads column for the entered agent"})
                        continue
                    logger.info("Valid assign_followup_lead: " + str(assign_followup_lead),
                                extra={'AppName': 'EasyAssist'})

            user = None
            password = generate_password(cobrowse_access_token.password_prefix)

            try:
                user = User.objects.get(username=user_email)
                user.first_name = user_name
                user.email = user_email
                user.save()
            except Exception:
                user = User.objects.create(first_name=user_name,
                                           email=user_email,
                                           username=user_email,
                                           status="2",
                                           role="bot_builder")
                user.set_password(password)
                user.save()

            platform_url = settings.EASYCHAT_HOST_URL

            thread = threading.Thread(target=send_password_over_email, args=(
                user_email, user_name, password, platform_url, ), daemon=True)
            thread.start()

            user.set_password(password)
            user.save()

            cobrowse_agent = CobrowseAgent.objects.create(user=user,
                                                          mobile_number=user_mobile,
                                                          role=user_type,
                                                          support_level=support_level,
                                                          access_token=cobrowse_access_token)

            update_agents_supervisors_creation_count(active_agent, user_type)

            if cobrowse_access_token.enable_followup_leads_tab and user_type == "agent":
                assign_lead_value = True

                if assign_followup_lead == "no":
                    assign_lead_value = False

                cobrowse_agent.assign_followup_leads = assign_lead_value
                cobrowse_agent.save()

            if cobrowse_agent.role == "admin_ally":
                selected_supervisor_pk_list = ["-1"]
                add_selected_supervisor(
                    selected_supervisor_pk_list, active_agent, cobrowse_agent, CobrowseAgent)
            else:
                supervisor_obj.agents.add(cobrowse_agent)
                supervisor_obj.save()
            add_supported_language(
                cobrowse_agent, selected_language_pk_list, LanguageSupport)
            add_supported_product(
                cobrowse_agent, product_category_pk_list, ProductCategory)
            logger.info("Added user user_name: " + str(user_name),
                        extra={'AppName': 'EasyAssist'})

        if total_rows <= 1:
            response["status"] = 301
            response["message"] = "File is empty."
        elif wrong_data_list:
            response["status"] = 302
            response["message"] = "Some entries are incorrect. Please check the below given excel link for more information."
            file_path = create_excel_wrong_user_data(
                wrong_data_list)
            file_access_management_key = create_file_access_management_obj(
                CobrowsingFileAccessManagement, cobrowse_access_token, "/" + file_path)
            response["file_path"] = 'easy-assist/download-file/' + \
                file_access_management_key
            logger.info("Some wrong data found %s", str(
                wrong_data_list), extra={'AppName': 'EasyAssist'})
        else:
            response["status"] = 200
            response["message"] = "users created"
            logger.info("Successfully exited from add_users_from_excel_document", extra={
                        'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_users_from_excel_document %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return response


def check_cogno_meet_status(meeting_io):
    try:
        start_date = meeting_io.meeting_start_date
        current_date = datetime.datetime.now()
        meeting_start_time = meeting_io.meeting_start_time
        current_time = current_date.strftime("%H:%M:%S")
        if start_date > current_date.date():
            return 'waiting'
        elif start_date < current_date.date():
            meeting_io.is_expired = True
            meeting_io.save()
        elif str(meeting_start_time) > str(current_time):
            return 'waiting'
        else:
            meeting_end_time = meeting_io.meeting_end_time
            if str(meeting_end_time) < str(current_time):
                meeting_io.is_expired = True
                meeting_io.save()
        return True
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error check_cogno_meet_status %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return False


def is_meeting_expired(meeting_io):
    try:
        start_date = meeting_io.meeting_start_date
        current_date = datetime.datetime.now()

        meeting_start_time = meeting_io.meeting_start_time
        current_time = current_date.strftime("%H:%M:%S")

        if start_date > current_date.date():
            return False
        elif start_date < current_date.date():
            return True
        elif str(meeting_start_time) >= str(current_time):
            return False
        else:
            meeting_end_time = meeting_io.meeting_end_time
            if str(meeting_end_time) < str(current_time):
                return True
        return False
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error is_meeting_expired %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return True


def send_agent_customer_connected_notification(agent, cobrowse_io):
    try:
        if agent == None:
            return

        if cobrowse_io == None:
            return

        product_name = "Cogno Cobrowse"
        cobrowse_config_obj = get_developer_console_cobrowsing_settings()
        if cobrowse_config_obj:
            product_name = cobrowse_config_obj.cobrowsing_title_text

        socket_response = {}
        notification_message = "Hi, " + agent.user.username + \
            "! A customer has connected with you on " + product_name + "."
        notification_list = [{
            "notification_message": notification_message,
            "product_name": product_name
        }]
        socket_response["status"] = 200
        socket_response["message"] = "success"
        socket_response["notification_list"] = notification_list
        send_data_from_server_to_client(
            "notification", socket_response, agent.user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error send_agent_customer_connected_notification %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return False


def send_notification_to_agent(agent, NotificationManagement):
    try:
        notification_objs = NotificationManagement.objects.filter(
            show_notification=True, agent=agent)

        response = {}
        notification_list = []

        for notification_obj in notification_objs:
            # notification_list.append(
            #     notification_obj.notification_message)
            notification_list.append({
                "notification_message": notification_obj.notification_message,
                "product_name": notification_obj.product_name
            })
            notification_obj.delete()

        response["status"] = 200
        response["message"] = "success"
        response["notification_list"] = notification_list

        send_data_from_server_to_client(
            "notification", response, agent.user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error send_notification_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return False


def send_page_refresh_request_to_agent(agent):
    try:
        send_data_from_server_to_client("page_refresh", None, agent.user)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error send_notification_to_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return False


def create_app_cobrowsing_session(cobrowse_io, user_type, AppCobrowsingSessionManagement):
    try:
        total_sessions = AppCobrowsingSessionManagement.objects.filter(
            cobrowse_io=cobrowse_io, user_type=user_type).count()

        user_alias = str(user_type) + "-" + str(total_sessions + 1)

        return AppCobrowsingSessionManagement.objects.create(cobrowse_io=cobrowse_io,
                                                             user_type=user_type,
                                                             user_alias=user_alias)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_app_cobrowsing_session %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def file_download(file_key, CobrowsingFileAccessManagement, SupportDocument, is_download_required=True):
    file_access_management_obj = None
    try:
        file_access_management_obj = CobrowsingFileAccessManagement.objects.get(
            key=file_key)

        path_to_file = file_access_management_obj.file_path

        original_file_name = file_access_management_obj.original_file_name

        support_document_obj = SupportDocument.objects.filter(
            file_access_management_key=file_key)
        if support_document_obj.count() > 0:
            filename = support_document_obj[0].file_name
        else:
            filename = path_to_file.split("/")[-1]

        if original_file_name == None:
            original_file_name = filename

        path_to_file = settings.BASE_DIR + path_to_file
        mime_type, _ = mimetypes.guess_type(path_to_file)

        if os.path.exists(path_to_file):
            with open(path_to_file, 'rb') as fh:
                response = HttpResponse(
                    fh.read(), status=200, content_type=mime_type)

                if is_download_required:
                    response['Content-Disposition'] = 'attachment; filename="%s"' % smart_str(
                        str(original_file_name))
                return response
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error FileDownload %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return HttpResponse(status=404)


def add_support_history_data_into_workbook(workbook, cobrowse_agent, start_date, end_date, cobrowse_io_objs=None):
    try:
        from EasyAssistApp.models import CobrowseIO, CobrowseIOInvitedAgentsDetails, CobrowseIOTransferredAgentsLogs

        def create_new_sheet():
            global comment_col, subremarks_col, language_col, type_of_lead_col, supervisor_col, auto_assigned_agent_col, \
                agents_invited_col, invited_agents_connected_col, agents_transferred_col, transferred_agents_connected_col, transferred_agents_rejected_col, session_start_datetime_col, session_end_datetime_col, time_spent_col, \
                lead_status_col, title_col, session_id_col, session_started_by_col, session_ended_by_col, agent_remarks_col, session_initiated_by, \
                assigned_agents_count_col, request_datetime_col, wait_time_col, drop_title_col, self_assign_time_col, product_col, \
                nps_rating_col, customer_remarks_col, product_title_col
            nonlocal wb_index

            if wb_index > 1:
                sheet = workbook.add_sheet("Attended Leads - " + str(wb_index))
            else:
                sheet = workbook.add_sheet("Attended Leads")

            wb_index += 1

            sheet.write(0, 0, CUSTOMER_NAME)
            sheet.col(0).width = 256 * 20
            sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
            sheet.col(1).width = 256 * 25
            sheet.write(0, 2, "Primary Agent")
            sheet.col(2).width = 256 * 35

            prev_col = 2

            if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
                auto_assigned_agent_col = prev_col + 1
                sheet.write(0, auto_assigned_agent_col, "Auto-assigned Agent")
                sheet.col(auto_assigned_agent_col).width = 256 * 35
                prev_col = auto_assigned_agent_col

                assigned_agents_count_col = prev_col + 1
                sheet.write(0, assigned_agents_count_col,
                            "Assigned Agents Count")
                sheet.col(assigned_agents_count_col).width = 256 * 35
                prev_col = assigned_agents_count_col

            session_initiated_by = prev_col + 1
            sheet.write(0, session_initiated_by, "Session Initiated By")
            sheet.col(session_initiated_by).width = 256 * 35
            prev_col = session_initiated_by

            if access_token_obj is not None and access_token_obj.enable_invite_agent_in_cobrowsing:
                agents_invited_col = prev_col + 1
                sheet.write(0, agents_invited_col, "Agents Invited")
                sheet.col(agents_invited_col).width = 256 * 35
                prev_col = agents_invited_col

                invited_agents_connected_col = prev_col + 1
                sheet.write(0, invited_agents_connected_col,
                            "Invited Agent Connected")
                sheet.col(invited_agents_connected_col).width = 256 * 35
                prev_col = invited_agents_connected_col

                if access_token_obj.enable_session_transfer_in_cobrowsing:
                    agents_transferred_col = prev_col + 1
                    sheet.write(0, agents_transferred_col,
                                "Transferred Agents")
                    sheet.col(agents_transferred_col).width = 256 * 35
                    prev_col = agents_transferred_col

                    transferred_agents_connected_col = prev_col + 1
                    sheet.write(0, transferred_agents_connected_col,
                                "Transferred Agents Connected")
                    sheet.col(transferred_agents_connected_col).width = 256 * 35
                    prev_col = transferred_agents_connected_col

                    transferred_agents_rejected_col = prev_col + 1
                    sheet.write(0, transferred_agents_rejected_col,
                                "Transferred Agents Not Connected")
                    sheet.col(transferred_agents_rejected_col).width = 256 * 35
                    prev_col = transferred_agents_rejected_col

            request_datetime_col = prev_col + 1
            sheet.write(0, request_datetime_col,
                        "Co-Browsing Request Date & Time")
            sheet.col(request_datetime_col).width = 256 * 40
            prev_col = request_datetime_col

            if access_token_obj. enable_request_in_queue:
                self_assign_time_col = prev_col + 1
                sheet.write(0, self_assign_time_col, "Self Assign Time")
                sheet.col(self_assign_time_col).width = 256 * 40
                prev_col = self_assign_time_col

            wait_time_col = prev_col + 1
            sheet.write(0, wait_time_col, "Wait Time")
            sheet.col(wait_time_col).width = 256 * 25
            prev_col = wait_time_col

            session_start_datetime_col = prev_col + 1
            sheet.write(0, session_start_datetime_col,
                        "Cobrowsing Start Date & Time")
            sheet.col(session_start_datetime_col).width = 256 * 25
            prev_col = session_start_datetime_col

            session_end_datetime_col = prev_col + 1
            sheet.write(0, session_end_datetime_col, "Session End Date & Time")
            sheet.col(session_end_datetime_col).width = 256 * 25
            prev_col = session_end_datetime_col

            time_spent_col = prev_col + 1
            sheet.write(0, time_spent_col, "Time Spent")
            sheet.col(time_spent_col).width = 256 * 20
            prev_col = time_spent_col

            lead_status_col = prev_col + 1
            sheet.write(0, lead_status_col, "Lead Status")
            sheet.col(lead_status_col).width = 256 * 15
            prev_col = lead_status_col

            product_title_col = prev_col + 1
            sheet.write(0, product_title_col, "Product Title")
            sheet.col(product_title_col).width = 256 * 60
            prev_col = product_title_col

            title_col = prev_col + 1
            sheet.write(0, title_col, "Session Start")
            sheet.col(title_col).width = 256 * 60
            prev_col = title_col

            drop_title_col = prev_col + 1
            sheet.write(0, drop_title_col, "Session End")
            sheet.col(drop_title_col).width = 256 * 60
            prev_col = drop_title_col

            session_id_col = prev_col + 1
            sheet.write(0, session_id_col, "Session ID")
            sheet.col(session_id_col).width = 256 * 45
            prev_col = session_id_col

            session_started_by_col = prev_col + 1
            sheet.write(0, session_started_by_col, "Session Started By")
            sheet.col(session_started_by_col).width = 256 * 20
            prev_col = session_started_by_col

            session_ended_by_col = prev_col + 1
            sheet.write(0, session_ended_by_col, "Session Ended By")
            sheet.col(session_ended_by_col).width = 256 * 20
            prev_col = session_ended_by_col

            agent_remarks_col = prev_col + 1
            sheet.write(0, agent_remarks_col, AGENT_REMARKS)
            sheet.col(agent_remarks_col).width = 256 * 80
            prev_col = agent_remarks_col

            if access_token_obj is not None and access_token_obj.enable_predefined_remarks and access_token_obj.enable_predefined_subremarks:
                subremarks_col = prev_col + 1
                sheet.write(0, subremarks_col, "Agent Sub-remarks")
                sheet.col(subremarks_col).width = 256 * 80
                prev_col = subremarks_col

            if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
                comment_col = prev_col + 1
                sheet.write(0, comment_col, "Agent Comments")
                sheet.col(comment_col).width = 256 * 80
                prev_col = comment_col

            customer_remarks_col = prev_col + 1
            sheet.write(0, customer_remarks_col, "Customer Remarks")
            sheet.col(customer_remarks_col).width = 256 * 80
            prev_col = customer_remarks_col

            nps_rating_col = prev_col + 1
            sheet.write(0, nps_rating_col, "NPS")
            sheet.col(nps_rating_col).width = 256 * 20
            prev_col = nps_rating_col
            
            if access_token_obj is not None and access_token_obj.allow_language_support:
                language_col = prev_col + 1
                sheet.write(0, language_col, "Language")
                sheet.col(language_col).width = 256 * 25
                prev_col = language_col

            if access_token_obj is not None and access_token_obj.enable_tag_based_assignment_for_outbound or access_token_obj.choose_product_category:
                product_col = prev_col + 1
                sheet.write(0, product_col, "Category")
                sheet.col(product_col).width = 256 * 25
                prev_col = product_col

            type_of_lead_col = prev_col + 1
            sheet.write(0, type_of_lead_col, "Type of Lead")
            sheet.col(type_of_lead_col).width = 256 * 20
            prev_col = type_of_lead_col

            if cobrowse_agent.role == "admin_ally":
                supervisor_col = prev_col + 1
                sheet.write(0, supervisor_col, "Supervisors")
                sheet.col(supervisor_col).width = 256 * 20
                prev_col = supervisor_col

            return sheet

        agents = get_list_agents_under_admin(
            cobrowse_agent, is_active=None, is_account_active=None)
        access_token_obj = cobrowse_agent.get_access_token_obj()

        if not cobrowse_io_objs:
            cobrowse_io_objs = CobrowseIO.objects.filter(
                is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
                request_datetime__date__gte=start_date, request_datetime__date__lte=end_date, agent__in=agents).filter(
                ~Q(cobrowsing_start_datetime=None))

            if cobrowse_agent.role == "supervisor" and access_token_obj.enable_invite_agent_in_cobrowsing:
                invited_agent_details_objs = CobrowseIOInvitedAgentsDetails.objects.select_related(
                    'cobrowse_io').filter(support_agents_joined__in=agents)
                if invited_agent_details_objs:
                    invited_session_ids = []
                    for invited_agent_details_obj in invited_agent_details_objs.iterator():
                        invited_session_ids.append(
                            invited_agent_details_obj.cobrowse_io.session_id)
                    invited_sessions_cobrowse_objs = CobrowseIO.objects.filter(session_id__in=invited_session_ids, is_test=False, access_token=access_token_obj,
                                                                               request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
                        request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(~Q(cobrowsing_start_datetime=None))
                    cobrowse_io_objs = cobrowse_io_objs | invited_sessions_cobrowse_objs

                if access_token_obj.enable_session_transfer_in_cobrowsing:
                    transferred_agent_details_objs = CobrowseIOTransferredAgentsLogs.objects.filter(
                        transferred_agent__in=agents, cobrowse_request_type="transferred")
                    if transferred_agent_details_objs:
                        transferred_session_ids = []
                        for transferred_agent_details_obj in transferred_agent_details_objs.iterator():
                            transferred_session_ids.append(
                                transferred_agent_details_obj.cobrowse_io.session_id)
                        transferred_sessions_cobrowse_objs = CobrowseIO.objects.filter(session_id__in=transferred_session_ids, is_test=False, access_token=access_token_obj,
                                                                                       request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
                            request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(~Q(cobrowsing_start_datetime=None))
                        cobrowse_io_objs = cobrowse_io_objs | transferred_sessions_cobrowse_objs

            if not cobrowse_io_objs:
                return False

        wb_index = 1
        sheet = create_new_sheet()

        style = easyxf('font: colour blue;')

        index = 1
        for cobrowse_io_obj in cobrowse_io_objs.iterator():

            if index > 50000:
                index = 1
                sheet = create_new_sheet()

            if not cobrowse_io_obj.is_lead:
                sheet.write(index, 0, cobrowse_io_obj.full_name)
                sheet.write(index, 1, get_masked_data_if_hashed(
                    cobrowse_io_obj.mobile_number))
            else:
                sheet.write(
                    index, 0, cobrowse_io_obj.get_sync_data_value("name"))
                sheet.write(
                    index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

            if cobrowse_io_obj.main_primary_agent:
                primary_agent_name = cobrowse_io_obj.main_primary_agent.user.username
            else:
                primary_agent_name = cobrowse_io_obj.agent.user.username

            sheet.write(index, 2, primary_agent_name)

            if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
                sheet.write(index, auto_assigned_agent_col,
                            cobrowse_io_obj.get_auto_assigned_agents())
                assigned_count_value = "-"
                lead_transfer_count = len(
                    cobrowse_io_obj.get_unattended_lead_transfer_audit_trail())
                if lead_transfer_count:
                    assigned_count_value = lead_transfer_count
                sheet.write(index, assigned_agents_count_col,
                            assigned_count_value)

            lead_initiated_by = "-"
            if cobrowse_io_obj.lead_initiated_by == 'greeting_bubble':
                lead_initiated_by = "Greeting Bubble"
            elif cobrowse_io_obj.lead_initiated_by == 'floating_button':
                lead_initiated_by = "Floating Button"
            elif cobrowse_io_obj.lead_initiated_by == 'icon':
                lead_initiated_by = "Icon"
            elif cobrowse_io_obj.lead_initiated_by == 'exit_intent':
                lead_initiated_by = "Exit Intent"
            elif cobrowse_io_obj.lead_initiated_by == 'inactivity_popup':
                lead_initiated_by = "Inactivity Pop-up"
            sheet.write(index, session_initiated_by, lead_initiated_by)

            if access_token_obj is not None and access_token_obj.enable_invite_agent_in_cobrowsing:

                agents_invited_str = "-"
                agents_invited_connected_str = "-"
                if cobrowse_io_obj.get_cobrowse_support_agent_details():
                    agents_invited = cobrowse_io_obj.get_cobrowse_support_agent_details(
                    ).support_agents_invited.all()
                    if agents_invited:
                        agents_invited_str = ""
                        for agent in agents_invited.iterator():
                            agents_invited_str = agents_invited_str + agent.user.username + ", "
                        agents_invited_str = agents_invited_str[:-2]

                    agents_invited_connected = cobrowse_io_obj.get_cobrowse_support_agent_details(
                    ).support_agents_joined.all()
                    if agents_invited_connected:
                        agents_invited_connected_str = ""
                        for agent in agents_invited_connected.iterator():
                            agents_invited_connected_str = agents_invited_connected_str + \
                                agent.user.username + ", "
                        agents_invited_connected_str = agents_invited_connected_str[
                            :-2]

                sheet.write(index, agents_invited_col, agents_invited_str)
                sheet.write(index, invited_agents_connected_col,
                            agents_invited_connected_str)

                if access_token_obj.enable_session_transfer_in_cobrowsing:
                    agents_transferred_str = cobrowse_io_obj.get_cobrowse_transferred_agent_list()
                    agents_transferred_connected_str = cobrowse_io_obj.get_cobrowse_transferred_agent_connected_list()
                    agents_transferred_rejected_str = cobrowse_io_obj.get_cobrowse_transferred_agent_rejected_list()

                    sheet.write(index, agents_transferred_col,
                                agents_transferred_str)
                    sheet.write(index, transferred_agents_connected_col,
                                agents_transferred_connected_str)
                    sheet.write(index, transferred_agents_rejected_col,
                                agents_transferred_rejected_str)

            sheet.write(index, request_datetime_col, str(cobrowse_io_obj.request_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))

            if cobrowse_io_obj.access_token.enable_request_in_queue:
                if cobrowse_io_obj.get_self_assign_time():
                    sheet.write(index, self_assign_time_col, str(
                        cobrowse_io_obj.get_self_assign_time()))
                else:
                    sheet.write(index, self_assign_time_col, "-")

            wait_time = cobrowse_io_obj.customer_wait_time_seconds()
            minutes, seconds = divmod(wait_time, 60)
            hours, minutes = divmod(minutes, 60)
            wait_time = "{:02d}:{:02d}:{:02d}".format(hours, minutes, seconds)

            sheet.write(index, wait_time_col, wait_time)
            sheet.write(index, session_start_datetime_col, str(cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            sheet.write(index, session_end_datetime_col, str(cobrowse_io_obj.last_update_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            sheet.write(index, time_spent_col, str(
                cobrowse_io_obj.total_time_spent()))

            if cobrowse_io_obj.is_archived == True and cobrowse_io_obj.is_helpful == True:
                sheet.write(index, lead_status_col, "Converted")
            else:
                sheet.write(index, lead_status_col, "Not Converted")

            if cobrowse_io_obj.product_name():
                sheet.write(index, product_title_col, cobrowse_io_obj.product_name())
            else:
                sheet.write(index, product_title_col, "-")

            if cobrowse_io_obj.product_url():
                sheet.write(index, title_col, cobrowse_io_obj.product_url())
            else:
                sheet.write(index, title_col, "-")

            if not cobrowse_io_obj.is_reverse_cobrowsing:
                sheet.write(index, drop_title_col, cobrowse_io_obj.drop_off_url())
            else:
                sheet.write(index, drop_title_col, '-', style=style)
            sheet.write(index, session_id_col, str(cobrowse_io_obj.session_id))

            sheet.write(index, session_started_by_col, cobrowse_io_obj.get_session_started_by())

            if cobrowse_io_obj.session_archived_cause:
                if cobrowse_io_obj.session_archived_cause in ["AGENT_ENDED", "AGENT_INACTIVITY"]:
                    sheet.write(index, session_ended_by_col, "Agent")
                elif cobrowse_io_obj.session_archived_cause in ["CLIENT_ENDED", "CLIENT_INACTIVITY"]:
                    sheet.write(index, session_ended_by_col, "Customer")
                else:
                    sheet.write(index, session_ended_by_col, cobrowse_io_obj.session_archived_cause[0].upper(
                    ) + cobrowse_io_obj.session_archived_cause[1:].lower())
            else:
                sheet.write(index, session_ended_by_col, "-")
            if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                sheet.write(index, agent_remarks_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                ).first().agent_comments)
            else:
                sheet.write(index, agent_remarks_col, NOT_PROVIDED)

            if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
                if access_token_obj.enable_predefined_subremarks:
                    if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                        if len(cobrowse_io_obj.get_cobrowsing_session_closing_comments().first().agent_subcomments) > 0:
                            sheet.write(index, subremarks_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                            ).first().agent_subcomments)
                        else:
                            sheet.write(index, subremarks_col, "-")
                    else:
                        sheet.write(index, subremarks_col, NOT_PROVIDED)

                if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                    if len(cobrowse_io_obj.get_cobrowsing_session_closing_comments().first().comment_desc) > 0:
                        sheet.write(index, comment_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                        ).first().comment_desc)
                    else:
                        sheet.write(index, comment_col, "-")
                else:
                    sheet.write(index, comment_col, NOT_PROVIDED)

            customer_remarks = "-"
            if cobrowse_io_obj.client_comments:
                customer_remarks = cobrowse_io_obj.client_comments
            sheet.write(index, customer_remarks_col, customer_remarks)

            nps_rating = "Not provided"
            if cobrowse_io_obj.agent_rating != None:
                nps_rating = str(cobrowse_io_obj.agent_rating)
            sheet.write(index, nps_rating_col, nps_rating)
            
            if access_token_obj is not None and access_token_obj.allow_language_support:
                if cobrowse_io_obj.supported_language is not None:
                    sheet.write(
                        index, language_col, cobrowse_io_obj.supported_language.title)
                else:
                    sheet.write(index, language_col, "-")

            if access_token_obj is not None and access_token_obj.enable_tag_based_assignment_for_outbound or access_token_obj.choose_product_category:
                if cobrowse_io_obj.product_category is not None:
                    sheet.write(
                        index, product_col, cobrowse_io_obj.product_category.title)
                else:
                    sheet.write(index, product_col, "-")

            lead_type = cobrowse_io_obj.get_lead_type()
            sheet.write(index, type_of_lead_col, lead_type)

            if cobrowse_agent.role == "admin_ally":
                sheet.write(index, supervisor_col,
                            cobrowse_io_obj.agent.get_supervisors())

            index += 1
        
        return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("In add_support_history_data_into_workbook! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
      
        return False
        

def add_live_chat_history_data_into_workbook(workbook, cobrowse_agent, start_date, end_date, CobrowseIO, CobrowseIOInvitedAgentsDetails, CobrowseIOTransferredAgentsLogs, CobrowseChatHistory, cobrowse_io_objs=None):

    def get_absolute_url(url):
        if url and (not url.startswith('http')):
            if url[0] == '/':
                url = settings.EASYCHAT_HOST_URL + url
            else:
                url = settings.EASYCHAT_HOST_URL + '/' + url
        return url
    
    def create_new_sheet():

        nonlocal wb_index

        if wb_index > 1:
            sheet = workbook.add_sheet("Chat History - " + str(wb_index))
        else:
            sheet = workbook.add_sheet("Chat History")

        wb_index += 1
        
        heading_style = easyxf(HEADING_FONT_STYLE)
        
        sheet.write(0, 0, "Session ID", style=heading_style)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, "Chat initiated by", style=heading_style)
        sheet.col(1).width = 256 * 20
        sheet.write(0, 2, "Request Time", style=heading_style)
        sheet.col(2).width = 256 * 20
        sheet.write(0, 3, "Start Time", style=heading_style)
        sheet.col(3).width = 256 * 20
        sheet.write(0, 4, "Customer Name", style=heading_style)
        sheet.col(4).width = 256 * 20
        sheet.write(0, 5, "Number", style=heading_style)
        sheet.col(5).width = 256 * 20
        sheet.write(0, 6, "Sender", style=heading_style)
        sheet.col(6).width = 256 * 20
        sheet.write(0, 7, "Message", style=heading_style)
        sheet.col(7).width = 256 * 20
        sheet.write(0, 8, "Message Timestamp", style=heading_style)
        sheet.col(8).width = 256 * 20
        sheet.write(0, 9, "Attachment", style=heading_style)
        sheet.col(9).width = 256 * 20
        
        return sheet

    wb_index = 1
    sheet = create_new_sheet()
    
    sheet.set_panes_frozen('1')
    sheet.set_horz_split_pos(1)  
    sheet.set_vert_split_pos(1)
    
    style = easyxf('font: colour blue; align: vert top;')
    alignment_style = easyxf(ALIGNMENT_STYLE)

    index = 1
    
    agents = get_list_agents_under_admin(
        cobrowse_agent, is_active=None, is_account_active=None)
    access_token_obj = cobrowse_agent.get_access_token_obj()

    if not cobrowse_io_objs:
        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
            request_datetime__date__gte=start_date, request_datetime__date__lte=end_date, agent__in=agents).filter(
            ~Q(cobrowsing_start_datetime=None))

        if cobrowse_agent.role == "supervisor" and access_token_obj.enable_invite_agent_in_cobrowsing:
            invited_agent_details_objs = CobrowseIOInvitedAgentsDetails.objects.select_related(
                'cobrowse_io').filter(support_agents_joined__in=agents)
            if invited_agent_details_objs:
                invited_session_ids = []
                for invited_agent_details_obj in invited_agent_details_objs.iterator():
                    invited_session_ids.append(
                        invited_agent_details_obj.cobrowse_io.session_id)
                invited_sessions_cobrowse_objs = CobrowseIO.objects.filter(session_id__in=invited_session_ids, is_test=False, access_token=access_token_obj,
                                                                           request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
                    request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(~Q(cobrowsing_start_datetime=None))
                cobrowse_io_objs = cobrowse_io_objs | invited_sessions_cobrowse_objs

            if access_token_obj.enable_session_transfer_in_cobrowsing:
                transferred_agent_details_objs = CobrowseIOTransferredAgentsLogs.objects.filter(
                    transferred_agent__in=agents, cobrowse_request_type="transferred")
                if transferred_agent_details_objs:
                    transferred_session_ids = []
                    for transferred_agent_details_obj in transferred_agent_details_objs.iterator():
                        transferred_session_ids.append(
                            transferred_agent_details_obj.cobrowse_io.session_id)
                    transferred_sessions_cobrowse_objs = CobrowseIO.objects.filter(session_id__in=transferred_session_ids, is_test=False, access_token=access_token_obj,
                                                                                   request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
                        request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(~Q(cobrowsing_start_datetime=None))
                    cobrowse_io_objs = cobrowse_io_objs | transferred_sessions_cobrowse_objs
        
        if not cobrowse_io_objs:
            return False
    
    cobrowse_io_objs = cobrowse_io_objs.order_by("-request_datetime")
    
    for cobrowse_io_obj in cobrowse_io_objs.iterator():

        if index > 50000:
            index = 1
            sheet = create_new_sheet()
         
        cobrowsing_chat_history_objs = CobrowseChatHistory.objects.filter(cobrowse_io=cobrowse_io_obj, chat_type="chat_message").order_by('datetime')
        
        if cobrowsing_chat_history_objs:
            sheet.write(index, 0, str(cobrowse_io_obj.session_id), style=style)            
            
            if cobrowsing_chat_history_objs[0].sender:
                sheet.write(index, 1, "Agent", style=alignment_style)
            else:
                sheet.write(index, 1, "Customer", style=alignment_style)

            sheet.write(index, 2, str(cobrowse_io_obj.request_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_7)), style=alignment_style)
            sheet.write(index, 3, str(cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_7)), style=alignment_style)
                
            if not cobrowse_io_obj.is_lead:
                sheet.write(index, 4, cobrowse_io_obj.full_name, style=alignment_style)
                sheet.write(index, 5, get_masked_data_if_hashed(
                    cobrowse_io_obj.mobile_number), style=alignment_style)
            else:
                sheet.write(index, 4, cobrowse_io_obj.get_sync_data_value("name"), style=alignment_style)
                sheet.write(index, 5, cobrowse_io_obj.get_sync_data_value("mobile"), style=alignment_style)
                
            agent_style = easyxf(AGENT_STYLE)
            invited_agent_style = easyxf(INVITED_AGENT_STYLE)
            customer_style = easyxf(CUSTOMER_STYLE)
            
            primary_agent_name = cobrowse_io_obj.agent.name()
            
            for cobrowsing_chat_history_obj in cobrowsing_chat_history_objs.iterator():
            
                sender_name = "Customer"
                sender_role = "Customer"
                if cobrowsing_chat_history_obj.sender:
                    sender_name = cobrowsing_chat_history_obj.sender.name()
                    sender_role = "AdminAgent"
                    if sender_name != primary_agent_name:
                        sender_role = "InvitedAgent"
                else:
                    if cobrowse_io_obj.full_name:
                        sender_name = cobrowse_io_obj.full_name
                    
                sheet.write(index, 6, sender_name, style=alignment_style)
                if sender_role == "AdminAgent":
                    sheet.write(index, 7, get_masked_data_if_hashed(cobrowsing_chat_history_obj.message), style=agent_style)
                elif sender_role == "Customer":
                    sheet.write(index, 7, get_masked_data_if_hashed(cobrowsing_chat_history_obj.message), style=customer_style)
                else:
                    sheet.write(index, 7, get_masked_data_if_hashed(cobrowsing_chat_history_obj.message), style=invited_agent_style) 
                
                sheet.write(index, 8, str(cobrowsing_chat_history_obj.datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_7)), style=alignment_style)
                
                attachment_url = get_absolute_url(cobrowsing_chat_history_obj.attachment)
                if attachment_url:
                    sheet.write(index, 9, Formula('HYPERLINK("{}")'.format(attachment_url)), style=style)
                else:
                    sheet.write(index, 9, "-", style=alignment_style)
                index += 1
                
            index += 1

    return True  
        

def get_custom_support_history(requested_data, cobrowse_agent, CobrowseIO, is_public=False, cobrowse_io_objs=None):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/support-history/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/support-history/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "support_history_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/support-history/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:

        logger.info("Inside get_custom_support_history",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        is_data_present = add_support_history_data_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date, cobrowse_io_objs)
        if is_data_present:
            add_agent_session_count_attended_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

            support_history_wb.save(absolute_file_path)
            if get_save_in_s3_bucket_status():
                file_name = get_file_name(".xls")
                key = s3_bucket_upload_file_by_file_path(
                    relative_file_path, file_name)
                s3_file_path = s3_bucket_download_file(
                    key, 'EasyAssistApp/support-history/', '.xls')
                return s3_file_path.split("EasyChat/", 1)[1]

            logger.info("Returning file path " + str(relative_file_path),
                        extra={'AppName': 'EasyAssist'})
            return relative_file_path
        return NO_DATA
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_support_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def get_custom_live_chat_history(requested_data, cobrowse_agent, CobrowseIO, CobrowseIOInvitedAgentsDetails, CobrowseIOTransferredAgentsLogs, CobrowseChatHistory, is_public=False, cobrowse_io_objs=None):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/live-chat-history/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/live-chat-history/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "live_chat_history_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/live-chat-history/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:

        logger.info("Inside get_custom_live_chat_history",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        is_data_present = add_live_chat_history_data_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date, CobrowseIO, 
            CobrowseIOInvitedAgentsDetails, CobrowseIOTransferredAgentsLogs, CobrowseChatHistory, cobrowse_io_objs)

        if is_data_present:
            support_history_wb.save(absolute_file_path)
            if get_save_in_s3_bucket_status():
                file_name = get_file_name(".xls")
                key = s3_bucket_upload_file_by_file_path(
                    relative_file_path, file_name)
                s3_file_path = s3_bucket_download_file(
                    key, 'EasyAssistApp/live-chat-history/', '.xls')
                return s3_file_path.split("EasyChat/", 1)[1]

            logger.info("Returning file path " + str(relative_file_path),
                        extra={'AppName': 'EasyAssist'})
            return relative_file_path

        return NO_DATA

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_live_chat_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def add_manually_converted_leads_data_into_workbook(workbook, cobrowse_agent, start_date, end_date, cobrowse_io_objs):

    from EasyAssistApp.models import CobrowseIO, CobrowseAgentComment

    def create_new_sheet():
        global auto_assigned_agent_col, language_col, product_category_col, prev_agent_remarks_col, prev_subremarks_col, \
            prev_comment_col, agent_remarks_col, subremarks_col, comment_col, lead_col, lead_type_col, lead_status_col, lead_converted_by_col
        nonlocal wb_index

        if wb_index > 1:
            sheet = workbook.add_sheet(
                "Manually Converted Leads - " + str(wb_index))
        else:
            sheet = workbook.add_sheet("Manually Converted Leads")

        wb_index += 1

        sheet.write(0, 0, CUSTOMER_NAME)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, "Date & Time")
        sheet.col(2).width = 256 * 25
        sheet.write(0, 3, "Session Initiated By")
        sheet.col(3).width = 256 * 25
        sheet.write(0, 4, "Product Title")
        sheet.col(4).width = 256 * 25
        sheet.write(0, 5, "Session Start")
        sheet.col(5).width = 256 * 60
        sheet.write(0, 6, "Agent Email ID")
        sheet.col(6).width = 256 * 35

        prev_col = 6

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            auto_assigned_agent_col = prev_col + 1
            sheet.write(0, auto_assigned_agent_col, "Auto-assigned Agent")
            sheet.col(auto_assigned_agent_col).width = 256 * 35
            prev_col = auto_assigned_agent_col

        if access_token_obj is not None and access_token_obj.allow_language_support:
            language_col = prev_col + 1
            sheet.write(0, language_col, "Language")
            sheet.col(language_col).width = 256 * 25
            prev_col = language_col

        if access_token_obj is not None and access_token_obj.choose_product_category:
            product_category_col = prev_col + 1
            sheet.write(0, product_category_col, "Category")
            sheet.col(product_category_col).width = 256 * 25
            prev_col = product_category_col

        prev_agent_remarks_col = prev_col + 1
        sheet.write(0, prev_agent_remarks_col, "Agent Previous Remarks")
        sheet.col(prev_agent_remarks_col).width = 256 * 25
        prev_col = prev_agent_remarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks and access_token_obj.enable_predefined_subremarks:
            prev_subremarks_col = prev_col + 1
            sheet.write(0, prev_subremarks_col, "Agent Previous Sub-remarks")
            sheet.col(prev_subremarks_col).width = 256 * 25
            prev_col = prev_subremarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            prev_comment_col = prev_col + 1
            sheet.write(0, prev_comment_col, "Agent Previous Comments")
            sheet.col(prev_comment_col).width = 256 * 25
            prev_col = prev_comment_col

        agent_remarks_col = prev_col + 1
        sheet.write(0, agent_remarks_col, "Agent Remarks")
        sheet.col(agent_remarks_col).width = 256 * 25
        prev_col = agent_remarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks and access_token_obj.enable_predefined_subremarks:
            subremarks_col = prev_col + 1
            sheet.write(0, subremarks_col, "Agent Sub-remarks")
            sheet.col(subremarks_col).width = 256 * 25
            prev_col = subremarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            comment_col = prev_col + 1
            sheet.write(0, comment_col, "Agent Comments")
            sheet.col(comment_col).width = 256 * 25
            prev_col = comment_col

        lead_col = prev_col + 1
        sheet.write(0, lead_col, "Lead")
        sheet.col(lead_col).width = 256 * 25
        prev_col = lead_col

        lead_type_col = prev_col + 1
        sheet.write(0, lead_type_col, "Lead Type")
        sheet.col(lead_type_col).width = 256 * 25
        prev_col = lead_type_col

        lead_status_col = prev_col + 1
        sheet.write(0, lead_status_col, "Lead Status")
        sheet.col(lead_status_col).width = 256 * 25
        prev_col = lead_status_col

        lead_converted_by_col = prev_col + 1
        sheet.write(0, lead_converted_by_col, "Lead Manually Converted")
        sheet.col(lead_converted_by_col).width = 256 * 35
        prev_col = lead_converted_by_col

        return sheet

    access_token_obj = cobrowse_agent.get_access_token_obj()

    if not cobrowse_io_objs:

        if cobrowse_agent.role in ["admin", "admin_ally"]:
            agent_objs = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)
        elif cobrowse_agent.role == "supervisor":
            agent_objs = list(cobrowse_agent.agents.all())

        session_archive_cause = ["UNATTENDED"]
        if access_token_obj.enable_followup_leads_tab:
            session_archive_cause = ["UNASSIGNED", "FOLLOWUP", "UNATTENDED"]

        cobrowse_io_objs = CobrowseIO.objects.filter(
            access_token=access_token_obj, is_test=False, request_datetime__date__gte=access_token_obj.go_live_date, agent__in=agent_objs,
            is_archived=True, cobrowsing_start_datetime=None, session_archived_cause__in=session_archive_cause, is_lead_manually_converted=True).filter(
            request_datetime__date__gte=start_date, request_datetime__date__lte=end_date)

        if not cobrowse_io_objs:
            return False

    wb_index = 1
    sheet = create_new_sheet()
    
    index = 1
    for cobrowse_io_obj in cobrowse_io_objs.iterator():

        if index > 50000:
            index = 1
            sheet = create_new_sheet()

        if not cobrowse_io_obj.is_lead:
            sheet.write(index, 0, cobrowse_io_obj.full_name)
            sheet.write(index, 1, get_masked_data_if_hashed(
                cobrowse_io_obj.mobile_number))
        else:
            sheet.write(
                index, 0, cobrowse_io_obj.get_sync_data_value("name"))
            sheet.write(
                index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

        sheet.write(index, 2, str(cobrowse_io_obj.request_datetime.astimezone(
            pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))

        lead_initiated_by = "-"
        if cobrowse_io_obj.lead_initiated_by == 'greeting_bubble':
            lead_initiated_by = "Greeting Bubble"
        elif cobrowse_io_obj.lead_initiated_by == 'floating_button':
            lead_initiated_by = "Floating Button"
        elif cobrowse_io_obj.lead_initiated_by == 'icon':
            lead_initiated_by = "Icon"
        elif cobrowse_io_obj.lead_initiated_by == 'inactivity_popup':
            lead_initiated_by = "Inactivity Pop-up"
        elif cobrowse_io_obj.lead_initiated_by == 'exit_intent':
            lead_initiated_by = "Exit Intent"

        sheet.write(index, 3, lead_initiated_by)

        if cobrowse_io_obj.product_name():
            sheet.write(index, 4, cobrowse_io_obj.product_name())
        else:
            sheet.write(index, 4, "-")

        if cobrowse_io_obj.product_url():
            sheet.write(index, 5, cobrowse_io_obj.product_url())
        else:
            sheet.write(index, 5, "-")

        if cobrowse_io_obj.main_primary_agent:
            primary_agent_name = cobrowse_io_obj.main_primary_agent.user.username
        elif cobrowse_io_obj.agent:
            primary_agent_name = cobrowse_io_obj.agent.user.username
        else:
            primary_agent_name = "No agent"

        sheet.write(index, 6, primary_agent_name)

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            sheet.write(index, auto_assigned_agent_col,
                        cobrowse_io_obj.get_auto_assigned_agents())

        if access_token_obj is not None and access_token_obj.allow_language_support:
            if cobrowse_io_obj.supported_language is not None:
                sheet.write(
                    index, language_col, cobrowse_io_obj.supported_language.title)
            else:
                sheet.write(index, language_col, "-")

        if access_token_obj is not None and access_token_obj.choose_product_category:
            if cobrowse_io_obj.product_category is not None:
                sheet.write(
                    index, product_category_col, cobrowse_io_obj.product_category.title)
            else:
                sheet.write(index, product_category_col, "-")

        agent_prev_remarks = "-"
        agent_prev_subremarks = "-"
        agent_prev_comment_desc = "-"
        agent_remarks = "-"
        agent_subremarks = "-"
        agent_comment_desc = "-"
        agent_comments_objs = CobrowseAgentComment.objects.filter(
            cobrowse_io=cobrowse_io_obj).order_by("datetime")
        if agent_comments_objs:
            if cobrowse_io_obj.session_archived_cause != "UNATTENDED":
                agent_comment = agent_comments_objs.first()
                if agent_comment.agent_comments:
                    agent_remarks = agent_comment.agent_comments
                if agent_comment.agent_subcomments:
                    agent_subremarks = agent_comment.agent_subcomments
                if agent_comment.comment_desc:
                    agent_comment_desc = agent_comment.comment_desc
            else:
                if agent_comments_objs.count() > 1:
                    agent_comments = agent_comments_objs[:2]

                    if agent_comments[0].agent_comments:
                        agent_prev_remarks = agent_comments[0].agent_comments
                    if agent_comments[0].agent_subcomments:
                        agent_prev_subremarks = agent_comments[
                            0].agent_subcomments
                    if agent_comments[0].comment_desc:
                        agent_prev_comment_desc = agent_comments[
                            0].comment_desc

                    if agent_comments[1].agent_comments:
                        agent_remarks = agent_comments[1].agent_comments
                    if agent_comments[1].agent_subcomments:
                        agent_subremarks = agent_comments[1].agent_subcomments
                    if agent_comments[1].comment_desc:
                        agent_comment_desc = agent_comments[1].comment_desc

                elif agent_comments_objs.count() == 1:
                    agent_comments = agent_comments_objs[0]

                    if agent_comments.agent_comments:
                        agent_remarks = agent_comments.agent_comments
                    if agent_comments.agent_subcomments:
                        agent_subremarks = agent_comments.agent_subcomments
                    if agent_comments.comment_desc:
                        agent_comment_desc = agent_comments.comment_desc

        sheet.write(index, prev_agent_remarks_col, agent_prev_remarks)
        sheet.write(index, agent_remarks_col, agent_remarks)

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks and access_token_obj.enable_predefined_subremarks:
            sheet.write(index, prev_subremarks_col, agent_prev_subremarks)
            sheet.write(index, subremarks_col, agent_subremarks)

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            sheet.write(index, prev_comment_col, agent_prev_comment_desc)
            sheet.write(index, comment_col, agent_comment_desc)

        if cobrowse_io_obj.session_archived_cause == "UNATTENDED":
            sheet.write(index, lead_col, "Unattended")
        else:
            sheet.write(index, lead_col, "Follow-up")

        lead_type = cobrowse_io_obj.get_lead_type()
        sheet.write(index, lead_type_col, lead_type)

        if cobrowse_io_obj.is_helpful:
            sheet.write(index, lead_status_col, "Converted")
        else:
            sheet.write(index, lead_status_col, "Not Converted")

        sheet.write(index, lead_converted_by_col,
                    cobrowse_io_obj.agent_manually_converted_lead.user.username)

        index += 1

    return True


def get_custom_manually_converted_leads_history(requested_data, cobrowse_agent, CobrowseIO, is_public=False, cobrowse_io_objs=None):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/manually-converted-leads/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/manually-converted-leads/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "manually_converted_leads_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = get_folder_path()
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:

        logger.info("Inside get_custom_manually_converted_leads_history",
                    extra={'AppName': 'EasyAssist'})

        if cobrowse_agent.role == "agent":
            return "None"

        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        manually_converted_leads_wb = Workbook()

        is_data_present = add_manually_converted_leads_data_into_workbook(
            manually_converted_leads_wb, cobrowse_agent, start_date, end_date, cobrowse_io_objs)
        
        if is_data_present:
            manually_converted_leads_wb.save(absolute_file_path)
            if get_save_in_s3_bucket_status():
                file_name = get_file_name(".xls")
                key = s3_bucket_upload_file_by_file_path(
                    relative_file_path, file_name)
                s3_file_path = s3_bucket_download_file(
                    key, 'EasyAssistApp/manually-converted-leads/', '.xls')
                return s3_file_path.split("EasyChat/", 1)[1]

            logger.info("Returning file path " + str(relative_file_path),
                        extra={'AppName': 'EasyAssist'})
            return relative_file_path
        
        return NO_DATA
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_manually_converted_leads_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


'''
    same filter logic as get_custom_support_history
'''


def add_agent_session_count_attended_into_workbook(workbook, cobrowse_agent, start_date, end_date):

    from EasyAssistApp.models import CobrowseIO

    agents = get_list_agents_under_admin(
        cobrowse_agent, is_active=None, is_account_active=None)
    access_token_obj = cobrowse_agent.get_access_token_obj()

    cobrowse_io_objs = CobrowseIO.objects.filter(
        is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
        request_datetime__date__gte=start_date, request_datetime__date__lte=end_date, agent__in=agents).filter(
        ~Q(cobrowsing_start_datetime=None))

    if not cobrowse_io_objs:
        return False

    sheet = workbook.add_sheet("Agent Wise Bifurcation")

    sheet.write(0, 0, "Agent Email ID")
    sheet.col(0).width = 256 * 20
    sheet.write(0, 1, "Attended Leads")
    sheet.col(1).width = 256 * 20

    index = 1
    for agent in agents:
        sheet.write(index, 0, agent.user.username)
        sheet.write(index, 1, cobrowse_io_objs.filter(agent=agent).count())
        index += 1

    return True


def get_agent_session_count_attended(requested_data, cobrowse_agent, CobrowseIO):

    def get_folder_path():
        folder_path = "EasyAssistApp/support-history/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "agent_wise_bifurcation_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        relative_file_path = SECURED_FILES_PATH + \
            get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        logger.info("Inside get_agent_session_count_attended",
                    extra={'AppName': 'EasyAssist'})

        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        absolute_file_path = get_absolute_file_path()

        support_history_wb = Workbook()

        add_agent_session_count_attended_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date)

        support_history_wb.save(absolute_file_path)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_session_count_attended! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_custom_meeting_support_history(requested_data, cobrowse_agent, CobrowseVideoConferencing, CobrowseVideoAuditTrail, CobrowsingFileAccessManagement, audit_trail_objs=None):

    def get_folder_path():
        folder_path = "EasyAssistApp/MeetingSupportHistory/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "meeting_support_history_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        relative_file_path = SECURED_FILES_PATH + \
            get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    def create_new_sheet():
        global supervisor_col, agents_invited_col, invited_agents_connected_col, meeting_start_datetime_col, \
            meeting_end_datetime_col, time_spent_col, meeting_status_col, meeting_id_col, meeting_notes_col, \
            customer_location_col, attachment_col, auto_assigned_agent_col, session_initiated_by, meeting_initiated_by_col, \
            assigned_agents_count_col
        nonlocal wb_index

        sheet = support_history_wb.add_sheet(
            "Meeting Support History - " + str(wb_index))
        wb_index += 1

        sheet.write(0, 0, CUSTOMER_NAME)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, "Agent Email ID")
        sheet.col(2).width = 256 * 35

        prev_col = 2

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            auto_assigned_agent_col = prev_col + 1
            sheet.write(0, auto_assigned_agent_col, "Auto-assigned Agent")
            sheet.col(auto_assigned_agent_col).width = 256 * 35
            prev_col = auto_assigned_agent_col

            assigned_agents_count_col = prev_col + 1
            sheet.write(0, assigned_agents_count_col, "Assigned Agents Count")
            sheet.col(assigned_agents_count_col).width = 256 * 35
            prev_col = assigned_agents_count_col

        session_initiated_by = prev_col + 1
        sheet.write(0, session_initiated_by, "Session Initiated By")
        sheet.col(session_initiated_by).width = 256 * 35
        prev_col = session_initiated_by

        if access_token_obj is not None and (access_token_obj.enable_invite_agent_in_cobrowsing or access_token_obj.enable_invite_agent_in_meeting):
            agents_invited_col = prev_col + 1
            sheet.write(0, agents_invited_col, "Support Agents")
            sheet.col(agents_invited_col).width = 256 * 35
            prev_col = agents_invited_col

            invited_agents_connected_col = prev_col + 1
            sheet.write(0, invited_agents_connected_col,
                        "Support Agents Connected")
            sheet.col(invited_agents_connected_col).width = 256 * 35
            prev_col = invited_agents_connected_col

        meeting_initiated_by_col = prev_col + 1
        sheet.write(0, meeting_initiated_by_col, "Meeting Initiated By")
        sheet.col(meeting_initiated_by_col).width = 256 * 25
        prev_col = meeting_initiated_by_col

        meeting_start_datetime_col = prev_col + 1
        sheet.write(0, meeting_start_datetime_col, "Meeting Start Date & Time")
        sheet.col(meeting_start_datetime_col).width = 256 * 25
        prev_col = meeting_start_datetime_col

        meeting_end_datetime_col = prev_col + 1
        sheet.write(0, meeting_end_datetime_col, "Meeting End Date & Time")
        sheet.col(meeting_end_datetime_col).width = 256 * 25
        prev_col = meeting_end_datetime_col

        time_spent_col = prev_col + 1
        sheet.write(0, time_spent_col, "Time Spent")
        sheet.col(time_spent_col).width = 256 * 15
        prev_col = time_spent_col

        meeting_status_col = prev_col + 1
        sheet.write(0, meeting_status_col, "Meeting Status")
        sheet.col(meeting_status_col).width = 256 * 15
        prev_col = meeting_status_col

        meeting_id_col = prev_col + 1
        sheet.write(0, meeting_id_col, "Meeting ID")
        sheet.col(meeting_id_col).width = 256 * 45
        prev_col = meeting_id_col

        meeting_notes_col = prev_col + 1
        sheet.write(0, meeting_notes_col, "Meeting Notes")
        sheet.col(meeting_notes_col).width = 256 * 45
        prev_col = meeting_notes_col

        customer_location_col = prev_col + 1
        sheet.write(0, customer_location_col, "Customer Location")
        sheet.col(customer_location_col).width = 256 * 90
        prev_col = customer_location_col

        attachment_col = prev_col + 1
        sheet.write(0, attachment_col, "Attachment")
        sheet.col(attachment_col).width = 256 * 45
        prev_col = attachment_col

        if cobrowse_agent.role == "admin_ally":
            supervisor_col = prev_col + 1
            sheet.write(0, supervisor_col, "Supervisors")
            sheet.col(supervisor_col).width = 256 * 20
            prev_col = supervisor_col

        return sheet

    try:
        logger.info("Inside get_custom_meeting_support_history",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        support_history_wb = Workbook()

        access_token_obj = cobrowse_agent.get_access_token_obj()

        if not audit_trail_objs:
            agents = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)
            cobrowse_video_objs = CobrowseVideoConferencing.objects.filter(
                agent__in=agents, meeting_start_date__gte=start_date, meeting_start_date__lte=end_date)

            if cobrowse_agent.role == "supervisor" and (access_token_obj.enable_invite_agent_in_meeting or access_token_obj.enable_invite_agent_in_cobrowsing):
                invited_agent_video_audit_objs = CobrowseVideoAuditTrail.objects.filter(
                    meeting_agents__in=agents)
                if invited_agent_video_audit_objs:
                    invited_session_ids = []
                    for invited_agent_video_audit_obj in invited_agent_video_audit_objs.iterator():
                        invited_session_ids.append(
                            invited_agent_video_audit_obj.cobrowse_video.meeting_id)
                    invited_sessions_cobrowse_objs = CobrowseVideoConferencing.objects.filter(
                        meeting_id__in=invited_session_ids, meeting_start_date__gte=start_date, meeting_start_date__lte=end_date)
                    cobrowse_video_objs = cobrowse_video_objs | invited_sessions_cobrowse_objs

            audit_trail_objs = CobrowseVideoAuditTrail.objects.filter(
                cobrowse_video__in=cobrowse_video_objs).order_by('-pk')

            if not audit_trail_objs:
                return False

        wb_index = 1
        sheet = create_new_sheet()

        index = 1
        for audit_trail_obj in audit_trail_objs.iterator():

            if index > 50000:
                index = 1
                sheet = create_new_sheet()

            client_location_details = json.loads(
                audit_trail_obj.client_location_details)
            client_address = []
            for location_obj in client_location_details['items']:
                client_name = location_obj['client_name']
                longitude = location_obj['longitude']
                latitude = location_obj['latitude']
                if latitude == 'None' or longitude == 'None':
                    continue

                if "address" in location_obj and location_obj["address"] != "None":
                    client_address.append(
                        client_name + " - " + location_obj["address"])
                else:
                    try:
                        location = geocoder.google(
                            [latitude, longitude], method="reverse", key=GOOGLE_GEOCODER_KEY)
                        client_address.append(
                            client_name + " - " + location.address)
                        location_obj["address"] = location.address
                    except Exception:
                        logger.warning("Error in get_custom_meeting_support_history, client's address not found",
                                       extra={'AppName': 'EasyAssist'})

            screenshots = json.loads(audit_trail_obj.meeting_screenshot)
            audit_trail_obj.client_location_details = json.dumps(
                client_location_details)
            audit_trail_obj.save()

            attachment_link = ""
            if len(screenshots['items']) > 0:
                zip_file_path = "secured_files/EasyAssistApp/MeetingSupportHistory/attachment_" + \
                    uuid.uuid4().hex + ".zip"

                zip_obj = ZipFile(zip_file_path, 'w')

                for screenshot_obj in screenshots['items']:
                    file_key = screenshot_obj['screenshot']
                    try:
                        file_access_management_obj = CobrowsingFileAccessManagement.objects.get(
                            key=file_key, is_public=False)

                        path_to_file = file_access_management_obj.file_path
                        file_name = path_to_file.split(os.sep)[-1]
                        zip_obj.write(path_to_file[1:], file_name)

                        if get_save_in_s3_bucket_status():
                            s3_bucket_upload_file_by_file_path(
                                zip_file_path, file_name)
                    except Exception:
                        logger.warning("Error in get_custom_meeting_support_history, Invalid file_key",
                                       extra={'AppName': 'EasyAssist'})
                zip_obj.close()

                zip_file_path = "/" + zip_file_path
                zip_file_obj = CobrowsingFileAccessManagement.objects.create(
                    file_path=zip_file_path, is_public=False, access_token=access_token_obj)
                attachment_link = settings.EASYCHAT_HOST_URL + \
                    '/easy-assist/download-file/' + str(zip_file_obj.key)

            sheet.write(index, 0, audit_trail_obj.cobrowse_video.full_name)
            sheet.write(
                index, 1, get_masked_data_if_hashed(audit_trail_obj.cobrowse_video.mobile_number))

            if audit_trail_obj.get_meeting_main_primary_agent():
                meeting_agent_name = audit_trail_obj.get_meeting_main_primary_agent()
            else:
                meeting_agent_name = audit_trail_obj.cobrowse_video.agent.user.username

            sheet.write(
                index, 2, meeting_agent_name)

            if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
                sheet.write(index, auto_assigned_agent_col,
                            audit_trail_obj.get_meeting_auto_assigned_agents())
                
                assigned_count_value = "-"
                lead_transfer_count = len(audit_trail_obj.get_meeting_assigned_agents_list())
                if lead_transfer_count:
                    assigned_count_value = lead_transfer_count   
                sheet.write(index, assigned_agents_count_col, assigned_count_value)

            lead_initiated_by = audit_trail_obj.get_lead_initiated_by()
            sheet.write(index, session_initiated_by, lead_initiated_by)

            if access_token_obj is not None and (access_token_obj.enable_invite_agent_in_cobrowsing or access_token_obj.enable_invite_agent_in_meeting):
                agents_invited_str = "-"
                agents_invited_connected_str = "-"

                agents_invited_list = audit_trail_obj.get_meeting_invited_agents()
                if len(agents_invited_list):
                    agents_invited_str = ""
                    agents_invited_str = ", ".join(agents_invited_list)

                agents_invited_connected_list = audit_trail_obj.get_meeting_agents()
                if len(agents_invited_connected_list):
                    agents_invited_connected_str = ""
                    agents_invited_connected_str = ", ".join(
                        agents_invited_connected_list)

                sheet.write(index, agents_invited_col, agents_invited_str)
                sheet.write(index, invited_agents_connected_col,
                            agents_invited_connected_str)

            sheet.write(index, meeting_initiated_by_col,
                        audit_trail_obj.get_meeting_initiated_by())
            sheet.write(index, meeting_start_datetime_col, str(audit_trail_obj.agent_joined.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%d-%m-%Y - %I:%M %p")))
            sheet.write(index, meeting_end_datetime_col, str(audit_trail_obj.meeting_ended.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%d-%m-%Y - %I:%M %p")))
            sheet.write(index, time_spent_col, str(
                audit_trail_obj.get_readable_meeting_duration()))

            if audit_trail_obj.cobrowse_video.is_expired == True:
                sheet.write(index, meeting_status_col, "Completed")
            else:
                sheet.write(index, meeting_status_col, "Scheduled")
            sheet.write(index, meeting_id_col, str(
                audit_trail_obj.cobrowse_video.meeting_id))

            if audit_trail_obj.agent_notes == None:
                sheet.write(index, meeting_notes_col, "")
            else:
                sheet.write(index, meeting_notes_col,
                            audit_trail_obj.agent_notes)

            sheet.write(index, customer_location_col,
                        " | ".join(client_address))
            sheet.write(index, attachment_col, attachment_link)

            if cobrowse_agent.role == "admin_ally":
                sheet.write(index, supervisor_col,
                            audit_trail_obj.cobrowse_video.agent.get_supervisors())

            index += 1

        support_history_wb.save(absolute_file_path)
        if get_save_in_s3_bucket_status():
            key = s3_bucket_upload_file_by_file_path(
                relative_file_path, file_name)
            s3_file_path = s3_bucket_download_file(
                key, 'EasyAssistApp/MeetingSupportHistory/', '.xls')
            return s3_file_path.split("EasyChat/", 1)[1]

        logger.info("Returning file path " + str(relative_file_path),
                    extra={'AppName': 'EasyAssist'})
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_meeting_support_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def add_unattended_leads_history_into_workbook(workbook, cobrowse_agent, start_date, end_date, cobrowse_io_objs=None):

    from EasyAssistApp.models import CobrowseIO

    def create_new_sheet():
        global comment_col, subremarks_col, language_col, type_of_lead_col, supervisor_col, \
            auto_assigned_agent_col, session_request_datetime_col, title_col, agent_remarks_col, \
            session_initiated_by, assigned_agents_count_col, self_assign_time_col, product_col, \
            product_title_col
        nonlocal wb_index

        if wb_index > 1:
            sheet = workbook.add_sheet("Unattended Leads - " + str(wb_index))
        else:
            sheet = workbook.add_sheet("Unattended Leads")

        wb_index += 1

        sheet.write(0, 0, CUSTOMER_NAME)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, "Agent Email ID")
        sheet.col(2).width = 256 * 35

        prev_col = 2

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            auto_assigned_agent_col = prev_col + 1
            sheet.write(0, auto_assigned_agent_col, "Auto-assigned Agent")
            sheet.col(auto_assigned_agent_col).width = 256 * 35
            prev_col = auto_assigned_agent_col

            assigned_agents_count_col = prev_col + 1
            sheet.write(0, assigned_agents_count_col, "Assigned Agents Count")
            sheet.col(assigned_agents_count_col).width = 256 * 35
            prev_col = assigned_agents_count_col

        session_initiated_by = prev_col + 1
        sheet.write(0, session_initiated_by, "Session Initiated By")
        sheet.col(session_initiated_by).width = 256 * 35
        prev_col = session_initiated_by

        session_request_datetime_col = prev_col + 1
        sheet.write(0, session_request_datetime_col,
                    COBROWSING_REQUEST_DATETIME)
        sheet.col(session_request_datetime_col).width = 256 * 30
        prev_col = session_request_datetime_col

        if access_token_obj.enable_request_in_queue:
            self_assign_time_col = prev_col + 1
            sheet.write(0, self_assign_time_col, "Self Assign Time")
            sheet.col(self_assign_time_col).width = 256 * 35
            prev_col = self_assign_time_col
        
        product_title_col = prev_col + 1
        sheet.write(0, product_title_col, "Product Title")
        sheet.col(product_title_col).width = 256 * 60
        prev_col = product_title_col

        title_col = prev_col + 1
        sheet.write(0, title_col, "Session Start")
        sheet.col(title_col).width = 256 * 60
        prev_col = title_col

        agent_remarks_col = prev_col + 1
        sheet.write(0, agent_remarks_col, AGENT_REMARKS)
        sheet.col(agent_remarks_col).width = 256 * 80
        prev_col = agent_remarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            subremarks_col = prev_col + 1
            sheet.write(0, subremarks_col, "Agent Sub-remarks")
            sheet.col(subremarks_col).width = 256 * 80
            prev_col = subremarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            comment_col = prev_col + 1
            sheet.write(0, comment_col, "Agent Comments")
            sheet.col(comment_col).width = 256 * 80
            prev_col = comment_col

        if access_token_obj is not None and access_token_obj.allow_language_support:
            language_col = prev_col + 1
            sheet.write(0, language_col, "Language")
            sheet.col(language_col).width = 256 * 25
            prev_col = language_col
        
        if access_token_obj is not None and access_token_obj.enable_tag_based_assignment_for_outbound or access_token_obj.choose_product_category:
            product_col = prev_col + 1
            sheet.write(0, product_col, "Category")
            sheet.col(product_col).width = 256 * 25
            prev_col = product_col

        type_of_lead_col = prev_col + 1
        sheet.write(0, type_of_lead_col, "Type of Lead")
        sheet.col(type_of_lead_col).width = 256 * 20
        prev_col = type_of_lead_col

        if cobrowse_agent.role == "admin_ally":
            supervisor_col = prev_col + 1
            sheet.write(0, supervisor_col, "Supervisors")
            sheet.col(supervisor_col).width = 256 * 20
            prev_col = supervisor_col

        return sheet

    access_token_obj = cobrowse_agent.get_access_token_obj()
    
    if not cobrowse_io_objs:
        agents = get_list_agents_under_admin(
            cobrowse_agent, is_active=None, is_account_active=None)

        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True, is_lead_manually_converted=False).filter(
            request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(
            agent__in=agents).filter(
            Q(cobrowsing_start_datetime=None)).filter(
            ~Q(allow_agent_cobrowse="false")).filter(session_archived_cause="UNATTENDED")

        if not cobrowse_io_objs:
            return False
    wb_index = 1
    sheet = create_new_sheet()
    
    index = 1
    for cobrowse_io_obj in cobrowse_io_objs.iterator():

        if index > 50000:
            index = 1
            sheet = create_new_sheet()

        if not cobrowse_io_obj.is_lead:
            sheet.write(index, 0, cobrowse_io_obj.full_name)
            sheet.write(index, 1, get_masked_data_if_hashed(
                cobrowse_io_obj.mobile_number))
        else:
            sheet.write(
                index, 0, cobrowse_io_obj.get_sync_data_value("name"))
            sheet.write(
                index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

        if cobrowse_io_obj.main_primary_agent:
            primary_agent_name = cobrowse_io_obj.main_primary_agent.user.username
        else:
            primary_agent_name = cobrowse_io_obj.agent.user.username

        sheet.write(index, 2, primary_agent_name)

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            sheet.write(index, auto_assigned_agent_col,
                        cobrowse_io_obj.get_auto_assigned_agents())
            
            assigned_count_value = "-"
            lead_transfer_count = len(cobrowse_io_obj.get_unattended_lead_transfer_audit_trail())
            if lead_transfer_count:
                assigned_count_value = lead_transfer_count   
            sheet.write(index, assigned_agents_count_col, assigned_count_value)
            
        lead_initiated_by = "-"
        if cobrowse_io_obj.lead_initiated_by == 'greeting_bubble':
            lead_initiated_by = "Greeting Bubble"
        elif cobrowse_io_obj.lead_initiated_by == 'floating_button':
            lead_initiated_by = "Floating Button"
        elif cobrowse_io_obj.lead_initiated_by == 'icon':
            lead_initiated_by = "Icon"
        elif cobrowse_io_obj.lead_initiated_by == 'exit_intent':
            lead_initiated_by = "Exit Intent"
        elif cobrowse_io_obj.lead_initiated_by == 'inactivity_popup':
            lead_initiated_by = "Inactivity Pop-up"
        sheet.write(index, session_initiated_by, lead_initiated_by)

        sheet.write(index, session_request_datetime_col, str(cobrowse_io_obj.request_datetime.astimezone(
            pytz.timezone(ASIA_KOLKATA)).strftime(DATE_TIME_FORMAT_2)))
        
        if cobrowse_io_obj.product_name():
            sheet.write(index, product_title_col, cobrowse_io_obj.product_name())
        else:
            sheet.write(index, product_title_col, "-")

        if cobrowse_io_obj.product_url():
            sheet.write(index, title_col, cobrowse_io_obj.product_url())
        else:
            sheet.write(index, title_col, "-")

        if access_token_obj.enable_request_in_queue:
            if cobrowse_io_obj.get_self_assign_time():
                sheet.write(index, self_assign_time_col, cobrowse_io_obj.get_self_assign_time())
            else:
                sheet.write(index, self_assign_time_col, "-")
        if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
            sheet.write(index, agent_remarks_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
            ).first().agent_comments)
        else:
            sheet.write(index, agent_remarks_col, NOT_PROVIDED)

        if access_token_obj.enable_predefined_remarks:
            if access_token_obj.enable_predefined_subremarks:
                if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                    if len(cobrowse_io_obj.get_cobrowsing_session_closing_comments().first().agent_subcomments) > 0:
                        sheet.write(index, subremarks_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                        ).first().agent_subcomments)
                    else:
                        sheet.write(index, subremarks_col, "-")
                else:
                    sheet.write(index, subremarks_col, NOT_PROVIDED)

            if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                if len(cobrowse_io_obj.get_cobrowsing_session_closing_comments().first().comment_desc) > 0:
                    sheet.write(index, comment_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                    ).first().comment_desc)
                else:
                    sheet.write(index, comment_col, "-")
            else:
                sheet.write(index, comment_col, NOT_PROVIDED)

        if access_token_obj is not None and access_token_obj.allow_language_support:
            if cobrowse_io_obj.supported_language is not None:
                sheet.write(
                    index, language_col, cobrowse_io_obj.supported_language.title)
            else:
                sheet.write(index, language_col, "-")

        if access_token_obj is not None and access_token_obj.enable_tag_based_assignment_for_outbound or access_token_obj.choose_product_category:
            if cobrowse_io_obj.product_category is not None:
                sheet.write(
                    index, product_col, cobrowse_io_obj.product_category.title)
            else:
                sheet.write(index, product_col, "-")

        lead_type = cobrowse_io_obj.get_lead_type()
        sheet.write(index, type_of_lead_col, lead_type)

        if cobrowse_agent.role == "admin_ally":
            sheet.write(index, supervisor_col,
                        cobrowse_io_obj.agent.get_supervisors())

        index += 1

    return True


def get_custom_unattended_leads_history(requested_data, cobrowse_agent, CobrowseIO, is_public=False, cobrowse_io_objs=None):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/UnattendedLeads/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/UnattendedLeads/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "unattended_leads_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/UnattendedLeads/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        # get_agent_session_count_unattended(requested_data, cobrowse_agent, CobrowseIO)

        logger.info("Inside get_custom_unattended_leads_history",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        is_data_present = add_unattended_leads_history_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date, cobrowse_io_objs)
        if is_data_present:
            add_agent_session_count_unattended_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

            support_history_wb.save(absolute_file_path)

            if get_save_in_s3_bucket_status():
                file_name = get_file_name(".xls")
                key = s3_bucket_upload_file_by_file_path(
                    relative_file_path, file_name)
                s3_file_path = s3_bucket_download_file(
                    key, 'EasyAssistApp/UnattendedLeads/', '.xls')
                return s3_file_path.split("EasyChat/", 1)[1]

            logger.info("Returning file path " + str(relative_file_path),
                        extra={'AppName': 'EasyAssist'})
            return relative_file_path

        return NO_DATA

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_unattended_leads_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


'''
    same filter logic as get_custom_unattended_leads_history
'''


def add_agent_session_count_unattended_into_workbook(workbook, cobrowse_agent, start_date, end_date):

    from EasyAssistApp.models import CobrowseIO

    agents = get_list_agents_under_admin(
        cobrowse_agent, is_active=None, is_account_active=None)
    access_token_obj = cobrowse_agent.get_access_token_obj()

    cobrowse_io_objs = CobrowseIO.objects.filter(
        is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
        request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(
        agent__in=agents).filter(
        Q(cobrowsing_start_datetime=None)).filter(
        ~Q(allow_agent_cobrowse="false")).filter(session_archived_cause="UNATTENDED")

    sheet = workbook.add_sheet("Agent Wise Bifurcation")

    sheet.write(0, 0, "Agent Email ID")
    sheet.col(0).width = 256 * 20
    sheet.write(0, 1, "Unattended Leads")
    sheet.col(1).width = 256 * 20

    index = 1
    for agent in agents:

        sheet.write(index, 0, agent.user.username)
        sheet.write(index, 1, cobrowse_io_objs.filter(agent=agent).count())
        index += 1


def get_agent_session_count_unattended(requested_data, cobrowse_agent, CobrowseIO):

    def get_folder_path():
        folder_path = "EasyAssistApp/UnattendedLeads/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "agent_wise_bifurcation_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        relative_file_path = SECURED_FILES_PATH + \
            get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        logger.info("Inside get_agent_session_count_unattended",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        absolute_file_path = get_absolute_file_path()

        support_history_wb = Workbook()

        add_agent_session_count_unattended_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date)

        support_history_wb.save(absolute_file_path)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_session_count_unattended! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def add_declined_leads_history_into_workbook(workbook, cobrowse_agent, start_date, end_date, cobrowse_io_objs=None):

    from EasyAssistApp.models import CobrowseIO
    from datetime import datetime, timedelta

    def create_new_sheet():
        global comment_col, language_col, type_of_lead_col, supervisor_col, \
            auto_assigned_agent_col, session_request_datetime_col, title_col, drop_title_col, \
            agent_remarks_col, session_initiated_by, assigned_agents_count_col, self_assign_time_col, \
            product_col, product_title_col
        nonlocal wb_index

        if wb_index > 1:
            sheet = workbook.add_sheet("Declined Leads - " + str(wb_index))
        else:
            sheet = workbook.add_sheet("Declined Leads")

        wb_index += 1

        sheet.write(0, 0, CUSTOMER_NAME)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, "Agent Email ID")
        sheet.col(2).width = 256 * 35

        prev_col = 2

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            auto_assigned_agent_col = prev_col + 1
            sheet.write(0, auto_assigned_agent_col, "Auto-assigned Agent")
            sheet.col(auto_assigned_agent_col).width = 256 * 35
            prev_col = auto_assigned_agent_col

            assigned_agents_count_col = prev_col + 1
            sheet.write(0, assigned_agents_count_col, "Assigned Agents Count")
            sheet.col(assigned_agents_count_col).width = 256 * 35
            prev_col = assigned_agents_count_col

        session_initiated_by = prev_col + 1
        sheet.write(0, session_initiated_by, "Session Initiated By")
        sheet.col(session_initiated_by).width = 256 * 35
        prev_col = session_initiated_by

        session_request_datetime_col = prev_col + 1
        sheet.write(0, session_request_datetime_col,
                    COBROWSING_REQUEST_DATETIME)
        sheet.col(session_request_datetime_col).width = 256 * 30
        prev_col = session_request_datetime_col

        if access_token_obj.enable_request_in_queue:
            self_assign_time_col = prev_col + 1
            sheet.write(0, self_assign_time_col, "Self Assign Time")
            sheet.col(self_assign_time_col).width = 256 * 35
            prev_col = self_assign_time_col

        product_title_col = prev_col + 1
        sheet.write(0, product_title_col, "Product Title")
        sheet.col(product_title_col).width = 256 * 60
        prev_col = product_title_col

        title_col = prev_col + 1
        sheet.write(0, title_col, "Session Start")
        sheet.col(title_col).width = 256 * 60
        prev_col = title_col
        
        drop_title_col = prev_col + 1
        sheet.write(0, drop_title_col, "Session End")
        sheet.col(drop_title_col).width = 256 * 60
        prev_col = drop_title_col

        agent_remarks_col = prev_col + 1
        sheet.write(0, agent_remarks_col, AGENT_REMARKS)
        sheet.col(agent_remarks_col).width = 256 * 80
        prev_col = agent_remarks_col

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            comment_col = prev_col + 1
            sheet.write(0, comment_col, "Agent Comments")
            sheet.col(comment_col).width = 256 * 80
            prev_col = comment_col

        if access_token_obj is not None and access_token_obj.allow_language_support:
            language_col = prev_col + 1
            sheet.write(0, language_col, "Language")
            sheet.col(language_col).width = 256 * 25
            prev_col = language_col
        
        if access_token_obj is not None and access_token_obj.enable_tag_based_assignment_for_outbound or access_token_obj.choose_product_category:
            product_col = prev_col + 1
            sheet.write(0, product_col, "Category")
            sheet.col(product_col).width = 256 * 25
            prev_col = product_col

        type_of_lead_col = prev_col + 1
        sheet.write(0, type_of_lead_col, "Type of Lead")
        sheet.col(type_of_lead_col).width = 256 * 20
        prev_col = type_of_lead_col

        if cobrowse_agent.role == "admin_ally":
            supervisor_col = prev_col + 1
            sheet.write(0, supervisor_col, "Supervisors")
            sheet.col(supervisor_col).width = 256 * 20
            prev_col = supervisor_col

        return sheet

    access_token_obj = cobrowse_agent.get_access_token_obj()

    if not cobrowse_io_objs:
        agents = get_list_agents_under_admin(
            cobrowse_agent, is_active=None, is_account_active=None)

        time_threshold = datetime.now() - timedelta(minutes=5)

        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
            request_datetime__lte=time_threshold, request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(
            agent__in=agents).filter(
            Q(cobrowsing_start_datetime=None)).filter(
            allow_agent_cobrowse="false")
        
        if not cobrowse_io_objs:
            return False
    
    wb_index = 1
    sheet = create_new_sheet()
    
    index = 1
    for cobrowse_io_obj in cobrowse_io_objs.iterator():

        if index > 50000:
            index = 1
            sheet = create_new_sheet()

        if not cobrowse_io_obj.is_lead:
            sheet.write(index, 0, cobrowse_io_obj.full_name)
            sheet.write(index, 1, get_masked_data_if_hashed(
                cobrowse_io_obj.mobile_number))
        else:
            sheet.write(
                index, 0, cobrowse_io_obj.get_sync_data_value("name"))
            sheet.write(
                index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

        if cobrowse_io_obj.main_primary_agent:
            primary_agent_name = cobrowse_io_obj.main_primary_agent.user.username
        else:
            primary_agent_name = cobrowse_io_obj.agent.user.username

        sheet.write(index, 2, primary_agent_name)

        if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
            sheet.write(index, auto_assigned_agent_col,
                        cobrowse_io_obj.get_auto_assigned_agents())
            
            assigned_count_value = "-"
            lead_transfer_count = len(cobrowse_io_obj.get_unattended_lead_transfer_audit_trail())
            if lead_transfer_count:
                assigned_count_value = lead_transfer_count   
            sheet.write(index, assigned_agents_count_col, assigned_count_value)

        lead_initiated_by = "-"
        if cobrowse_io_obj.lead_initiated_by == 'greeting_bubble':
            lead_initiated_by = "Greeting Bubble"
        elif cobrowse_io_obj.lead_initiated_by == 'floating_button':
            lead_initiated_by = "Floating Button"
        elif cobrowse_io_obj.lead_initiated_by == 'icon':
            lead_initiated_by = "Icon"
        elif cobrowse_io_obj.lead_initiated_by == 'exit_intent':
            lead_initiated_by = "Exit Intent"
        elif cobrowse_io_obj.lead_initiated_by == 'inactivity_popup':
            lead_initiated_by = "Inactivity Pop-up"

        sheet.write(index, session_initiated_by, lead_initiated_by)

        sheet.write(index, session_request_datetime_col, str(cobrowse_io_obj.request_datetime.astimezone(
            pytz.timezone(ASIA_KOLKATA)).strftime(DATE_TIME_FORMAT_2)))

        if access_token_obj.enable_request_in_queue:
            if cobrowse_io_obj.get_self_assign_time():
                sheet.write(index, self_assign_time_col, cobrowse_io_obj.get_self_assign_time())
            else:
                sheet.write(index, self_assign_time_col, "-")
        
        if cobrowse_io_obj.product_name():
            sheet.write(index, product_title_col, cobrowse_io_obj.product_name())
        else:
            sheet.write(index, product_title_col, "-")

        if cobrowse_io_obj.product_url():
            sheet.write(index, title_col, cobrowse_io_obj.product_url())
        else:
            sheet.write(index, title_col, "-")
        
        if not cobrowse_io_obj.is_reverse_cobrowsing:
            sheet.write(index, drop_title_col, cobrowse_io_obj.drop_off_url())
        else:
            sheet.write(index, drop_title_col, '-')
        if cobrowse_io_obj.agent_comments:
            sheet.write(index, agent_remarks_col, cobrowse_io_obj.agent_comments)
        else:
            sheet.write(index, agent_remarks_col, "-")

        if access_token_obj is not None and access_token_obj.enable_predefined_remarks:
            if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                if len(cobrowse_io_obj.get_cobrowsing_session_closing_comments().first().comment_desc) > 0:
                    sheet.write(index, comment_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                    ).first().comment_desc)
                else:
                    sheet.write(index, comment_col, "-")
            else:
                sheet.write(index, comment_col, NOT_PROVIDED)

        if access_token_obj is not None and access_token_obj.allow_language_support:
            if cobrowse_io_obj.supported_language is not None:
                sheet.write(
                    index, language_col, cobrowse_io_obj.supported_language.title)
            else:
                sheet.write(index, language_col, "-")

        if access_token_obj is not None and access_token_obj.enable_tag_based_assignment_for_outbound or access_token_obj.choose_product_category:
            if cobrowse_io_obj.product_category is not None:
                sheet.write(
                    index, product_col, cobrowse_io_obj.product_category.title)
            else:
                sheet.write(index, product_col, "-")
        
        lead_type = cobrowse_io_obj.get_lead_type()
        sheet.write(index, type_of_lead_col, lead_type)

        if cobrowse_agent.role == "admin_ally":
            sheet.write(index, supervisor_col,
                        cobrowse_io_obj.agent.get_supervisors())

        index += 1

    return True


def get_custom_declined_leads_history(requested_data, cobrowse_agent, CobrowseIO, is_public=False, cobrowse_io_objs=None):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/DeclinedLeads/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/DeclinedLeads/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "declined_leads_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/DeclinedLeads/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        logger.info("Inside get_custom_declined_leads_history",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        is_data_present = add_declined_leads_history_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date, cobrowse_io_objs)

        if is_data_present:
            support_history_wb.save(absolute_file_path)

            if get_save_in_s3_bucket_status():
                file_name = get_file_name(".xls")
                key = s3_bucket_upload_file_by_file_path(
                    relative_file_path, file_name)
                s3_file_path = s3_bucket_download_file(
                    key, 'EasyAssistApp/DeclinedLeads/', '.xls')
                return s3_file_path.split("EasyChat/", 1)[1]

            logger.info("Returning file path " + str(relative_file_path),
                        extra={'AppName': 'EasyAssist'})
            return relative_file_path

        return NO_DATA
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_declined_leads_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def add_followup_leads_history_into_workbook(workbook, cobrowse_agent, start_date, end_date, cobrowse_io_objs=None):

    from EasyAssistApp.models import CobrowseIO

    def create_new_sheet():
        global category_col, language_col
        nonlocal wb_index

        if wb_index > 1:
            sheet = workbook.add_sheet("Followup Leads - " + str(wb_index))
        else:
            sheet = workbook.add_sheet("Followup Leads")

        wb_index += 1

        sheet.write(0, 0, CUSTOMER_NAME)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, COBROWSING_REQUEST_DATETIME)
        sheet.col(2).width = 256 * 30
        sheet.write(0, 3, "Session Initiated By")
        sheet.col(3).width = 256 * 30
        sheet.write(0, 4, "Product Title")
        sheet.col(4).width = 256 * 80
        sheet.write(0, 5, "Session Start")
        sheet.col(5).width = 256 * 80
        sheet.write(0, 6, "Assigned Agent")
        sheet.col(6).width = 256 * 25

        prev_col = 6

        if access_token_obj is not None and access_token_obj.choose_product_category:
            category_col = prev_col + 1
            sheet.write(0, category_col, "Category")
            sheet.col(category_col).width = 256 * 30
            prev_col = category_col

        if access_token_obj is not None and access_token_obj.allow_language_support:
            language_col = prev_col + 1
            sheet.write(0, language_col, "Language")
            sheet.col(language_col).width = 256 * 25
            prev_col = language_col

        return sheet

    access_token_obj = cobrowse_agent.get_access_token_obj()

    if not cobrowse_io_objs:        
        if cobrowse_agent.role in ["admin", "admin_ally"]:
            agent_objs = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)

        elif cobrowse_agent.role == "supervisor":
            agent_objs = list(cobrowse_agent.agents.all())

        cobrowse_io_objs = CobrowseIO.objects.filter(
            access_token=access_token_obj, is_test=False, request_datetime__date__gte=access_token_obj.go_live_date,
            agent__in=agent_objs, is_archived=True, is_lead=False, session_archived_cause__in=["UNASSIGNED", "FOLLOWUP"], is_lead_manually_converted=False).filter(
                request_datetime__date__gte=start_date, request_datetime__date__lte=end_date)

        if not cobrowse_io_objs:
            return False

    wb_index = 1
    sheet = create_new_sheet()
    
    index = 1
    for cobrowse_io_obj in cobrowse_io_objs.iterator():

        if index > 50000:
            index = 1
            sheet = create_new_sheet()

        if not cobrowse_io_obj.is_lead:
            sheet.write(index, 0, cobrowse_io_obj.full_name)
            sheet.write(index, 1, get_masked_data_if_hashed(
                cobrowse_io_obj.mobile_number))
        else:
            sheet.write(
                index, 0, cobrowse_io_obj.get_sync_data_value("name"))
            sheet.write(
                index, 1, cobrowse_io_obj.get_sync_data_value("mobile"))

        sheet.write(index, 2, str(cobrowse_io_obj.request_datetime.astimezone(
            pytz.timezone(ASIA_KOLKATA)).strftime(DATE_TIME_FORMAT_2)))

        lead_initiated_by = "-"
        if cobrowse_io_obj.lead_initiated_by == 'greeting_bubble':
            lead_initiated_by = "Greeting Bubble"
        elif cobrowse_io_obj.lead_initiated_by == 'floating_button':
            lead_initiated_by = "Floating Button"
        elif cobrowse_io_obj.lead_initiated_by == 'icon':
            lead_initiated_by = "Icon"
        elif cobrowse_io_obj.lead_initiated_by == 'exit_intent':
            lead_initiated_by = "Exit Intent"
        elif cobrowse_io_obj.lead_initiated_by == 'inactivity_popup':
            lead_initiated_by = "Inactivity Pop-up"
        sheet.write(index, 3, lead_initiated_by)

        if cobrowse_io_obj.product_name():
            sheet.write(index, 4, cobrowse_io_obj.product_name())
        else:
            sheet.write(index, 4, "-")

        if cobrowse_io_obj.product_url():
            sheet.write(index, 5, cobrowse_io_obj.product_url())
        else:
            sheet.write(index, 5, "-")

        assigned_agent_name = "-"
        if cobrowse_io_obj.agent != None:
            assigned_agent_name = cobrowse_io_obj.agent.user.username

        sheet.write(index, 6, assigned_agent_name)

        if access_token_obj is not None and access_token_obj.choose_product_category:
            if cobrowse_io_obj.product_category is not None:
                sheet.write(
                    index, category_col, cobrowse_io_obj.product_category.title)
            else:
                sheet.write(index, category_col, "")
        if access_token_obj is not None and access_token_obj.allow_language_support:
            if cobrowse_io_obj.supported_language is not None:
                sheet.write(
                    index, language_col, cobrowse_io_obj.supported_language.title)
            else:
                sheet.write(index, language_col, "")
        index += 1

    return True


def get_custom_followup_leads_history(requested_data, cobrowse_agent, CobrowseIO, is_public=False, cobrowse_io_objs=None):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/FollowupLeads/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/FollowupLeads/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "followup_leads_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/FollowupLeads/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        logger.info("Inside get_custom_followup_leads_history",
                    extra={'AppName': 'EasyAssist'})

        if cobrowse_agent.role == "agent":
            return "None"

        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        is_data_present = add_followup_leads_history_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date, cobrowse_io_objs)

        if is_data_present:
            support_history_wb.save(absolute_file_path)

            if get_save_in_s3_bucket_status():
                file_name = get_file_name(".xls")
                key = s3_bucket_upload_file_by_file_path(
                    relative_file_path, file_name)
                s3_file_path = s3_bucket_download_file(
                    key, 'EasyAssistApp/FollowupLeads/', '.xls')
                return s3_file_path.split("EasyChat/", 1)[1]

            logger.info("Returning file path " + str(relative_file_path),
                        extra={'AppName': 'EasyAssist'})
            return relative_file_path
        return NO_DATA
        
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_followup_leads_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def get_custom_screen_recording_history(requested_data, cobrowse_agent, CobrowseScreenRecordingAuditTrail, screen_recording_objs=None):

    def get_folder_path():
        folder_path = "EasyAssistApp/ScreenRecordingHistory/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "screen_recording_history_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        relative_file_path = SECURED_FILES_PATH + \
            get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    def create_new_sheet():
        global supervisor_col
        nonlocal wb_index

        sheet = support_history_wb.add_sheet(
            "Screen Recording History - " + str(wb_index))
        wb_index += 1

        sheet.write(0, 0, CUSTOMER_NAME)
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, CUSTOMER_MOBILE_NUMBER)
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, "Agent Email ID")
        sheet.col(2).width = 256 * 35
        sheet.write(0, 3, "Recording Start Date & Time")
        sheet.col(3).width = 256 * 25
        sheet.write(0, 4, "Recording Duration")
        sheet.col(4).width = 256 * 20

        prev_col = 4

        if cobrowse_agent.role == "admin_ally":
            supervisor_col = prev_col + 1
            sheet.write(0, supervisor_col, "Supervisors")
            sheet.col(supervisor_col).width = 256 * 20
            prev_col = supervisor_col

        return sheet

    try:
        logger.info("Inside get_custom_screen_recording_history",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        if not screen_recording_objs:
            agents = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)

            screen_recording_objs = CobrowseScreenRecordingAuditTrail.objects.filter(
                agent__in=agents, recording_started__date__gte=start_date, recording_started__date__lte=end_date)

        wb_index = 1
        sheet = create_new_sheet()

        index = 1
        for screen_recording_obj in screen_recording_objs.iterator():

            if index > 50000:
                index = 1
                sheet = create_new_sheet()

            sheet.write(index, 0, screen_recording_obj.cobrowse_io.full_name)
            sheet.write(
                index, 1, get_masked_data_if_hashed(screen_recording_obj.cobrowse_io.mobile_number))

            sheet.write(
                index, 2, screen_recording_obj.agent.user.username)
            sheet.write(index, 3, str(screen_recording_obj.recording_started.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%d-%m-%Y - %I:%M %p")))
            sheet.write(index, 4, str(
                screen_recording_obj.get_screen_recording_duration()))

            if cobrowse_agent.role == "admin_ally":
                sheet.write(index, supervisor_col,
                            screen_recording_obj.agent.get_supervisors())

            index += 1

        support_history_wb.save(absolute_file_path)

        if get_save_in_s3_bucket_status():
            file_name = get_file_name(".xls")
            key = s3_bucket_upload_file_by_file_path(
                relative_file_path, file_name)
            s3_file_path = s3_bucket_download_file(
                key, 'EasyAssistApp/ScreenRecordingHistory/', '.xls')
            return s3_file_path.split("EasyChat/", 1)[1]

        logger.info("Returning file path " + str(relative_file_path),
                    extra={'AppName': 'EasyAssist'})
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_custom_screen_recording_history! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def get_audit_trail_dump(requested_data, cobrowse_agent, CobrowsingAuditTrail, audit_trail_objs=None):

    def get_folder_path():
        folder_path = "EasyAssistApp/AuditTrail/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "audit_trail_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        relative_file_path = SECURED_FILES_PATH + \
            get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    def create_new_sheet():
        nonlocal wb_index

        sheet = support_history_wb.add_sheet("Audit Trail - " + str(wb_index))
        wb_index += 1

        sheet.write(0, 0, "User Email ID")
        sheet.col(0).width = 256 * 25
        sheet.write(0, 1, "Date & Time")
        sheet.col(1).width = 256 * 20
        sheet.write(0, 2, "Action")
        sheet.col(2).width = 256 * 35
        sheet.write(0, 3, "Description")
        sheet.col(3).width = 256 * 35

        return sheet

    try:
        logger.info("Inside get_audit_trail_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        if not audit_trail_objs:
            agents = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)

            agents += cobrowse_agent.agents.all().filter(
                role="supervisor")

            if cobrowse_agent not in agents:
                agents += [cobrowse_agent]
            audit_trail_objs = CobrowsingAuditTrail.objects.filter(
                agent__in=agents, datetime__date__gte=start_date, datetime__date__lte=end_date)

        wb_index = 1
        sheet = create_new_sheet()

        index = 1
        audit_trail_objs = audit_trail_objs.order_by("-datetime")
        for audit_trail_obj in audit_trail_objs.iterator():

            if index > 50000:
                index = 1
                sheet = create_new_sheet()

            sheet.write(index, 0, audit_trail_obj.agent.user.username)
            sheet.write(index, 1, str(audit_trail_obj.datetime.astimezone(
                pytz.timezone(ASIA_KOLKATA)).strftime(DATE_TIME_FORMAT_2)))
            sheet.write(index, 2, audit_trail_obj.get_action())
            sheet.write(index, 3, audit_trail_obj.action_description)
            index += 1

        support_history_wb.save(absolute_file_path)

        if get_save_in_s3_bucket_status():
            file_name = get_file_name(".xls")
            key = s3_bucket_upload_file_by_file_path(
                relative_file_path, file_name)
            s3_file_path = s3_bucket_download_file(
                key, 'EasyAssistApp/AuditTrail/', '.xls')
            return s3_file_path.split("EasyChat/", 1)[1]
        logger.info("Returning file path " + str(relative_file_path),
                    extra={'AppName': 'EasyAssist'})
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_audit_trail_dump! %s %s", str(e), str(
            exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return "None"


def get_agent_audit_trail_dump(requested_data, cobrowse_agent, CobrowsingAuditTrail, CobrowseAgentWorkAuditTrail, CobrowseAgentOnlineAuditTrail, agent_audit_trail_list=None):

    def get_folder_path():
        folder_path = "EasyAssistApp/AgentAuditTrail/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "agent_audit_trail_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        relative_file_path = SECURED_FILES_PATH + \
            get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    def create_new_sheet():
        nonlocal wb_index

        sheet = support_history_wb.add_sheet(
            "Agent Audit Trail - " + str(wb_index))
        wb_index += 1

        sheet.write(0, 0, "Agent Email ID")
        sheet.col(0).width = 256 * 25
        sheet.write(0, 1, "Supervisors")
        sheet.col(1).width = 256 * 45
        sheet.write(0, 2, "Last Login time")
        sheet.col(2).width = 256 * 20
        sheet.write(0, 3, "Last Logout time")
        sheet.col(3).width = 256 * 20
        sheet.write(0, 4, "Date")
        sheet.col(4).width = 256 * 20
        sheet.write(0, 5, "Total Online Duration")
        sheet.col(5).width = 256 * 20
        sheet.write(0, 6, "Idle Time")
        sheet.col(6).width = 256 * 20
        sheet.write(0, 7, "Offline Duration")
        sheet.col(7).width = 256 * 20

        return sheet

    try:
        logger.info("Inside get_agent_audit_trail_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        relative_file_path = get_relative_file_path()
        absolute_file_path = get_absolute_file_path()

        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path

        support_history_wb = Workbook()

        agents = []
        if cobrowse_agent.role == "supervisor" or (cobrowse_agent.role == "agent" and cobrowse_agent.is_switch_allowed):
            agents = list(cobrowse_agent.agents.all())
        elif cobrowse_agent.role in ["admin", "admin_ally"] or (cobrowse_agent.role == "agent" and cobrowse_agent.is_switch_allowed):
            agents = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)
        else:
            agents = []

        if cobrowse_agent.is_switch_allowed and (cobrowse_agent not in agents):
            agents += [cobrowse_agent]
        if not agent_audit_trail_list:
            agent_audit_trail_list = CobrowsingAuditTrail.objects.filter(
                agent__in=agents, action__in=[COBROWSING_LOGIN_ACTION, COBROWSING_LOGOUT_ACTION], datetime__date__gte=start_date, datetime__date__lte=end_date)
        
        agent_online_audit_trail_objs = CobrowseAgentOnlineAuditTrail.objects.filter(
            agent__in=agents, last_online_start_datetime__date__gte=start_date, last_online_start_datetime__date__lte=end_date)

        agent_wise_audit_trail = get_agent_wise_audit_trail(
            agent_audit_trail_list, cobrowse_agent, agents, agent_online_audit_trail_objs, CobrowseAgentWorkAuditTrail)

        wb_index = 1
        sheet = create_new_sheet()

        index = 1
        for audit_trail_obj in agent_wise_audit_trail:

            if index > 50000:
                index = 1
                sheet = create_new_sheet()

            sheet.write(index, 0, audit_trail_obj["agent_username"])
            sheet.write(index, 1, audit_trail_obj["agent_supervisors"])
            sheet.write(index, 2, audit_trail_obj["last_login_time"])
            sheet.write(index, 3, audit_trail_obj["last_logout_time"])
            sheet.write(index, 4, audit_trail_obj["date"])
            sheet.write(index, 5, audit_trail_obj["duration"])
            sheet.write(index, 6, audit_trail_obj["idle_time"])
            sheet.write(index, 7, audit_trail_obj["offline_duration"])
            index += 1

        support_history_wb.save(absolute_file_path)

        if get_save_in_s3_bucket_status():
            file_name = get_file_name(".xls")
            key = s3_bucket_upload_file_by_file_path(
                relative_file_path, file_name)
            s3_file_path = s3_bucket_download_file(
                key, 'EasyAssistApp/AgentAuditTrail/', '.xls')
            return s3_file_path.split("EasyChat/", 1)[1]
        logger.info("Returning file path " + str(relative_file_path),
                    extra={'AppName': 'EasyAssist'})
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_audit_trail_dump! %s %s", str(e), str(
            exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return "None"


def get_agent_online_audit_trail_dump(requested_data, cobrowse_agent, CobrowseAgentWorkAuditTrail, CobrowseAgentOnlineAuditTrail, agent_online_audit_trail_objs=None):

    def get_folder_path():
        folder_path = "EasyAssistApp/AgentOnlineAuditTrail/" + \
            str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name():
        file_name = "agent_online_audit_trail_" + \
            str(start_date) + "-" + str(end_date) + ".xls"
        return file_name

    def get_relative_file_path():
        relative_file_path = SECURED_FILES_PATH + get_folder_path() + get_file_name()
        return relative_file_path

    def get_absolute_file_path():
        absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
        if not path.exists(absolute_folder_path):
            os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name()
        return absolute_file_path

    def create_new_sheet():
        nonlocal wb_index

        sheet = support_history_wb.add_sheet(
            "Agent online audit trail - " + str(wb_index))
        wb_index += 1

        sheet.write(0, 0, "Date")
        sheet.col(0).width = 256 * 20
        sheet.write(0, 1, "Agent Email ID")
        sheet.col(1).width = 256 * 25
        sheet.write(0, 2, "Supervisors")
        sheet.col(2).width = 256 * 45
        sheet.write(0, 3, "Online time")
        sheet.col(3).width = 256 * 20
        sheet.write(0, 4, "Offline time")
        sheet.col(4).width = 256 * 20
        sheet.write(0, 5, "Online Duration")
        sheet.col(5).width = 256 * 20
        sheet.write(0, 6, "Idle Time")
        sheet.col(6).width = 256 * 20

        return sheet

    try:
        logger.info("Inside get_agent_online_audit_trail_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()
        yesterdays_date = (datetime.now() - timedelta(days=1)).date()

        agents = []
        if cobrowse_agent.role == "supervisor" or (cobrowse_agent.role == "agent" and cobrowse_agent.is_switch_allowed):
            agents = list(cobrowse_agent.agents.all())
        elif cobrowse_agent.role in ["admin", "admin_ally"] or (cobrowse_agent.role == "agent" and cobrowse_agent.is_switch_allowed):
            agents = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)
        else:
            agents = []

        if cobrowse_agent.is_switch_allowed and (cobrowse_agent not in agents):
            agents += [cobrowse_agent]
        
        if not agent_online_audit_trail_objs:
            agent_online_audit_trail_objs = CobrowseAgentOnlineAuditTrail.objects.filter(
                agent__in=agents,
                last_online_start_datetime__date__gte=start_date,
                last_online_end_datetime__date__lte=end_date)
        
        absolute_file_path = get_absolute_file_path()
        relative_file_path = get_relative_file_path()

        if path.exists(absolute_file_path) and (end_date < yesterdays_date):
            return relative_file_path    

        agent_online_audit_trail_objs = agent_online_audit_trail_objs.order_by(
            '-last_online_start_datetime')
        agent_wise_audit_trail = get_agent_wise_online_audit_trail(
            cobrowse_agent, agents, agent_online_audit_trail_objs, CobrowseAgentWorkAuditTrail)

        support_history_wb = Workbook()

        wb_index = 1
        sheet = create_new_sheet()

        index = 1
        for audit_trail_obj in agent_wise_audit_trail:

            if index > 50000:
                index = 1
                sheet = create_new_sheet()

            sheet.write(index, 0, audit_trail_obj["date"])
            sheet.write(index, 1, audit_trail_obj["agent_username"])
            sheet.write(index, 2, audit_trail_obj["agent_supervisors"])
            sheet.write(index, 3, audit_trail_obj["online_time"])
            sheet.write(index, 4, audit_trail_obj["offline_time"])
            sheet.write(index, 5, audit_trail_obj["online_duration"])
            sheet.write(index, 6, audit_trail_obj["idle_time"])
            index += 1

        support_history_wb.save(absolute_file_path)

        if get_save_in_s3_bucket_status():
            file_name = get_file_name(".xls")
            key = s3_bucket_upload_file_by_file_path(
                relative_file_path, file_name)
            s3_file_path = s3_bucket_download_file(
                key, 'EasyAssistApp/AgentOnlineAuditTrail/', '.xls')
            return s3_file_path.split("EasyChat/", 1)[1]
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_online_audit_trail_dump! %s %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return "None"


def auto_assign_unattended_leads(CobrowseAccessToken, CobrowseIO, CobrowseAgent, User, UnattendedLeadTransferAuditTrail):
    try:
        cobrowse_access_token_objs = CobrowseAccessToken.objects.all()

        allowed_agent_support_levels = ["L1"]

        for access_token_obj in cobrowse_access_token_objs.iterator():
            if access_token_obj.enable_auto_assign_unattended_lead:

                admin_agent = access_token_obj.agent
                allow_video_meeting_only = access_token_obj.allow_video_meeting_only

                if not access_token_obj.assign_agent_under_same_supervisor:
                    active_agents = get_list_agents_under_admin(
                        admin_agent, is_active=True)

                max_lead_reassignment_count = access_token_obj.unattended_lead_auto_assignment_counter
                
                if allow_video_meeting_only:
                    cobrowse_io_objs = CobrowseIO.objects.filter(access_token=access_token_obj, agent__isnull=False, meeting_start_datetime=None,
                                                                 is_lead=False, is_reverse_cobrowsing=False, is_droplink_lead=False, is_archived=False, is_agent_request_for_cobrowsing=False,
                                                                 lead_reassignment_counter__lt=max_lead_reassignment_count).exclude(cobrowsing_type="modified-inbound")
                else:
                    cobrowse_io_objs = CobrowseIO.objects.filter(access_token=access_token_obj, agent__isnull=False, cobrowsing_start_datetime=None,
                                                                 is_lead=False, is_reverse_cobrowsing=False, is_droplink_lead=False, is_archived=False, is_agent_request_for_cobrowsing=False,
                                                                 lead_reassignment_counter__lt=max_lead_reassignment_count).exclude(cobrowsing_type="modified-inbound")

                for cobrowse_io in cobrowse_io_objs.iterator():
                    if cobrowse_io.is_auto_assign_timer_exhausted():

                        is_agent_filtering_required = True
                        
                        if access_token_obj.assign_agent_under_same_supervisor:
                            prev_selected_agent = cobrowse_io.agent
                            # fetch agents mapped only under admin
                            admin_agents = admin_agent.agents.filter(role="agent", is_active=True)
                            if admin_agent.is_switch_allowed and admin_agent.is_active and admin_agent not in admin_agents:
                                admin_agents += [admin_agent]
                            if prev_selected_agent == admin_agent or prev_selected_agent in admin_agents:
                                # filter free active agents under admin
                                filtered_admin_agents = filter_free_active_agent(
                                    admin_agents, cobrowse_io, for_meeting=allow_video_meeting_only, support_levels=allowed_agent_support_levels)
                                
                                if len(filtered_admin_agents) and prev_selected_agent in filtered_admin_agents:
                                    filtered_admin_agents.remove(prev_selected_agent)
                                
                                if len(filtered_admin_agents):
                                    active_agents = filtered_admin_agents
                                    is_agent_filtering_required = False
                                else:
                                    # fetch all agents throughout organization
                                    active_agents = get_list_agents_under_admin(
                                        admin_agent, is_active=True)
                            else:
                                # fetch agents of supervisor
                                selected_supervisor = None
                                supervisors_list = admin_agent.agents.filter(role="supervisor")
                                for supervisor in supervisors_list.iterator():
                                    if prev_selected_agent in supervisor.agents.all():
                                        selected_supervisor = supervisor
                                        break
                                if selected_supervisor:
                                    supervisor_agents = selected_supervisor.agents.filter(role="agent", is_active=True)
                                    filtered_supervisor_agents = filter_free_active_agent(
                                        supervisor_agents, cobrowse_io, for_meeting=allow_video_meeting_only, support_levels=allowed_agent_support_levels)
                                    
                                    if len(filtered_supervisor_agents) and prev_selected_agent in filtered_supervisor_agents:
                                        filtered_supervisor_agents.remove(prev_selected_agent)
                                    
                                    if len(filtered_supervisor_agents):
                                        active_agents = filtered_supervisor_agents
                                        is_agent_filtering_required = False
                                    else:
                                        active_agents = get_list_agents_under_admin(
                                            admin_agent, is_active=True)
                                else:
                                    active_agents = get_list_agents_under_admin(
                                        admin_agent, is_active=True)

                        if is_agent_filtering_required:
                            active_agents = filter_free_active_agent(
                                active_agents, cobrowse_io, for_meeting=allow_video_meeting_only, support_levels=allowed_agent_support_levels)

                        if cobrowse_io.lead_reassignment_counter == 0:
                            cobrowse_io.main_primary_agent = cobrowse_io.agent

                        if len(active_agents) != 0:
                            if cobrowse_io.agent in active_agents:
                                active_agents.remove(cobrowse_io.agent)

                        # Assign new agent and update list of leads assigned,
                        # if no agent then reassign current agent.
                        if len(active_agents) == 0:
                            cobrowse_io.agents_assigned_list.add(
                                cobrowse_io.agent)
                            update_unattended_lead_transfer_audit_trail(cobrowse_io, cobrowse_io.agent, UnattendedLeadTransferAuditTrail)
                        else:
                            new_cobrowse_agent = get_agent_for_reassignment(
                                active_agents, cobrowse_io, CobrowseIO, CobrowseAgent, User)
                            if new_cobrowse_agent:
                                cobrowse_io.agent = new_cobrowse_agent
                                cobrowse_io.agents_assigned_list.add(
                                    new_cobrowse_agent)
                                update_unattended_lead_transfer_audit_trail(cobrowse_io, new_cobrowse_agent, UnattendedLeadTransferAuditTrail)
                            else:
                                cobrowse_io.agents_assigned_list.add(
                                    cobrowse_io.agent)
                                update_unattended_lead_transfer_audit_trail(cobrowse_io, cobrowse_io.agent, UnattendedLeadTransferAuditTrail)

                        # if customer should not be notified if the same agent is reassigned then shift the below line to the place
                        # where new agent is being assigned also set this to
                        # True where current agent is being reassigned
                        cobrowse_io.is_customer_notified = False
                        cobrowse_io.last_agent_assignment_datetime = timezone.now()
                        cobrowse_io.lead_reassignment_counter = cobrowse_io.lead_reassignment_counter + 1
                        cobrowse_io.save()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("auto_assign_unattended_leads %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_agent_for_reassignment(active_agents, cobrowse_io, CobrowseIO, CobrowseAgent, User):
    try:
        agent_dict = {}
        for agent in active_agents:
            agent_dict[agent.user.username] = []
            agent_dict[agent.user.username].append(0)
            if agent.last_lead_assigned_datetime != None:
                agent_dict[agent.user.username].append(
                    datetime.datetime.timestamp(agent.last_lead_assigned_datetime))
            else:
                agent_dict[agent.user.username].append(0)

        if not agent_dict:
            logger.info("Active Agents are not available for reassignment to an unattended lead.",
                        extra={'AppName': 'EasyAssist'})
            return None

        waiting_cobrowse_io_objs = CobrowseIO.objects.filter(
            agent__in=active_agents, is_archived=False)

        for waiting_cobrowse_io_obj in waiting_cobrowse_io_objs:
            if waiting_cobrowse_io_obj.is_active_timer() == False:
                continue

            if waiting_cobrowse_io_obj.agent.user.username in agent_dict:
                agent_dict[waiting_cobrowse_io_obj.agent.user.username][0] += 1

        agent_min_lead_entry = min(agent_dict.items(),
                                   key=lambda item: (item[1][0], item[1][1]))

        agent_username = agent_min_lead_entry[0]
        agent_active_leads_count = agent_min_lead_entry[1][0]

        logger.info("Agent with min lead eligible for reassignment for an unattended lead: " + str(agent_username) + ":" +
                    str(agent_active_leads_count), extra={'AppName': 'EasyAssist'})

        if cobrowse_io.access_token.maximum_active_leads and agent_active_leads_count >= cobrowse_io.access_token.maximum_active_leads_threshold:
            logger.info("All active agents have atleast %s active leads and thus the lead cannot be reassigned to a new agent.",
                        str(cobrowse_io.access_token.maximum_active_leads_threshold), extra={'AppName': 'EasyAssist'})
            return None

        user = User.objects.get(username=agent_username)
        selected_agent = CobrowseAgent.objects.get(user=user)
        if selected_agent.is_active == False:
            return None

        logger.warning("Details of active agent selected for reassignment to unattended lead are - %s : %s",
                       str(cobrowse_io.session_id), selected_agent.user.username, extra={'AppName': 'EasyAssist'})
        selected_agent.last_lead_assigned_datetime = timezone.now()
        selected_agent.save()
        return selected_agent
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_for_reassignment %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return None


def archive_cobrowse_objects(CobrowseAccessToken, CobrowseIO, SystemAuditTrail):
    try:

        cobrowse_access_token_objs = CobrowseAccessToken.objects.all()

        for cobrowse_access_token_obj in cobrowse_access_token_objs.iterator():

            threshold_minutes = cobrowse_access_token_obj.archive_on_unassigned_time_threshold
            unassigned_time_threshold = (datetime.datetime.today() - datetime.timedelta(
                minutes=threshold_minutes)).astimezone(pytz.timezone(settings.TIME_ZONE))

            # Leads which are not assigned
            cobrowse_io_objs = CobrowseIO.objects.filter(access_token=cobrowse_access_token_obj).filter(
                agent=None, is_archived=False, request_datetime__lt=unassigned_time_threshold)

            for cobrowse_io_obj in cobrowse_io_objs.iterator():
                cobrowse_io_obj.is_archived = True

                if cobrowse_io_obj.session_archived_cause == None:
                    cobrowse_io_obj.session_archived_cause = "UNASSIGNED"
                    cobrowse_io_obj.session_archived_datetime = timezone.now()

                cobrowse_io_obj.save()

                category = "session_closed"
                description = "Session is archived, agant not available for cobrowsing"
                save_system_audit_trail(category, description, cobrowse_io_obj,
                                        cobrowse_io_obj.access_token, SystemAuditTrail, sender=None)

            # Leads which are assigned and inactive

            inactive_archive_time = cobrowse_access_token_obj.archive_on_common_inactivity_threshold
            inactive_archive_time_threshold = (datetime.datetime.now() - datetime.timedelta(
                minutes=inactive_archive_time)).astimezone(pytz.timezone(settings.TIME_ZONE))

            cobrowse_io_objs = CobrowseIO.objects.filter(access_token=cobrowse_access_token_obj).filter(~Q(cobrowsing_start_datetime=None)).filter(
                is_archived=False).filter(Q(last_update_datetime__lte=inactive_archive_time_threshold) | Q(last_agent_update_datetime__lte=inactive_archive_time_threshold))
            for cobrowse_io_obj in cobrowse_io_objs.iterator():
                update_virtual_agent_code(cobrowse_io_obj)
                cobrowse_io_obj.is_archived = True
                cobrowse_io_obj.is_active = False
                cobrowse_io_obj.save()

                category = "session_closed"

                if (cobrowse_io_obj.last_update_datetime).astimezone(pytz.timezone(settings.TIME_ZONE)) < inactive_archive_time_threshold:
                    total_min = (
                        timezone.now() - cobrowse_io_obj.last_update_datetime).total_seconds() // 60
                    description = "Session is archived on inactivity of client for " + \
                        str(total_min) + " minutes"
                    if cobrowse_io_obj.session_archived_cause == None:
                        cobrowse_io_obj.session_archived_cause = "CLIENT_INACTIVITY"
                        cobrowse_io_obj.session_archived_datetime = timezone.now()
                        cobrowse_io_obj.save()
                elif (cobrowse_io_obj.last_agent_update_datetime).astimezone(pytz.timezone(settings.TIME_ZONE)) < inactive_archive_time_threshold:
                    total_min = (timezone.now(
                    ) - cobrowse_io_obj.last_agent_update_datetime).total_seconds() // 60
                    description = "Session is archived on inactivity of agent for " + \
                        str(total_min) + " minutes"
                    if cobrowse_io_obj.session_archived_cause == None:
                        cobrowse_io_obj.session_archived_cause = "AGENT_INACTIVITY"
                        cobrowse_io_obj.session_archived_datetime = timezone.now()
                        cobrowse_io_obj.save()
                else:
                    description = "Session is archived on inactivity"
                    logger.warning("archive_cobrowse_objects should not run", extra={
                                   'AppName': 'EasyAssist'})

                save_system_audit_trail(category, description, cobrowse_io_obj,
                                        cobrowse_io_obj.access_token, SystemAuditTrail, sender=None)
                send_page_refresh_request_to_agent(cobrowse_io_obj.agent)

            # Leads that are assigned and still on agent's dahboard and agent does not attended

            current_archive_time = (datetime.datetime.today() - datetime.timedelta(
                minutes=cobrowse_access_token_obj.auto_archive_cobrowsing_session_timer)).astimezone(pytz.timezone(settings.TIME_ZONE))

            cobrowse_io_objs = CobrowseIO.objects.filter(access_token=cobrowse_access_token_obj, is_archived=False, cobrowsing_start_datetime=None,
                                                         last_update_datetime__lte=current_archive_time)

            for cobrowse_io_obj in cobrowse_io_objs:
                cobrowse_io_obj.is_archived = True
                cobrowse_io_obj.is_active = False

                if not cobrowse_io_obj.session_archived_cause:
                    cobrowse_io_obj.session_archived_cause = "UNATTENDED"
                    cobrowse_io_obj.session_archived_datetime = timezone.now()

                cobrowse_io_obj.save()

                category = "session_closed"
                description = "Session is archived due to agent inactivity"
                save_system_audit_trail(
                    category, description, cobrowse_io_obj, cobrowse_access_token_obj, SystemAuditTrail, sender=None)

                send_page_refresh_request_to_agent(cobrowse_io_obj.agent)
            
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error archive_cobrowse_objects %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def assign_followup_leads_to_agents():

    # logger.info("Into assign_followup_leads_to_agents",
    #                 extra={'AppName': 'EasyAssist'})

    try:

        from EasyAssistApp.models import CobrowseIO, CobrowseAccessToken, CobrowseAgent
        from EasyChatApp.models import User

        access_token_objs = CobrowseAccessToken.objects.filter(is_active=True)

        for access_token_obj in access_token_objs.iterator():

            if access_token_obj.enable_followup_leads_tab == False:
                continue

            cobrowse_agent_admin = access_token_obj.agent

            active_agents = get_list_agents_under_admin(
                cobrowse_agent_admin, is_active=None)

            active_agents = get_agents_with_followup_enabled(active_agents)

            # get all cobrowseio objects which are unassigned or for follow up
            cobrowse_io_objs = CobrowseIO.objects.filter(
                access_token=access_token_obj, is_test=False, request_datetime__date__gte=access_token_obj.go_live_date, agent=None, is_archived=True, is_lead=False)

            if(len(cobrowse_io_objs) != 0):

                for cobrowse_io_obj in cobrowse_io_objs.iterator():

                    if cobrowse_io_obj.supported_language != None or cobrowse_io_obj.product_category != None:
                        agent_followup_list = filter_agents_for_followup(
                            active_agents, cobrowse_io_obj)
                    else:
                        agent_followup_list = active_agents

                    if(len(agent_followup_list) != 0):
                        # a dictionary which stores the agent username as a key
                        # and the value would be the number of sessions alloted
                        # to him for followup
                        agent_session_alloted_count = {}

                        for agent in agent_followup_list:

                            # getting the count of followup/unassigned leads
                            # for the current agent
                            number_of_cobrowseio_alloted = CobrowseIO.objects.filter(
                                access_token=access_token_obj, is_test=False, request_datetime__date__gte=access_token_obj.go_live_date,
                                agent=agent, is_archived=True, is_lead=False, session_archived_cause__in=["UNASSIGNED", "FOLLOWUP"]).count()

                            agent_session_alloted_count[
                                agent.user.username] = number_of_cobrowseio_alloted

                        # sorting the dictionary so that the agent with least
                        # number of alloted leads comes first
                        agent_session_alloted_count = {k: v for k, v in sorted(
                            agent_session_alloted_count.items(), key=lambda item: item[1])}

                        agent_username = list(
                            agent_session_alloted_count.keys())[0]

                        user = User.objects.get(username=agent_username)
                        selected_agent = CobrowseAgent.objects.get(user=user)
                        selected_agent.last_lead_assigned_datetime = timezone.now()
                        selected_agent.save()

                        cobrowse_io_obj.agent = selected_agent
                        cobrowse_io_obj.save()

                    # else:
                    #     logger.info("Agents list for assigning followup/unassigned lead is empty",
                    #                 extra={'AppName': 'EasyAssist'})

            # else:
            #     logger.info("No followup or unassigned leads available",
            #                 extra={'AppName': 'EasyAssist'})

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error assign_followup_leads_to_agents %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_agents_with_followup_enabled(agents_list):
    filtered_agents_list = []
    for agent in agents_list:
        if agent.assign_followup_leads:
            filtered_agents_list.append(agent)

    return filtered_agents_list


def filter_agents_for_followup(agent_list, cobrowse_io_obj, for_meeting=False, support_levels=None):
    agent_list_for_followup = []
    try:
        for agent in agent_list:
            if support_levels and (agent.support_level not in support_levels):
                continue

            if cobrowse_io_obj.access_token.allow_language_support:
                if cobrowse_io_obj.supported_language in agent.supported_language.filter(is_deleted=False):
                    if cobrowse_io_obj.access_token.choose_product_category:
                        if cobrowse_io_obj.product_category in agent.product_category.filter(is_deleted=False):
                            agent_list_for_followup.append(agent)
                    else:
                        agent_list_for_followup.append(agent)
            else:
                if cobrowse_io_obj.access_token.choose_product_category:
                    if cobrowse_io_obj.product_category in agent.product_category.filter(is_deleted=False):
                        agent_list_for_followup.append(agent)
                else:
                    agent_list_for_followup.append(agent)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error filter_free_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        # agent_list_for_followup = agent_list
        agent_list_for_followup = []

    return agent_list_for_followup


def mask_cobrowseio_customer_details(CobrowseDropLink, CobrowseChatHistory):

    try:
        drop_link_objs = CobrowseDropLink.objects.filter(
            is_pii_data_masked=False)
        for drop_link_obj in drop_link_objs.iterator():
            drop_link_obj.mask_customer_details()

        chat_history_objs = CobrowseChatHistory.objects.filter(
            is_pii_data_masked=False)
        for chat_history_obj in chat_history_objs.iterator():
            chat_history_obj.mask_customer_chat()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error mask_cobrowseio_customer_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def archive_meeting_objects(CobrowseVideoConferencing):
    try:
        meeting_io_objs = CobrowseVideoConferencing.objects.filter(
            is_expired=False)
        for meeting_io in meeting_io_objs:
            check_cogno_meet_status(meeting_io)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error archive_meeting_objects %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def mask_cogno_meet_customer_details(CobrowseVideoAuditTrail):
    try:
        audit_trail_objs = CobrowseVideoAuditTrail.objects.filter(
            is_pii_data_masked=False, cobrowse_video__is_expired=True)

        for audit_trail_obj in audit_trail_objs.iterator():
            audit_trail_obj.mask_video_conferencing_data()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error mask_cogno_meet_customer_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def update_resend_password_counter(cobrowse_agent):
    try:
        logger.info("Inside update_resend_password_counter", extra={
            'AppName': 'EasyAssist'})
        current_date = timezone.now().date()

        if cobrowse_agent.last_password_resend_date != current_date:
            cobrowse_agent.last_password_resend_date = timezone.now()
            cobrowse_agent.resend_password_count = RESEND_PASSWORD_THRESHOLD

        if cobrowse_agent.resend_password_count >= 0:
            cobrowse_agent.resend_password_count -= 1
        cobrowse_agent.save()

        logger.info("Successfully Exiting update_resend_password_counter", extra={
            'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error archive_cobrowse_objects %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def copy_files_base_to_token(file_obj_list):
    try:
        for file_obj in file_obj_list:
            src_filepath = "EasyAssistApp/static/EasyAssistApp" + \
                "/" + file_obj["type"] + "/" + file_obj["name"]
            dst_filepath = file_obj["relative_path"]
            static_dst_filepath = get_similar_static_files_path(dst_filepath)
            if os.path.exists(src_filepath):
                content = read_file_content(src_filepath)
                if not os.path.exists(dst_filepath):
                    write_file_content(dst_filepath, content)
                if not os.path.exists(static_dst_filepath):
                    write_file_content(static_dst_filepath, content)
            else:
                logger.error("src_filepath not exist : %s",
                             src_filepath, extra={'AppName': 'EasyAssist'})
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error copy_files_base_to_token %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def create_directory(directory_path, remove_first=False):
    if remove_first and os.path.exists(directory_path):
        try:
            shutil.rmtree(directory_path)
        except OSError as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_directory %s at %s", str(e), str(
                exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)


def create_static_files_access_token_specific(cobrowse_access_token, reset=False):
    try:
        # create folder for token if not exists
        js_token_dir = "EasyAssistApp/static/EasyAssistApp/js/" + \
            str(cobrowse_access_token.key)
        create_directory(js_token_dir, remove_first=reset)
        create_directory(get_similar_static_files_path(
            js_token_dir), remove_first=reset)

        css_token_dir = "EasyAssistApp/static/EasyAssistApp/css/" + \
            str(cobrowse_access_token.key)
        create_directory(css_token_dir, remove_first=reset)
        create_directory(get_similar_static_files_path(
            css_token_dir), remove_first=reset)

        img_token_dir = "EasyAssistApp/static/EasyAssistApp/img/" + \
            str(cobrowse_access_token.key)
        create_directory(img_token_dir, remove_first=reset)
        create_directory(get_similar_static_files_path(
            img_token_dir), remove_first=reset)

        # copy dynamic changing files
        js_file_list = cobrowse_access_token.get_static_file_token_wise_list_specific_type(
            "js")
        copy_files_base_to_token(js_file_list)

        css_file_list = cobrowse_access_token.get_static_file_token_wise_list_specific_type(
            "css")
        copy_files_base_to_token(css_file_list)

        img_file_list = cobrowse_access_token.get_static_file_token_wise_list_specific_type(
            "img")
        copy_files_base_to_token(img_file_list)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_static_files_access_token_specific %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def save_static_file_change_audit_trail(cobrowse_access_token_obj, source_file_path, backup_file_path, active_agent, StaticFileChangeLogger):
    try:
        static_file_change_logger_objs = StaticFileChangeLogger.objects.filter(
            source_file_path=source_file_path).order_by("-update_time")
        max_save_limit = 4
        remove_objs = static_file_change_logger_objs[max_save_limit:]
        for remove_obj in remove_objs:
            if remove_obj.is_deleted == False and os.path.exists(remove_obj.backup_file_path):
                os.remove(remove_obj.backup_file_path)
                remove_obj.is_deleted = True
                remove_obj.save()
        StaticFileChangeLogger.objects.create(agent=active_agent, access_token=cobrowse_access_token_obj,
                                              source_file_path=source_file_path, backup_file_path=backup_file_path)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error save_static_file_change_audit_trail %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def backup_static_file(source_file_path, cobrowse_access_token_obj, active_agent, StaticFileChangeLogger):
    try:
        current_timestamp = str(datetime.datetime.now().timestamp())
        source_file_name = source_file_path.split("/")[-1]

        extention_index = source_file_name.rfind(".")
        extention = source_file_name[extention_index + 1:]
        backup_file_name = source_file_name[:extention_index] + \
            "_" + current_timestamp + "." + extention
        backup_directory = "files/EasyAssistApp/static_backup/" + \
            extention + "/" + str(cobrowse_access_token_obj.key)
        if not os.path.exists(backup_directory):
            os.makedirs(backup_directory)
        backup_file_path = backup_directory + "/" + backup_file_name

        content = read_file_content(source_file_path)
        write_file_content(backup_file_path, content)

        save_static_file_change_audit_trail(
            cobrowse_access_token_obj, source_file_path, backup_file_path, active_agent, StaticFileChangeLogger)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error backup_static_file %s at %s", str(e), str(
            exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def delete_static_file_directory(directory_path):
    try:
        shutil.rmtree(directory_path)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error delete_static_file_directory %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def delete_access_token_specific_static_file(access_token_key):
    try:
        # create folder for token if not exists
        js_token_dir = "EasyAssistApp/static/EasyAssistApp/js/" + \
            str(access_token_key)
        delete_static_file_directory(js_token_dir)
        delete_static_file_directory(
            get_similar_static_files_path(js_token_dir))

        css_token_dir = "EasyAssistApp/static/EasyAssistApp/css/" + \
            str(access_token_key)
        delete_static_file_directory(css_token_dir)
        delete_static_file_directory(
            get_similar_static_files_path(css_token_dir))

        img_token_dir = "EasyAssistApp/static/EasyAssistApp/img/" + \
            str(access_token_key)
        delete_static_file_directory(img_token_dir)
        delete_static_file_directory(
            get_similar_static_files_path(img_token_dir))

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error delete_access_token_specific_static_file %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def read_file_content(file_path):
    file_obj = open(file_path, "r")
    file_content = file_obj.read()
    file_obj.close()
    return file_content


def write_file_content(file_path, content):
    with open(file_path, "w") as file_obj:
        file_obj.write(content)
        file_obj.close()


def get_similar_static_files_path(path):
    static_index = path.find("static")
    new_path = path[static_index:]
    return new_path


def generate_drop_link_with_data(request, client_page_link, active_agent, customer_name, customer_email_id, customer_mobile_number, CobrowseDropLink):
    try:
        form_obj = forms.URLField()
        try:
            form_obj.clean(client_page_link)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error generate_drop_link_with_data %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return None

        cobrowse_drop_link_obj = CobrowseDropLink.objects.create(
            client_page_link=client_page_link,
            agent=active_agent,
            customer_name=customer_name)
        enable_masking_pii_data = active_agent.get_access_token_obj().enable_masking_pii_data
        generated_link = settings.EASYCHAT_HOST_URL + \
            "/easy-assist/initiate-droplink-session" + \
            "?eadKey=" + str(cobrowse_drop_link_obj.key)
        # shorten_tiny_url_obj = UrlShortenTinyurl()
        # generated_link = shorten_tiny_url_obj.shorten(generated_link)

        if customer_email_id != "":
            agent_name = str(active_agent.user.first_name) + " " + str(active_agent.user.last_name)
            if agent_name.strip() == "" or agent_name == "None":
                agent_name = str(active_agent.user.username)
            thread = threading.Thread(target=send_drop_link_over_email, args=(
                customer_email_id, customer_name, agent_name, generated_link, ), daemon=True)
            thread.start()
        
            if enable_masking_pii_data:
                customer_email_id = hashlib.md5(
                    customer_email_id.encode()).hexdigest()
        
        if not enable_masking_pii_data:
            cobrowse_drop_link_obj.is_pii_data_masked = True

        cobrowse_drop_link_obj.customer_mobile = customer_mobile_number
        cobrowse_drop_link_obj.customer_email = customer_email_id
        cobrowse_drop_link_obj.generated_link = generated_link
        cobrowse_drop_link_obj.generate_datetime = timezone.now()
        cobrowse_drop_link_obj.save()
        return cobrowse_drop_link_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error generate_drop_link_with_data %s at %s", str(
            e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return None


def get_authorization_token(client_id, CRMIntegrationModel):
    crm_integration_model = CRMIntegrationModel.objects.filter(
        client_id=client_id).first()
    if crm_integration_model:
        return crm_integration_model.client_key, crm_integration_model.client_iv
    else:
        return None, None


def docx_replace_regex(doc_obj, regex, replace):

    for para in doc_obj.paragraphs:
        if regex.search(para.text):
            inline = para.runs
            # Loop added to work with runs (strings with same style)
            for index in range(len(inline)):
                if regex.search(inline[index].text):
                    text = regex.sub(replace, inline[index].text)
                    inline[index].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


class CRMDocumentGenerator():

    def __init__(self, crm_integration_model, CobrowsingFileAccessManagement):
        self.crm_integration_model = crm_integration_model
        self.CobrowsingFileAccessManagement = CobrowsingFileAccessManagement

    def create_inbound_document(self):
        try:
            # Inbound api doc generation
            inbound_doc_src = "files/templates/easyassist-crm-template/CRM_Inbound.docx"
            document_obj_inbound = Document(inbound_doc_src)

            crm_inbound_url = settings.EASYCHAT_HOST_URL + \
                "/easy-assist/crm/inbound/" + self.crm_integration_model.client_id
            docx_replace_regex(document_obj_inbound, re.compile(
                r"crm_integration_url"), crm_inbound_url)
            crm_integration_token = self.crm_integration_model.auth_token
            docx_replace_regex(document_obj_inbound, re.compile(
                r"crm_integration_token"), crm_integration_token)
            crm_integration_client_name = self.crm_integration_model.client_id
            docx_replace_regex(document_obj_inbound, re.compile(
                r"crm_integration_client_name"), crm_integration_client_name)
            docx_replace_regex(document_obj_inbound, re.compile(
                r"crm_integration_hostname"), settings.EASYCHAT_HOST_URL)

            target_document_folder = CRM_DOCUMENTS_PATH + \
                str(self.crm_integration_model.access_token.agent.user.username)
            create_directory(target_document_folder)

            target_document_path = target_document_folder + "/CRM_Inbound_Api_Doc.docx"
            document_obj_inbound.save(target_document_path)

            file_access_management_obj = self.CobrowsingFileAccessManagement.objects.create(
                file_path="/" + target_document_path, is_public=False, access_token=self.crm_integration_model.access_token)
            inbound_doc = DOWNLOAD_FILE_PATH + \
                str(file_access_management_obj.key)
            return inbound_doc
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_inbound_document %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return ''

    def create_search_lead_document(self):
        try:
            # Searchlead api doc generation
            search_lead_doc_src = "files/templates/easyassist-crm-template/CRM_Search_Lead.docx"
            document_obj_search_lead = Document(search_lead_doc_src)

            crm_search_lead_url = settings.EASYCHAT_HOST_URL + \
                "/easy-assist/crm/search-lead/" + self.crm_integration_model.client_id
            docx_replace_regex(document_obj_search_lead, re.compile(
                r"crm_integration_url"), crm_search_lead_url)
            crm_integration_token = self.crm_integration_model.auth_token
            docx_replace_regex(document_obj_search_lead, re.compile(
                r"crm_integration_token"), crm_integration_token)
            crm_integration_client_name = self.crm_integration_model.client_id
            docx_replace_regex(document_obj_search_lead, re.compile(
                r"crm_integration_client_name"), crm_integration_client_name)
            docx_replace_regex(document_obj_search_lead, re.compile(
                r"crm_integration_hostname"), settings.EASYCHAT_HOST_URL)

            target_document_folder = CRM_DOCUMENTS_PATH + \
                str(self.crm_integration_model.access_token.agent.user.username)
            create_directory(target_document_folder)

            target_document_path = target_document_folder + "/CRM_Search_Lead_Api_Doc.docx"
            document_obj_search_lead.save(target_document_path)

            file_access_management_obj = self.CobrowsingFileAccessManagement.objects.create(
                file_path="/" + target_document_path, is_public=False, access_token=self.crm_integration_model.access_token)
            search_lead_doc = DOWNLOAD_FILE_PATH + \
                str(file_access_management_obj.key)
            return search_lead_doc
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_search_lead_document %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return ''

    def create_drop_link_document(self):
        try:
            # Droplink api doc generation
            drop_link_doc_src = "files/templates/easyassist-crm-template/CRM_Drop_Link.docx"
            document_obj_drop_link = Document(drop_link_doc_src)

            crm_drop_link_url = settings.EASYCHAT_HOST_URL + \
                "/easy-assist/crm/drop-link/" + self.crm_integration_model.client_id
            docx_replace_regex(document_obj_drop_link, re.compile(
                r"crm_integration_url"), crm_drop_link_url)
            crm_integration_token = self.crm_integration_model.auth_token
            docx_replace_regex(document_obj_drop_link, re.compile(
                r"crm_integration_token"), crm_integration_token)
            crm_integration_client_name = self.crm_integration_model.client_id
            docx_replace_regex(document_obj_drop_link, re.compile(
                r"crm_integration_client_name"), crm_integration_client_name)
            docx_replace_regex(document_obj_drop_link, re.compile(
                r"crm_integration_hostname"), settings.EASYCHAT_HOST_URL)

            target_document_folder = CRM_DOCUMENTS_PATH + \
                str(self.crm_integration_model.access_token.agent.user.username)
            create_directory(target_document_folder)

            target_document_path = target_document_folder + "/CRM_Drop_Link_Api_Doc.docx"
            document_obj_drop_link.save(target_document_path)

            file_access_management_obj = self.CobrowsingFileAccessManagement.objects.create(
                file_path="/" + target_document_path, is_public=False, access_token=self.crm_integration_model.access_token)
            drop_link_doc = DOWNLOAD_FILE_PATH + \
                str(file_access_management_obj.key)
            return drop_link_doc
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_drop_link_document %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return ''

    def create_support_history_document(self):
        try:
            # Droplink api doc generation
            support_history_doc_src = "files/templates/easyassist-crm-template/CRM_Support_History.docx"
            document_obj_support_history = Document(support_history_doc_src)

            crm_support_history_url = settings.EASYCHAT_HOST_URL + \
                "/easy-assist/crm/support-history/" + self.crm_integration_model.client_id
            docx_replace_regex(document_obj_support_history, re.compile(
                r"crm_integration_url"), crm_support_history_url)
            crm_integration_token = self.crm_integration_model.auth_token
            docx_replace_regex(document_obj_support_history, re.compile(
                r"crm_integration_token"), crm_integration_token)
            crm_integration_client_name = self.crm_integration_model.client_id
            docx_replace_regex(document_obj_support_history, re.compile(
                r"crm_integration_client_name"), crm_integration_client_name)
            docx_replace_regex(document_obj_support_history, re.compile(
                r"crm_integration_hostname"), settings.EASYCHAT_HOST_URL)

            target_document_folder = CRM_DOCUMENTS_PATH + \
                str(self.crm_integration_model.access_token.agent.user.username)
            create_directory(target_document_folder)

            target_document_path = target_document_folder + \
                "/CRM_Support_History_Api_Doc.docx"
            document_obj_support_history.save(target_document_path)

            file_access_management_obj = self.CobrowsingFileAccessManagement.objects.create(
                file_path="/" + target_document_path, is_public=False, access_token=self.crm_integration_model.access_token)
            support_history_doc = DOWNLOAD_FILE_PATH + \
                str(file_access_management_obj.key)
            return support_history_doc
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_support_history_document %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return ''

    def create_chat_history_document(self):
        try:
            # Droplink api doc generation
            chat_history_doc_src = "files/templates/easyassist-crm-template/CRM_Chat_History.docx"
            document_obj_chat_history = Document(chat_history_doc_src)

            crm_chat_history_url = settings.EASYCHAT_HOST_URL + \
                "/easy-assist/crm/chat-history/" + self.crm_integration_model.client_id
            docx_replace_regex(document_obj_chat_history, re.compile(
                r"crm_integration_url"), crm_chat_history_url)
            crm_integration_token = self.crm_integration_model.auth_token
            docx_replace_regex(document_obj_chat_history, re.compile(
                r"crm_integration_token"), crm_integration_token)
            crm_integration_client_name = self.crm_integration_model.client_id
            docx_replace_regex(document_obj_chat_history, re.compile(
                r"crm_integration_client_name"), crm_integration_client_name)
            docx_replace_regex(document_obj_chat_history, re.compile(
                r"crm_integration_hostname"), settings.EASYCHAT_HOST_URL)

            target_document_folder = CRM_DOCUMENTS_PATH + \
                str(self.crm_integration_model.access_token.agent.user.username)
            create_directory(target_document_folder)

            target_document_path = target_document_folder + "/CRM_Chat_History_Api_Doc.docx"
            document_obj_chat_history.save(target_document_path)

            file_access_management_obj = self.CobrowsingFileAccessManagement.objects.create(
                file_path="/" + target_document_path, is_public=False, access_token=self.crm_integration_model.access_token)
            chat_history_doc = DOWNLOAD_FILE_PATH + \
                str(file_access_management_obj.key)
            return chat_history_doc
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error create_chat_history_document %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
            return ''

    def generate_crm_api_document(self):
        try:
            file_dictionary = {
                "inbound_doc": self.create_inbound_document(),
                "search_lead_doc": self.create_search_lead_document(),
                "drop_link_doc": self.create_drop_link_document(),
                "support_history": self.create_support_history_document(),
                "chat_history": self.create_chat_history_document(),
            }
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            logger.error("Error generate_crm_api_document %s at %s", str(
                e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return json.dumps(file_dictionary)


def analytics_ongoing_meeting_count(cogno_vid_objs, CobrowseVideoAuditTrail):
    current_date = timezone.localtime(timezone.now()).date()
    current_time = timezone.localtime(timezone.now()).time()

    total_ongoing_meeting = 0
    ongoing_meetings = cogno_vid_objs.filter(
        meeting_start_date=current_date, meeting_start_time__lte=current_time, is_expired=False)
    for ongoing_meeting in ongoing_meetings:
        cogno_vid_audit_trail = CobrowseVideoAuditTrail.objects.filter(
            cobrowse_video=ongoing_meeting.meeting_id).first()
        if cogno_vid_audit_trail != None and cogno_vid_audit_trail.is_meeting_ended == False:
            total_ongoing_meeting += 1
    return total_ongoing_meeting


def calculate_agent_online_and_work_time(online_audit_trail_objs, work_audit_trail_objs):

    try:
        common_time_list = []
        for online_audit_trail_obj in online_audit_trail_objs:
            common_time_list.append([
                online_audit_trail_obj.last_online_start_datetime,
                online_audit_trail_obj.last_online_end_datetime,
                "online"
            ])

        for work_audit_trail_obj in work_audit_trail_objs:
            common_time_list.append([
                work_audit_trail_obj.session_start_datetime,
                work_audit_trail_obj.session_end_datetime,
                "work"
            ])

        return calcuate_agent_online_session_common_time(common_time_list)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error calculate_agent_online_and_work_time %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return 0


def calcuate_agent_online_session_common_time(common_time_list):
    try:
        if len(common_time_list) == 0:
            return 0

        common_time_list.sort()
        total_common_time = 0
        past_interval = common_time_list[0]
        for active_interval in common_time_list[1:]:
            is_past_interval_online = (past_interval[2] == "online")
            past_interval_start_time = past_interval[0]
            past_interval_end_time = past_interval[1]

            is_active_interval_online = (active_interval[2] == "online")
            active_interval_start_time = active_interval[0]
            active_interval_end_time = active_interval[1]

            if is_past_interval_online == is_active_interval_online:
                past_interval = active_interval
            else:
                if active_interval_start_time > past_interval_end_time:
                    past_interval = active_interval
                elif active_interval_end_time <= past_interval_end_time:
                    total_common_time += get_common_online_and_work_time(
                        past_interval_start_time, past_interval_end_time,
                        active_interval_start_time, active_interval_end_time)
                else:
                    total_common_time += get_common_online_and_work_time(
                        past_interval_start_time, past_interval_end_time,
                        active_interval_start_time, active_interval_end_time)

                    past_interval = active_interval

        return int(total_common_time)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error calcuate_agent_online_session_common_time %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return 0


def get_common_online_and_work_time(past_interval_start_time, past_interval_end_time, active_interval_start_time, active_interval_end_time):
    try:
        if active_interval_start_time > past_interval_end_time:
            return 0
        elif active_interval_end_time <= past_interval_end_time:
            return (active_interval_end_time - active_interval_start_time).total_seconds()
        elif active_interval_start_time <= past_interval_end_time:
            return (past_interval_end_time - active_interval_start_time).total_seconds()
        return 0
    except Exception:
        return 0


def calculate_agent_cobrowsing_meeting_session_duration(work_audit_trail_objs):
    try:
        total_session_time = 0

        for work_audit_trail_obj in work_audit_trail_objs:
            total_session_time += (work_audit_trail_obj.session_end_datetime -
                                   work_audit_trail_obj.session_start_datetime).total_seconds()
        return int(total_session_time)
    except Exception:
        return 0


def calculate_agent_online_duration(online_audit_trail_objs):
    try:
        total_online_time = 0

        for online_audit_trail_obj in online_audit_trail_objs:
            total_online_time += (online_audit_trail_obj.last_online_end_datetime -
                                  online_audit_trail_obj.last_online_start_datetime).total_seconds()
        return int(total_online_time)
    except Exception:
        return 0


def get_agent_wise_audit_trail(agent_audit_trail_list, cobrowse_agent, agents, agent_online_audit_trail_objs, CobrowseAgentWorkAuditTrail):
    try:
        agent_date_wise_duration = {}
        est = pytz.timezone(settings.TIME_ZONE)

        for audit_trail_obj in agent_audit_trail_list:
            date_string = audit_trail_obj.datetime.astimezone(
                est).strftime(DATE_TIME_FORMAT_3)

            agent_username = str(audit_trail_obj.agent.user.username)

            if date_string not in agent_date_wise_duration:
                agent_date_wise_duration[date_string] = {}

            if agent_username not in agent_date_wise_duration[date_string]:
                agent_date_wise_duration[date_string][agent_username] = {
                    "last_login_time": "-",
                    "last_logout_time": "-",
                    "total_online_duration": 0,
                    "idle_time": 0,
                    "offline_duration": 0,
                }

            if audit_trail_obj.action == COBROWSING_LOGIN_ACTION:
                if agent_date_wise_duration[date_string][agent_username]["last_login_time"] == "-" or agent_date_wise_duration[date_string][agent_username]["last_login_time"] < audit_trail_obj.datetime:
                    agent_date_wise_duration[date_string][agent_username][
                        "last_login_time"] = audit_trail_obj.datetime

            if audit_trail_obj.action == COBROWSING_LOGOUT_ACTION:
                if agent_date_wise_duration[date_string][agent_username]["last_logout_time"] == "-" or agent_date_wise_duration[date_string][agent_username]["last_logout_time"] < audit_trail_obj.datetime:
                    agent_date_wise_duration[date_string][agent_username][
                        "last_logout_time"] = audit_trail_obj.datetime

        agent_work_start_time = cobrowse_agent.get_access_token_obj().start_time
        agent_work_end_time = cobrowse_agent.get_access_token_obj().end_time
        agent_work_start_time = datetime.datetime.combine(
            datetime.date.today(), agent_work_start_time)
        agent_work_end_time = datetime.datetime.combine(
            datetime.date.today(), agent_work_end_time)

        current_time = datetime.datetime.now()

        for date_string in agent_date_wise_duration:
            date_obj = datetime.datetime.strptime(
                date_string, DATE_TIME_FORMAT_3).date()

            if current_time.date() == date_obj:
                if agent_work_end_time > current_time:
                    agent_session_work_end_time = current_time
                else:
                    agent_session_work_end_time = agent_work_end_time
            else:
                agent_session_work_end_time = agent_work_end_time

            agent_working_duration = int(
                (agent_session_work_end_time - agent_work_start_time).total_seconds())

            for agent in agents:
                agent_username = agent.user.username
                agent_supervisors = agent.get_supervisors()

                if agent.role == "admin" or not agent_supervisors:
                    agent_supervisors = "-"
                
                if agent_username not in agent_date_wise_duration[date_string]:
                    continue

                online_audit_trail_objs = agent_online_audit_trail_objs.filter(
                    last_online_start_datetime__date=date_obj, agent=agent)

                # Total Online Duration
                total_online_duration = calculate_agent_online_duration(
                    online_audit_trail_objs)

                # Agent Idle time duration
                work_audit_trail_objs = CobrowseAgentWorkAuditTrail.objects.filter(
                    session_start_datetime__date=date_obj, agent=agent)

                total_common_time = calculate_agent_online_and_work_time(
                    online_audit_trail_objs, work_audit_trail_objs)

                agent_idle_time = total_online_duration - total_common_time

                # Agent offline duration
                online_audit_trail_objs = online_audit_trail_objs.filter(
                    last_online_start_datetime__time__range=[agent_work_start_time.time(), agent_work_end_time.time()]).filter(
                    last_online_end_datetime__time__range=[agent_work_start_time.time(), agent_work_end_time.time()])

                work_audit_trail_objs = work_audit_trail_objs.filter(
                    session_start_datetime__time__range=[agent_work_start_time.time(), agent_work_end_time.time()]).filter(
                    session_end_datetime__time__range=[agent_work_start_time.time(), agent_work_end_time.time()])
                total_common_time = calculate_agent_online_and_work_time(
                    online_audit_trail_objs, work_audit_trail_objs)
                total_work_time_online_duration = calculate_agent_online_duration(
                    online_audit_trail_objs)
                total_session_duration = calculate_agent_cobrowsing_meeting_session_duration(
                    work_audit_trail_objs)
                agent_offline_duration = agent_working_duration - \
                    (total_work_time_online_duration + total_session_duration) + \
                    total_common_time

                last_login_time = agent_date_wise_duration[
                    date_string][agent_username]["last_login_time"]
                last_logout_time = agent_date_wise_duration[
                    date_string][agent_username]["last_logout_time"]

                if last_login_time != "-" and last_logout_time != "-" and last_logout_time < last_login_time:
                    last_logout_time = "-"

                if last_login_time != "-":
                    last_login_time = last_login_time.astimezone(
                        est).strftime(DATE_TIME_FORMAT_4)

                if last_logout_time != "-":
                    last_logout_time = last_logout_time.astimezone(
                        est).strftime(DATE_TIME_FORMAT_4)

                agent_date_wise_duration[date_string][
                    agent_username]["idle_time"] = agent_idle_time
                agent_date_wise_duration[date_string][agent_username][
                    "offline_duration"] = agent_offline_duration
                agent_date_wise_duration[date_string][agent_username]["total_online_duration"] = total_online_duration + \
                    total_session_duration - total_common_time
                agent_date_wise_duration[date_string][agent_username][
                    "last_login_time"] = last_login_time
                agent_date_wise_duration[date_string][agent_username][
                    "last_logout_time"] = last_logout_time
                agent_date_wise_duration[date_string][agent_username][
                    "agent_supervisors"] = agent_supervisors

        agent_wise_audit_trail = []
        for date_string in agent_date_wise_duration:
            date_obj = datetime.datetime.strptime(
                date_string, DATE_TIME_FORMAT_3).date()
            display_date_obj = datetime.datetime.strftime(date_obj, "%d-%b-%Y")
            date_obj = datetime.datetime.strftime(date_obj, DATE_TIME_FORMAT)
            for agent_username in agent_date_wise_duration[date_string]:
                agent_wise_audit_trail.append({
                    "agent_username": agent_username,
                    "date": display_date_obj,
                    "duration": convert_seconds_to_hours_minutes(agent_date_wise_duration[date_string][agent_username]["total_online_duration"]),
                    "idle_time": convert_seconds_to_hours_minutes(agent_date_wise_duration[date_string][agent_username]["idle_time"]),
                    "offline_duration": convert_seconds_to_hours_minutes(agent_date_wise_duration[date_string][agent_username]["offline_duration"]),
                    "last_login_time": agent_date_wise_duration[date_string][agent_username]["last_login_time"],
                    "last_logout_time": agent_date_wise_duration[date_string][agent_username]["last_logout_time"],
                    "formatted_date": date_obj,
                    "agent_supervisors": agent_date_wise_duration[date_string][agent_username]["agent_supervisors"],
                })
        return agent_wise_audit_trail
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_agent_wise_audit_trail %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return []


def get_online_duration_within_agent_login_session(login_datetime, logout_datetime, online_start_datetime, online_end_datetime):

    try:
        if logout_datetime < online_start_datetime:
            return 0

        if login_datetime > online_end_datetime:
            return 0

        if login_datetime > online_start_datetime:
            online_start_datetime = login_datetime

        if logout_datetime < online_end_datetime:
            online_end_datetime = logout_datetime

        return int((online_end_datetime - online_start_datetime).total_seconds())
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_online_duration_within_agent_login_session %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return 0


def get_session_duration_within_agent_login_session(login_datetime, logout_datetime, session_start_datetime, session_end_datetime):

    try:
        if logout_datetime < session_start_datetime:
            return 0

        if login_datetime > session_end_datetime:
            return 0

        if login_datetime > session_start_datetime:
            session_start_datetime = login_datetime

        if logout_datetime < session_end_datetime:
            session_end_datetime = logout_datetime

        return int((session_end_datetime - session_start_datetime).total_seconds())
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_session_duration_within_agent_login_session %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return 0


def get_time_duration_with_login_session(login_datetime, logout_datetime, start_datetime, end_datetime):
    try:
        if logout_datetime < start_datetime:
            return None

        if login_datetime > end_datetime:
            return None

        if login_datetime > start_datetime:
            start_datetime = login_datetime

        if logout_datetime < end_datetime:
            end_datetime = logout_datetime

        return [start_datetime, end_datetime]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_time_duration_with_login_session %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return None


def get_agent_wise_online_audit_trail(cobrowse_agent, agents, agent_online_audit_trail_objs, CobrowseAgentWorkAuditTrail):
    try:
        agent_date_wise_duration = {}
        est = pytz.timezone(settings.TIME_ZONE)

        for online_audit_trail_obj in agent_online_audit_trail_objs:
            date_string = online_audit_trail_obj.last_online_start_datetime.astimezone(
                est).strftime("%d-%b-%Y")

            agent_username = str(online_audit_trail_obj.agent.user.username)

            agent_supervisors = online_audit_trail_obj.agent.get_supervisors()

            if online_audit_trail_obj.agent.role == "admin" or not agent_supervisors:
                agent_supervisors = "-"

            online_time = online_audit_trail_obj.last_online_start_datetime

            offline_time = online_audit_trail_obj.last_online_end_datetime

            online_duration = calculate_agent_online_duration(
                [online_audit_trail_obj])

            idle_time = online_audit_trail_obj.idle_time

            if date_string not in agent_date_wise_duration:
                agent_date_wise_duration[date_string] = {}

            if agent_username not in agent_date_wise_duration[date_string]:
                agent_date_wise_duration[date_string][agent_username] = []

            agent_date_wise_duration[date_string][agent_username].append({
                "online_time": online_time,
                "offline_time": offline_time,
                "online_duration": online_duration,
                "idle_time": idle_time,
                "agent_supervisors": agent_supervisors,
            })

        agent_wise_online_audit_trail = []
        for date_string in agent_date_wise_duration:
            for agent_username in agent_date_wise_duration[date_string]:
                audit_trail_objs = agent_date_wise_duration[
                    date_string][agent_username]

                for audit_trail_obj in audit_trail_objs:
                    online_time = audit_trail_obj["online_time"]
                    offline_time = audit_trail_obj["offline_time"]
                    online_duration = audit_trail_obj["online_duration"]
                    idle_time = audit_trail_obj["idle_time"]
                    agent_supervisors = audit_trail_obj["agent_supervisors"]

                    online_time = online_time.astimezone(
                        est).strftime(DATE_TIME_FORMAT_4)

                    offline_time = offline_time.astimezone(
                        est).strftime(DATE_TIME_FORMAT_4)

                    online_duration = get_formatted_time_from_seconds(
                        online_duration)
                    idle_time = get_formatted_time_from_seconds(idle_time)

                    agent_wise_online_audit_trail.append({
                        "date": date_string,
                        "agent_username": agent_username,
                        "online_time": online_time,
                        "offline_time": offline_time,
                        "online_duration": online_duration,
                        "idle_time": idle_time,
                        "agent_supervisors": agent_supervisors,
                    })

        return agent_wise_online_audit_trail
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_agent_wise_online_audit_trail %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return []


def update_online_audit_trail_idle_time(agent, online_audit_trail_obj, CobrowseAgentWorkAuditTrail):
    try:
        audit_trail_obj_date = online_audit_trail_obj.last_online_start_datetime.date()
        work_audit_trail_objs = CobrowseAgentWorkAuditTrail.objects.filter(
            agent=agent,
            session_start_datetime__date=audit_trail_obj_date)

        online_session_common_time = calculate_agent_online_and_work_time(
            [online_audit_trail_obj], work_audit_trail_objs)

        total_online_duration = calculate_agent_online_duration(
            [online_audit_trail_obj])

        online_audit_trail_obj.idle_time = total_online_duration - online_session_common_time
        online_audit_trail_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error update_online_audit_trail_idle_time %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_hashed_data(message):
    final_string = message
    try:
        if message == None:
            return message

        # pan
        pan_regex = r"\b[a-zA-Z]{3}[pP][a-zA-Z][1-9][0-9]{3}[a-zA-Z]\b"
        if re.match(pan_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        # email id
        email_regex = r"\S+@\S+"
        if re.match(email_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        # dates
        # Regex date format: dd/mm/yyyy
        date_regex = r"[\d]{1,2}/[\d]{1,2}/[\d]{2,4}"
        if re.match(date_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        date_regex = r"[\d]{1,2}-[\d]{1,2}-[\d]{2}"
        if re.match(date_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        date_regex = r"[\d]{1,2} [ADFJMNOS]\w* [\d]{2,4}"
        if re.match(date_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        date_regex = r"[ADFJMNOS]\w* [\d]{1,2} [\d]{2,4}"
        if re.match(date_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        # mobile number, account number, aadhar number (10-12 digits number
        # should have space before and after)
        num_regex = r"\b[0-9]{10,12}\b"
        if re.match(num_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        # age and address (1-3 digits number
        # should have space before and after)
        age_regex = r"\b[0-9]{1,3}\b"
        if re.match(age_regex, message):
            message = hashlib.sha256(message.encode()).hexdigest()

        # customer_id (any string that contains atleast 1 digit)
        id_regex = r"\b[A-Za-z0-9]*\d[A-Za-z0-9]*\b"
        if re.match(id_regex, message) and len(message) != 64:
            message = hashlib.sha256(message.encode()).hexdigest()

        final_string = message
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_hashed_data: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return final_string


def hash_crucial_info_in_data(message):
    final_string = message
    try:
        if message == None:
            return message

        # pan
        pan_pattern = re.findall(
            r"\b[a-zA-Z]{3}[pP][a-zA-Z][1-9][0-9]{3}[a-zA-Z]\b", message)
        for item in pan_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # email id
        email_pattern = re.findall(r"\S+@\S+", message)
        for item in email_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # dates
        # Regex date format: dd/mm/yyyy
        date_format_ddmmyyyy = re.findall(
            r"[\d]{1,2}/[\d]{1,2}/[\d]{2,4}", message)
        for item in date_format_ddmmyyyy:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # Regex date format: dd-mm-yy
        date_format_ddmmyyyy_two = re.findall(
            r"[\d]{1,2}-[\d]{1,2}-[\d]{2}", message)
        for item in date_format_ddmmyyyy_two:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # Regex date format: dd AUG YYYY
        date_format_ddmonthnameyyyy = re.findall(
            r"[\d]{1,2} [ADFJMNOS]\w* [\d]{2,4}", message)
        for item in date_format_ddmonthnameyyyy:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # Regex date format: AUG dd YYYY
        date_format_monthnameddyyyy = re.findall(
            r"[ADFJMNOS]\w* [\d]{1,2} [\d]{2,4}", message)
        for item in date_format_monthnameddyyyy:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # mobile number, account number, aadhar number (10-12 digits number
        # should have space before and after)
        mobile_pattern = re.findall(r"\b[0-9]{10,12}\b", message)
        for item in mobile_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # age and address (1-3 digits number
        # should have space before and after)
        age_pattern = re.findall(r"\b[0-9]{1,3}\b", message)
        for item in age_pattern:
            message = message.replace(
                item, hashlib.sha256(item.encode()).hexdigest())

        # customer_id (any string that contains atleast 1 digit)
        id_pattern = re.findall(r"\b[A-Za-z0-9]*\d[A-Za-z0-9]*\b", message)
        for item in id_pattern:
            if len(item) == 64:
                continue
            reg = r"\b" + str(item) + r"\b"
            message = re.sub(reg, hashlib.sha256(
                item.encode()).hexdigest(), message)

        final_string = message
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("hash_crucial_info_in_data: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return final_string


def get_masked_data_if_hashed(message):
    final_string = message
    try:
        reg_pattern = r"[a-fA-F0-9]{64}"
        hash_pattern = re.findall(reg_pattern, message)
        for item in hash_pattern:
            message = message.replace(
                item, mask_hashed_data(item))

        final_string = message
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_masked_data_if_hashed: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return final_string


def mask_hashed_data(hash_message):
    final_string = hash_message
    try:
        final_string = hash_message[0:5] + "***" + hash_message[-2:]
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("mask_hashed_data: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return final_string


def disable_inbound_cobrowsing_features(cobrowse_access_token):
    try:
        if cobrowse_access_token.enable_inbound == False:
            cobrowse_access_token.show_floating_easyassist_button = False
            cobrowse_access_token.show_easyassist_connect_agent_icon = False
            cobrowse_access_token.choose_product_category = False
            cobrowse_access_token.field_stuck_event_handler = False
            cobrowse_access_token.allow_popup_on_browser_leave = False
            cobrowse_access_token.enable_followup_leads_tab = False
            cobrowse_access_token.enable_low_bandwidth_cobrowsing = False
            cobrowse_access_token.enable_manual_switching = False
            cobrowse_access_token.allow_connect_with_virtual_agent_code = False
            cobrowse_access_token.connect_with_virtual_agent_code_mandatory = False
            cobrowse_access_token.save()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("disable_inbound_cobrowsing_features: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def parse_cobrowse_io_data(cobrowse_io_objs, agents, cobrowse_agent):
    from EasyAssistApp.models import CobrowseIOTransferredAgentsLogs
    try:
        cobrowse_io_obj_details = []
        unattended_lead_transfer_audit_trail = []
        for cobrowse_io_obj in cobrowse_io_objs:
            is_last_agent = False
            is_lead_reassigned = False
            
            transfer_lead_time_details = ""
            if cobrowse_io_obj.is_agent_connected:
                is_lead_reassigned = False
            
            if cobrowse_io_obj.access_token.enable_session_transfer_in_cobrowsing:
                if cobrowse_io_obj.is_transfer_request_timer_exhausted():
                    continue                    
                else:
                    transfer_lead_time_details = cobrowse_io_obj.get_transfer_request_remaining_time()
            cobrowse_transferred_agent_logs = None
            cobrowse_agent_logs = None
            if cobrowse_io_obj.access_token.enable_session_transfer_in_cobrowsing:
                cobrowse_transferred_agent_logs = CobrowseIOTransferredAgentsLogs.objects.filter(
                    cobrowse_io=cobrowse_io_obj, cobrowse_request_type="transferred", transferred_agent__in=agents, transferred_status="").order_by("-log_request_datetime").first()
            if cobrowse_io_obj.access_token.enable_invite_agent_in_cobrowsing:
                cobrowse_agent_logs = CobrowseIOTransferredAgentsLogs.objects.filter(
                    cobrowse_io=cobrowse_io_obj, cobrowse_request_type="invited", transferred_agent__in=agents, invited_status="").order_by("-log_request_datetime").first()
            lead_type = ""
            inviting_agent = ""
            invited_agent = ""
            transferred_agent = ""
            if cobrowse_transferred_agent_logs:
                lead_type = "transferred"
                is_lead_reassigned = False
                transferred_agent = cobrowse_transferred_agent_logs.transferred_agent.user.username
            elif cobrowse_agent_logs:
                lead_type = "invited"
                is_lead_reassigned = False
                inviting_agent = cobrowse_agent_logs.inviting_agent.user.username
                invited_agent = cobrowse_agent_logs.transferred_agent.user.username
            
            if cobrowse_io_obj.access_token.enable_auto_assign_unattended_lead and cobrowse_io_obj.is_lead == False and \
                cobrowse_io_obj.is_reverse_cobrowsing == False and cobrowse_io_obj.is_droplink_lead == False and \
                    cobrowse_io_obj.cobrowsing_type != "modified-inbound":

                if cobrowse_io_obj.access_token.enable_auto_assign_to_one_agent:
                    if cobrowse_io_obj.is_unattended_lead_timer_elapsed():
                        continue
                
                max_lead_reassignment_count = cobrowse_io_obj.access_token.unattended_lead_auto_assignment_counter

                if cobrowse_io_obj.lead_reassignment_counter == max_lead_reassignment_count or \
                        cobrowse_io_obj.is_agent_request_for_cobrowsing:
                    is_last_agent = True

                if cobrowse_io_obj.is_auto_assign_timer_exhausted():
                    if cobrowse_io_obj.lead_reassignment_counter != max_lead_reassignment_count and \
                        cobrowse_io_obj.is_agent_request_for_cobrowsing == False and not \
                            cobrowse_io_obj.is_cobrowse_session_initiated():
                        continue
                
                unattended_lead_transfer_audit_trail = cobrowse_io_obj.get_unattended_lead_transfer_audit_trail()
                if len(unattended_lead_transfer_audit_trail) > 1:
                    is_lead_reassigned = True

            est = pytz.timezone(settings.TIME_ZONE)

            last_update_datetime = None
            if cobrowse_io_obj.last_update_datetime:
                last_update_datetime = cobrowse_io_obj.last_update_datetime.astimezone(
                    est).strftime(TIME_FORMAT)

            request_datetime = cobrowse_io_obj.request_datetime.astimezone(
                est).strftime("%d-%b-%Y %I:%M %p")

            cobrowsing_start_datetime = None
            if cobrowse_io_obj.cobrowsing_start_datetime:
                cobrowsing_start_datetime = cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                    est).strftime("%d-%b-%Y %I:%M %p")

            meeting_start_datetime = None
            if cobrowse_io_obj.meeting_start_datetime:
                meeting_start_datetime = cobrowse_io_obj.meeting_start_datetime.astimezone(
                    est).strftime(TIME_FORMAT)

            is_active = cobrowse_io_obj.is_active_timer() and cobrowse_io_obj.is_active

            show_verification_code_modal = cobrowse_io_obj.access_token.show_verification_code_modal

            OTP = "-"
            if is_active and cobrowse_io_obj.otp_validation != None and show_verification_code_modal == True:
                OTP = cobrowse_io_obj.otp_validation

            product_category_title = None
            product_category_pk = None
            supported_language_title = None
            supported_language_pk = None

            if cobrowse_io_obj.product_category:
                product_category_title = cobrowse_io_obj.product_category.title
                product_category_pk = cobrowse_io_obj.product_category.pk

            if cobrowse_io_obj.supported_language:
                supported_language_title = cobrowse_io_obj.supported_language.title
                supported_language_pk = cobrowse_io_obj.supported_language.pk

            agent_username = "-"
            if cobrowse_io_obj.agent:
                agent_username = cobrowse_io_obj.agent.user.username
                agent_supervisors = cobrowse_io_obj.agent.get_supervisors()

            if is_last_agent or cobrowse_io_obj.is_cobrowse_session_initiated() or cobrowse_io_obj.cobrowsing_type == "modified-inbound":
                auto_assign_time_details = {
                    "remaining_time_str": "-",
                    "is_highlight_required": False
                }
            else:
                auto_assign_time_details = cobrowse_io_obj.get_readable_auto_assign_remaining_time()

            lead_initiated_by = "-"
            if cobrowse_io_obj.lead_initiated_by == 'greeting_bubble':
                lead_initiated_by = "Greeting Bubble"
            elif cobrowse_io_obj.lead_initiated_by == 'floating_button':
                lead_initiated_by = "Floating Button"
            elif cobrowse_io_obj.lead_initiated_by == 'icon':
                lead_initiated_by = "Icon"
            elif cobrowse_io_obj.lead_initiated_by == 'inactivity_popup':
                lead_initiated_by = "Inactivity Pop-up"
            elif cobrowse_io_obj.lead_initiated_by == 'exit_intent':
                lead_initiated_by = "Exit Intent"

            cobrowse_io_obj_details.append({
                "otp": OTP,
                "is_active": is_active,
                "is_lead": cobrowse_io_obj.is_lead,
                "session_id": str(cobrowse_io_obj.session_id),
                "full_name": cobrowse_io_obj.full_name,
                "mobile_number": get_masked_data_if_hashed(
                    cobrowse_io_obj.mobile_number),
                "get_sync_data": cobrowse_io_obj.get_sync_data(),
                "request_datetime": request_datetime,
                "last_update_datetime": last_update_datetime,
                "cobrowsing_start_datetime": cobrowsing_start_datetime,
                "agent_username": agent_username,
                "agent_email": cobrowse_io_obj.agent.user.email,
                "product_url": cobrowse_io_obj.product_url(),
                "product_name": cobrowse_io_obj.product_name(),
                "product_category": product_category_title,
                "product_category_pk": product_category_pk,
                "supported_language": supported_language_title,
                "supported_language_pk": supported_language_pk,
                "total_time_spent": cobrowse_io_obj.total_time_spent(),
                "agent_assistant_request_status": cobrowse_io_obj.agent_assistant_request_status,
                "agent_meeting_request_status": cobrowse_io_obj.agent_meeting_request_status,
                "allow_agent_meeting": cobrowse_io_obj.allow_agent_meeting,
                "allow_agent_cobrowse": cobrowse_io_obj.allow_agent_cobrowse,
                "meeting_start_datetime": meeting_start_datetime,
                "is_droplink_lead": cobrowse_io_obj.is_droplink_lead,
                "is_agent_request_for_cobrowsing": cobrowse_io_obj.is_agent_request_for_cobrowsing,
                "agent_notified_count": cobrowse_io_obj.agent_notified_count,
                "is_reverse_cobrowsing": cobrowse_io_obj.is_reverse_cobrowsing,
                "share_client_session": cobrowse_io_obj.share_client_session,
                "show_verification_code_modal": show_verification_code_modal,
                "allow_agent_to_customer_cobrowsing": cobrowse_io_obj.access_token.allow_agent_to_customer_cobrowsing,
                "choose_product_category": cobrowse_io_obj.access_token.choose_product_category,
                "enable_tag_based_assignment_for_outbound": cobrowse_io_obj.access_token.enable_tag_based_assignment_for_outbound,
                "allow_language_support": cobrowse_io_obj.access_token.allow_language_support,
                "enable_verification_code_popup": cobrowse_io_obj.access_token.enable_verification_code_popup,
                "allow_cobrowsing_meeting": cobrowse_io_obj.access_token.allow_cobrowsing_meeting,
                "allow_video_meeting_only": cobrowse_io_obj.access_token.allow_video_meeting_only,
                "allow_screen_sharing_cobrowse": cobrowse_io_obj.access_token.allow_screen_sharing_cobrowse,
                "agent_supervisors": agent_supervisors,
                "enable_auto_assign_unattended_lead": cobrowse_io_obj.access_token.enable_auto_assign_unattended_lead,
                "auto_assign_time_details": auto_assign_time_details,
                "unattended_lead_transfer_audit_trail": unattended_lead_transfer_audit_trail,
                "lead_initiated_by": lead_initiated_by,
                "enable_session_transfer_in_cobrowsing": cobrowse_io_obj.access_token.enable_session_transfer_in_cobrowsing,
                "enable_lead_status": cobrowse_io_obj.access_token.enable_lead_status,
                "lead_type": lead_type,
                "is_lead_reassigned": is_lead_reassigned,
                "inviting_agent": inviting_agent,
                "invited_agent": invited_agent,
                "transfer_lead_time_details": transfer_lead_time_details,
                "transferred_agent": transferred_agent,
                "agent_role": cobrowse_agent.role,
                # "show_verification_code": cobrowse_io_obj.access_token.enable_verification_code_popup,
            })

        return cobrowse_io_obj_details
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_cobrowse_io_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return []


def parse_cobrowse_io_screenshot_info(cobrowse_io_obj, CobrowsingFileAccessManagement):

    meta_information_list = []
    try:
        cobrowse_meta_objs = cobrowse_io_obj.get_cobrowsing_session_meta_data().order_by('-datetime')

        for cobrowse_meta_obj in cobrowse_meta_objs:

            est = pytz.timezone(settings.TIME_ZONE)
            datetime = cobrowse_meta_obj.datetime.astimezone(
                est).strftime(TIME_FORMAT)

            file_path = cobrowse_meta_obj.content
            file_name = file_path.split("/")[-1]

            file_access_management_obj = CobrowsingFileAccessManagement.objects.filter(
                file_path="/" + file_path, is_public=True, original_file_name=file_name).first()

            if file_access_management_obj == None:
                file_access_management_obj = CobrowsingFileAccessManagement.objects.create(
                    file_path="/" + file_path, is_public=True, original_file_name=file_name, access_token=cobrowse_io_obj.agent.get_access_token_obj())

            public_path = settings.EASYCHAT_HOST_URL + \
                '/easy-assist/download-file/' + \
                str(file_access_management_obj.pk)

            meta_information_list.append({
                "url": public_path,
                "datetime": datetime
            })
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_cobrowse_io_screenshot_info %s at %s for session_id %s",
                     str(e), str(exc_tb.tb_lineno), str(cobrowse_io_obj.session_id), extra={'AppName': 'EasyAssist'})
        meta_information_list = []

    return meta_information_list


def parse_cobrowse_io_data_public_api(cobrowse_io_obj, CobrowsingFileAccessManagement):

    session_data = {}
    try:

        access_token = cobrowse_io_obj.access_token

        est = pytz.timezone(settings.TIME_ZONE)

        request_datetime = cobrowse_io_obj.request_datetime.astimezone(
            est).strftime(TIME_FORMAT)

        cobrowsing_start_datetime = None
        if cobrowse_io_obj.cobrowsing_start_datetime:
            cobrowsing_start_datetime = cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                est).strftime(TIME_FORMAT)

        last_agent_update_datetime = None
        if cobrowse_io_obj.last_agent_update_datetime:
            last_agent_update_datetime = cobrowse_io_obj.last_agent_update_datetime.astimezone(
                est).strftime(TIME_FORMAT)

        product_category_title = ''
        supported_language_title = ''

        if cobrowse_io_obj.product_category:
            product_category_title = cobrowse_io_obj.product_category.title

        if cobrowse_io_obj.supported_language:
            supported_language_title = cobrowse_io_obj.supported_language.title

        customer_name = ''
        mobile_number = ''
        session_started_by = ''

        if cobrowse_io_obj.is_lead == False:
            customer_name = cobrowse_io_obj.full_name
            mobile_number = get_masked_data_if_hashed(
                cobrowse_io_obj.mobile_number)
            session_started_by = 'Customer'
        else:
            customer_name = cobrowse_io_obj.get_sync_data_value("name")
            mobile_number = get_masked_data_if_hashed(
                cobrowse_io_obj.get_sync_data_value("mobile"))
            session_started_by = 'Agent'

        agent_id = ''
        agent_name = ''
        if cobrowse_io_obj.agent:
            agent_id = cobrowse_io_obj.agent.name()
            agent_name = cobrowse_io_obj.agent.agent_name()

        session_ended_by = ''
        if cobrowse_io_obj.session_archived_cause:
            if cobrowse_io_obj.session_archived_cause in ["AGENT_ENDED", "AGENT_INACTIVITY"]:
                session_ended_by = "Agent"
            elif cobrowse_io_obj.session_archived_cause in ["CLIENT_ENDED", "CLIENT_INACTIVITY"]:
                session_ended_by = "Customer"
            else:
                session_ended_by = cobrowse_io_obj.session_archived_cause[0].upper(
                ) + cobrowse_io_obj.session_archived_cause[1:].lower()

        agent_nps = 'Not provided'
        if cobrowse_io_obj.agent_rating != None:
            agent_nps = cobrowse_io_obj.agent_rating

        customer_comment = 'Not provided'
        if cobrowse_io_obj.client_comments:
            customer_comment = cobrowse_io_obj.client_comments

        agent_remarks = 'Not provided'
        if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
            agent_remarks = cobrowse_io_obj.get_cobrowsing_session_closing_comments(
            ).first().agent_comments

        agent_comments = 'Not provided'
        if access_token is not None and access_token.enable_predefined_remarks:
            agent_comments = ''

        lead_type = cobrowse_io_obj.get_lead_type()
        lead_status = ''
        if cobrowse_io_obj.is_archived and cobrowse_io_obj.is_helpful:
            lead_status = 'Converted'
        else:
            lead_status = 'Not Converted'

        screenshots = parse_cobrowse_io_screenshot_info(
            cobrowse_io_obj, CobrowsingFileAccessManagement)

        session_data = {
            "session_id": str(cobrowse_io_obj.session_id),
            "agent_id": agent_id,
            "agent_name": agent_name,
            "customer_name": customer_name,
            "customer_mobile_number": mobile_number,
            "product_url": cobrowse_io_obj.product_url(),
            "product_title": cobrowse_io_obj.product_name(),
            "product_category": product_category_title,
            "support_language": supported_language_title,
            "request_datetime": request_datetime,
            "session_start_datetime": cobrowsing_start_datetime,
            "session_end_datetime": last_agent_update_datetime,
            "customer_wait_time_seconds": cobrowse_io_obj.customer_wait_time_seconds(),
            "total_time_spent_seconds": cobrowse_io_obj.total_time_spent_seconds(),
            "session_started_by": session_started_by,
            "session_ended_by": session_ended_by,
            "agent_nps": agent_nps,
            "customer_comment": customer_comment,
            "agent_remarks": agent_remarks,
            "agent_comments": agent_comments,
            "lead_type": lead_type,
            "lead_status": lead_status,
            "screenshots": screenshots,
        }

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_cobrowse_io_data_public_api %s at %s for session_id %s",
                     str(e), str(exc_tb.tb_lineno), str(cobrowse_io_obj.session_id), extra={'AppName': 'EasyAssist'})
        session_data = {}

    return session_data


def parse_cobrowse_io_chat_history(cobrowse_io_obj):

    def get_absolute_url(url):
        if url and (not url.startswith('http')):
            if url[0] == '/':
                url = settings.EASYCHAT_HOST_URL + url
            else:
                url = settings.EASYCHAT_HOST_URL + '/' + url
        return url

    chat_history = []
    try:
        cobrowsing_chat_history_objs = cobrowse_io_obj.get_cobrowsing_chat_history()

        chat_history = []
        for cobrowsing_chat_history_obj in cobrowsing_chat_history_objs:

            est = pytz.timezone(settings.TIME_ZONE)
            datetime = cobrowsing_chat_history_obj.datetime.astimezone(
                est).strftime(TIME_FORMAT)

            sender_id = "customer"
            sender_name = "customer"
            sender_role = "customer"
            if cobrowsing_chat_history_obj.sender != None:
                sender_id = cobrowsing_chat_history_obj.sender.name()
                sender_name = cobrowsing_chat_history_obj.sender.agent_name()
                sender_role = "AdminAgent"
                if sender_id != cobrowse_io_obj.agent.name():
                    sender_role = "InvitedAgent"

            chat_history.append({
                "chat_type": cobrowsing_chat_history_obj.chat_type,
                "sender_id": sender_id,
                "sender_name": sender_name,
                "sender_role": sender_role,
                "message": get_masked_data_if_hashed(cobrowsing_chat_history_obj.message),
                "datetime": datetime,
                "attachment": get_absolute_url(cobrowsing_chat_history_obj.attachment),
                "attachment_file_name": cobrowsing_chat_history_obj.attachment_file_name,
            })

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_cobrowse_io_chat_history %s at %s for session_id %s",
                     str(e), str(exc_tb.tb_lineno), str(cobrowse_io_obj.session_id), extra={'AppName': 'EasyAssist'})
        chat_history = []

    return chat_history


def update_inbound_analytics_model(CobrowseDateWiseInboundAnalytics, CobrowseIO, EasyAssistCustomer, CobrowseAccessToken, date):
    try:
        access_token_objs = CobrowseAccessToken.objects.all()

        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False, is_archived=True, is_reverse_cobrowsing=False, is_lead=False, request_datetime__date=date).exclude(
            cobrowsing_type="outbound-proxy-cobrowsing")

        request_not_initiated_objs = EasyAssistCustomer.objects.filter(
            request_datetime__date=date, cobrowse_io=None)

        for access_token_obj in access_token_objs.iterator():
            admin_agent = access_token_obj.agent
            if admin_agent == None:
                continue

            agent_objs = get_list_agents_under_admin(
                admin_agent, is_active=None, is_account_active=None)
            agent_objs += get_list_supervisor_under_admin(admin_agent, None)

            filtered_cobrowse_io_objs = cobrowse_io_objs.filter(
                access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date)

            total_request_not_initiated = request_not_initiated_objs.filter(
                access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date).count()
            for agent in agent_objs:
                analytics_data = get_agent_wise_inbound_analytics(
                    filtered_cobrowse_io_objs, agent)

                if analytics_data == None:
                    continue

                cobrowse_analytics_obj = CobrowseDateWiseInboundAnalytics.objects.filter(
                    date=date, agent=agent).first()
                if cobrowse_analytics_obj == None:
                    cobrowse_analytics_obj = CobrowseDateWiseInboundAnalytics.objects.create(
                        agent=agent,
                        date=date,
                        request_initiated=analytics_data["request_initiated"],
                        request_initiated_by_floating_button_icon=analytics_data[
                            "request_initiated_by_floating_button_icon"],
                        request_initiated_by_greeting_bubble=analytics_data[
                            "request_initiated_by_greeting_bubble"],
                        request_initiated_by_exit_intent=analytics_data[
                            "request_initiated_by_exit_intent"],
                        request_initiated_by_inactivity_popup=analytics_data[
                            "request_initiated_by_inactivity_popup"],
                        request_attended=analytics_data["request_attended"],
                        request_unattended=analytics_data[
                            "request_unattended"],
                        followup_leads=analytics_data["followup_leads"],
                        customers_converted=analytics_data[
                            "customers_converted"],
                        customers_converted_by_url=analytics_data[
                            "customers_converted_by_url"],
                        declined_leads=analytics_data["declined_leads"],
                        total_session_time=analytics_data[
                            "total_session_time"],
                        attended_leads_total_wait_time=analytics_data[
                            "attended_leads_total_wait_time"],
                        unattended_leads_total_wait_time=analytics_data[
                            "unattended_leads_total_wait_time"],
                        request_assistance_total_wait_time=analytics_data[
                            "request_assistance_total_wait_time"],
                        group_cobrowse_request_initiated=analytics_data[
                            "group_cobrowse_request_initiated"],
                        group_cobrowse_request_received=analytics_data[
                            "group_cobrowse_request_received"],
                        group_cobrowse_request_connected=analytics_data["group_cobrowse_request_connected"],
                        transfer_requests_received=analytics_data["transfer_requests_received"],
                        transfer_requests_connected=analytics_data["transfer_requests_connected"],
                        transfer_requests_rejected=analytics_data["transfer_requests_rejected"],
                        total_self_assign_time=analytics_data["total_self_assign_time"],
                        self_assign_sessions=analytics_data["self_assign_sessions"])
                else:
                    cobrowse_analytics_obj.request_initiated = analytics_data[
                        "request_initiated"]
                    cobrowse_analytics_obj.request_initiated_by_floating_button_icon = analytics_data[
                        "request_initiated_by_floating_button_icon"]
                    cobrowse_analytics_obj.request_initiated_by_greeting_bubble = analytics_data[
                        "request_initiated_by_greeting_bubble"]
                    cobrowse_analytics_obj.request_initiated_by_exit_intent = analytics_data[
                        "request_initiated_by_exit_intent"]
                    cobrowse_analytics_obj.request_initiated_by_inactivity_popup = analytics_data[
                        "request_initiated_by_inactivity_popup"]
                    cobrowse_analytics_obj.request_attended = analytics_data[
                        "request_attended"]
                    cobrowse_analytics_obj.request_unattended = analytics_data[
                        "request_unattended"]
                    cobrowse_analytics_obj.followup_leads = analytics_data[
                        "followup_leads"]
                    cobrowse_analytics_obj.customers_converted = analytics_data[
                        "customers_converted"]
                    cobrowse_analytics_obj.customers_converted_by_url = analytics_data[
                        "customers_converted_by_url"]
                    cobrowse_analytics_obj.declined_leads = analytics_data[
                        "declined_leads"]
                    cobrowse_analytics_obj.total_session_time = analytics_data[
                        "total_session_time"]
                    cobrowse_analytics_obj.attended_leads_total_wait_time = analytics_data[
                        "attended_leads_total_wait_time"]
                    cobrowse_analytics_obj.unattended_leads_total_wait_time = analytics_data[
                        "unattended_leads_total_wait_time"]
                    cobrowse_analytics_obj.request_assistance_total_wait_time = analytics_data[
                        "request_assistance_total_wait_time"]
                    cobrowse_analytics_obj.group_cobrowse_request_initiated = analytics_data[
                        "group_cobrowse_request_initiated"]
                    cobrowse_analytics_obj.group_cobrowse_request_received = analytics_data[
                        "group_cobrowse_request_received"]
                    cobrowse_analytics_obj.group_cobrowse_request_connected = analytics_data[
                        "group_cobrowse_request_connected"]
                    cobrowse_analytics_obj.transfer_requests_received = analytics_data["transfer_requests_received"]
                    cobrowse_analytics_obj.transfer_requests_connected = analytics_data["transfer_requests_connected"]
                    cobrowse_analytics_obj.transfer_requests_rejected = analytics_data["transfer_requests_rejected"]
                    cobrowse_analytics_obj.total_self_assign_time = analytics_data["total_self_assign_time"]
                    cobrowse_analytics_obj.self_assign_sessions = analytics_data["self_assign_sessions"]
                    cobrowse_analytics_obj.save()

                if agent.role in ["admin", "supervisor"] or agent.is_switch_allowed:
                    cobrowse_analytics_obj.request_not_initiated = total_request_not_initiated
                    cobrowse_analytics_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_inbound_analytics_model %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_agent_wise_inbound_analytics(cobrowse_io_objs, agent):
    from EasyAssistApp.models import CobrowseIOInvitedAgentsDetails, CobrowseIOTransferredAgentsLogs
    try:
        total_sr = 0
        total_sr_floating_button_icon = 0
        total_sr_greeting_bubble = 0
        total_sr_exit_intent = 0
        total_sr_inactivity_popup = 0
        total_sr_closed = 0
        total_sr_closed_by_url = 0
        total_sr_attended = 0
        total_session_duration = 0
        total_wait_time = 0
        total_wait_time_unattended = 0
        browsing_time_before_connect_click = 0
        group_cobrowse_request_initiated = 0
        group_cobrowse_request_received = 0
        group_cobrowse_request_connected = 0
        transfer_requests_received = 0
        transfer_requests_connected = 0
        transfer_requests_rejected = 0
        self_assign_sessions = 0
        total_self_assign_time = 0
        total_followup_leads = 0

        group_cobrowse_request_received = CobrowseIOInvitedAgentsDetails.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, support_agents_invited__in=[agent]).count()
        group_cobrowse_request_connected = CobrowseIOInvitedAgentsDetails.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, support_agents_joined__in=[agent]).count()
        transfer_requests_received = CobrowseIOTransferredAgentsLogs.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, cobrowse_request_type="transferred", transferred_agent__in=[agent]).count()
        transfer_requests_connected = CobrowseIOTransferredAgentsLogs.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, transferred_agent__in=[agent], cobrowse_request_type="transferred", transferred_status="accepted").count()
        transfer_requests_rejected = transfer_requests_received - transfer_requests_connected
        cobrowse_io_objs = cobrowse_io_objs.filter(agent=agent)
        
        # total cobrowsing request Initiated/Notinitiated
        total_sr = cobrowse_io_objs.count()
        if total_sr > 0:
            # total cobrowsing request initiated by floating button and icon
            total_sr_floating_button_icon = cobrowse_io_objs.filter(
                lead_initiated_by__in=['floating_button', 'icon']).count()

            # total cobrowsing request initiated by greeting bubble
            total_sr_greeting_bubble = cobrowse_io_objs.filter(
                lead_initiated_by='greeting_bubble').count()

            # total cobrowsing request initiated by exit intent
            total_sr_exit_intent = cobrowse_io_objs.filter(
                lead_initiated_by='exit_intent').count()

            # total cobrowsing request initiated by Inactivity pop-up
            total_sr_inactivity_popup = cobrowse_io_objs.filter(
                lead_initiated_by='inactivity_popup').count()

            # total cobrowsing request attended
            total_sr_attended = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request initiated and converted successfully
            total_sr_closed = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=False).filter(~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request initiated and converted successfully by
            # landing at specified URL
            total_sr_closed_by_url = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=True).filter(~Q(cobrowsing_start_datetime=None)).count()

            cobrowse_io_initiated = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None))
            
            self_assign_session_obj = cobrowse_io_objs.filter(~Q(self_assign_time=None))
            self_assign_sessions = self_assign_session_obj.count()
            total_self_assign_time = get_total_self_assign_time(self_assign_session_obj)

            for cobrowse_io in cobrowse_io_initiated.iterator():
                total_session_duration += cobrowse_io.session_time_in_seconds()
                total_wait_time += cobrowse_io.customer_wait_time_in_seconds()

            cobrowse_io_not_initiated = cobrowse_io_objs.filter(
                Q(cobrowsing_start_datetime=None)).filter(
                ~Q(allow_agent_cobrowse="false")).exclude(session_archived_cause__in=["FOLLOWUP", "UNASSIGNED"])
            # Customer request not initiated
            total_sr_unattended = cobrowse_io_not_initiated.count()

            for cobrowse_io in cobrowse_io_not_initiated.iterator():
                total_wait_time_unattended += cobrowse_io.customer_wait_time_in_seconds()

            browsing_time_before_connect_click = cobrowse_io_objs.aggregate(
                Sum('browsing_time_before_connect_click'))

            if browsing_time_before_connect_click['browsing_time_before_connect_click__sum']:
                browsing_time_before_connect_click = browsing_time_before_connect_click[
                    'browsing_time_before_connect_click__sum']
            else:
                browsing_time_before_connect_click = 0

            total_declined_leads = cobrowse_io_objs.filter(
                cobrowsing_start_datetime=None, allow_agent_cobrowse="false").count()

            total_followup_leads = cobrowse_io_objs.filter(
                is_archived=True, session_archived_cause__in=["UNASSIGNED", "FOLLOWUP"]).count()
            
            group_cobrowse_request_initiated = CobrowseIOInvitedAgentsDetails.objects.filter(
                cobrowse_io__in=cobrowse_io_objs).count()

            return {
                "request_initiated": total_sr,
                "request_initiated_by_floating_button_icon": total_sr_floating_button_icon,
                "request_initiated_by_greeting_bubble": total_sr_greeting_bubble,
                "request_initiated_by_exit_intent": total_sr_exit_intent,
                "request_initiated_by_inactivity_popup": total_sr_inactivity_popup,
                "request_attended": total_sr_attended,
                "request_unattended": total_sr_unattended,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "declined_leads": total_declined_leads,
                "followup_leads": total_followup_leads,
                "total_session_time": total_session_duration,
                "attended_leads_total_wait_time": total_wait_time,
                "unattended_leads_total_wait_time": total_wait_time_unattended,
                "request_assistance_total_wait_time": browsing_time_before_connect_click,
                "group_cobrowse_request_initiated": group_cobrowse_request_initiated,
                "group_cobrowse_request_received": group_cobrowse_request_received,
                "group_cobrowse_request_connected": group_cobrowse_request_connected,
                "transfer_requests_received": transfer_requests_received,
                "transfer_requests_connected": transfer_requests_connected,
                "transfer_requests_rejected": transfer_requests_rejected,
                "total_self_assign_time": total_self_assign_time,
                "self_assign_sessions": self_assign_sessions,
            }

        else:
            return {
                "request_initiated": total_sr,
                "request_initiated_by_floating_button_icon": total_sr_floating_button_icon,
                "request_initiated_by_greeting_bubble": total_sr_greeting_bubble,
                "request_initiated_by_exit_intent": total_sr_exit_intent,
                "request_initiated_by_inactivity_popup": total_sr_inactivity_popup,
                "request_attended": total_sr_attended,
                "request_unattended": 0,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "declined_leads": 0,
                "followup_leads": 0,
                "total_session_time": total_session_duration,
                "attended_leads_total_wait_time": total_wait_time,
                "unattended_leads_total_wait_time": total_wait_time_unattended,
                "request_assistance_total_wait_time": browsing_time_before_connect_click,
                "group_cobrowse_request_initiated": group_cobrowse_request_initiated,
                "group_cobrowse_request_received": group_cobrowse_request_received,
                "group_cobrowse_request_connected": group_cobrowse_request_connected,
                "transfer_requests_received": transfer_requests_received,
                "transfer_requests_connected": transfer_requests_connected,
                "transfer_requests_rejected": transfer_requests_rejected,
                "total_self_assign_time": total_self_assign_time,
                "self_assign_sessions": self_assign_sessions,
            }

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_wise_inbound_analytics %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

    return None


def get_readable_average_time(total_time, factor):
    try:
        average_time = 0
        if factor > 0:
            average_time = total_time / factor

        if average_time < 60:
            average_time = str(round(average_time)) + " sec"
        else:
            average_time = average_time / 60
            average_time = str(round(average_time)) + " min"

        return average_time
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_wise_inbound_analytics %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
    return ZERO_SEC


def easyassist_get_average_time(total_time, factor):
    try:
        average_time = 0
        if factor > 0:
            average_time = round(total_time / factor)

        minutes, seconds = divmod(average_time, 60)
        hours, minutes = divmod(minutes, 60)
        return (hours, minutes, seconds)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("easyassist_get_average_time %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
    return (0, 0, 0)


def add_inbound_analytics_in_excel_sheet(analytics_workbook, requested_data, cobrowse_agent, CobrowseDateWiseInboundAnalytics, CobrowseDateWiseOutboundDroplinkAnalytics):
    try:
        def create_new_sheet(access_token_obj):
            global invited_agents_connected_col, transfer_requests_col, transferred_agents_connected_col, transferred_agents_rejected_col, \
                self_assign_time_col, average_self_assign_time_col, invited_agents_requested_col, invited_agents_col

            sheet = analytics_workbook.add_sheet(
                "Agent Wise Analytics - Inbound")
            sheet.write(0, 0, "Agent Email ID")
            sheet.col(0).width = 256 * 20
            sheet.write(0, 1, "Request Initiated")
            sheet.col(1).width = 256 * 25
            sheet.write(0, 2, "Request Initiated by Floating Button/Icon")
            sheet.col(2).width = 256 * 25
            sheet.write(0, 3, "Request Initiated by Greeting Bubble")
            sheet.col(3).width = 256 * 30
            sheet.write(0, 4, "Request Initiated by Exit Intent")
            sheet.col(4).width = 256 * 25
            sheet.write(0, 5, "Request Initiated by Inactivity Pop-up")
            sheet.col(5).width = 256 * 30
            sheet.write(0, 6, "Request Attended")
            sheet.col(6).width = 256 * 30
            sheet.write(0, 7, "Links Generated")
            sheet.col(7).width = 256 * 30
            sheet.write(0, 8, CUSTOMERS_CONVERTED_BY_AGENT)
            sheet.col(8).width = 256 * 35
            sheet.write(0, 9, CUSTOMERS_CONVERTED_THROUGH_URL)
            sheet.col(9).width = 256 * 35
            sheet.write(0, 10, "Request Unattended")
            sheet.col(10).width = 256 * 35
            sheet.write(0, 11, AVERAGE_SESSION_TIME)
            sheet.col(11).width = 256 * 20
            sheet.write(0, 12, AVERAGE_WAITING_TIME_FOR_ATTENDED_LEADS)
            sheet.col(12).width = 256 * 40
            sheet.write(0, 13, AVERAGE_WAITING_TIME_FOR_UNATTENDED_LEADS)
            sheet.col(13).width = 256 * 40
            prev_col = 13
            if access_token_obj and access_token_obj.enable_request_in_queue:
                self_assign_time_col = prev_col + 1
                sheet.write(0, self_assign_time_col, "Self Assigned Sessions")
                sheet.col(self_assign_time_col).width = 256 * 40
                average_self_assign_time_col = self_assign_time_col + 1
                sheet.write(0, average_self_assign_time_col,
                            "Average Self Assign Time")
                sheet.col(average_self_assign_time_col).width = 256 * 40
                prev_col = average_self_assign_time_col

            if access_token_obj and access_token_obj.enable_invite_agent_in_cobrowsing:
                invited_agents_requested_col = prev_col + 1
                sheet.write(0, invited_agents_requested_col,
                            "Group Cobrowse Request Initiated")
                sheet.col(invited_agents_requested_col).width = 256 * 40
                invited_agents_col = invited_agents_requested_col + 1
                sheet.write(0, invited_agents_col,
                            "Group Cobrowse Request Received")
                sheet.col(invited_agents_col).width = 256 * 40
                invited_agents_connected_col = invited_agents_col + 1
                sheet.write(0, invited_agents_connected_col,
                            "Group Cobrowse Request Connected")
                sheet.col(invited_agents_connected_col).width = 256 * 40
                prev_col = invited_agents_connected_col

                if access_token_obj.enable_session_transfer_in_cobrowsing:
                    transfer_requests_col = prev_col + 1
                    sheet.write(0, transfer_requests_col,
                                "Transferred Requests Received")
                    sheet.col(transfer_requests_col).width = 256 * 45
                    transferred_agents_connected_col = transfer_requests_col + 1
                    sheet.write(0, transferred_agents_connected_col,
                                "Transferred Requests Connected")
                    sheet.col(transferred_agents_connected_col).width = 256 * 45
                    transferred_agents_rejected_col = transferred_agents_connected_col + 1
                    sheet.write(0, transferred_agents_rejected_col,
                                "Transferred Agent Not Connected")
                    sheet.col(transferred_agents_rejected_col).width = 256 * 45

            return sheet

        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT
        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        report_period_days = (end_date - start_date).days + 1

        temp_date = start_date

        agents = get_list_agents_under_admin(
            cobrowse_agent, is_active=None, is_account_active=None)

        inbound_analytics_objs = CobrowseDateWiseInboundAnalytics.objects.filter(
            date__gte=start_date, date__lte=end_date, agent__in=agents)
        access_token_obj = cobrowse_agent.get_access_token_obj()

        drop_link_analytics_objs = CobrowseDateWiseOutboundDroplinkAnalytics.objects.filter(
            date__gte=start_date, date__lte=end_date, agent__in=agents)

        date_wise_analytics_data = {}
        agent_wise_analytics_data = {}
        analytics_data_summary = {
            "request_initiated": 0,
            "request_initiated_by_floating_button_icon": 0,
            "request_initiated_by_greeting_bubble": 0,
            "request_initiated_by_exit_intent": 0,
            "request_initiated_by_inactivity_popup": 0,
            "request_attended": 0,
            "customers_converted": 0,
            "customers_converted_by_url": 0,
            "total_session_time": 0,
            "request_not_initiated": 0,
            "attended_leads_total_wait_time": 0,
            "unattended_leads_total_wait_time": 0,
            "request_assistance_total_wait_time": 0,
            "total_links_generated": 0,
        }

        for analytics_obj in inbound_analytics_objs.iterator():
            date = analytics_obj.date.strftime(DATE_TIME_FORMAT)
            agent_username = analytics_obj.agent.user.username

            if date not in date_wise_analytics_data:
                date_wise_analytics_data[date] = {
                    "request_initiated": 0,
                    "request_initiated_by_floating_button_icon": 0,
                    "request_initiated_by_greeting_bubble": 0,
                    "request_initiated_by_exit_intent": 0,
                    "request_initiated_by_inactivity_popup": 0,
                    "request_attended": 0,
                    "request_unattended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "total_session_time": 0,
                    "request_not_initiated": 0,
                    "attended_leads_total_wait_time": 0,
                    "unattended_leads_total_wait_time": 0,
                    "request_assistance_total_wait_time": 0,
                    "total_links_generated": 0,
                }

            if agent_username not in agent_wise_analytics_data:
                agent_wise_analytics_data[agent_username] = {
                    "request_initiated": 0,
                    "request_initiated_by_floating_button_icon": 0,
                    "request_initiated_by_greeting_bubble": 0,
                    "request_initiated_by_exit_intent": 0,
                    "request_initiated_by_inactivity_popup": 0,
                    "request_attended": 0,
                    "request_unattended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "total_session_time": 0,
                    "request_not_initiated": 0,
                    "attended_leads_total_wait_time": 0,
                    "unattended_leads_total_wait_time": 0,
                    "request_assistance_total_wait_time": 0,
                    "group_cobrowse_request_initiated": 0,
                    "group_cobrowse_request_received": 0,
                    "group_cobrowse_request_connected": 0,
                    "transfer_requests_received": 0,
                    "transfer_requests_connected": 0,
                    "transfer_requests_rejected": 0,
                    "self_assign_sessions": 0,
                    "total_self_assign_time": 0,
                    "total_links_generated": 0,
                }

            date_wise_analytics_data[date][
                "request_initiated"] += analytics_obj.request_initiated
            date_wise_analytics_data[date][
                "request_initiated_by_floating_button_icon"] += analytics_obj.request_initiated_by_floating_button_icon
            date_wise_analytics_data[date][
                "request_initiated_by_greeting_bubble"] += analytics_obj.request_initiated_by_greeting_bubble
            date_wise_analytics_data[date][
                "request_initiated_by_exit_intent"] += analytics_obj.request_initiated_by_exit_intent
            date_wise_analytics_data[date][
                "request_initiated_by_inactivity_popup"] += analytics_obj.request_initiated_by_inactivity_popup
            date_wise_analytics_data[date][
                "request_attended"] += analytics_obj.request_attended
            date_wise_analytics_data[date][
                "request_unattended"] += analytics_obj.request_unattended
            date_wise_analytics_data[date][
                "customers_converted"] += analytics_obj.customers_converted
            date_wise_analytics_data[date][
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url
            date_wise_analytics_data[date][
                "total_session_time"] += analytics_obj.total_session_time
            date_wise_analytics_data[date][
                "attended_leads_total_wait_time"] += analytics_obj.attended_leads_total_wait_time
            date_wise_analytics_data[date][
                "unattended_leads_total_wait_time"] += analytics_obj.unattended_leads_total_wait_time
            date_wise_analytics_data[date][
                "request_assistance_total_wait_time"] += analytics_obj.request_assistance_total_wait_time

            if analytics_obj.agent == cobrowse_agent:
                date_wise_analytics_data[date][
                    "request_not_initiated"] += analytics_obj.request_not_initiated
                analytics_data_summary["request_not_initiated"] += analytics_obj.request_not_initiated

            agent_wise_analytics_data[agent_username][
                "request_initiated"] += analytics_obj.request_initiated
            agent_wise_analytics_data[agent_username][
                "request_initiated_by_floating_button_icon"] += analytics_obj.request_initiated_by_floating_button_icon
            agent_wise_analytics_data[agent_username][
                "request_initiated_by_greeting_bubble"] += analytics_obj.request_initiated_by_greeting_bubble
            agent_wise_analytics_data[agent_username][
                "request_initiated_by_exit_intent"] += analytics_obj.request_initiated_by_exit_intent
            agent_wise_analytics_data[agent_username][
                "request_initiated_by_inactivity_popup"] += analytics_obj.request_initiated_by_inactivity_popup
            agent_wise_analytics_data[agent_username][
                "request_attended"] += analytics_obj.request_attended
            agent_wise_analytics_data[agent_username][
                "request_unattended"] += analytics_obj.request_unattended
            agent_wise_analytics_data[agent_username][
                "customers_converted"] += analytics_obj.customers_converted
            agent_wise_analytics_data[agent_username][
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url
            agent_wise_analytics_data[agent_username][
                "total_session_time"] += analytics_obj.total_session_time
            agent_wise_analytics_data[agent_username][
                "attended_leads_total_wait_time"] += analytics_obj.attended_leads_total_wait_time
            agent_wise_analytics_data[agent_username][
                "unattended_leads_total_wait_time"] += analytics_obj.unattended_leads_total_wait_time
            agent_wise_analytics_data[agent_username][
                "group_cobrowse_request_initiated"] += analytics_obj.group_cobrowse_request_initiated
            agent_wise_analytics_data[agent_username][
                "group_cobrowse_request_received"] += analytics_obj.group_cobrowse_request_received
            agent_wise_analytics_data[agent_username][
                "group_cobrowse_request_connected"] += analytics_obj.group_cobrowse_request_connected
            agent_wise_analytics_data[agent_username][
                "transfer_requests_received"] += analytics_obj.transfer_requests_received
            agent_wise_analytics_data[agent_username][
                "transfer_requests_connected"] += analytics_obj.transfer_requests_connected
            agent_wise_analytics_data[agent_username][
                "transfer_requests_rejected"] += analytics_obj.transfer_requests_rejected
            agent_wise_analytics_data[agent_username][
                "self_assign_sessions"] += analytics_obj.self_assign_sessions
            agent_wise_analytics_data[agent_username][
                "total_self_assign_time"] += analytics_obj.total_self_assign_time

            analytics_data_summary["request_initiated"] += analytics_obj.request_initiated
            analytics_data_summary["request_initiated_by_floating_button_icon"] += analytics_obj.request_initiated_by_floating_button_icon
            analytics_data_summary["request_initiated_by_greeting_bubble"] += analytics_obj.request_initiated_by_greeting_bubble
            analytics_data_summary["request_initiated_by_exit_intent"] += analytics_obj.request_initiated_by_exit_intent
            analytics_data_summary["request_initiated_by_inactivity_popup"] += analytics_obj.request_initiated_by_inactivity_popup
            analytics_data_summary["request_attended"] += analytics_obj.request_attended
            analytics_data_summary["customers_converted"] += analytics_obj.customers_converted
            analytics_data_summary["customers_converted_by_url"] += analytics_obj.customers_converted_by_url

        for drop_link_analytics_obj in drop_link_analytics_objs.iterator():
            date = drop_link_analytics_obj.date.strftime(DATE_TIME_FORMAT)
            agent_username = drop_link_analytics_obj.agent.user.username

            if date not in date_wise_analytics_data:
                date_wise_analytics_data[date] = {
                    "request_initiated": 0,
                    "request_initiated_by_floating_button_icon": 0,
                    "request_initiated_by_greeting_bubble": 0,
                    "request_initiated_by_exit_intent": 0,
                    "request_initiated_by_inactivity_popup": 0,
                    "request_attended": 0,
                    "request_unattended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "total_session_time": 0,
                    "request_not_initiated": 0,
                    "attended_leads_total_wait_time": 0,
                    "unattended_leads_total_wait_time": 0,
                    "request_assistance_total_wait_time": 0,
                    "total_links_generated": 0,
                }

            if agent_username not in agent_wise_analytics_data:
                agent_wise_analytics_data[agent_username] = {
                    "request_initiated": 0,
                    "request_initiated_by_floating_button_icon": 0,
                    "request_initiated_by_greeting_bubble": 0,
                    "request_initiated_by_exit_intent": 0,
                    "request_initiated_by_inactivity_popup": 0,
                    "request_attended": 0,
                    "request_unattended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "total_session_time": 0,
                    "request_not_initiated": 0,
                    "attended_leads_total_wait_time": 0,
                    "unattended_leads_total_wait_time": 0,
                    "request_assistance_total_wait_time": 0,
                    "group_cobrowse_request_initiated": 0,
                    "group_cobrowse_request_received": 0,
                    "group_cobrowse_request_connected": 0,
                    "transfer_requests_received": 0,
                    "transfer_requests_connected": 0,
                    "transfer_requests_rejected": 0,
                    "self_assign_sessions": 0,
                    "total_self_assign_time": 0,
                    "total_links_generated": 0,
                }

            date_wise_analytics_data[date][
                "total_links_generated"] += drop_link_analytics_obj.total_droplinks_generated

            agent_wise_analytics_data[agent_username][
                "total_links_generated"] += drop_link_analytics_obj.total_droplinks_generated

            analytics_data_summary["total_links_generated"] += drop_link_analytics_obj.total_droplinks_generated
        
        sheet1 = analytics_workbook.add_sheet("Inbound")

        sheet1.write(0, 0, "Date")
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, "Cobrowsing Request Initiated")
        sheet1.col(1).width = 256 * 25
        sheet1.write(0, 2, "Cobrowsing Request Initiated by Floating Button/Icon")
        sheet1.col(2).width = 256 * 25
        sheet1.write(0, 3, "Cobrowsing Request Initiated by Greeting Bubble")
        sheet1.col(3).width = 256 * 25
        sheet1.write(0, 4, "Cobrowsing Request Initiated by Exit Intent")
        sheet1.col(4).width = 256 * 25
        sheet1.write(0, 5, "Cobrowsing Request Initiated by Inactivity Pop-up")
        sheet1.col(5).width = 256 * 25
        sheet1.write(0, 6, "Cobrowsing Request Attended")
        sheet1.col(6).width = 256 * 30
        sheet1.write(0, 7, "Links Generated")
        sheet1.col(7).width = 256 * 30
        sheet1.write(0, 8, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet1.col(8).width = 256 * 35
        sheet1.write(0, 9, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet1.col(9).width = 256 * 35
        sheet1.write(0, 10, "Cobrowsing Request Not Initiated")
        sheet1.col(10).width = 256 * 35
        sheet1.write(0, 11, AVERAGE_SESSION_TIME)
        sheet1.col(11).width = 256 * 20
        sheet1.write(0, 12, AVERAGE_WAITING_TIME_FOR_ATTENDED_LEADS)
        sheet1.col(12).width = 256 * 40
        sheet1.write(0, 13, AVERAGE_WAITING_TIME_FOR_UNATTENDED_LEADS)
        sheet1.col(13).width = 256 * 40
        sheet1.write(0, 14, "Average Waiting Time for Requesting Assistance")
        sheet1.col(14).width = 256 * 45

        index = 1
        while temp_date <= end_date:

            date = temp_date.strftime(DATE_TIME_FORMAT)
            excel_sheet_date_format = temp_date.strftime(DATE_TIME_FORMAT_6)
            temp_date = temp_date + timedelta(1)

            if date not in date_wise_analytics_data:
                total_request = 0
                total_request_floating_button_icon = 0
                total_request_greeting_bubble = 0
                total_request_exit_intent = 0
                total_request_inactivity_popup = 0
                total_request_attended = 0
                total_customers_converted = 0
                total_customers_converted_by_url = 0
                total_request_not_initiated = 0
                average_session_time = ZERO_SEC
                attended_leads_avg_wait_time = ZERO_SEC
                unattended_leads_avg_wait_time = ZERO_SEC
                request_assistance_avg_wait_time = ZERO_SEC
                total_links_generated = 0
            else:
                total_links_generated = date_wise_analytics_data[
                    date]["total_links_generated"]

                total_session_time = date_wise_analytics_data[
                    date]["total_session_time"]

                total_request_attended = date_wise_analytics_data[
                    date]["request_attended"]

                attended_leads_total_wait_time = date_wise_analytics_data[
                    date]["attended_leads_total_wait_time"]

                unattended_leads_total_wait_time = date_wise_analytics_data[
                    date]["unattended_leads_total_wait_time"]

                total_request_unattended = date_wise_analytics_data[
                    date]["request_unattended"]

                request_assistance_total_wait_time = date_wise_analytics_data[
                    date]["request_assistance_total_wait_time"]

                total_request = date_wise_analytics_data[
                    date]["request_initiated"]

                total_request_floating_button_icon = date_wise_analytics_data[
                    date]["request_initiated_by_floating_button_icon"]

                total_request_greeting_bubble = date_wise_analytics_data[
                    date]["request_initiated_by_greeting_bubble"]

                total_request_exit_intent = date_wise_analytics_data[
                    date]["request_initiated_by_exit_intent"]

                total_request_inactivity_popup = date_wise_analytics_data[
                    date]["request_initiated_by_inactivity_popup"]

                total_customers_converted = date_wise_analytics_data[
                    date]["customers_converted"]

                total_customers_converted_by_url = date_wise_analytics_data[
                    date]["customers_converted_by_url"]

                total_request_not_initiated = date_wise_analytics_data[
                    date]["request_not_initiated"]

                average_session_time = get_readable_average_time(
                    total_session_time, total_request_attended)

                attended_leads_avg_wait_time = get_readable_average_time(
                    attended_leads_total_wait_time, total_request_attended)

                unattended_leads_avg_wait_time = get_readable_average_time(
                    unattended_leads_total_wait_time, total_request_unattended)

                request_assistance_avg_wait_time = get_readable_average_time(
                    request_assistance_total_wait_time, total_request)

            analytics_data_summary["total_session_time"] += easyassist_time_in_seconds(average_session_time)
            analytics_data_summary["attended_leads_total_wait_time"] += easyassist_time_in_seconds(attended_leads_avg_wait_time)
            analytics_data_summary["unattended_leads_total_wait_time"] += easyassist_time_in_seconds(unattended_leads_avg_wait_time)
            analytics_data_summary["request_assistance_total_wait_time"] += easyassist_time_in_seconds(request_assistance_avg_wait_time)

            sheet1.write(index, 0, excel_sheet_date_format)
            sheet1.write(index, 1, total_request)
            sheet1.write(index, 2, total_request_floating_button_icon)
            sheet1.write(index, 3, total_request_greeting_bubble)
            sheet1.write(index, 4, total_request_exit_intent)
            sheet1.write(index, 5, total_request_inactivity_popup)
            sheet1.write(index, 6, total_request_attended)
            sheet1.write(index, 7, total_links_generated)
            sheet1.write(index, 8, total_customers_converted)
            sheet1.write(index, 9, total_customers_converted_by_url)
            sheet1.write(index, 10, total_request_not_initiated)
            sheet1.write(index, 11, average_session_time)
            sheet1.write(index, 12, attended_leads_avg_wait_time)
            sheet1.write(index, 13, unattended_leads_avg_wait_time)
            sheet1.write(index, 14, request_assistance_avg_wait_time)

            index += 1

        sheet2 = create_new_sheet(access_token_obj)

        index = 1
        for agent in agents:
            agent_username = agent.user.username

            if agent_username not in agent_wise_analytics_data:
                total_request = 0
                total_request_floating_button_icon = 0
                total_request_greeting_bubble = 0
                total_request_exit_intent = 0
                total_request_inactivity_popup = 0
                total_request_attended = 0
                total_customers_converted = 0
                total_customers_converted_by_url = 0
                total_request_unattended = 0
                average_session_time = ZERO_SEC
                attended_leads_avg_wait_time = ZERO_SEC
                unattended_leads_avg_wait_time = ZERO_SEC
                group_cobrowse_request_initiated = 0
                group_cobrowse_request_received = 0
                group_cobrowse_request_connected = 0
                transfer_requests_received = 0
                transfer_requests_connected = 0
                transfer_requests_rejected = 0
                self_assign_sessions = 0
                total_self_assign_time = 0
                average_self_assign_time = ""
                total_links_generated = 0
            else:

                total_links_generated = agent_wise_analytics_data[
                    agent_username]["total_links_generated"]

                total_request = agent_wise_analytics_data[
                    agent_username]["request_initiated"]

                total_request_floating_button_icon = agent_wise_analytics_data[
                    agent_username]["request_initiated_by_floating_button_icon"]

                total_request_greeting_bubble = agent_wise_analytics_data[
                    agent_username]["request_initiated_by_greeting_bubble"]

                total_request_exit_intent = agent_wise_analytics_data[
                    agent_username]["request_initiated_by_exit_intent"]

                total_request_inactivity_popup = agent_wise_analytics_data[
                    agent_username]["request_initiated_by_inactivity_popup"]

                total_customers_converted = agent_wise_analytics_data[
                    agent_username]["customers_converted"]

                total_customers_converted_by_url = agent_wise_analytics_data[
                    agent_username]["customers_converted_by_url"]

                total_session_time = agent_wise_analytics_data[
                    agent_username]["total_session_time"]

                total_request_attended = agent_wise_analytics_data[
                    agent_username]["request_attended"]

                attended_leads_total_wait_time = agent_wise_analytics_data[
                    agent_username]["attended_leads_total_wait_time"]

                unattended_leads_total_wait_time = agent_wise_analytics_data[
                    agent_username]["unattended_leads_total_wait_time"]

                total_request_unattended = agent_wise_analytics_data[
                    agent_username]["request_unattended"]

                average_session_time = get_readable_average_time(
                    total_session_time, total_request_attended)

                attended_leads_avg_wait_time = get_readable_average_time(
                    attended_leads_total_wait_time, total_request_attended)

                unattended_leads_avg_wait_time = get_readable_average_time(
                    unattended_leads_total_wait_time, total_request_unattended)

                group_cobrowse_request_initiated = agent_wise_analytics_data[
                    agent_username]["group_cobrowse_request_initiated"]

                group_cobrowse_request_received = agent_wise_analytics_data[
                    agent_username]["group_cobrowse_request_received"]

                group_cobrowse_request_connected = agent_wise_analytics_data[
                    agent_username]["group_cobrowse_request_connected"]
                
                transfer_requests_received = agent_wise_analytics_data[
                    agent_username]["transfer_requests_received"]

                transfer_requests_connected = agent_wise_analytics_data[
                    agent_username]["transfer_requests_connected"]

                transfer_requests_rejected = agent_wise_analytics_data[
                    agent_username]["transfer_requests_rejected"]

                self_assign_sessions = agent_wise_analytics_data[
                    agent_username]["self_assign_sessions"]

                total_self_assign_time = agent_wise_analytics_data[
                    agent_username]["total_self_assign_time"]

                average_self_time = total_self_assign_time // self_assign_sessions if self_assign_sessions else 0
                average_self_assign_time = get_time_in_hours_mins_secs(
                    average_self_time)
            
            sheet2.write(index, 0, agent_username)
            sheet2.write(index, 1, total_request)
            sheet2.write(index, 2, total_request_floating_button_icon)
            sheet2.write(index, 3, total_request_greeting_bubble)
            sheet2.write(index, 4, total_request_exit_intent)
            sheet2.write(index, 5, total_request_inactivity_popup)
            sheet2.write(index, 6, total_request_attended)
            sheet2.write(index, 7, total_links_generated)
            sheet2.write(index, 8, total_customers_converted)
            sheet2.write(index, 9, total_customers_converted_by_url)
            sheet2.write(index, 10, total_request_unattended)
            sheet2.write(index, 11, average_session_time)
            sheet2.write(index, 12, attended_leads_avg_wait_time)
            sheet2.write(index, 13, unattended_leads_avg_wait_time)
            
            if access_token_obj and access_token_obj.enable_request_in_queue:
                if self_assign_sessions:
                    sheet2.write(index, self_assign_time_col, self_assign_sessions)
                else:
                    sheet2.write(index, self_assign_time_col, "-")
                if average_self_assign_time != " 0 sec ":
                    sheet2.write(index, average_self_assign_time_col, average_self_assign_time)
                else:
                    sheet2.write(index, average_self_assign_time_col, "-")

            if access_token_obj and access_token_obj.enable_invite_agent_in_cobrowsing:
                sheet2.write(index, invited_agents_requested_col, group_cobrowse_request_initiated)
                sheet2.write(index, invited_agents_col, group_cobrowse_request_received)
                sheet2.write(index, invited_agents_connected_col, group_cobrowse_request_connected)

                if access_token_obj.enable_session_transfer_in_cobrowsing:
                    sheet2.write(index, transfer_requests_col, transfer_requests_received)
                    sheet2.write(index, transferred_agents_connected_col, transfer_requests_connected)
                    sheet2.write(index, transferred_agents_rejected_col, transfer_requests_rejected)

            index += 1

        sheet3 = analytics_workbook.add_sheet("Inbound Summary")

        sheet3.write(0, 0, "Co-browsing Requests Initiated")
        sheet3.col(0).width = 256 * 35
        sheet3.write(
            0, 1, "Co-browsing Requests Initiated by Floating Button/Icon")
        sheet3.col(1).width = 256 * 35
        sheet3.write(0, 2, "Co-browsing Requests Initiated by Greeting Bubble")
        sheet3.col(2).width = 256 * 35
        sheet3.write(0, 3, "Co-browsing Requests Initiated by Exit Intent")
        sheet3.col(3).width = 256 * 35
        sheet3.write(
            0, 4, "Co-browsing Requests Initiated by Inactivity Pop-up")
        sheet3.col(4).width = 256 * 35
        sheet3.write(0, 5, "Co-browsing Requests Attended")
        sheet3.col(5).width = 256 * 30
        sheet3.write(0, 6, "Links Generated")
        sheet3.col(6).width = 256 * 30
        sheet3.write(0, 7, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet3.col(7).width = 256 * 35
        sheet3.write(0, 8, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet3.col(8).width = 256 * 35
        sheet3.write(0, 9, "Co-browsing Requests Not Initiated")
        sheet3.col(9).width = 256 * 35
        sheet3.write(0, 10, AVERAGE_SESSION_TIME)
        sheet3.col(10).width = 256 * 35
        sheet3.write(0, 11, AVERAGE_WAITING_TIME_FOR_ATTENDED_LEADS)
        sheet3.col(11).width = 256 * 40
        sheet3.write(0, 12, AVERAGE_WAITING_TIME_FOR_UNATTENDED_LEADS)
        sheet3.col(12).width = 256 * 40
        sheet3.write(0, 13, "Average Waiting Time for Requesting Assistance")
        sheet3.col(13).width = 256 * 45

        sheet3.write(1, 0, analytics_data_summary["request_initiated"])
        sheet3.write(
            1, 1, analytics_data_summary["request_initiated_by_floating_button_icon"])
        sheet3.write(
            1, 2, analytics_data_summary["request_initiated_by_greeting_bubble"])
        sheet3.write(
            1, 3, analytics_data_summary["request_initiated_by_exit_intent"])
        sheet3.write(
            1, 4, analytics_data_summary["request_initiated_by_inactivity_popup"])
        sheet3.write(1, 5, analytics_data_summary["request_attended"])
        sheet3.write(1, 6, analytics_data_summary["total_links_generated"])
        sheet3.write(1, 7, analytics_data_summary["customers_converted"])
        sheet3.write(
            1, 8, analytics_data_summary["customers_converted_by_url"])
        sheet3.write(
            1, 9, analytics_data_summary["request_not_initiated"])

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["total_session_time"], report_period_days)

        average_session_time = readable_time_format(hours, minutes, seconds)
        sheet3.write(1, 10, average_session_time)

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["attended_leads_total_wait_time"], report_period_days)

        attended_leads_avg_wait_time = readable_time_format(
            hours, minutes, seconds)
        sheet3.write(1, 11, attended_leads_avg_wait_time)

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["unattended_leads_total_wait_time"], report_period_days)

        unattended_leads_avg_wait_time = readable_time_format(
            hours, minutes, seconds)
        sheet3.write(1, 12, unattended_leads_avg_wait_time)

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["request_assistance_total_wait_time"], report_period_days)

        request_assistance_avg_wait_time = readable_time_format(
            hours, minutes, seconds)
        sheet3.write(1, 13, request_assistance_avg_wait_time)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("add_inbound_analytics_in_excel_sheet %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_inbound_analytics_data_dump(requested_data, cobrowse_agent, CobrowseDateWiseInboundAnalytics, CobrowseDateWiseOutboundDroplinkAnalytics):
    try:
        logger.info("Inside get_inbound_analytics_data_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]
        
        date_format = DATE_TIME_FORMAT
        #               for email in email_str.split(",") if email != ""]

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()
        file_directory = settings.SECURE_MEDIA_ROOT + \
            "EasyAssistApp/InboundAnalytics/" + \
            str(cobrowse_agent.user.username)
        if not path.exists(file_directory):
            os.makedirs(file_directory)
        file_path = file_directory + "/inbound_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"
        absolute_file_path = "/secured_files/EasyAssistApp/InboundAnalytics/" + \
            str(cobrowse_agent.user.username) + \
            "/inbound_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"

        if path.exists(file_path):
            return absolute_file_path
        analytics_workbook = Workbook()
        
        add_inbound_analytics_in_excel_sheet(
            analytics_workbook, requested_data, cobrowse_agent, CobrowseDateWiseInboundAnalytics, CobrowseDateWiseOutboundDroplinkAnalytics)

        analytics_workbook.save(file_path)
        return absolute_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_inbound_analytics_data_dump %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def update_outbound_analytics_model(CobrowseDateWiseOutboundAnalytics, CobrowseIO, CobrowseAccessToken, date):
    try:
        access_token_objs = CobrowseAccessToken.objects.all()

        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False, is_archived=True, is_reverse_cobrowsing=False, is_lead=True, request_datetime__date=date)

        for access_token_obj in access_token_objs.iterator():
            admin_agent = access_token_obj.agent
            if admin_agent == None:
                continue

            agent_objs = get_list_agents_under_admin(
                admin_agent, is_active=None, is_account_active=None)
            agent_objs += get_list_supervisor_under_admin(admin_agent, None)

            filtered_cobrowse_io_objs = cobrowse_io_objs.filter(
                access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date)

            total_sr_lead_captured = cobrowse_io_objs.filter(
                access_token=access_token_obj).count()

            for agent in agent_objs:
                analytics_data = get_agent_wise_outbound_analytics(
                    filtered_cobrowse_io_objs, agent)

                if analytics_data == None:
                    continue

                cobrowse_analytics_obj = CobrowseDateWiseOutboundAnalytics.objects.filter(
                    date=date, agent=agent).first()

                if cobrowse_analytics_obj == None:
                    cobrowse_analytics_obj = CobrowseDateWiseOutboundAnalytics.objects.create(
                        agent=agent,
                        date=date,
                        searched_leads=analytics_data["searched_leads"],
                        request_attended=analytics_data["request_attended"],
                        request_unattended=analytics_data[
                            "request_unattended"],
                        customers_converted=analytics_data[
                            "customers_converted"],
                        customers_converted_by_url=analytics_data[
                            "customers_converted_by_url"],
                        requests_denied_by_customers=analytics_data[
                            "requests_denied_by_customers"],
                        total_session_time=analytics_data[
                            "total_session_time"],
                        attended_leads_total_wait_time=analytics_data[
                            "attended_leads_total_wait_time"],
                        unattended_leads_total_wait_time=analytics_data[
                            "unattended_leads_total_wait_time"],
                        group_cobrowse_request_initiated=analytics_data[
                            "group_cobrowse_request_initiated"],
                        group_cobrowse_request_received=analytics_data[
                            "group_cobrowse_request_received"],
                        group_cobrowse_request_connected=analytics_data["group_cobrowse_request_connected"],
                        transfer_requests_received=analytics_data["transfer_requests_received"],
                        transfer_requests_connected=analytics_data["transfer_requests_connected"],
                        transfer_requests_rejected=analytics_data["transfer_requests_rejected"],)
                else:
                    cobrowse_analytics_obj.searched_leads = analytics_data[
                        "searched_leads"]
                    cobrowse_analytics_obj.request_attended = analytics_data[
                        "request_attended"]
                    cobrowse_analytics_obj.request_unattended = analytics_data[
                        "request_unattended"]
                    cobrowse_analytics_obj.customers_converted = analytics_data[
                        "customers_converted"]
                    cobrowse_analytics_obj.customers_converted_by_url = analytics_data[
                        "customers_converted_by_url"]
                    cobrowse_analytics_obj.requests_denied_by_customers = analytics_data[
                        "requests_denied_by_customers"]
                    cobrowse_analytics_obj.total_session_time = analytics_data[
                        "total_session_time"]
                    cobrowse_analytics_obj.attended_leads_total_wait_time = analytics_data[
                        "attended_leads_total_wait_time"]
                    cobrowse_analytics_obj.unattended_leads_total_wait_time = analytics_data[
                        "unattended_leads_total_wait_time"]
                    cobrowse_analytics_obj.group_cobrowse_request_initiated = analytics_data[
                        "group_cobrowse_request_initiated"]
                    cobrowse_analytics_obj.group_cobrowse_request_received = analytics_data[
                        "group_cobrowse_request_received"]
                    cobrowse_analytics_obj.group_cobrowse_request_connected = analytics_data[
                        "group_cobrowse_request_connected"]
                    cobrowse_analytics_obj.transfer_requests_received = analytics_data["transfer_requests_received"]
                    cobrowse_analytics_obj.transfer_requests_connected = analytics_data["transfer_requests_connected"]
                    cobrowse_analytics_obj.transfer_requests_rejected = analytics_data["transfer_requests_rejected"]
                    
                    cobrowse_analytics_obj.save()

                if agent.role in ["admin", "supervisor"] or agent.is_switch_allowed:
                    cobrowse_analytics_obj.captured_leads = total_sr_lead_captured
                    cobrowse_analytics_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_outbound_analytics_model %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_agent_wise_outbound_analytics(cobrowse_io_objs, agent):
    from EasyAssistApp.models import CobrowseIOInvitedAgentsDetails, CobrowseIOTransferredAgentsLogs
    try:
        total_sr = 0
        total_sr_closed = 0
        total_sr_closed_by_url = 0
        total_sr_attended = 0
        total_session_duration = 0
        total_wait_time = 0
        total_wait_time_unattended = 0
        group_cobrowse_request_initiated = 0
        group_cobrowse_request_received = 0
        group_cobrowse_request_connected = 0
        transfer_requests_connected = 0
        transfer_requests_rejected = 0
        transfer_requests_received = 0

        group_cobrowse_request_received = CobrowseIOInvitedAgentsDetails.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, support_agents_invited__in=[agent]).count()
        group_cobrowse_request_connected = CobrowseIOInvitedAgentsDetails.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, support_agents_joined__in=[agent]).count()

        transfer_requests_received = CobrowseIOTransferredAgentsLogs.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, cobrowse_request_type="transferred", transferred_agent__in=[agent]).count()
        transfer_requests_connected = CobrowseIOTransferredAgentsLogs.objects.filter(
            cobrowse_io__in=cobrowse_io_objs, transferred_agent__in=[agent], cobrowse_request_type="transferred", transferred_status="accepted").count()
        transfer_requests_rejected = transfer_requests_received - transfer_requests_connected
        
        cobrowse_io_objs = cobrowse_io_objs.filter(agent=agent)

        # total cobrowsing request Initiated/Notinitiated
        total_sr = cobrowse_io_objs.count()

        if total_sr > 0:

            # total cobrowsing request initiated and converted successfully
            total_sr_closed = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=False).filter(~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request initiated and converted successfully by
            # landing at specified URL
            total_sr_closed_by_url = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=True).filter(~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request attended
            total_sr_attended = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request denied by customer
            total_sr_customer_denied = cobrowse_io_objs.filter(
                consent_cancel_count__gte=1, allow_agent_cobrowse="false", cobrowsing_start_datetime=None).count()

            cobrowse_io_initiated = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None))
            for cobrowse_io in cobrowse_io_initiated.iterator():
                total_session_duration += cobrowse_io.session_time_in_seconds()
                total_wait_time += cobrowse_io.customer_wait_time_in_seconds()

            cobrowse_io_not_initiated = cobrowse_io_objs.filter(
                Q(cobrowsing_start_datetime=None)).filter(
                ~Q(allow_agent_cobrowse="false"))
            total_sr_unattended = cobrowse_io_not_initiated.count()

            for cobrowse_io in cobrowse_io_not_initiated.iterator():
                total_wait_time_unattended += cobrowse_io.customer_wait_time_in_seconds()

            group_cobrowse_request_initiated = CobrowseIOInvitedAgentsDetails.objects.filter(
                cobrowse_io__in=cobrowse_io_objs).count()

            return {
                "searched_leads": total_sr,
                "request_attended": total_sr_attended,
                "request_unattended": total_sr_unattended,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "requests_denied_by_customers": total_sr_customer_denied,
                "total_session_time": total_session_duration,
                "attended_leads_total_wait_time": total_wait_time,
                "unattended_leads_total_wait_time": total_wait_time_unattended,
                "group_cobrowse_request_initiated": group_cobrowse_request_initiated,
                "group_cobrowse_request_received": group_cobrowse_request_received,
                "group_cobrowse_request_connected": group_cobrowse_request_connected,
                "transfer_requests_received": transfer_requests_received,
                "transfer_requests_connected": transfer_requests_connected,
                "transfer_requests_rejected": transfer_requests_rejected,
            }

        else:
            return {
                "searched_leads": total_sr,
                "request_attended": total_sr_attended,
                "request_unattended": 0,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "requests_denied_by_customers": 0,
                "total_session_time": total_session_duration,
                "attended_leads_total_wait_time": total_wait_time,
                "unattended_leads_total_wait_time": total_wait_time_unattended,
                "group_cobrowse_request_initiated": group_cobrowse_request_initiated,
                "group_cobrowse_request_received": group_cobrowse_request_received,
                "group_cobrowse_request_connected": group_cobrowse_request_connected,
                "transfer_requests_received": transfer_requests_received,
                "transfer_requests_connected": transfer_requests_connected,
                "transfer_requests_rejected": transfer_requests_rejected,
            }

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_wise_outbound_analytics %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

    return None


def add_outbound_analytics_in_excel_sheet(analytics_workbook, requested_data, cobrowse_agent, CobrowseDateWiseOutboundAnalytics):
    try:
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        report_period_days = (end_date - start_date).days + 1

        temp_date = start_date

        agents = get_list_agents_under_admin(
            cobrowse_agent, is_active=None, is_account_active=None)

        outbound_analytics_objs = CobrowseDateWiseOutboundAnalytics.objects.filter(
            date__gte=start_date, date__lte=end_date, agent__in=agents)

        access_token_obj = cobrowse_agent.get_access_token_obj()

        date_wise_analytics_data = {}
        agent_wise_analytics_data = {}
        analytics_data_summary = {
            "captured_leads": 0,
            "searched_leads": 0,
            "request_attended": 0,
            "customers_converted": 0,
            "customers_converted_by_url": 0,
            "requests_denied_by_customers": 0,
            "total_session_time": 0,
            "conversion_rate": 0,
            "attended_leads_total_wait_time": 0,
            "unattended_leads_total_wait_time": 0,
        }

        for analytics_obj in outbound_analytics_objs.iterator():
            date = analytics_obj.date.strftime(DATE_TIME_FORMAT)
            agent_username = analytics_obj.agent.user.username

            if date not in date_wise_analytics_data:
                date_wise_analytics_data[date] = {
                    "captured_leads": 0,
                    "searched_leads": 0,
                    "request_attended": 0,
                    "request_unattended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "requests_denied_by_customers": 0,
                    "total_session_time": 0,
                    "attended_leads_total_wait_time": 0,
                    "unattended_leads_total_wait_time": 0,
                }

            if agent_username not in agent_wise_analytics_data:
                agent_wise_analytics_data[agent_username] = {
                    "searched_leads": 0,
                    "request_attended": 0,
                    "request_unattended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "total_session_time": 0,
                    "attended_leads_total_wait_time": 0,
                    "unattended_leads_total_wait_time": 0,
                    "group_cobrowse_request_initiated": 0,
                    "group_cobrowse_request_received": 0,
                    "group_cobrowse_request_connected": 0,
                    "transfer_requests_received": 0,
                    "transfer_requests_connected": 0,
                    "transfer_requests_rejected": 0,
                }

            date_wise_analytics_data[date][
                "searched_leads"] += analytics_obj.searched_leads
            date_wise_analytics_data[date][
                "request_attended"] += analytics_obj.request_attended
            date_wise_analytics_data[date][
                "request_unattended"] += analytics_obj.request_unattended
            date_wise_analytics_data[date][
                "customers_converted"] += analytics_obj.customers_converted
            date_wise_analytics_data[date][
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url
            date_wise_analytics_data[date][
                "requests_denied_by_customers"] += analytics_obj.requests_denied_by_customers
            date_wise_analytics_data[date][
                "total_session_time"] += analytics_obj.total_session_time
            date_wise_analytics_data[date][
                "attended_leads_total_wait_time"] += analytics_obj.attended_leads_total_wait_time
            date_wise_analytics_data[date][
                "unattended_leads_total_wait_time"] += analytics_obj.unattended_leads_total_wait_time

            if analytics_obj.agent == cobrowse_agent:
                date_wise_analytics_data[date][
                    "captured_leads"] += analytics_obj.captured_leads
                analytics_data_summary[
                    "captured_leads"] += analytics_obj.captured_leads

            agent_wise_analytics_data[agent_username][
                "searched_leads"] += analytics_obj.searched_leads
            agent_wise_analytics_data[agent_username][
                "request_attended"] += analytics_obj.request_attended
            agent_wise_analytics_data[agent_username][
                "request_unattended"] += analytics_obj.request_unattended
            agent_wise_analytics_data[agent_username][
                "customers_converted"] += analytics_obj.customers_converted
            agent_wise_analytics_data[agent_username][
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url
            agent_wise_analytics_data[agent_username][
                "total_session_time"] += analytics_obj.total_session_time
            agent_wise_analytics_data[agent_username][
                "attended_leads_total_wait_time"] += analytics_obj.attended_leads_total_wait_time
            agent_wise_analytics_data[agent_username][
                "unattended_leads_total_wait_time"] += analytics_obj.unattended_leads_total_wait_time
            agent_wise_analytics_data[agent_username][
                "group_cobrowse_request_initiated"] += analytics_obj.group_cobrowse_request_initiated
            agent_wise_analytics_data[agent_username][
                "group_cobrowse_request_received"] += analytics_obj.group_cobrowse_request_received
            agent_wise_analytics_data[agent_username][
                "group_cobrowse_request_connected"] += analytics_obj.group_cobrowse_request_connected
            agent_wise_analytics_data[agent_username][
                "transfer_requests_received"] += analytics_obj.transfer_requests_received
            agent_wise_analytics_data[agent_username][
                "transfer_requests_connected"] += analytics_obj.transfer_requests_connected
            agent_wise_analytics_data[agent_username][
                "transfer_requests_rejected"] += analytics_obj.transfer_requests_rejected

            analytics_data_summary[
                "searched_leads"] += analytics_obj.searched_leads
            analytics_data_summary[
                "request_attended"] += analytics_obj.request_attended
            analytics_data_summary[
                "customers_converted"] += analytics_obj.customers_converted
            analytics_data_summary[
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url
            analytics_data_summary[
                "requests_denied_by_customers"] += analytics_obj.requests_denied_by_customers

        sheet1 = analytics_workbook.add_sheet("Outbound")

        sheet1.write(0, 0, "Date")
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, "Captured Leads")
        sheet1.col(1).width = 256 * 20
        sheet1.write(0, 2, "Leads Searched")
        sheet1.col(2).width = 256 * 25
        sheet1.write(0, 3, "Cobrowsing Request Attended")
        sheet1.col(3).width = 256 * 30
        sheet1.write(0, 4, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet1.col(4).width = 256 * 35
        sheet1.write(0, 5, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet1.col(5).width = 256 * 35
        sheet1.write(0, 6, "Cobrowsing Request denied by Customers")
        sheet1.col(6).width = 256 * 35
        sheet1.write(0, 7, AVERAGE_SESSION_TIME)
        sheet1.col(7).width = 256 * 20
        sheet1.write(0, 8, "Conversion Rate")
        sheet1.col(8).width = 256 * 25
        sheet1.write(0, 9, AVERAGE_WAITING_TIME_FOR_ATTENDED_LEADS)
        sheet1.col(9).width = 256 * 40
        sheet1.write(0, 10, AVERAGE_WAITING_TIME_FOR_UNATTENDED_LEADS)
        sheet1.col(10).width = 256 * 40

        index = 1
        while temp_date <= end_date:

            date = temp_date.strftime(DATE_TIME_FORMAT)
            excel_sheet_date_format = temp_date.strftime(DATE_TIME_FORMAT_6)
            temp_date = temp_date + timedelta(1)

            if date not in date_wise_analytics_data:
                total_captured_leads = 0
                total_searched_leads = 0
                total_request_attended = 0
                total_customers_converted = 0
                total_customers_converted_by_url = 0
                total_requests_denied_by_customers = 0
                average_session_time = ZERO_SEC
                conversion_rate = "0%"
                attended_leads_avg_wait_time = ZERO_SEC
                unattended_leads_avg_wait_time = ZERO_SEC
            else:
                total_session_time = date_wise_analytics_data[
                    date]["total_session_time"]

                total_request_attended = date_wise_analytics_data[
                    date]["request_attended"]

                attended_leads_total_wait_time = date_wise_analytics_data[
                    date]["attended_leads_total_wait_time"]

                unattended_leads_total_wait_time = date_wise_analytics_data[
                    date]["unattended_leads_total_wait_time"]

                total_customers_converted = date_wise_analytics_data[
                    date]["customers_converted"]

                total_customers_converted_by_url = date_wise_analytics_data[
                    date]["customers_converted_by_url"]

                total_request_unattended = date_wise_analytics_data[
                    date]["request_unattended"]

                total_captured_leads = date_wise_analytics_data[
                    date]["captured_leads"]

                total_searched_leads = date_wise_analytics_data[
                    date]["searched_leads"]

                total_requests_denied_by_customers = date_wise_analytics_data[
                    date]["requests_denied_by_customers"]

                average_session_time = get_readable_average_time(
                    total_session_time, total_request_attended)

                attended_leads_avg_wait_time = get_readable_average_time(
                    attended_leads_total_wait_time, total_request_attended)

                unattended_leads_avg_wait_time = get_readable_average_time(
                    unattended_leads_total_wait_time, total_request_unattended)

                conversion_rate = "0%"
                if total_request_attended > 0:
                    conversion_rate = (
                        total_customers_converted / total_request_attended) * 100
                    conversion_rate = round(conversion_rate)
                    analytics_data_summary["conversion_rate"] += conversion_rate
                    conversion_rate = str(conversion_rate) + "%"

            analytics_data_summary[
                "total_session_time"] += easyassist_time_in_seconds(average_session_time)
            analytics_data_summary[
                "attended_leads_total_wait_time"] += easyassist_time_in_seconds(attended_leads_avg_wait_time)
            analytics_data_summary[
                "unattended_leads_total_wait_time"] += easyassist_time_in_seconds(unattended_leads_avg_wait_time)

            sheet1.write(index, 0, excel_sheet_date_format)
            sheet1.write(index, 1, total_captured_leads)
            sheet1.write(index, 2, total_searched_leads)
            sheet1.write(index, 3, total_request_attended)
            sheet1.write(index, 4, total_customers_converted)
            sheet1.write(index, 5, total_customers_converted_by_url)
            sheet1.write(index, 6, total_requests_denied_by_customers)
            sheet1.write(index, 7, average_session_time)
            sheet1.write(index, 8, conversion_rate)
            sheet1.write(index, 9, attended_leads_avg_wait_time)
            sheet1.write(index, 10, unattended_leads_avg_wait_time)

            index += 1

        sheet2 = analytics_workbook.add_sheet(
            "Agent Wise Analytics - Outbound")
        sheet2.write(0, 0, "Agent Email ID")
        sheet2.col(0).width = 256 * 20
        sheet2.write(0, 1, "Leads Searched")
        sheet2.col(1).width = 256 * 25
        sheet2.write(0, 2, "Request Attended")
        sheet2.col(2).width = 256 * 30
        sheet2.write(0, 3, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet2.col(3).width = 256 * 35
        sheet2.write(0, 4, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet2.col(4).width = 256 * 35
        sheet2.write(0, 5, "Request Unattended")
        sheet2.col(5).width = 256 * 35
        sheet2.write(0, 6, AVERAGE_SESSION_TIME)
        sheet2.col(6).width = 256 * 20
        sheet2.write(0, 7, AVERAGE_WAITING_TIME_FOR_ATTENDED_LEADS)
        sheet2.col(7).width = 256 * 40
        sheet2.write(0, 8, AVERAGE_WAITING_TIME_FOR_UNATTENDED_LEADS)
        sheet2.col(8).width = 256 * 40

        if access_token_obj and access_token_obj.enable_invite_agent_in_cobrowsing:
            sheet2.write(0, 9, "Group Cobrowse Request Initiated")
            sheet2.col(9).width = 256 * 40
            sheet2.write(0, 10, "Group Cobrowse Request Received")
            sheet2.col(10).width = 256 * 40
            sheet2.write(0, 11, "Group Cobrowse Request Connected")
            sheet2.col(11).width = 256 * 40

            if access_token_obj.enable_session_transfer_in_cobrowsing:
                sheet2.write(0, 12, "Transfer Requests Received")
                sheet2.col(12).width = 256 * 45
                sheet2.write(0, 13, "Transfer Requests Connected")
                sheet2.col(13).width = 256 * 45
                sheet2.write(0, 14, "Transferred Agent Not Connected")
                sheet2.col(14).width = 256 * 45

        index = 1
        for agent in agents:
            agent_username = agent.user.username
            if agent_username not in agent_wise_analytics_data:
                total_searched_leads = 0
                total_request_attended = 0
                total_customers_converted = 0
                total_customers_converted_by_url = 0
                total_request_unattended = 0
                average_session_time = ZERO_SEC
                attended_leads_avg_wait_time = ZERO_SEC
                unattended_leads_avg_wait_time = ZERO_SEC
                group_cobrowse_request_initiated = 0
                group_cobrowse_request_received = 0
                group_cobrowse_request_connected = 0
                transfer_requests_connected = 0
                transfer_requests_rejected = 0
                transfer_requests_received = 0
            else:

                total_session_time = agent_wise_analytics_data[
                    agent_username]["total_session_time"]

                total_request_attended = agent_wise_analytics_data[
                    agent_username]["request_attended"]

                attended_leads_total_wait_time = agent_wise_analytics_data[
                    agent_username]["attended_leads_total_wait_time"]

                unattended_leads_total_wait_time = agent_wise_analytics_data[
                    agent_username]["unattended_leads_total_wait_time"]

                total_customers_converted = agent_wise_analytics_data[
                    agent_username]["customers_converted"]

                total_customers_converted_by_url = agent_wise_analytics_data[
                    agent_username]["customers_converted_by_url"]

                total_request_unattended = agent_wise_analytics_data[
                    agent_username]["request_unattended"]

                total_searched_leads = agent_wise_analytics_data[
                    agent_username]["searched_leads"]

                average_session_time = get_readable_average_time(
                    total_session_time, total_request_attended)

                attended_leads_avg_wait_time = get_readable_average_time(
                    attended_leads_total_wait_time, total_request_attended)

                unattended_leads_avg_wait_time = get_readable_average_time(
                    unattended_leads_total_wait_time, total_request_unattended)

                group_cobrowse_request_initiated = agent_wise_analytics_data[
                    agent_username]["group_cobrowse_request_initiated"]

                group_cobrowse_request_received = agent_wise_analytics_data[
                    agent_username]["group_cobrowse_request_received"]

                group_cobrowse_request_connected = agent_wise_analytics_data[
                    agent_username]["group_cobrowse_request_connected"]

                transfer_requests_received = agent_wise_analytics_data[
                    agent_username]["transfer_requests_received"]

                transfer_requests_connected = agent_wise_analytics_data[
                    agent_username]["transfer_requests_connected"]

                transfer_requests_rejected = agent_wise_analytics_data[
                    agent_username]["transfer_requests_rejected"]

            sheet2.write(index, 0, agent_username)
            sheet2.write(index, 1, total_searched_leads)
            sheet2.write(index, 2, total_request_attended)
            sheet2.write(index, 3, total_customers_converted)
            sheet2.write(index, 4, total_customers_converted_by_url)
            sheet2.write(index, 5, total_request_unattended)
            sheet2.write(index, 6, average_session_time)
            sheet2.write(index, 7, attended_leads_avg_wait_time)
            sheet2.write(index, 8, unattended_leads_avg_wait_time)

            if access_token_obj and access_token_obj.enable_invite_agent_in_cobrowsing:
                sheet2.write(index, 9, group_cobrowse_request_initiated)
                sheet2.write(index, 10, group_cobrowse_request_received)
                sheet2.write(index, 11, group_cobrowse_request_connected)

                if access_token_obj.enable_session_transfer_in_cobrowsing:
                    sheet2.write(index, 12, transfer_requests_received)
                    sheet2.write(index, 13, transfer_requests_connected)
                    sheet2.write(index, 14, transfer_requests_rejected)

            index += 1

        sheet3 = analytics_workbook.add_sheet("Outbound Summary")

        sheet3.write(0, 0, "Captured Leads")
        sheet3.col(0).width = 256 * 25
        sheet3.write(
            0, 1, "Leads Searched")
        sheet3.col(1).width = 256 * 25
        sheet3.write(0, 2, "Co-browsing Requests Attended")
        sheet3.col(2).width = 256 * 35
        sheet3.write(0, 3, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet3.col(3).width = 256 * 35
        sheet3.write(0, 4, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet3.col(4).width = 256 * 35
        sheet3.write(0, 5, "Co-browsing Requests Denied by Customers")
        sheet3.col(5).width = 256 * 35
        sheet3.write(0, 6, AVERAGE_SESSION_TIME)
        sheet3.col(6).width = 256 * 35
        sheet3.write(0, 7, "Conversion Rate")
        sheet3.col(7).width = 256 * 35
        sheet3.write(0, 8, AVERAGE_WAITING_TIME_FOR_ATTENDED_LEADS)
        sheet3.col(8).width = 256 * 35
        sheet3.write(0, 9, AVERAGE_WAITING_TIME_FOR_UNATTENDED_LEADS)
        sheet3.col(9).width = 256 * 40

        sheet3.write(1, 0, analytics_data_summary["captured_leads"])
        sheet3.write(1, 1, analytics_data_summary["searched_leads"])
        sheet3.write(1, 2, analytics_data_summary["request_attended"])
        sheet3.write(1, 3, analytics_data_summary["customers_converted"])
        sheet3.write(
            1, 4, analytics_data_summary["customers_converted_by_url"])
        sheet3.write(
            1, 5, analytics_data_summary["requests_denied_by_customers"])

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["total_session_time"], report_period_days)

        average_session_time = readable_time_format(hours, minutes, seconds)
        sheet3.write(1, 6, average_session_time)

        sheet3.write(
            1, 7, "{:.1f} %".format(analytics_data_summary["conversion_rate"] / report_period_days))

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["attended_leads_total_wait_time"], report_period_days)

        attended_leads_avg_wait_time = readable_time_format(
            hours, minutes, seconds)
        sheet3.write(1, 8, attended_leads_avg_wait_time)

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["unattended_leads_total_wait_time"], report_period_days)

        unattended_leads_avg_wait_time = readable_time_format(
            hours, minutes, seconds)
        sheet3.write(1, 9, unattended_leads_avg_wait_time)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("add_outbound_analytics_in_excel_sheet %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_outbound_analytics_data_dump(requested_data, cobrowse_agent, CobrowseDateWiseOutboundAnalytics):
    try:
        logger.info("Inside get_outbound_analytics_data_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT
        #               for email in email_str.split(",") if email != ""]

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        file_directory = settings.SECURE_MEDIA_ROOT + \
            "EasyAssistApp/OutboundAnalytics/" + \
            str(cobrowse_agent.user.username)
        if not path.exists(file_directory):
            os.makedirs(file_directory)

        file_path = file_directory + "/outbound_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"
        absolute_file_path = "/secured_files/EasyAssistApp/OutboundAnalytics/" + \
            str(cobrowse_agent.user.username) + \
            "/outbound_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"

        if path.exists(file_path):
            return absolute_file_path

        analytics_workbook = Workbook()

        add_outbound_analytics_in_excel_sheet(
            analytics_workbook, requested_data, cobrowse_agent, CobrowseDateWiseOutboundAnalytics)

        analytics_workbook.save(file_path)
        return absolute_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_outbound_analytics_data_dump %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def update_reverse_analytics_model(CobrowseDateWiseReverseAnalytics, CobrowseIO, CobrowseAccessToken, date):
    try:
        access_token_objs = CobrowseAccessToken.objects.all()

        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False, is_archived=True, is_reverse_cobrowsing=True, is_lead=False, request_datetime__date=date)

        for access_token_obj in access_token_objs.iterator():
            admin_agent = access_token_obj.agent
            if admin_agent == None:
                continue

            agent_objs = get_list_agents_under_admin(
                admin_agent, is_active=None, is_account_active=None)
            agent_objs += get_list_supervisor_under_admin(admin_agent, None)

            filtered_cobrowse_io_objs = cobrowse_io_objs.filter(
                access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date)

            for agent in agent_objs:
                analytics_data = get_agent_wise_reverse_analytics(
                    filtered_cobrowse_io_objs, agent)

                if analytics_data == None:
                    continue

                cobrowse_analytics_obj = CobrowseDateWiseReverseAnalytics.objects.filter(
                    date=date, agent=agent).first()

                if cobrowse_analytics_obj == None:
                    cobrowse_analytics_obj = CobrowseDateWiseReverseAnalytics.objects.create(
                        agent=agent,
                        date=date,
                        links_generated=analytics_data["links_generated"],
                        request_attended=analytics_data["request_attended"],
                        customers_converted=analytics_data[
                            "customers_converted"],
                        customers_converted_by_url=analytics_data[
                            "customers_converted_by_url"],
                        total_session_time=analytics_data["total_session_time"])
                else:
                    cobrowse_analytics_obj.links_generated = analytics_data[
                        "links_generated"]
                    cobrowse_analytics_obj.request_attended = analytics_data[
                        "request_attended"]
                    cobrowse_analytics_obj.customers_converted = analytics_data[
                        "customers_converted"]
                    cobrowse_analytics_obj.customers_converted_by_url = analytics_data[
                        "customers_converted_by_url"]
                    cobrowse_analytics_obj.total_session_time = analytics_data[
                        "total_session_time"]
                    cobrowse_analytics_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_reverse_analytics_model %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_agent_wise_reverse_analytics(cobrowse_io_objs, agent):
    try:
        total_sr = 0
        total_sr_closed = 0
        total_sr_closed_by_url = 0
        total_sr_attended = 0
        total_session_duration = 0

        cobrowse_io_objs = cobrowse_io_objs.filter(agent=agent)

        # total cobrowsing request Initiated/Notinitiated
        total_sr = cobrowse_io_objs.count()

        if total_sr > 0:

            total_sr_closed = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=False).filter(~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request initiated and converted successfully by
            # landing at specified URL
            total_sr_closed_by_url = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=True).filter(~Q(cobrowsing_start_datetime=None)).count()

            total_sr_attended = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None)).count()

            cobrowse_io_initiated = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None))

            for cobrowse_io in cobrowse_io_initiated.iterator():
                total_session_duration += cobrowse_io.session_time_in_seconds()

            return {
                "links_generated": total_sr,
                "request_attended": total_sr_attended,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "total_session_time": total_session_duration,
            }
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_wise_outbound_analytics %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

    return None


def get_reverse_analytics_data_dump(requested_data, cobrowse_agent, CobrowseDateWiseReverseAnalytics):
    try:
        logger.info("Inside get_reverse_analytics_data_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT
        #               for email in email_str.split(",") if email != ""]

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        report_period_days = (end_date - start_date).days + 1

        file_directory = settings.SECURE_MEDIA_ROOT + \
            "EasyAssistApp/ReverseAnalytics/" + \
            str(cobrowse_agent.user.username)
        if not path.exists(file_directory):
            os.makedirs(file_directory)

        file_path = file_directory + "/reverse_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"
        absolute_file_path = "/secured_files/EasyAssistApp/ReverseAnalytics/" + \
            str(cobrowse_agent.user.username) + \
            "/reverse_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"

        if path.exists(file_path):
            return absolute_file_path

        temp_date = start_date

        agents = get_list_agents_under_admin(
            cobrowse_agent, is_active=None, is_account_active=None)

        reverse_analytics_objs = CobrowseDateWiseReverseAnalytics.objects.filter(
            date__gte=start_date, date__lte=end_date, agent__in=agents)

        date_wise_analytics_data = {}
        analytics_data_summary = {
            "links_generated": 0,
            "request_attended": 0,
            "customers_converted": 0,
            "customers_converted_by_url": 0,
            "total_session_time": 0,
        }

        for analytics_obj in reverse_analytics_objs.iterator():
            date = analytics_obj.date.strftime(DATE_TIME_FORMAT)

            if date not in date_wise_analytics_data:
                date_wise_analytics_data[date] = {
                    "links_generated": 0,
                    "request_attended": 0,
                    "customers_converted": 0,
                    "customers_converted_by_url": 0,
                    "total_session_time": 0,
                }

            date_wise_analytics_data[date][
                "links_generated"] += analytics_obj.links_generated
            date_wise_analytics_data[date][
                "request_attended"] += analytics_obj.request_attended
            date_wise_analytics_data[date][
                "customers_converted"] += analytics_obj.customers_converted
            date_wise_analytics_data[date][
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url
            date_wise_analytics_data[date][
                "total_session_time"] += analytics_obj.total_session_time

            analytics_data_summary[
                "links_generated"] += analytics_obj.links_generated
            analytics_data_summary[
                "request_attended"] += analytics_obj.request_attended
            analytics_data_summary[
                "customers_converted"] += analytics_obj.customers_converted
            analytics_data_summary[
                "customers_converted_by_url"] += analytics_obj.customers_converted_by_url

        analytics_workbook = Workbook()
        sheet1 = analytics_workbook.add_sheet("Reverse")

        sheet1.write(0, 0, "Date")
        sheet1.col(0).width = 256 * 20
        sheet1.write(0, 1, "Links Generated")
        sheet1.col(1).width = 256 * 20
        sheet1.write(0, 2, "Session Joined by the Customer")
        sheet1.col(2).width = 256 * 25
        sheet1.write(0, 3, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet1.col(3).width = 256 * 35
        sheet1.write(0, 4, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet1.col(4).width = 256 * 35
        sheet1.write(0, 5, AVERAGE_SESSION_TIME)
        sheet1.col(5).width = 256 * 20

        index = 1
        while temp_date <= end_date:

            date = temp_date.strftime(DATE_TIME_FORMAT)
            excel_sheet_date_format = temp_date.strftime(DATE_TIME_FORMAT_6)
            temp_date = temp_date + timedelta(1)

            if date not in date_wise_analytics_data:
                total_links_generated = 0
                total_customers_converted = 0
                total_customers_converted_by_url = 0
                total_request_attended = 0
                average_session_time = ZERO_SEC
            else:
                total_links_generated = date_wise_analytics_data[
                    date]["links_generated"]
                total_session_time = date_wise_analytics_data[
                    date]["total_session_time"]
                total_customers_converted = date_wise_analytics_data[
                    date]["customers_converted"]
                total_customers_converted_by_url = date_wise_analytics_data[
                    date]["customers_converted_by_url"]
                total_request_attended = date_wise_analytics_data[
                    date]["request_attended"]
                average_session_time = 0
                if total_request_attended > 0:
                    average_session_time = total_session_time / total_request_attended

                if average_session_time < 60:
                    average_session_time = str(
                        round(average_session_time)) + " sec"
                else:
                    average_session_time = average_session_time / 60
                    average_session_time = str(
                        round(average_session_time)) + " min"

            analytics_data_summary[
                "total_session_time"] += easyassist_time_in_seconds(average_session_time)

            sheet1.write(index, 0, excel_sheet_date_format)
            sheet1.write(index, 1, total_links_generated)
            sheet1.write(index, 2, total_request_attended)
            sheet1.write(index, 3, total_customers_converted)
            sheet1.write(index, 4, total_customers_converted_by_url)
            sheet1.write(index, 5, average_session_time)

            index += 1

        sheet2 = analytics_workbook.add_sheet("Reverse Summary")

        sheet2.write(0, 0, "Links Generated")
        sheet2.col(0).width = 256 * 20
        sheet2.write(0, 1, "Sessions Joined by the Customer")
        sheet2.col(1).width = 256 * 30
        sheet2.write(0, 2, CUSTOMERS_CONVERTED_BY_AGENT)
        sheet2.col(2).width = 256 * 30
        sheet2.write(0, 3, CUSTOMERS_CONVERTED_THROUGH_URL)
        sheet2.col(3).width = 256 * 30
        sheet2.write(0, 4, AVERAGE_SESSION_TIME)
        sheet2.col(4).width = 256 * 20

        sheet2.write(1, 0, analytics_data_summary["links_generated"])
        sheet2.write(1, 1, analytics_data_summary["request_attended"])
        sheet2.write(1, 2, analytics_data_summary["customers_converted"])
        sheet2.write(
            1, 3, analytics_data_summary["customers_converted_by_url"])

        hours, minutes, seconds = easyassist_get_average_time(
            analytics_data_summary["total_session_time"], report_period_days)

        total_session_time = readable_time_format(hours, minutes, seconds)
        sheet2.write(1, 4, total_session_time)

        analytics_workbook.save(file_path)
        return absolute_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_reverse_analytics_data_dump %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def get_general_analytics_data_dump(requested_data, cobrowse_agent, CobrowseDateWiseInboundAnalytics, CobrowseDateWiseOutboundAnalytics, CobrowseDateWiseOutboundDroplinkAnalytics):
    try:
        logger.info("Inside get_general_analytics_data_dump",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT
        #               for email in email_str.split(",") if email != ""]

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        file_directory = settings.SECURE_MEDIA_ROOT + \
            "EasyAssistApp/GeneralAnalytics/" + \
            str(cobrowse_agent.user.username)
        if not path.exists(file_directory):
            os.makedirs(file_directory)

        file_path = file_directory + "/general_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"
        absolute_file_path = "/secured_files/EasyAssistApp/GeneralAnalytics/" + \
            str(cobrowse_agent.user.username) + \
            "/general_analytics_" + \
            str(start_date) + "-" + str(end_date) + ".xls"

        if path.exists(file_path):
            return absolute_file_path

        analytics_workbook = Workbook()

        add_inbound_analytics_in_excel_sheet(
            analytics_workbook, requested_data, cobrowse_agent, CobrowseDateWiseInboundAnalytics, CobrowseDateWiseOutboundDroplinkAnalytics)

        add_outbound_analytics_in_excel_sheet(
            analytics_workbook, requested_data, cobrowse_agent, CobrowseDateWiseOutboundAnalytics)

        analytics_workbook.save(file_path)
        return absolute_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_reverse_analytics_data_dump %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

    return "None"


def update_outbound_analytics_droplink_model(CobrowseDateWiseOutboundDroplinkAnalytics, CobrowseIO, CobrowseAccessToken, CobrowseDropLink, date):
    try:
        access_token_objs = CobrowseAccessToken.objects.all()

        cobrowse_io_objs = CobrowseIO.objects.filter(
            is_test=False,
            is_archived=True,
            is_lead=False,
            is_droplink_lead=True,
            is_reverse_cobrowsing=False,
            request_datetime__date=date)

        for access_token_obj in access_token_objs.iterator():
            admin_agent = access_token_obj.agent
            if admin_agent == None:
                continue

            agent_objs = get_list_agents_under_admin(
                admin_agent, is_active=None, is_account_active=None)
            agent_objs += get_list_supervisor_under_admin(admin_agent, None)

            filtered_cobrowse_io_objs = cobrowse_io_objs.filter(
                access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date)

            drop_link_objs = CobrowseDropLink.objects.filter(generate_datetime__date__gte=access_token_obj.go_live_date).filter(
                agent__in=agent_objs, generate_datetime__date=date, proxy_cobrowse_io=None)

            for agent in agent_objs:
                analytics_data = get_agent_wise_outbound_droplink_analytics(
                    filtered_cobrowse_io_objs, drop_link_objs, agent)
                if analytics_data == None:
                    continue

                cobrowse_analytics_obj = CobrowseDateWiseOutboundDroplinkAnalytics.objects.filter(
                    date=date, agent=agent).first()

                if cobrowse_analytics_obj == None:
                    cobrowse_analytics_obj = CobrowseDateWiseOutboundDroplinkAnalytics.objects.create(
                        agent=agent,
                        date=date,
                        request_initiated=analytics_data["request_initiated"],
                        request_attended=analytics_data["request_attended"],
                        request_unattended=analytics_data[
                            "request_unattended"],
                        customers_converted=analytics_data[
                            "customers_converted"],
                        customers_converted_by_url=analytics_data[
                            "customers_converted_by_url"],
                        declined_leads=analytics_data["declined_leads"],
                        total_session_time=analytics_data["total_session_time"],
                        total_droplinks_generated=analytics_data["total_links_generated"])
                else:
                    cobrowse_analytics_obj.request_initiated = analytics_data[
                        "request_initiated"]
                    cobrowse_analytics_obj.request_attended = analytics_data[
                        "request_attended"]
                    cobrowse_analytics_obj.request_unattended = analytics_data[
                        "request_unattended"]
                    cobrowse_analytics_obj.customers_converted = analytics_data[
                        "customers_converted"]
                    cobrowse_analytics_obj.customers_converted_by_url = analytics_data[
                        "customers_converted_by_url"]
                    cobrowse_analytics_obj.declined_leads = analytics_data[
                        "declined_leads"]
                    cobrowse_analytics_obj.total_session_time = analytics_data[
                        "total_session_time"]
                    cobrowse_analytics_obj.total_droplinks_generated = analytics_data[
                        "total_links_generated"]
                    cobrowse_analytics_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_outbound_analytics_model %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def get_agent_wise_outbound_droplink_analytics(cobrowse_io_objs, droplink_objs, agent):
    try:
        total_sr = 0
        total_sr_closed = 0
        total_sr_closed_by_url = 0
        total_sr_attended = 0
        total_wait_time = 0
        total_session_duration = 0
        total_wait_time_unattended = 0
        total_sr_unattended = 0
        declined_leads = 0

        cobrowse_io_objs = cobrowse_io_objs.filter(agent=agent)

        # total cobrowsing request Initiated/Notinitiated
        total_sr = cobrowse_io_objs.count()

        total_links_generated = droplink_objs.filter(agent=agent).count()

        if total_sr > 0:

            # total cobrowsing request initiated and converted successfully
            total_sr_closed = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=False).filter(~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request initiated and converted successfully by
            # landing at specified URL
            total_sr_closed_by_url = cobrowse_io_objs.filter(
                is_helpful=True, is_archived=True, is_lead_converted_by_url=True).filter(~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request attended
            total_sr_attended = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None)).count()

            # total cobrowsing request denied by customer
            declined_leads = cobrowse_io_objs.filter(
                consent_cancel_count__gte=1, allow_agent_cobrowse="false", cobrowsing_start_datetime=None).count()

            cobrowse_io_initiated = cobrowse_io_objs.filter(
                ~Q(cobrowsing_start_datetime=None))

            for cobrowse_io in cobrowse_io_initiated.iterator():
                total_session_duration += cobrowse_io.session_time_in_seconds()
                total_wait_time += cobrowse_io.customer_wait_time_in_seconds()

            cobrowse_io_not_initiated = cobrowse_io_objs.filter(
                Q(cobrowsing_start_datetime=None)).filter(
                ~Q(allow_agent_cobrowse="false")).exclude(session_archived_cause__in=["FOLLOWUP", "UNASSIGNED"])
            total_sr_unattended = cobrowse_io_not_initiated.count()

            for cobrowse_io in cobrowse_io_not_initiated.iterator():
                total_wait_time_unattended += cobrowse_io.customer_wait_time_in_seconds()

            return {
                "request_initiated": total_sr,
                "request_attended": total_sr_attended,
                "request_unattended": total_sr_unattended,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "declined_leads": declined_leads,
                "total_session_time": total_session_duration,
                "total_wait_time_unattended": total_wait_time_unattended,
                "total_links_generated": total_links_generated
            }

        if total_links_generated:
            return {
                "request_initiated": total_sr,
                "request_attended": total_sr_attended,
                "request_unattended": total_sr_unattended,
                "customers_converted": total_sr_closed,
                "customers_converted_by_url": total_sr_closed_by_url,
                "declined_leads": declined_leads,
                "total_session_time": total_session_duration,
                "total_wait_time_unattended": total_wait_time_unattended,
                "total_links_generated": total_links_generated
            }

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_wise_outbound_analytics %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

    return None


def get_requested_data_for_daily():
    # Last day data
    date_format = DATE_TIME_FORMAT

    start_date = datetime.datetime.now() - datetime.timedelta(days=1)
    start_date = start_date.strftime(date_format)
    end_date = start_date

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def get_requested_data_for_week():
    # Last 7 days data
    date_format = DATE_TIME_FORMAT

    start_date = datetime.datetime.now() - datetime.timedelta(days=7)
    start_date = start_date.strftime(date_format)
    end_date = datetime.datetime.now() - datetime.timedelta(days=1)
    end_date = end_date.strftime(date_format)

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def get_requested_data_for_month():
    # Last 30 days data
    date_format = DATE_TIME_FORMAT

    start_date = datetime.datetime.now() - datetime.timedelta(days=30)
    start_date = start_date.strftime(date_format)
    end_date = datetime.datetime.now() - datetime.timedelta(days=1)
    end_date = end_date.strftime(date_format)

    requested_data = {
        "startdate": start_date,
        "enddate": end_date,
    }
    return requested_data


def add_agent_wise_bifurcation_into_workbook(workbook, cobrowse_agent, start_date, end_date):

    from EasyAssistApp.models import CobrowseIO

    agents = get_list_agents_under_admin(
        cobrowse_agent, is_active=None, is_account_active=None)
    access_token_obj = cobrowse_agent.get_access_token_obj()

    cobrowse_io_objs_attended = CobrowseIO.objects.filter(
        is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
        request_datetime__date__gte=start_date, request_datetime__date__lte=end_date).filter(
        agent__in=agents).filter(
        ~Q(cobrowsing_start_datetime=None))

    cobrowse_io_objs_unattended = CobrowseIO.objects.filter(
        is_test=False, access_token=access_token_obj, request_datetime__date__gte=access_token_obj.go_live_date, is_archived=True).filter(
        request_datetime__date__gte=start_date, request_datetime__date__lte=end_date, agent__in=agents).filter(
        Q(cobrowsing_start_datetime=None)).filter(~Q(allow_agent_cobrowse="false")).exclude(session_archived_cause__in=["FOLLOWUP", "UNASSIGNED"])

    if not cobrowse_io_objs_attended and not cobrowse_io_objs_unattended:
        return False

    sheet = workbook.add_sheet("Agent Wise Bifurcation")

    sheet.write(0, 0, "Agent Email ID")
    sheet.col(0).width = 256 * 20
    sheet.write(0, 1, "Attended Leads")
    sheet.col(1).width = 256 * 20
    sheet.write(0, 2, "Unattended Leads")
    sheet.col(2).width = 256 * 20

    index = 1
    for agent in agents:

        sheet.write(index, 0, agent.user.username)
        sheet.write(index, 1, cobrowse_io_objs_attended.filter(
            agent=agent).count())
        sheet.write(index, 2, cobrowse_io_objs_unattended.filter(
            agent=agent).count())
        index += 1
    
    return True


def get_agent_wise_bifurcation(requested_data, cobrowse_agent, CobrowseIO, is_public=False):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/support-history/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/support-history/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "agent_wise_bifurcation_common_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/support-history/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        logger.info("Inside get_agent_wise_bifurcation",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        absolute_file_path = get_absolute_file_path()
        relative_file_path = get_relative_file_path()

        support_history_wb = Workbook()

        is_data_present = add_agent_wise_bifurcation_into_workbook(
            support_history_wb, cobrowse_agent, start_date, end_date)

        if is_data_present:
            support_history_wb.save(absolute_file_path)
            return relative_file_path

        return NO_DATA
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_agent_wise_bifurcation! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
        return None


def get_data_dump_in_one_file(requested_data, cobrowse_agent, history_type_list, is_public=False):

    def get_folder_path():
        if is_public:
            folder_path = "files/EasyAssistApp/CommonDataDump/" + \
                str(cobrowse_agent.user.username) + "/"
        else:
            folder_path = "EasyAssistApp/CommonDataDump/" + \
                str(cobrowse_agent.user.username) + "/"
        return folder_path

    def get_file_name(extention):
        file_name = "data_dump_" + \
            str(start_date) + "_to_" + str(end_date) + extention
        return file_name

    def get_relative_file_path(extention=".xls"):
        if is_public:
            relative_file_path = get_folder_path() + get_file_name(extention)
        else:
            relative_file_path = SECURED_FILES_PATH + \
                get_folder_path() + get_file_name(extention)
        return relative_file_path

    def get_absolute_file_path(extention=".xls"):
        if is_public:
            absolute_folder_path = settings.MEDIA_ROOT + \
                "EasyAssistApp/CommonDataDump/" + \
                str(cobrowse_agent.user.username) + "/"
            if not path.exists(absolute_folder_path):
                os.makedirs(absolute_folder_path)
        else:
            absolute_folder_path = settings.SECURE_MEDIA_ROOT + get_folder_path()
            if not path.exists(absolute_folder_path):
                os.mkdir(absolute_folder_path)
        absolute_file_path = absolute_folder_path + get_file_name(extention)
        return absolute_file_path

    try:
        logger.info("Inside get_data_dump_in_one_file",
                    extra={'AppName': 'EasyAssist'})
        from datetime import datetime, timedelta
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]

        date_format = DATE_TIME_FORMAT

        start_date = datetime.strptime(start_date, date_format).date()
        end_date = datetime.strptime(end_date, date_format).date()

        absolute_file_path = get_absolute_file_path()
        relative_file_path = get_relative_file_path()

        support_history_wb = Workbook()
        is_attended_present = True
        is_unattended_leads_present = True
        is_declined_leads_present = True
        is_follow_upleads_present = True
        is_agent_wise_data_present = True
        if "attended_leads" in history_type_list:
            is_attended_present = add_support_history_data_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

        if "unattended_leads" in history_type_list:
            is_unattended_leads_present = add_unattended_leads_history_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

        if "declined_leads" in history_type_list:
            is_declined_leads_present = add_declined_leads_history_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

        if "followup_leads" in history_type_list:
            is_follow_upleads_present = add_followup_leads_history_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

        if "agent_wise_reports" in history_type_list:
            is_agent_wise_data_present = add_agent_wise_bifurcation_into_workbook(
                support_history_wb, cobrowse_agent, start_date, end_date)

        is_data_present = is_attended_present or is_follow_upleads_present or is_agent_wise_data_present or \
            is_declined_leads_present or is_unattended_leads_present
        if is_data_present:
            support_history_wb.save(absolute_file_path)
            return relative_file_path

        return NO_DATA
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_data_dump_in_one_file! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
        return None


def generate_crm_api_document(access_token):

    def docx_replace_regex(doc_obj, regex, replace_text):
        for para in doc_obj.paragraphs:
            if regex.search(para.text):
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for index in range(len(inline)):
                    if regex.search(inline[index].text):
                        text = regex.sub(replace_text, inline[index].text)
                        inline[index].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace_text)

    def edit_document(document_path, target_path, metadata_dict):
        document_obj = Document(document_path)

        for metadata in metadata_dict:
            docx_replace_regex(document_obj, re.compile(
                r"" + metadata + ""), metadata_dict[metadata])

        document_obj.save(target_path)

    def get_metadata_dict():
        crm_integration_url = settings.EASYCHAT_HOST_URL + \
            "/easy-assist/crm/" + document_details["url_suffix"] + "/"

        metadata_dict = {
            "crm_integration_url": crm_integration_url
        }

        return metadata_dict

    def get_base_document_path(document_details):
        original_file_name = document_details["original_file_name"]
        base_document_path = "files/templates/easyassist-crm-template/" + original_file_name
        return base_document_path

    def get_document_folder_path():
        target_document_folder = CRM_DOCUMENTS_PATH + \
            str(access_token.key)
        create_directory(target_document_folder)
        return target_document_folder

    def get_target_document_path(document_details):
        target_document_folder = get_document_folder_path()
        original_file_name = document_details["original_file_name"]
        target_document_path = target_document_folder + "/" + original_file_name
        return target_document_path

    try:
        for document_type, document_details in CRM_DOCUMENTS.items():
            base_document_path = get_base_document_path(document_details)
            target_document_path = get_target_document_path(document_details)
            metadata_dict = get_metadata_dict()
            edit_document(base_document_path,
                          target_document_path, metadata_dict)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("generate_crm_api_document! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})
        return None


def get_crm_integration_model(request, CRMIntegrationModel):
    try:
        authorization_header = request.META['HTTP_AUTHORIZATION']
        auth_token = authorization_header.replace("Bearer ", "")

        min_datetime = datetime.datetime.now() - datetime.timedelta(hours=1)
        integration_obj = CRMIntegrationModel.objects.filter(
            auth_token=auth_token, is_expired=False, datetime__gte=min_datetime).first()
    except Exception:
        integration_obj = None

    return integration_obj


def change_agent_is_active_flag(status, cobrowse_agent, CobrowseAgentOnlineAuditTrail, CobrowseAgentWorkAuditTrail):

    if status == True:
        time_threshold = datetime.datetime.now() - datetime.timedelta(minutes=1)
        agent_online_audit_trail_obj = CobrowseAgentOnlineAuditTrail.objects.filter(
            agent=cobrowse_agent,
            last_online_end_datetime__gte=time_threshold).order_by(
                '-last_online_start_datetime').first()

        if agent_online_audit_trail_obj != None:
            agent_online_audit_trail_obj.last_online_end_datetime = timezone.now()
            agent_online_audit_trail_obj.save()
        else:
            agent_online_audit_trail_obj = CobrowseAgentOnlineAuditTrail.objects.create(
                agent=cobrowse_agent,
                last_online_start_datetime=timezone.now(),
                last_online_end_datetime=timezone.now())
            agent_online_audit_trail_obj.save()

        update_online_audit_trail_idle_time(
            cobrowse_agent, agent_online_audit_trail_obj, CobrowseAgentWorkAuditTrail)

    cobrowse_agent.is_active = status
    cobrowse_agent.last_agent_active_datetime = timezone.now()
    cobrowse_agent.is_cobrowsing_active = False
    cobrowse_agent.save()


def change_agent_is_account_active_flag(activate, active_agent, cobrowse_agent, request, data, CobrowsingAuditTrail):

    cobrowse_agent.is_account_active = activate
    cobrowse_agent.is_active = cobrowse_agent.is_active and activate

    if not activate:
        description = cobrowse_agent.user.username + \
            " (" + cobrowse_agent.role + ") was marked as inactive through CRM Api"
        save_audit_trail(
            active_agent, COBROWSING_DEACTIVATEUSER_ACTION, description, CobrowsingAuditTrail)
        cobrowse_agent.agent_deactivation_datetime = timezone.now()
        cobrowse_agent.save()
        add_audit_trail(
            "EASYASSISTAPP",
            active_agent.user,
            "Deactivate-User",
            description,
            json.dumps(data),
            request.META.get("PATH_INFO"),
            request.META.get('HTTP_X_FORWARDED_FOR')
        )
    else:
        description = cobrowse_agent.user.username + \
            " (" + cobrowse_agent.role + ") was marked as active through CRM Api"
        save_audit_trail(
            active_agent, COBROWSING_ACTIVATEUSER_ACTION, description, CobrowsingAuditTrail)
        cobrowse_agent.agent_activation_datetime = timezone.now()
        cobrowse_agent.save()
        add_audit_trail(
            "EASYASSISTAPP",
            active_agent.user,
            "Activate-User",
            description,
            json.dumps(data),
            request.META.get("PATH_INFO"),
            request.META.get('HTTP_X_FORWARDED_FOR')
        )


def parse_meeting_details_crm(cobrowse_video_obj, CobrowseVideoAuditTrail):
    response = {}

    meeting_audit_trail = CobrowseVideoAuditTrail.objects.filter(
        cobrowse_video=cobrowse_video_obj).first()

    feedback_rating = cobrowse_video_obj.feedback_rating
    agent_remarks = cobrowse_video_obj.agent_comments
    customer_remarks = cobrowse_video_obj.feedback_comment
    message_history = None
    agent_notes = None
    screenshots = None
    customer_name = cobrowse_video_obj.full_name
    mobile_number = cobrowse_video_obj.mobile_number

    agent_join_time = None
    meeting_status = "NotStarted"
    meeting_duration = None

    meeting_start_date_time = None
    if cobrowse_video_obj.meeting_start_date and cobrowse_video_obj.meeting_start_time:
        meeting_start_date = cobrowse_video_obj.meeting_start_date.strftime(
            DATE_TIME_FORMAT)
        meeting_start_time = str(
            cobrowse_video_obj.meeting_start_time.strftime("%I:%M:%S %p"))
        meeting_start_date_time = meeting_start_date + " " + meeting_start_time

    if meeting_audit_trail:
        message_history = meeting_audit_trail.get_meeting_chats()
        agent_notes = meeting_audit_trail.agent_notes
        screenshots = meeting_audit_trail.get_meeting_screenshot_links()

        meeting_duration = meeting_audit_trail.get_meeting_duration_in_seconds()

        meeting_status = None
        if meeting_audit_trail.is_meeting_ended == True:
            meeting_status = "Completed"
        else:
            meeting_status = "Ongoing"

        agent_join_time = None
        if meeting_audit_trail.agent_joined:
            agent_join_time = meeting_audit_trail.agent_joined.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%Y-%m-%d %I:%M:%S %p")

    response["nps"] = feedback_rating
    response["customer_name"] = customer_name
    response["mobile_number"] = mobile_number
    response["meeting_start_date_time"] = meeting_start_date_time
    response["agent_join_time"] = agent_join_time
    response["meeting_status"] = meeting_status
    response["meeting_duration"] = meeting_duration
    response["agent_remarks"] = agent_remarks
    response["customer_remarks"] = customer_remarks
    response["message_history"] = message_history
    response["agent_notes"] = agent_notes
    response["screenshots"] = screenshots

    return response


def update_sandbox_user_cobrowsing_settings(sandbox_user_obj, access_token_obj):
    try:
        if sandbox_user_obj.enable_cobrowsing:
            access_token_obj.allow_generate_meeting = True
            access_token_obj.allow_video_meeting_only = False
            if sandbox_user_obj.enable_reverse_cobrowsing:
                access_token_obj.allow_agent_to_customer_cobrowsing = True
                access_token_obj.enable_inbound = False
                access_token_obj.lead_generation = False

            else:
                if sandbox_user_obj.enable_inbound:
                    access_token_obj.allow_agent_to_customer_cobrowsing = False
                    access_token_obj.enable_inbound = True

                if sandbox_user_obj.enable_outbound:
                    access_token_obj.allow_agent_to_customer_cobrowsing = False
                    access_token_obj.lead_generation = True

        if sandbox_user_obj.enable_video_meeting:
            access_token_obj.allow_generate_meeting = True
            access_token_obj.allow_video_meeting_only = True
            access_token_obj.enable_inbound = False
            access_token_obj.lead_generation = False
            access_token_obj.allow_agent_to_customer_cobrowsing = False
            access_token_obj.enable_inbound = False
            access_token_obj.allow_agent_to_customer_cobrowsing = False
            access_token_obj.lead_generation = False

        access_token_obj.whitelisted_domain = "*"
        access_token_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("update_sandbox_user_cobrowsing_settings! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})


def easyassist_check_for_admin_expired_credentials(user_obj, CobrowseAgent, CobrowseSandboxUser):

    try:
        cobrowse_agent = CobrowseAgent.objects.filter(user=user_obj)

        if cobrowse_agent.count() == 0:
            return False

        cobrowse_agent = cobrowse_agent.first()

        admin_agent = get_admin_from_active_agent(
            cobrowse_agent, CobrowseAgent)

        user_obj = admin_agent.user

        if easyassist_check_for_expired_credentials(user_obj, CobrowseSandboxUser):
            return True

        return False

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("easyassist_check_for_expired_credentials! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return False


def easyassist_check_for_expired_credentials(user_obj, CobrowseSandboxUser):
    try:
        sandbox_user_obj = CobrowseSandboxUser.objects.filter(
            user=user_obj).first()

        if not sandbox_user_obj:
            return False

        if sandbox_user_obj.is_expired:
            return True

        current_date = timezone.now().date()
        expiration_date = sandbox_user_obj.expire_date
        if current_date > expiration_date:
            sandbox_user_obj.is_expired = True
            sandbox_user_obj.save()
            return True

        return False
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("easyassist_check_for_expired_credentials! %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return False


"""
function: get_number_of_day
input params:
    year:
    month:
expected output:
    1. Provides number of days in a month of a perticular year.
"""


def get_number_of_day(year, month):
    year = int(year)
    month = int(month)
    leap = 0
    if year % 400 == 0:
        leap = 1
    elif year % 100 == 0:
        leap = 0
    elif year % 4 == 0:
        leap = 1
    if month == 2:
        return 28 + leap
    list = [1, 3, 5, 7, 8, 10, 12]
    if month in list:
        return 31
    return 30


"""
function: check_and_update_based_on_event_type
input params:
    user:
    calender_obj:
    selected_event:
    start_time:
    end_time:
    description: description of event
    auto_response: auto response will be given to livechat customer on this day.
expected output:
    1. If user is admin, then this event will be updated based on event type.
"""


def check_and_update_based_on_event_type(user, calender_obj, selected_event, start_time, end_time, description, auto_response):
    try:
        if user.role == "admin":
            if str(selected_event) == "1":
                calender_obj.start_time = start_time
                calender_obj.end_time = end_time
                calender_obj.event_type = "1"
            else:
                calender_obj.description = description
                calender_obj.auto_response = auto_response
                calender_obj.event_type = "2"
            calender_obj.modified_by = user
            calender_obj.modified_date = timezone.now()
            calender_obj.save()

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside check_and_update_based_on_event_type: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def enable_livechat_cobrowse_settings(access_token_obj):
    try:
        access_token_obj.enable_chat_functionality = False
        access_token_obj.enable_auto_assign_unattended_lead = False
        access_token_obj.show_floating_easyassist_button = False
        access_token_obj.show_easyassist_connect_agent_icon = False
        access_token_obj.enable_invite_agent_in_cobrowsing = True
        access_token_obj.enable_agent_connect_message = True
        access_token_obj.show_verification_code_modal = False
        access_token_obj.enable_verification_code_popup = False
        access_token_obj.field_stuck_event_handler = False
        access_token_obj.save()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside enable_livechat_cobrowse_settings: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def is_system_command_present(code):
    try:
        for cmd in SYSTEM_COMMANDS_LIST:
            if cmd in code:
                return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside is_system_command_present: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return False


def create_excel_easyassist_unique_customers(cobrowsing_type, active_agent, cobrowse_io_objs, access_token_obj, requested_data):
    from datetime import datetime
    filename = None
    try:
        logger.info("create_excel_easyassist_unique_customers",
                    extra={'AppName': 'EasyAssist'})

        easyassist_unique_customers_wb = Workbook()

        def create_new_sheet():
            nonlocal wb_index

            global auto_assigned_agent_col, agents_invited_col, invited_agents_connected_col, session_start_datetime_col, session_end_datetime_col, time_spent_col, \
                lead_status_col, nps_col, session_id_col, session_started_by_col, session_ended_by_col, agent_remarks_col, session_initiated_by, url_col, session_duration_col

            sheet_name = cobrowsing_type.title() + "_Unique_Customers-" + str(wb_index)
            sheet = easyassist_unique_customers_wb.add_sheet(sheet_name)

            sheet.write(0, 0, "Customer Name")
            sheet.col(0).width = 256 * 25
            sheet.write(0, 1, "Customer Mobile Number")
            sheet.col(1).width = 256 * 45
            sheet.write(0, 2, "Location")
            sheet.col(2).width = 256 * 50
            sheet.write(0, 3, "Primary Agent")
            sheet.col(3).width = 256 * 30

            prev_col = 3

            if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
                auto_assigned_agent_col = prev_col + 1
                sheet.write(0, auto_assigned_agent_col, "Auto-assigned Agent")
                sheet.col(auto_assigned_agent_col).width = 256 * 35
                prev_col = auto_assigned_agent_col

            session_initiated_by = prev_col + 1
            sheet.write(0, session_initiated_by, "Session Initiated By")
            sheet.col(session_initiated_by).width = 256 * 35
            prev_col = session_initiated_by

            if access_token_obj is not None and access_token_obj.enable_invite_agent_in_cobrowsing:
                agents_invited_col = prev_col + 1
                sheet.write(0, agents_invited_col, "Agents Invited")
                sheet.col(agents_invited_col).width = 256 * 35
                prev_col = agents_invited_col

                invited_agents_connected_col = prev_col + 1
                sheet.write(0, invited_agents_connected_col,
                            "Invited Agent Connected")
                sheet.col(invited_agents_connected_col).width = 256 * 35
                prev_col = invited_agents_connected_col

            time_spent_col = prev_col + 1
            sheet.write(0, time_spent_col, "Wait Time")
            sheet.col(time_spent_col).width = 256 * 20
            prev_col = time_spent_col

            session_start_datetime_col = prev_col + 1
            sheet.write(0, session_start_datetime_col,
                        "Cobrowsing Start Date & Time")
            sheet.col(session_start_datetime_col).width = 256 * 25
            prev_col = session_start_datetime_col

            session_end_datetime_col = prev_col + 1
            sheet.write(0, session_end_datetime_col, "Session End Date & Time")
            sheet.col(session_end_datetime_col).width = 256 * 25
            prev_col = session_end_datetime_col

            session_duration_col = prev_col + 1
            sheet.write(0, session_duration_col, "Session Duration")
            sheet.col(session_duration_col).width = 256 * 15
            prev_col = session_duration_col

            lead_status_col = prev_col + 1
            sheet.write(0, lead_status_col, "Lead Status")
            sheet.col(lead_status_col).width = 256 * 15
            prev_col = lead_status_col

            url_col = prev_col + 1
            sheet.write(0, url_col, "Captured from (Webpage)")
            sheet.col(url_col).width = 256 * 60
            prev_col = url_col

            nps_col = prev_col + 1
            sheet.write(0, nps_col, "NPS")
            sheet.col(nps_col).width = 256 * 60
            prev_col = nps_col

            session_id_col = prev_col + 1
            sheet.write(0, session_id_col, "Session ID")
            sheet.col(session_id_col).width = 256 * 45
            prev_col = session_id_col

            session_started_by_col = prev_col + 1
            sheet.write(0, session_started_by_col, "Session Started By")
            sheet.col(session_started_by_col).width = 256 * 20
            prev_col = session_started_by_col

            session_ended_by_col = prev_col + 1
            sheet.write(0, session_ended_by_col, "Session Ended By")
            sheet.col(session_ended_by_col).width = 256 * 20
            prev_col = session_ended_by_col

            agent_remarks_col = prev_col + 1
            sheet.write(0, agent_remarks_col, "Agent Remarks")
            sheet.col(agent_remarks_col).width = 256 * 80
            prev_col = agent_remarks_col

            return sheet

        index = 1
        wb_index = 1
        sheet = create_new_sheet()
        cobrowse_io_objs = cobrowse_io_objs.order_by('request_datetime')
        for cobrowse_io_obj in cobrowse_io_objs.iterator():
            if index > 50000:
                wb_index += 1
                sheet = create_new_sheet()
                index = 1

            sheet.write(index, 0, cobrowse_io_obj.full_name)
            sheet.write(index, 1, get_masked_data_if_hashed(
                cobrowse_io_obj.mobile_number))

            longitude = cobrowse_io_obj.longitude
            latitude = cobrowse_io_obj.latitude

            if latitude == 'None' or longitude == 'None':
                sheet.write(index, 2, "-")
            elif latitude == 'null' or longitude == 'null':
                sheet.write(index, 2, "-")
            else:
                try:
                    location = geocoder.google(
                        [latitude, longitude], method="reverse", key=GOOGLE_GEOCODER_KEY)
                    if location.address != None:
                        sheet.write(index, 2, location.address)
                    else:
                        sheet.write(index, 2, "-")
                except Exception:
                    sheet.write(index, 2, "-")
                    logger.warning("Error in create_excel_easyassist_uniue_customers, client's address not found",
                                   extra={'AppName': 'EasyAssist'})

            if cobrowse_io_obj.main_primary_agent:
                primary_agent_name = cobrowse_io_obj.main_primary_agent.user.username
            else:
                primary_agent_name = cobrowse_io_obj.agent.user.username

            sheet.write(index, 3, primary_agent_name)

            if access_token_obj is not None and access_token_obj.enable_auto_assign_unattended_lead:
                sheet.write(index, auto_assigned_agent_col,
                            cobrowse_io_obj.get_auto_assigned_agents())

            sheet.write(index, session_initiated_by,
                        cobrowse_io_obj.get_lead_initiated_by())

            if access_token_obj is not None and access_token_obj.enable_invite_agent_in_cobrowsing:
                agents_invited_str = "-"
                agents_invited_connected_str = "-"
                if cobrowse_io_obj.get_cobrowse_support_agent_details():
                    agents_invited = cobrowse_io_obj.get_cobrowse_support_agent_details(
                    ).support_agents_invited.all()
                    if agents_invited:
                        agents_invited_str = ""
                        for agent in agents_invited.iterator():
                            agents_invited_str = agents_invited_str + agent.user.username + ", "
                        agents_invited_str = agents_invited_str[:-2]

                    agents_invited_connected = cobrowse_io_obj.get_cobrowse_support_agent_details(
                    ).support_agents_joined.all()
                    if agents_invited_connected:
                        agents_invited_connected_str = ""
                        for agent in agents_invited_connected.iterator():
                            agents_invited_connected_str = agents_invited_connected_str + \
                                agent.user.username + ", "
                        agents_invited_connected_str = agents_invited_connected_str[:-2]

                sheet.write(index, agents_invited_col, agents_invited_str)
                sheet.write(index, invited_agents_connected_col,
                            agents_invited_connected_str)

            if cobrowse_io_obj.customer_wait_time_seconds():
                sheet.write(index, time_spent_col,
                            cobrowse_io_obj.customer_wait_time_seconds())
            else:
                sheet.write(index, time_spent_col, "-")

            if(cobrowse_io_obj.cobrowsing_start_datetime):
                sheet.write(index, session_start_datetime_col, str(cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
                sheet.write(index, session_end_datetime_col, str(cobrowse_io_obj.last_update_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            else:
                sheet.write(index, session_start_datetime_col, "-")
                sheet.write(index, session_end_datetime_col, "-")
            sheet.write(index, session_duration_col, str(
                cobrowse_io_obj.total_time_spent()))

            if cobrowse_io_obj.is_helpful:
                sheet.write(index, lead_status_col, "Converted")
            else:
                sheet.write(index, lead_status_col, "Not Converted")

            sheet.write(index, url_col, cobrowse_io_obj.title)

            if cobrowse_io_obj.agent_rating != None:
                sheet.write(index, nps_col, cobrowse_io_obj.agent_rating)
            else:
                sheet.write(index, nps_col, "-")

            sheet.write(index, session_id_col, str(cobrowse_io_obj.session_id))

            sheet.write(index, session_started_by_col, cobrowse_io_obj.get_session_started_by())

            if cobrowse_io_obj.session_archived_cause:
                if cobrowse_io_obj.session_archived_cause in ["AGENT_ENDED", "AGENT_INACTIVITY"]:
                    sheet.write(index, session_ended_by_col, "Agent")
                elif cobrowse_io_obj.session_archived_cause in ["CLIENT_ENDED", "CLIENT_INACTIVITY"]:
                    sheet.write(index, session_ended_by_col, "Customer")
                else:
                    sheet.write(index, session_ended_by_col, cobrowse_io_obj.session_archived_cause[0].upper(
                    ) + cobrowse_io_obj.session_archived_cause[1:].lower())
            else:
                sheet.write(index, session_ended_by_col, "-")

            if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                sheet.write(index, agent_remarks_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                ).first().agent_comments)
            else:
                sheet.write(index, agent_remarks_col, "-")

            index += 1

        try:
            start_date = datetime.strptime(
                requested_data["start_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
            end_date = datetime.strptime(
                requested_data["end_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
        except Exception:
            start_date = requested_data["start_date"]
            end_date = requested_data["end_date"]

        file_directory = EASYASSISTAPP_SECURED_FILES_PATH + "UniqueCustomers/Inbound/" + \
            active_agent.user.username + "/"

        if not os.path.exists(file_directory):
            os.makedirs(file_directory)

        file_name = 'Inbound_Unique_Customers_'

        if cobrowsing_type == "reverse":
            file_name = 'Reverse_Unique_Customers_'

        filename_prefix = file_directory + file_name + \
            str(start_date) + '-' + str(end_date) + "_" + \
            active_agent.user.username

        filename = filename_prefix + ".xls"
        easyassist_unique_customers_wb.save(filename)
        return filename
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_unique_customers %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_easyassist_repeated_customers(cobrowsing_type, cobrowse_io_objs, active_agent, requested_data):
    from datetime import datetime
    filename = None
    try:
        logger.info("create_excel_easyassist_repeated_customers",
                    extra={'AppName': 'EasyAssist'})

        easyassist_repeated_customers_wb = Workbook()

        def create_new_sheet():
            nonlocal wb_index
            sheet_name = cobrowsing_type.title() + "_Repeated_Customers-" + str(wb_index)
            sheet = easyassist_repeated_customers_wb.add_sheet(
                sheet_name)

            sheet.write(0, 0, "Customer Name")
            sheet.col(0).width = 256 * 25
            sheet.write(0, 1, "Customer Mobile Number")
            sheet.col(1).width = 256 * 45
            sheet.write(0, 2, "Requests Initiated")
            sheet.col(2).width = 256 * 25
            sheet.write(0, 3, "Co-browsing Sessions attended")
            sheet.col(3).width = 256 * 25
            sheet.write(0, 4, "Average Session Duration")
            sheet.col(4).width = 256 * 25
            sheet.write(0, 5, "Average Wait Time")
            sheet.col(5).width = 256 * 45

            return sheet

        wb_index = 1
        sheet = create_new_sheet()
        index = 1
        total_time_spent = 0
        prev_value = 0
        cobrowse_io_objs = cobrowse_io_objs.order_by("mobile_number")
        for cobrowse_io_obj in cobrowse_io_objs.iterator():
            if index > 50000:
                wb_index += 1
                sheet = create_new_sheet()
                index = 1
            if prev_value != cobrowse_io_obj.mobile_number:
                sheet.write(index, 0, cobrowse_io_obj.full_name)
                sheet.write(index, 1, get_masked_data_if_hashed(
                    cobrowse_io_obj.mobile_number))
                prev_value = cobrowse_io_obj.mobile_number
            else:
                continue

            total_request_from_mobile_number = cobrowse_io_objs.filter(
                mobile_number=prev_value).count()
            sheet.write(index, 2, total_request_from_mobile_number)

            customer_count = cobrowse_io_objs.filter(mobile_number=prev_value).filter(
                ~Q(cobrowsing_start_datetime=None)).count()
            sheet.write(index, 3, customer_count)

            mobile_number_details = cobrowse_io_objs.filter(
                mobile_number=prev_value)

            total_time_spent = 0
            total_wait_time = 0
            for cust_obj in mobile_number_details.iterator():
                total_time_spent += cust_obj.total_time_spent_seconds()
                total_wait_time += cust_obj.customer_wait_time_seconds()
            if mobile_number_details.count() != 0:
                avg_time_spent = total_time_spent / mobile_number_details.count()
                sheet.write(index, 4, round(avg_time_spent, 2))
            else:
                sheet.write(index, 4, " - ")
            if customer_count != 0:
                avg_wait_time = total_wait_time / customer_count
                sheet.write(index, 5, round(avg_wait_time, 2))
            else:
                sheet.write(index, 5, "-")

            index += 1

        try:
            start_date = datetime.strptime(
                requested_data["start_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
            end_date = datetime.strptime(
                requested_data["end_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
        except Exception:
            start_date = requested_data["start_date"]
            end_date = requested_data["end_date"]

        file_directory = EASYASSISTAPP_SECURED_FILES_PATH + "RepeatedCustomers/Inbound/" + \
            active_agent.user.username + "/"

        if not os.path.exists(file_directory):
            os.makedirs(file_directory)

        file_name = 'Inbound_Repeated_Customers_'

        if cobrowsing_type == "reverse":
            file_name = 'Reverse_Repeated_Customers_'

        filename_prefix = file_directory + file_name + \
            str(start_date) + '-' + str(end_date) + "_" + \
            active_agent.user.username

        filename = filename_prefix + ".xls"
        easyassist_repeated_customers_wb.save(filename)
        return filename
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_repeated_customers %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_easyassist_unique_customers_outbound(cobrowse_io_objs, active_agent, access_token_obj, requested_data):
    from EasyAssistApp.models import CobrowseCapturedLeadData, CobrowseLeadHTMLField
    from datetime import datetime
    filename = None
    try:
        logger.info("create_excel_easyassist_unique_customers_outbound",
                    extra={'AppName': 'EasyAssist'})

        easyassist_unique_customers_wb = Workbook()

        def create_new_sheet():
            nonlocal wb_index
            global agents_invited_col, invited_agents_connected_col, session_start_datetime_col, session_end_datetime_col, \
                lead_status_col, nps_col, session_id_col, agent_remarks_col, url_col, session_duration_col, location_col, \
                agent_col

            prev_col = 0
            sheet = easyassist_unique_customers_wb.add_sheet(
                "Outbound_Unique_Customers-" + str(wb_index), cell_overwrite_ok=True)

            sheet.write(0, prev_col, "Customer Details")
            sheet.col(prev_col).width = 256 * 50
            location_col = prev_col + 1
            sheet.write(0, location_col, "Location")
            sheet.col(location_col).width = 256 * 35
            agent_col = location_col + 1
            sheet.write(0, agent_col, "Primary Agent")
            sheet.col(agent_col).width = 256 * 35
            prev_col = agent_col
            if access_token_obj is not None and access_token_obj.enable_invite_agent_in_cobrowsing:
                agents_invited_col = prev_col + 1
                sheet.write(0, agents_invited_col, "Agents Invited")
                sheet.col(agents_invited_col).width = 256 * 35
                prev_col = agents_invited_col

                invited_agents_connected_col = prev_col + 1
                sheet.write(0, invited_agents_connected_col,
                            "Invited Agent Connected")
                sheet.col(invited_agents_connected_col).width = 256 * 35
                prev_col = invited_agents_connected_col

            session_start_datetime_col = prev_col + 1
            sheet.write(0, session_start_datetime_col,
                        "Cobrowsing Start Date & Time")
            sheet.col(session_start_datetime_col).width = 256 * 25
            prev_col = session_start_datetime_col

            session_end_datetime_col = prev_col + 1
            sheet.write(0, session_end_datetime_col, "Session End Date & Time")
            sheet.col(session_end_datetime_col).width = 256 * 25
            prev_col = session_end_datetime_col

            session_duration_col = prev_col + 1
            sheet.write(0, session_duration_col, "Session Duration")
            sheet.col(session_duration_col).width = 256 * 15
            prev_col = session_duration_col

            lead_status_col = prev_col + 1
            sheet.write(0, lead_status_col, "Lead Status")
            sheet.col(lead_status_col).width = 256 * 15
            prev_col = lead_status_col

            url_col = prev_col + 1
            sheet.write(0, url_col, "Captured from (Webpage)")
            sheet.col(url_col).width = 256 * 60
            prev_col = url_col

            nps_col = prev_col + 1
            sheet.write(0, nps_col, "NPS")
            sheet.col(nps_col).width = 256 * 15
            prev_col = nps_col

            session_id_col = prev_col + 1
            sheet.write(0, session_id_col, "Session ID")
            sheet.col(session_id_col).width = 256 * 45
            prev_col = session_id_col

            agent_remarks_col = prev_col + 1
            sheet.write(0, agent_remarks_col, "Agent Remarks")
            sheet.col(agent_remarks_col).width = 256 * 80
            prev_col = agent_remarks_col

            return sheet

        index = 1
        wb_index = 1
        sheet = create_new_sheet()
        cobrowse_io_objs = cobrowse_io_objs.order_by('request_datetime')
        for cobrowse_io_obj in cobrowse_io_objs.iterator():
            if index > 50000:
                wb_index += 1
                sheet = create_new_sheet()
                index = 1

            data_captured = CobrowseCapturedLeadData.objects.filter(
                session_id=cobrowse_io_obj.session_id, agent_searched=True).first()

            if data_captured:
                label_name = CobrowseLeadHTMLField.objects.filter(
                    id=data_captured.search_field_id).first().tag_label
                sheet.write(index, 0, label_name + " : " +
                            str(data_captured.primary_value))
            else:
                continue

            longitude = cobrowse_io_obj.longitude
            latitude = cobrowse_io_obj.latitude

            if latitude == 'None' or longitude == 'None':
                sheet.write(index, location_col, "-")
            else:
                try:
                    location = geocoder.google(
                        [latitude, longitude], method="reverse", key=GOOGLE_GEOCODER_KEY)
                    if location.address != None:
                        sheet.write(index, location_col, location.address)
                    else:
                        sheet.write(index, location_col, "-")
                except Exception:
                    sheet.write(index, location_col, "-")
                    logger.warning("Error in create_excel_easyassist_uniue_customers_outbound, client's address not found",
                                   extra={'AppName': 'EasyAssist'})

            if cobrowse_io_obj.main_primary_agent:
                primary_agent_name = cobrowse_io_obj.main_primary_agent.user.username
            else:
                primary_agent_name = cobrowse_io_obj.agent.user.username

            sheet.write(index, agent_col, primary_agent_name)

            if access_token_obj is not None and access_token_obj.enable_invite_agent_in_cobrowsing:
                agents_invited_str = "-"
                agents_invited_connected_str = "-"
                if cobrowse_io_obj.get_cobrowse_support_agent_details():
                    agents_invited = cobrowse_io_obj.get_cobrowse_support_agent_details(
                    ).support_agents_invited.all()
                    if agents_invited:
                        agents_invited_str = ""
                        for agent in agents_invited.iterator():
                            agents_invited_str = agents_invited_str + agent.user.username + ", "
                        agents_invited_str = agents_invited_str[:-2]

                    agents_invited_connected = cobrowse_io_obj.get_cobrowse_support_agent_details(
                    ).support_agents_joined.all()
                    if agents_invited_connected:
                        agents_invited_connected_str = ""
                        for agent in agents_invited_connected.iterator():
                            agents_invited_connected_str = agents_invited_connected_str + \
                                agent.user.username + ", "
                        agents_invited_connected_str = agents_invited_connected_str[:-2]

                sheet.write(index, agents_invited_col, agents_invited_str)
                sheet.write(index, invited_agents_connected_col,
                            agents_invited_connected_str)

            if(cobrowse_io_obj.cobrowsing_start_datetime):
                sheet.write(index, session_start_datetime_col, str(cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
                sheet.write(index, session_end_datetime_col, str(cobrowse_io_obj.last_update_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            else:
                sheet.write(index, session_start_datetime_col, "-")
                sheet.write(index, session_end_datetime_col, "-")
            sheet.write(index, session_duration_col, str(
                cobrowse_io_obj.total_time_spent()))

            if cobrowse_io_obj.is_helpful:
                sheet.write(index, lead_status_col, "Converted")
            else:
                sheet.write(index, lead_status_col, "Not Converted")
            sheet.write(index, url_col, cobrowse_io_obj.title)
            if cobrowse_io_obj.agent_rating != None:
                sheet.write(index, nps_col, cobrowse_io_obj.agent_rating)
            else:
                sheet.write(index, nps_col, "-")
            sheet.write(index, session_id_col, str(cobrowse_io_obj.session_id))

            if cobrowse_io_obj.get_cobrowsing_session_closing_comments():
                sheet.write(index, agent_remarks_col, cobrowse_io_obj.get_cobrowsing_session_closing_comments(
                ).first().agent_comments)
            else:
                sheet.write(index, agent_remarks_col, "-")

            index += 1

        try:
            start_date = datetime.strptime(
                requested_data["start_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
            end_date = datetime.strptime(
                requested_data["end_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
        except Exception:
            start_date = requested_data["start_date"]
            end_date = requested_data["end_date"]

        file_directory = EASYASSISTAPP_SECURED_FILES_PATH + "UniqueCustomers/Outbound/" + \
            active_agent.user.username + "/"

        if not os.path.exists(file_directory):
            os.makedirs(file_directory)

        filename_prefix = file_directory + "Outbound_Unique_Customers_" + \
            str(start_date) + '-' + str(end_date) + "_" + \
            active_agent.user.username

        filename = filename_prefix + ".xls"
        easyassist_unique_customers_wb.save(filename)
        return filename
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_unique_customers_outbound %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def create_excel_easyassist_repeated_customers_outbound(cobrowse_io_objs, active_agent, repeated_primary_value, captured_data, requested_data):
    from EasyAssistApp.models import CobrowseLeadHTMLField, CobrowseCapturedLeadData
    from datetime import datetime
    filename = None
    try:
        logger.info("create_excel_easyassist_repeated_customers_outbound",
                    extra={'AppName': 'EasyAssist'})

        easyassist_repeated_customers_wb = Workbook()

        def create_new_sheet():
            global initiated_col, session_col, avg_session_col
            nonlocal wb_index
            prev_col = 0
            sheet = easyassist_repeated_customers_wb.add_sheet(
                "Outbound_Repeated_Customers-" + str(wb_index), cell_overwrite_ok=True)
            sheet.write(0, prev_col, "Customer Details")
            sheet.col(prev_col).width = 256 * 50
            initiated_col = prev_col + 1
            sheet.write(0, initiated_col, "Requests Initiated")
            sheet.col(initiated_col).width = 256 * 45
            session_col = initiated_col + 1
            sheet.write(0, session_col, "Co-browsing Sessions attended")
            sheet.col(session_col).width = 256 * 25
            avg_session_col = session_col + 1
            sheet.write(0, avg_session_col, "Average Session Duration")
            sheet.col(avg_session_col).width = 256 * 25

            return sheet

        index = 1
        wb_index = 1
        sheet = create_new_sheet()
        total_time_spent = 0
        prev_value_list = {}
        prev_value = 0
        for cobrowse_io_obj in cobrowse_io_objs.iterator():
            if index > 50000:
                wb_index += 1
                sheet = create_new_sheet()
                index = 1

            data_captured = cobrowse_io_obj.captured_lead.filter(
                agent_searched=True).first()

            if data_captured and data_captured.primary_value not in prev_value_list and repeated_primary_value.filter(primary_value=data_captured.primary_value).count():
                label_name = CobrowseLeadHTMLField.objects.filter(
                    id=data_captured.search_field_id).first().tag_label
                sheet.write(index, 0,
                            label_name + ' : ' + str(data_captured.primary_value))
                prev_value_list[data_captured.primary_value] = 1
                prev_value = data_captured.primary_value
            else:
                continue

            total_request = captured_data.filter(
                primary_value=prev_value, agent_searched=True)
            sheet.write(index, initiated_col, total_request.count())
            session_attended = cobrowse_io_objs.filter(
                session_id__in=total_request).filter(~Q(cobrowsing_start_datetime=None))
            total_time_spent = 0
            sheet.write(index, session_col, session_attended.count())
            for cust_obj in session_attended.iterator():
                total_time_spent += cust_obj.total_time_spent_seconds()
            if session_attended.count() != 0:
                avg_time_spent = total_time_spent / session_attended.count()
                sheet.write(index, avg_session_col, round(avg_time_spent, 2))
            else:
                sheet.write(index, avg_session_col, "-")

            index += 1

        try:
            start_date = datetime.strptime(
                requested_data["start_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
            end_date = datetime.strptime(
                requested_data["end_date"], "%d/%m/%Y").strftime("%Y-%m-%d")
        except Exception:
            start_date = requested_data["start_date"]
            end_date = requested_data["end_date"]

        file_directory = EASYASSISTAPP_SECURED_FILES_PATH + "RepeatedCustomers/Outbound/" + \
            active_agent.user.username + "/"

        if not os.path.exists(file_directory):
            os.makedirs(file_directory)

        filename_prefix = file_directory + "Outbound_Repeated_Customers_" + \
            str(start_date) + '-' + str(end_date) + "_" + \
            active_agent.user.username

        filename = filename_prefix + ".xls"
        easyassist_repeated_customers_wb.save(filename)
        return filename
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_repeated_customers_outbound %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def generate_jaas_jwt_token(username, user_email, user_avatar, api_key, is_moderator, api_id, private_key):
    jaasJwt = JaaSJwtBuilder()
    token = jaasJwt.withDefaults().withApiKey(api_key).withUserName(username).withUserEmail(user_email).withModerator(is_moderator).withAppID(
        api_id).withUserAvatar(user_avatar).signWith(private_key)  # Pass the private here
    return token.decode("utf-8")


def populate_captured_lead_tag_value_in_sheet(sheet, index, client_data_captured, columns):
    for captured_data in client_data_captured.iterator():
        if captured_data.primary_value:
            sheet.write(index, columns[captured_data.search_field.tag_label], captured_data.primary_value)
        else:
            sheet.write(index, columns[captured_data.search_field.tag_label], "-")


def create_excel_easyassist_capture_leads(data, cobrowse_io_objs, access_token_obj):
    filename = None
    try:
        logger.info("create_excel_easyassist_capture_leads",
                    extra={'AppName': 'EasyAssist'})

        start_date = data["start_date"]
        end_date = data["end_date"]
        email_id = access_token_obj.agent.user.username

        cobrowseio_drop_off_leads = cobrowse_io_objs.filter(
            agent=None, cobrowsing_start_datetime=None, meeting_start_datetime=None).order_by('request_datetime')

        cobrowseio_attended_leads = cobrowse_io_objs.filter(
            ~Q(cobrowsing_start_datetime=None)).filter(~Q(agent=None)).order_by('request_datetime')

        cobrowseio_not_attended_leads = cobrowse_io_objs.filter(cobrowsing_start_datetime=None).filter(
            ~Q(agent=None)).exclude(session_archived_cause__in=["FOLLOWUP", "UNASSIGNED"]).order_by('request_datetime')

        easyassist_capture_leads_wb = Workbook()
        create_excel_easyassist_capture_leads_sheet1(
            easyassist_capture_leads_wb, cobrowseio_drop_off_leads, access_token_obj)
        create_excel_easyassist_capture_leads_sheet2(
            easyassist_capture_leads_wb, cobrowseio_attended_leads, access_token_obj)
        create_excel_easyassist_capture_leads_sheet3(
            easyassist_capture_leads_wb, cobrowseio_not_attended_leads, access_token_obj)

        file_directory = EASYASSISTAPP_SECURED_FILES_PATH + "CapturedLeadsAnalytics/" + email_id + "/"

        if not os.path.exists(file_directory):
            os.makedirs(file_directory)

        filename_prefix = file_directory + "EasyAssist_CapturedLeads_" + \
            str(start_date) + ' - ' + str(end_date) + \
            '_' + email_id

        filename = filename_prefix + ".xls"
        easyassist_capture_leads_wb.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_capture_leads %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None

    return filename


def create_excel_easyassist_capture_leads_sheet1(easyassist_capture_leads_wb, cobrowseio_drop_off_leads, access_token_obj):
    try:
        sheet = easyassist_capture_leads_wb.add_sheet(
            "Drop off leads", cell_overwrite_ok=True)

        custom_cols = ["Lead generated on",
                       "Captured from (Webpage)"]
        columns = populate_columns(access_token_obj, sheet, custom_cols)

        index = 1

        for cobrowse_io_obj in cobrowseio_drop_off_leads.iterator():
            sheet.write(index, columns["Lead generated on"], str(cobrowse_io_obj.request_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            sheet.write(
                index, columns["Captured from (Webpage)"], cobrowse_io_obj.title)

            for search_field in access_token_obj.search_fields.all():
                sheet.write(
                    index, columns[search_field.tag_label], "-")

            client_data_captured = cobrowse_io_obj.captured_lead.all()

            populate_captured_lead_tag_value_in_sheet(sheet, index, client_data_captured, columns)
            
            index += 1
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_capture_leads_sheet1 %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def populate_columns(access_token_obj, sheet, custom_cols=[]):
    column_number = 0
    columns = {}

    for column in access_token_obj.search_fields.all().order_by('id'):
        columns[column.tag_label] = column_number
        sheet.write(0, column_number, column.tag_label)
        sheet.col(column_number).width = 256 * 25
        column_number += 1

    for col in custom_cols:
        columns[col] = column_number
        sheet.write(0, column_number, col)
        sheet.col(column_number).width = 256 * 25
        column_number += 1

    return columns


def create_excel_easyassist_capture_leads_sheet2(easyassist_capture_leads_wb, cobrowseio_attended_leads, access_token_obj):

    def get_invited_and_connected_agents(access_token_obj, cobrowse_io_obj):
        agents_invited_str = "-"
        agents_invited_connected_str = "-"

        if access_token_obj is not None and (access_token_obj.enable_invite_agent_in_cobrowsing or access_token_obj.enable_invite_agent_in_meeting):
            if cobrowse_io_obj.get_cobrowse_support_agent_details():
                agents_invited = cobrowse_io_obj.get_cobrowse_support_agent_details(
                ).support_agents_invited.all()
                if agents_invited:
                    agents_invited_str = ""
                    for agent in agents_invited.iterator():
                        agents_invited_str = agents_invited_str + agent.user.username + ", "
                    agents_invited_str = agents_invited_str[:-2]

                agents_invited_connected = cobrowse_io_obj.get_cobrowse_support_agent_details(
                ).support_agents_joined.all()
                if agents_invited_connected:
                    agents_invited_connected_str = ""
                    for agent in agents_invited_connected.iterator():
                        agents_invited_connected_str = agents_invited_connected_str + \
                            agent.user.username + ", "
                    agents_invited_connected_str = agents_invited_connected_str[:-2]

        return agents_invited_str, agents_invited_connected_str

    try:
        sheet = easyassist_capture_leads_wb.add_sheet(
            "Attended Session", cell_overwrite_ok=True)

        custom_cols = ["Co-browsing Start Date & Time", "Session End Date Time", "Location", "Primary Agent",
                       "Session Duration", "Lead Status", "Captured from (Webpage)",
                       "NPS", "Session ID", "Agent Remarks"]

        if access_token_obj.enable_invite_agent_in_cobrowsing or access_token_obj.enable_invite_agent_in_meeting:
            custom_cols.insert(4, "Agents Invited")
            custom_cols.insert(5, "Invited agents connected")

        columns = populate_columns(access_token_obj, sheet, custom_cols)

        index = 1

        for cobrowse_io_obj in cobrowseio_attended_leads.iterator():
            sheet.write(
                index, columns["Co-browsing Start Date & Time"], str(cobrowse_io_obj.cobrowsing_start_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            sheet.write(
                index, columns["Session End Date Time"], str(cobrowse_io_obj.last_update_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))

            customer_location = "-"

            if cobrowse_io_obj.latitude != None and cobrowse_io_obj.longitude != None and \
                    cobrowse_io_obj.latitude != "None" and cobrowse_io_obj.longitude != "None":
                try:
                    location = geocoder.google(
                        [cobrowse_io_obj.latitude, cobrowse_io_obj.longitude], method="reverse", key=GOOGLE_GEOCODER_KEY)

                    customer_location = location.address
                except Exception as e:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    logger.error("Error create_excel_easyassist_capture_leads_sheet2 %s at %s",
                                 str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

            sheet.write(
                index, columns["Location"], customer_location)

            sheet.write(
                index, columns["Primary Agent"], cobrowse_io_obj.agent.user.username)

            if access_token_obj.enable_invite_agent_in_cobrowsing or access_token_obj.enable_invite_agent_in_meeting:
                agents_invited_str, agents_invited_connected_str = "-", "-"
                try:
                    agents_invited_str, agents_invited_connected_str = get_invited_and_connected_agents(
                        access_token_obj, cobrowse_io_obj)
                except Exception as e:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    logger.error("Error create_excel_easyassist_capture_leads_sheet2 %s at %s",
                                 str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

                sheet.write(
                    index, columns["Agents Invited"], agents_invited_str)
                sheet.write(
                    index, columns["Invited agents connected"], agents_invited_connected_str)

            sheet.write(
                index, columns["Session Duration"], str(cobrowse_io_obj.total_time_spent()))

            if cobrowse_io_obj.is_archived == True and cobrowse_io_obj.is_helpful == True:
                sheet.write(index, columns["Lead Status"], "Converted")
            else:
                sheet.write(index, columns["Lead Status"], "Not Converted")

            sheet.write(
                index, columns["Captured from (Webpage)"], cobrowse_io_obj.title)

            agent_nps = '-'
            if cobrowse_io_obj.agent_rating != None:
                agent_nps = cobrowse_io_obj.agent_rating

            sheet.write(
                index, columns["NPS"], agent_nps)

            sheet.write(
                index, columns["Session ID"], str(cobrowse_io_obj.session_id))

            agent_remarks = "-"
            if cobrowse_io_obj.get_cobrowsing_session_closing_comments().first():
                agent_remarks = cobrowse_io_obj.get_cobrowsing_session_closing_comments().first().agent_comments

            sheet.write(index, columns["Agent Remarks"], agent_remarks)

            for search_field in access_token_obj.search_fields.all().iterator():
                sheet.write(
                    index, columns[search_field.tag_label], "-")

            client_data_captured = cobrowse_io_obj.captured_lead.all()
            
            populate_captured_lead_tag_value_in_sheet(sheet, index, client_data_captured, columns)

            index += 1
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_capture_leads_sheet2 %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def create_excel_easyassist_capture_leads_sheet3(easyassist_capture_leads_wb, cobrowseio_not_attended_leads, access_token_obj):
    try:
        sheet = easyassist_capture_leads_wb.add_sheet(
            "Not Attended Session", cell_overwrite_ok=True)

        custom_cols = ["Co-browsing request Date & Time",
                       "Captured from (Webpage)"]
        columns = populate_columns(access_token_obj, sheet, custom_cols)

        index = 1

        for cobrowse_io_obj in cobrowseio_not_attended_leads.iterator():
            sheet.write(
                index, columns["Co-browsing request Date & Time"], str(cobrowse_io_obj.request_datetime.astimezone(
                    pytz.timezone(settings.TIME_ZONE)).strftime(DATE_TIME_FORMAT_2)))
            sheet.write(
                index, columns["Captured from (Webpage)"], cobrowse_io_obj.title)

            for search_field in access_token_obj.search_fields.all().iterator():
                sheet.write(
                    index, columns[search_field.tag_label], "-")

            client_data_captured = cobrowse_io_obj.captured_lead.all()
            
            populate_captured_lead_tag_value_in_sheet(sheet, index, client_data_captured, columns)
            
            index += 1
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_easyassist_capture_leads_sheet3 %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def easyassist_is_active_user(user_obj, CobrowseAgent):
    try:
        cobrowse_agent = CobrowseAgent.objects.filter(user=user_obj)

        if not cobrowse_agent:
            return True

        cobrowse_agent = cobrowse_agent.first()
        return cobrowse_agent.is_account_active

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error easyassist_is_active_user %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return False


def update_unattended_lead_transfer_audit_trail(cobrowse_io, cobrowse_agent, UnattendedLeadTransferAuditTrail):
    try:
        UnattendedLeadTransferAuditTrail.objects.create(cobrowse_io=cobrowse_io,
                                                        auto_assigned_agent=cobrowse_agent)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in update_unattended_lead_transfer_audit_trail %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def readable_time_format(hours, minutes, seconds):
    time_string = ""
    try:
        if hours:
            time_string += str(hours) + " hr "

        if minutes:
            time_string += str(minutes) + " min "

        if seconds:
            time_string += str(seconds) + " sec"

        time_string = time_string.strip()
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in readable_time_format %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return time_string if time_string else ZERO_SEC


def easyassist_time_in_seconds(time_in_string):
    try:
        if "min" in time_in_string:
            time_in_string = time_in_string.split(" ")
            return int(time_in_string[0]) * 60
        elif "sec" in time_in_string:
            time_in_string = time_in_string.split(" ")
            return int(time_in_string[0])
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in easyassist_time_in_seconds %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return 0


def check_access_token(request, file_key, CobrowseAgent, CobrowsingFileAccessManagement):
    try:
        active_agent = get_active_agent_obj(request, CobrowseAgent)
        file_access_management_obj = CobrowsingFileAccessManagement.objects.filter(
            key=file_key).first()
        return active_agent.get_access_token_obj() == file_access_management_obj.access_token
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in check_access_token %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return False


def create_file_access_management_obj(CobrowsingFileAccessManagement, access_token, file_path, is_public=False):
    try:
        file_access_management_obj = CobrowsingFileAccessManagement.objects.create(
            file_path=file_path, is_public=is_public, access_token=access_token)

        return str(file_access_management_obj.key)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in create_file_access_management_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
    return None


def parse_request_io_data(cobrowse_io_objs, cobrowse_agent):
    try:
        cobrowse_io_obj_details = []
        for cobrowse_io_obj in cobrowse_io_objs:
            if cobrowse_io_obj.access_token.enable_auto_assign_unattended_lead:
                continue

            est = pytz.timezone(settings.TIME_ZONE)

            request_datetime = cobrowse_io_obj.request_datetime.astimezone(
                est).strftime("%d-%b-%Y %I:%M %p")

            product_category_title = None
            product_category_pk = None
            supported_language_title = None
            supported_language_pk = None

            if cobrowse_io_obj.product_category:
                product_category_title = cobrowse_io_obj.product_category.title
                product_category_pk = cobrowse_io_obj.product_category.pk

            if cobrowse_io_obj.supported_language:
                supported_language_title = cobrowse_io_obj.supported_language.title
                supported_language_pk = cobrowse_io_obj.supported_language.pk

            cobrowse_io_obj_details.append({
                "session_id": str(cobrowse_io_obj.session_id),
                "full_name": cobrowse_io_obj.full_name,
                "mobile_number": get_masked_data_if_hashed(
                    cobrowse_io_obj.mobile_number),
                "request_datetime": request_datetime,
                "product_url": cobrowse_io_obj.product_url(),
                "product_name": cobrowse_io_obj.product_name(),
                "product_category": product_category_title,
                "queue_time": cobrowse_io_obj.lead_in_queue_time(),
                "product_category_pk": product_category_pk,
                "supported_language": supported_language_title,
                "supported_language_pk": supported_language_pk,
                "choose_product_category": cobrowse_io_obj.access_token.choose_product_category,
                "allow_language_support": cobrowse_io_obj.access_token.allow_language_support,
                "allow_video_meeting_only": cobrowse_io_obj.access_token.allow_video_meeting_only,
                "allow_agent_to_customer_cobrowsing": cobrowse_io_obj.access_token.allow_agent_to_customer_cobrowsing,
            })

        return cobrowse_io_obj_details
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_request_io_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return []


def get_time_in_hours_mins_secs(total_seconds, time_format=False):
    time_spent = " "
    time_list = []
    if total_seconds >= 3600:
        hours = round(total_seconds) // 3600
        time_list.append(str(hours))
        time_spent += str(hours) + " hr "
        total_seconds = total_seconds - hours * 3600
    if total_seconds >= 60:
        mins = round(total_seconds) // 60
        time_list.append(str(mins))
        time_spent += str(mins) + " min "
        total_seconds = total_seconds - mins * 60
    if total_seconds < 60:
        secs = round(total_seconds)
        time_list.append(str(secs))
        time_spent += str(secs) + " sec "
    if time_format:
        return time_list
    return time_spent


def get_average_self_assign_time(total_objects, assign_objs, time_format=False):
    try:
        total_time = 0
        for assign_obj in assign_objs:
            self_assign_time = assign_obj.self_assign_time
            request_datetime = assign_obj.request_datetime
            total_time += int((self_assign_time - request_datetime).total_seconds())

        if total_time:
            average_time = round((total_time / total_objects))
            average_time = get_time_in_hours_mins_secs(average_time, time_format)
            return average_time
        if time_format:
            return "0"
        return "0 sec"
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_average_self_assign_time %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return "0 sec"


def paginate(obj_list, items_per_page, page):
    try:
        page_object_count = obj_list.count()
    except Exception:
        page_object_count = len(obj_list)

    paginator = Paginator(obj_list, items_per_page)

    try:
        obj_list = paginator.page(page)
    except PageNotAnInteger:
        obj_list = paginator.page(1)
    except EmptyPage:
        obj_list = paginator.page(paginator.num_pages)

    if page:
        start_point = items_per_page * (int(page) - 1) + 1
        end_point = min(items_per_page *
                        int(page), page_object_count)
    else:
        start_point = 1
        end_point = min(items_per_page, page_object_count)

    return page_object_count, obj_list, start_point, end_point


def get_pagination_data(page_object: Page):
    try:
        pagination_data = {
            'has_other_pages': page_object.has_other_pages(),
            'page_range': page_object.paginator.page_range.stop,
            'number': page_object.number,
            'num_pages': page_object.paginator.num_pages,
        }
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_pagination_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        pagination_data = {
            'has_other_pages': False,
        }

    return pagination_data


def get_easyassist_popup_configurations_obj(access_token_obj, EasyAssistPopupConfigurations):
    try:
        popup_configurations_obj = EasyAssistPopupConfigurations.objects.filter(
            access_token=access_token_obj).first()

        if not popup_configurations_obj:
            popup_configurations_obj = EasyAssistPopupConfigurations.objects.create(
                access_token=access_token_obj)

        return popup_configurations_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error in get_easyassist_popup_configurations_obj %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def create_notification_objects(agent, cobrowse_io_obj, product_name, NotificationManagement, custom_message=""):
    
    notification_message = "Hi, " + agent.user.username + \
        "! A customer is waiting for you to connect on " + product_name + "."

    if custom_message:
        notification_message = custom_message

    NotificationManagement.objects.create(show_notification=True,
                                          agent=agent,
                                          notification_message=notification_message,
                                          cobrowse_io=cobrowse_io_obj,
                                          product_name=product_name)


def get_product_name():
    product_name = "Cogno Cobrowse"
    cobrowse_config_obj = get_developer_console_cobrowsing_settings()
    if cobrowse_config_obj:
        product_name = cobrowse_config_obj.cobrowsing_title_text

    return product_name


def check_for_supervisor_category_language_match(active_agent, user_type, selected_supervisor_pk_list, selected_language_pk_list,
                                                 selected_product_category_pk_list, CobrowseAgent):

    language_pk_set = set()
    product_category_pk_set = set()

    access_token_obj = active_agent.get_access_token_obj()
    language_support = access_token_obj.allow_language_support
    product_support = access_token_obj.choose_product_category or access_token_obj.enable_tag_based_assignment_for_outbound

    if not language_support and not product_support:
        return {}

    if language_support:
        for language_pk in selected_language_pk_list:
            language_pk_set.add(int(language_pk))

    if product_support:
        for product_category_pk in selected_product_category_pk_list:
            product_category_pk_set.add(int(product_category_pk))

    for current_supervisor_pk in selected_supervisor_pk_list:
        current_supervisor_pk = int(current_supervisor_pk)
        if current_supervisor_pk != -1:
            current_supervisor = CobrowseAgent.objects.get(
                pk=current_supervisor_pk)
            supervisor_supported_language_set = set(current_supervisor.supported_language.filter(
                is_deleted=False).values_list("pk", flat=True))
            supervisor_product_category_set = set(current_supervisor.product_category.filter(
                is_deleted=False).values_list("pk", flat=True))
            if user_type == "admin_ally":
                if language_support and not supervisor_supported_language_set.issubset(language_pk_set):
                    return {"matched_error": "language", "supervisor": current_supervisor.user.username}
                if product_support and not supervisor_product_category_set.issubset(product_category_pk_set):
                    return {"matched_error": "product", "supervisor": current_supervisor.user.username}
            else:
                if language_support and not language_pk_set.issubset(supervisor_supported_language_set):
                    return {"matched_error": "language", "supervisor": current_supervisor.user.username}
                if product_support and not product_category_pk_set.issubset(supervisor_product_category_set):
                    return {"matched_error": "product", "supervisor": current_supervisor.user.username}

    return {}


def get_total_self_assign_time(assign_objs):
    try:
        total_time = 0
        for assign_obj in assign_objs:
            self_assign_time = assign_obj.self_assign_time
            request_datetime = assign_obj.request_datetime
            total_time += int((self_assign_time -
                              request_datetime).total_seconds())

        return total_time
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_total_self_assign_time %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return 0


def check_for_details_match(active_agent, user_type, agent_email, selected_language_pk_list,
                            selected_product_category_pk_list, CobrowseAgent, User):

    language_pk_set = set()
    product_category_pk_set = set()
    message = ""
    access_token_obj = active_agent.get_access_token_obj()
    language_support = access_token_obj.allow_language_support
    product_support = access_token_obj.choose_product_category or access_token_obj.enable_tag_based_assignment_for_outbound

    if not language_support and not product_support:
        return message

    if language_support:
        for language_pk in selected_language_pk_list:
            language_pk_set.add(int(language_pk))

    if product_support:
        for product_category_pk in selected_product_category_pk_list:
            product_category_pk_set.add(int(product_category_pk))

    if user_type == "admin_ally":
        admin_ally_user = User.objects.get(username=agent_email)
        admin_ally = CobrowseAgent.objects.get(user=admin_ally_user)
        admin_ally_supervisors = get_list_supervisor_under_admin(
            admin_ally, None)
        for supervisor in admin_ally_supervisors:
            supervisor_supported_language_set = set(supervisor.supported_language.filter(
                is_deleted=False).values_list("pk", flat=True))
            supervisor_product_category_set = set(supervisor.product_category.filter(
                is_deleted=False).values_list("pk", flat=True))
            if product_support and not supervisor_product_category_set.issubset(product_category_pk_set):
                message = "The language/categories you have deselected are present with \
                    some supervisors in your team. Please remove those categories/languages from their account and try again."
            if language_support and not supervisor_supported_language_set.issubset(language_pk_set):
                message = "The language/categories you have deselected are present with \
                    some supervisors in your team. Please remove those categories/languages from their account and try again."
    elif active_agent.role == "admin" and user_type == "supervisor":
        supervisor_obj = get_cobrowse_agent_from_username(agent_email, User, CobrowseAgent)
        admin_allies = get_adminally_from_supervsior(supervisor_obj, CobrowseAgent)
        for admin_ally in admin_allies:
            admin_ally_supported_language_set = set(admin_ally.supported_language.filter(
                is_deleted=False).values_list("pk", flat=True))
            admin_ally_product_category_set = set(admin_ally.product_category.filter(
                is_deleted=False).values_list("pk", flat=True))
            if product_support and not product_category_pk_set.issubset(admin_ally_product_category_set):
                message = "The language/categories you have deselected are present with \
                    some admin allies in your team. Please remove those categories/languages from their account and try again."
            if language_support and not language_pk_set.issubset(admin_ally_supported_language_set):
                message = "The language/categories you have deselected are present with \
                    some admin allies in your team. Please remove those categories/languages from their account and try again."
    else:
        supervisor_obj = get_cobrowse_agent_from_username(
            agent_email, User, CobrowseAgent)
        supervisors_agent = get_list_agents_under_admin(
            supervisor_obj, is_active=None, is_account_active=None)
        for agent in supervisors_agent:
            supervisor_supported_language_set = set(agent.supported_language.filter(
                is_deleted=False).values_list("pk", flat=True))
            supervisor_product_category_set = set(agent.product_category.filter(
                is_deleted=False).values_list("pk", flat=True))

            if product_support and not supervisor_product_category_set.issubset(product_category_pk_set):
                message = "The language/categories you have deselected are present with \
                    some agents in your team. Please remove those categories/languages from their account and try again."
            if language_support and not supervisor_supported_language_set.issubset(language_pk_set):
                message = "The language/categories you have deselected are present with \
                    some agents in your team. Please remove those categories/languages from their account and try again."

    return message


def get_agents_under_cobrowse_agent(cobrowse_agent):
    try:
        agents = []
        if cobrowse_agent.role == "supervisor" or (cobrowse_agent.role == "agent" and cobrowse_agent.is_switch_allowed):
            agents = list(cobrowse_agent.agents.all())
        elif cobrowse_agent.role in ["admin", "admin_ally"] or (cobrowse_agent.role == "agent" and cobrowse_agent.is_switch_allowed):
            agents = get_list_agents_under_admin(
                cobrowse_agent, is_active=None, is_account_active=None)
        else:
            agents = []

        if cobrowse_agent.is_switch_allowed and (cobrowse_agent not in agents):
            agents += [cobrowse_agent]

        return agents
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_agents_under_cobrowse_agent: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return []


def validate_canned_keyword(keyword):
    try:
        reg = r'^[a-z0-9]+$'

        if re.match(reg, keyword):
            return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside validate_canned_keyword: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return False


def validate_canned_response(response):
    try:
        reg = r'^[a-zA-Z @.?!,0-9]+$'

        if re.match(reg, response.lower()):
            return True

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside validate_canned_response: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return False


def get_supervisors_of_active_agent(active_agent, CobrowseAgent):

    try:
        if active_agent.is_switch_allowed and active_agent.role == "admin":
            return active_agent

        parent_user = CobrowseAgent.objects.filter(
            agents__pk=active_agent.pk, role="supervisor", is_account_active=True)

        return list(parent_user)

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_supervisors_of_active_agent %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def get_canned_responses_data_dump(cobrowse_agent, canned_responses_obj):
    try:
        logger.info("Inside get_canned_responses_data_dump",
                    extra={'AppName': 'EasyAssist'})
        file_directory = settings.SECURE_MEDIA_ROOT + \
            "EasyAssistApp/CannedResponses/" + \
            str(cobrowse_agent.user.username)
  
        if not path.exists(file_directory):
            os.makedirs(file_directory)

        current_date_obj = datetime.date.today()
        current_date = current_date_obj.strftime("%d-%m-%Y")

        file_path = file_directory + "/Canned_responses_export_" + \
            str(current_date) + ".xls"
        relative_file_path = "/secured_files/EasyAssistApp/CannedResponses/" + \
            str(cobrowse_agent.user.username) + \
            "/Canned_responses_export_" + \
            str(current_date) + ".xls"

        canned_workbook = Workbook(style_compression=2)

        add_canned_response_in_excel_sheet(
            canned_workbook, canned_responses_obj)

        canned_workbook.save(file_path)
        return relative_file_path
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_canned_responses_data_dump %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


def add_canned_response_in_excel_sheet(canned_workbook, canned_responses_obj):
    try:
        def create_new_sheet():
            nonlocal wb_index
            sheet = canned_workbook.add_sheet(
                "Canned Responses-" + str(wb_index))

            style = easyxf(HEADING_FONT_STYLE)

            sheet.write(0, 0, "Sl No.", style=style)
            sheet.col(0).width = 256 * 6
            sheet.write(0, 1, "Keyword", style=style)
            sheet.col(1).width = 256 * 20
            sheet.write(0, 2, "Response", style=style)
            sheet.col(2).width = 256 * 20
            sheet.write(0, 3, "Added By", style=style)
            sheet.col(3).width = 256 * 20
            sheet.write(0, 4, "User Email ID", style=style)
            sheet.col(4).width = 256 * 20

            return sheet

        wb_index = 1
        sheet = create_new_sheet()
        index = 1
        for canned_response in canned_responses_obj.iterator():
            if index > 50000:
                wb_index += 1
                sheet = create_new_sheet()
                index = 1
            sheet.write(index, 0, str(index))
            sheet.write(index, 1, canned_response.keyword)
            sheet.write(index, 2, canned_response.response)
            sheet.write(index, 3, canned_response.agent.role.capitalize())
            sheet.write(index, 4, canned_response.agent.user.username)
            index += 1

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error add_canned_response_in_excel_sheet %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_canned_response_list(agent, LiveChatCannedResponse, CobrowseAgent):
    try:
        canned_response_obj = LiveChatCannedResponse.objects.filter(
            access_token=agent.get_access_token_obj(), is_deleted=False)
        return canned_response_obj
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_canned_response_list %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return []


def create_excel_wrong_canned_response_data(wrong_data_list):
    filename = None
    try:
        logger.info("In create_excel_wrong_canned_response_data",
                    extra={'AppName': 'EasyAssist'})

        wrong_data_workbook = Workbook()
        wrong_data_sheet = wrong_data_workbook.add_sheet(
            "Wrong Information Sheet")

        style = XFStyle()

        font = Font()
        font.bold = True
        style.font = font

        wrong_data_sheet.write(0, 0, "Sl. No.", style=style)
        wrong_data_sheet.col(0).width = 256 * 10
        wrong_data_sheet.write(0, 1, "Keyword", style=style)
        wrong_data_sheet.col(1).width = 256 * 10
        wrong_data_sheet.write(0, 2, "Responses", style=style)
        wrong_data_sheet.col(2).width = 256 * 100
        wrong_data_sheet.write(0, 3, "Error Message", style=style)
        wrong_data_sheet.col(3).width = 256 * 60

        row_number = 1
        for wrong_data in wrong_data_list:
            data_list = []
            data_list.append(wrong_data["row_num"])
            data_list.append(wrong_data["keyword"])
            data_list.append(wrong_data["response"])
            data_list.append(wrong_data["detail"])
            col_number = 0
            for data in data_list:
                wrong_data_sheet.write(row_number, col_number, data)
                col_number += 1
            row_number += 1

        folder_path = EASYASSISTAPP_SECURED_FILES_PATH + "EasyassistCannedResponseData/"

        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        filename = folder_path + "canned_responses_error_file" + \
            ".xls"
        wrong_data_workbook.save(filename)
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_excel_wrong_canned_response_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        filename = None
    return filename


def get_supervisors_list_under_admin(cobrowse_agent):
    try:
        supervisors_obj = get_list_supervisor_under_admin(
            admin_user=cobrowse_agent, is_active=None, is_account_active=None)
        admin_allies = get_list_admin_ally(
            cobrowse_agent, is_active=None, is_account_active=None)
        for admin_ally in admin_allies:
            supervisors_obj += get_list_supervisor_under_admin(
                admin_user=admin_ally, is_active=None, is_account_active=None)
        return supervisors_obj

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Inside get_supervisors_list_under_admin: %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return []


def get_adminally_from_supervsior(active_agent, CobrowseAgent):
    try:
        if active_agent.is_switch_allowed and active_agent.role == "admin":
            return active_agent

        parent_user = CobrowseAgent.objects.filter(
            agents__pk=active_agent.pk, role="admin_ally", is_account_active=True)
        return parent_user
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_adminally_from_supervsior %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def get_cobrowse_agent_from_username(user_email, User, CobrowseAgent):
    try:
        user_obj = User.objects.get(username=user_email)
        cobrowse_agent_obj = CobrowseAgent.objects.get(user=user_obj)
        return cobrowse_agent_obj

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error get_cobrowse_agent_from_username %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

    return None


def parse_canned_response_data(canned_response_list):
    try:
        canned_response_data = []
        for canned_response in canned_response_list:
            visibility = "Everyone"

            if canned_response.agent.role == "supervisor":
                visibility = "Team Only"

            canned_response_data.append({
                "pk": str(canned_response.pk),
                "keyword": canned_response.keyword,
                "response": canned_response.response,
                "added_by": canned_response.agent.user.username,
                "visibility": visibility,
                "agent_role": canned_response.agent.role,
            })

        return canned_response_data
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error parse_canned_response_data %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})
        return []


def update_virtual_agent_code(cobrowse_io):
    try:
        cobrowse_agent = cobrowse_io.agent
        cobrowse_access_token = cobrowse_io.access_token
        if cobrowse_io.is_lead or cobrowse_io.is_droplink_lead or cobrowse_io.is_reverse_cobrowsing or not cobrowse_agent:
            return
        if cobrowse_access_token.allow_connect_with_virtual_agent_code:
            uuid_str = str(uuid.uuid4())
            new_virtual_agent_code = uuid_str[-7:]
            if new_virtual_agent_code.find("-") != -1:
                new_virtual_agent_code = new_virtual_agent_code.replace("-", "a")
            
            cobrowse_agent.virtual_agent_code = new_virtual_agent_code
            cobrowse_agent.save(update_fields=["virtual_agent_code"])
        
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("An error occurred while update_virtual_agent_code %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_encrypted_response_packet(response, custom_encrypt_obj):
    response = json.dumps(response)
    encrypted_response = custom_encrypt_obj.encrypt(response)
    response = {"Response": encrypted_response}
    return response


def update_agents_supervisors_creation_count(active_agent, agent_role):

    if agent_role == "agent":
        active_agent.total_agents_created += 1
    if agent_role == "supervisor":
        active_agent.total_supervisors_created += 1

    active_agent.save()


def parse_cobrowse_agent_details(agents, access_token_obj):
    try:
        agent_details_list = []
        for agent in agents:
            agent_details_list.append({
                "pk": agent.pk,
                "is_account_active": agent.is_account_active,
                "first_name": agent.user.first_name,
                "email": agent.user.email,
                "mobile_number": agent.mobile_number,
                "product_categories": agent.get_product_categories(),
                "product_categories_limited": agent.get_product_categories_limited(),
                "supported_languages": agent.get_supported_languages(),
                "supported_languages_limited": agent.get_supported_languages_limited(),
                "is_active": agent.is_active,
                "is_online": agent.user.is_online,
                "allow_agent_to_customer_cobrowsing": access_token_obj.allow_agent_to_customer_cobrowsing,
                "is_cobrowsing_active": agent.is_cobrowsing_active,
                "is_cognomeet_active": agent.is_cognomeet_active,
                "supervisors": agent.get_supervisors(),
                "choose_product_category": access_token_obj.choose_product_category,
                "enable_tag_based_assignment_for_outbound": access_token_obj.enable_tag_based_assignment_for_outbound,
                "allow_language_support": access_token_obj.allow_language_support,
                "assign_followup_leads": agent.assign_followup_leads,
                "support_level": agent.support_level,
            })

        return agent_details_list
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("An error occurred while parse_cobrowse_agent_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return []


def get_encrypted_no_data_response_packet(custom_encrypt_obj):
    response = {}
    response["status"] = 302
    response["message"] = NO_DATA_FOUND
    return get_encrypted_response_packet(response, custom_encrypt_obj)


def check_user_online_status(cobrowse_agent, UserSession, Session, SecuredLogin, CobrowsingAuditTrail):
    from EasyChatApp.utils import save_audit_trail as save_audit_trail_easychat
    try:
        if not cobrowse_agent.is_cobrowsing_active and not cobrowse_agent.is_cognomeet_active:
            user_obj = cobrowse_agent.user
            if not user_obj.is_online:
                return False
            user_session_obj = UserSession.objects.filter(
                user=user_obj).first()
            if user_session_obj:
                time_zone = pytz.timezone(settings.TIME_ZONE)
                last_updated_time = user_session_obj.last_update_datetime
                if not last_updated_time:
                    return False
                last_updated_time = last_updated_time.astimezone(time_zone)
                current_time = datetime.datetime.now().astimezone(time_zone)
                inactive_time = (
                    current_time - last_updated_time).total_seconds()
                if inactive_time > 75:
                    session_objs = Session.objects.filter(
                        pk=user_session_obj.session_key)
                    user_session_obj.delete()
                    if session_objs:
                        session_objs.delete()
                    user_obj.is_online = False
                    user_obj.save(update_fields=["is_online"])
                    secured_login = SecuredLogin.objects.filter(
                        user=user_obj).first()
                    secured_login.failed_attempts = 0
                    secured_login.save(update_fields=["failed_attempts"])

                    audit_trail_data = json.dumps({
                        "user_id": user_obj.pk
                    })
                    save_audit_trail_easychat(user_obj, "7", audit_trail_data)

                    cobrowse_agent.is_active = False
                    cobrowse_agent.is_agent_connected = False
                    cobrowse_agent.save(
                        update_fields=["is_active", "is_agent_connected"])

                    save_audit_trail(cobrowse_agent, "2",
                                     "Logout from System", CobrowsingAuditTrail)
                    return False
                else:
                    return True
        return True
    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("An error occurred in check_user_online_status: %s at %s", e, str(
            exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})

        return False


def create_droplink_customer_details(analytics_workbook, drop_link_objs):
    try:
        logger.info("create_droplink_customer_details",
                    extra={'AppName': 'EasyAssist'})

        def create_new_sheet():
            global supervisors_col, account_status_col, last_active_date_col, creation_date_col,\
                deactivation_date_col, activation_date_col
            nonlocal wb_index
            sheet = analytics_workbook.add_sheet("Links Generated - " + str(wb_index))

            style = easyxf(HEADING_FONT_STYLE)

            sheet.write(0, 0, "Sl No.", style=style)
            sheet.col(0).width = 256 * 6
            sheet.write(0, 1, "Agent Name", style=style)
            sheet.col(1).width = 256 * 30
            sheet.write(0, 2, "Agent Email ID", style=style)
            sheet.col(2).width = 256 * 30
            sheet.write(0, 3, "Link Generated Date & Time", style=style)
            sheet.col(3).width = 256 * 25
            sheet.write(0, 4, "Webpage Link", style=style)
            sheet.col(4).width = 256 * 30
            sheet.write(0, 5, "Customer Name", style=style)
            sheet.col(5).width = 256 * 20
            sheet.write(0, 6, "Customer Mobile No.", style=style)
            sheet.col(6).width = 256 * 20
            sheet.write(0, 7, "Customer Email ID", style=style)
            sheet.col(7).width = 256 * 30

            return sheet

        index = 1
        wb_index = 1
        sheet = create_new_sheet()
        for drop_link_obj in drop_link_objs.iterator():
            if index > 50000:
                wb_index += 1
                index = 1
                sheet = create_new_sheet()

            sheet.write(index, 0, index)
            sheet.write(index, 1, drop_link_obj.agent.agent_full_name())
            sheet.write(index, 2, drop_link_obj.agent.user.username)
            sheet.write(index, 3, str(drop_link_obj.generate_datetime.astimezone(
                pytz.timezone(settings.TIME_ZONE)).strftime("%d-%m-%Y %I:%M %p")))
            
            if len(drop_link_obj.client_page_link) < 255:
                sheet.write(index, 4, Formula('HYPERLINK("{}")'.format(drop_link_obj.client_page_link)), style=easyxf('font: colour blue; align: vert top;'))
            else:
                sheet.write(index, 4, drop_link_obj.client_page_link)

            sheet.write(index, 5, drop_link_obj.customer_name)
            sheet.write(index, 6, drop_link_obj.customer_mobile)
            if drop_link_obj.customer_email:
                sheet.write(index, 7, drop_link_obj.customer_email)
            else:
                sheet.write(index, 7, "-")
            
            index += 1

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("Error create_droplink_customer_details %s at %s",
                     str(e), str(exc_tb.tb_lineno), extra={'AppName': 'EasyAssist'})


def get_generated_link_data_dump(requested_data, active_agent, CobrowseDropLink, drop_link_objs=None):
    try:
        logger.info("Inside get_generated_link_data_dump",
                    extra={'AppName': 'EasyAssist'})
        start_date = requested_data["startdate"]
        end_date = requested_data["enddate"]
        try:
            start_date_obj = datetime.datetime.strptime(start_date, DATE_TIME_FORMAT).date()
            end_date_obj = datetime.datetime.strptime(end_date, DATE_TIME_FORMAT).date()
        except:
            start_date_obj = datetime.datetime.strptime(start_date, "%d-%m-%Y").date()
            end_date_obj = datetime.datetime.strptime(end_date, "%d-%m-%Y").date()

        start_date_string = start_date_obj.strftime("%d-%m-%Y")
        end_date_string = end_date_obj.strftime("%d-%m-%Y")

        file_directory = settings.SECURE_MEDIA_ROOT + \
            "EasyAssistApp/InboundAnalytics/" + \
            str(active_agent.user.username)

        if not path.exists(file_directory):
            os.makedirs(file_directory)

        file_path = file_directory + "/GDL_Links_Generated_" + \
            str(active_agent.user.username) + "_" + str(start_date_string) + "_" + str(end_date_string) + ".xls"

        relative_file_path = "/secured_files/EasyAssistApp/InboundAnalytics/" + \
            str(active_agent.user.username) + \
            "/GDL_Links_Generated_" + str(active_agent.user.username) + "_" + str(start_date_string) + "_" + str(end_date_string) + ".xls"

        yesterdays_date_obj = (datetime.datetime.now() - datetime.timedelta(days=1)).date()

        if path.exists(file_path) and end_date_obj < yesterdays_date_obj:
            return relative_file_path
        if not drop_link_objs:
            if active_agent.role in ["admin", "admin_ally"]:
                agent_objs = get_list_agents_under_admin(
                    active_agent, is_active=None, is_account_active=None)
            elif active_agent.role == "supervisor":
                agent_objs = list(active_agent.agents.all())

            go_live_date = active_agent.get_access_token_obj().go_live_date
            drop_link_objs = CobrowseDropLink.objects.filter(agent__in=agent_objs, generate_datetime__date__gte=go_live_date).filter(
                generate_datetime__date__gte=start_date, generate_datetime__date__lte=end_date, proxy_cobrowse_io=None).order_by("generate_datetime")

            if not drop_link_objs:
                return NO_DATA

        analytics_workbook = Workbook(style_compression=2)

        create_droplink_customer_details(
            analytics_workbook, drop_link_objs)

        analytics_workbook.save(file_path)
        return relative_file_path

    except Exception as e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        logger.error("get_generated_link_data_dump %s %s", str(e), str(exc_tb.tb_lineno), extra={
                     'AppName': 'EasyAssist'})

        return "None"


"""
function: convert_time_to_24_hrs
input time_obj: Time in 12 hours format
expected output:
    Provides time in 24 hours format
"""


def convert_time_to_24_hrs(time_obj):
    if time_obj:
        raw_time = datetime.datetime.strptime(time_obj, "%I:%M %p")
        converted_time = datetime.datetime.strftime(raw_time, "%H:%M")
        return converted_time
    else:
        return time_obj
